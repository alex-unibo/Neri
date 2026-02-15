import sys
import json
import locale
import os
from glob import glob
from collections import defaultdict, Counter
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from functools import partial  # Per il metodo partial
from PyQt6.QtWidgets import QSpinBox  # Per il widget SpinBox
from datetime import datetime, timedelta
from collections import defaultdict
from PyQt6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QGridLayout, 
                           QLabel, QPushButton, QFrame, QScrollArea, QDialog,
                           QApplication, QMessageBox, QToolTip, QSizePolicy,
                           QHeaderView, QTableWidget, QTableWidgetItem)
from PyQt6.QtCore import Qt, QDate, pyqtSignal, QTimer, QRect
from PyQt6.QtGui import QFont, QPalette, QColor, QPainter, QLinearGradient

from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                           QHBoxLayout, QGridLayout, QLabel, QLineEdit, 
                           QPushButton, QCheckBox, QRadioButton, QButtonGroup,
                           QTabWidget, QScrollArea, QFrame, QMessageBox,
                           QTreeWidget, QTreeWidgetItem, QDialog, QDateEdit,
                           QGroupBox, QSplitter, QTextEdit, QComboBox,
                           QProgressBar, QSpacerItem, QSizePolicy, QToolBar,
                           QHeaderView, QTableWidget, QTableWidgetItem)
from PyQt6.QtCore import Qt, QDate, pyqtSignal, QSize
from PyQt6.QtGui import QFont, QPalette, QColor, QIcon, QAction
import matplotlib.pyplot as plt
import time
from PyQt6.QtCore import QThread, pyqtSignal, QTimer, QRect
from PyQt6.QtGui import QPixmap, QPainter, QLinearGradient, QPen
from PyQt6.QtWidgets import QSplashScreen, QProgressBar
# Aggiungi queste righe dopo gli import esistenti
from datetime import datetime, timedelta
from collections import defaultdict
# ========== AGGIUNGI QUESTE IMPORT ALL'INIZIO DEL FILE ==========
from PyQt6.QtCore import Qt, QTimer
from PyQt6.QtWidgets import QLineEdit, QHBoxLayout, QLabel, QFrame
from PyQt6.QtGui import QShortcut, QKeySequence
# Import necessari all'inizio
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
archivio_dialogs_aperti = []
# ========== GESTORE PERCORSI MULTIPLI ==========
class GestorePercorsi:
    """Gestisce automaticamente i percorsi su diversi PC aziendali"""
    
    def __init__(self):
        self.percorso_base = None
        self.percorsi_file = {}
        self.inizializza_percorsi()
        self.menu_extra_widgets = {}  # Per salvare i SpinBox dei menu extra
        self.evento_in_fiera = False
    
    def inizializza_percorsi(self):
        """Trova automaticamente il percorso corretto e configura tutti i percorsi"""
        
        # Lista dei percorsi base possibili
        percorsi_possibili = [
            "./",                                    # Cartella corrente (fallback)

            r"G:\.shortcut-targets-by-id\1tG5T-Pcv2j4b_f9flOCvQ3HwYwHyB_j5\ordini_drive",
            r"H:\.shortcut-targets-by-id\1tG5T-Pcv2j4b_f9flOCvQ3HwYwHyB_j5\ordini_drive",
            r"F:\.shortcut-targets-by-id\1tG5T-Pcv2j4b_f9flOCvQ3HwYwHyB_j5\ordini_drive",
            r"C:\.shortcut-targets-by-id\1tG5T-Pcv2j4b_f9flOCvQ3HwYwHyB_j5\ordini_drive",
            r"E:\.shortcut-targets-by-id\1tG5T-Pcv2j4b_f9flOCvQ3HwYwHyB_j5\ordini_drive",
            r"D:\.shortcut-targets-by-id\1tG5T-Pcv2j4b_f9flOCvQ3HwYwHyB_j5\ordini_drive",
            r"C:\Users\Alex\OneDrive - Alma Mater Studiorum Università di Bologna\Documenti\NeriCatering\dist", # Desktop utente
            r"C:\Users\Utente\Desktop\Condivisione", # Desktop utente (EN)
            r"C:\Condivisione",                      # C: drive alternativo
            r"C:\@condivisione\Condivisione",        # Sottocartella

            "./",                                    # Cartella corrente (fallback)
        ]
        
        print("🔍 Ricerca percorso catering su questo PC...")
        
        for percorso in percorsi_possibili:
            if os.path.exists(percorso):
                # Controlla se contiene i file necessari
                if self.verifica_cartella_catering(percorso):
                    self.percorso_base = percorso
                    print(f"✅ TROVATO percorso catering: {percorso}")
                    break
        
        if not self.percorso_base:
            # Fallback: usa cartella corrente
            self.percorso_base = "./"
            print(f"⚠️ Uso cartella corrente come fallback")
        
        # Configura tutti i percorsi specifici
        self.configura_percorsi_file()
        
        # Crea cartelle mancanti
        self.crea_cartelle_necessarie()
    
    def verifica_cartella_catering(self, percorso):
        """Verifica se la cartella contiene i file necessari per il catering"""
        file_necessari = [
            "catering_data.json",
            "referenze_cibi.json"
        ]
        
        for file_name in file_necessari:
            file_path = os.path.join(percorso, file_name)
            if not os.path.exists(file_path):
                return False
        
        return True
    
    def configura_percorsi_file(self):
        """Configura tutti i percorsi specifici basandosi sul percorso base"""
        base = self.percorso_base
        
        self.percorsi_file = {
            "catering_data": os.path.join(base, "catering_data.json"),
            "referenze_cibi": os.path.join(base,"referenze_cibi.json"),
            "ordini_json": os.path.join(base, "archivio_ordini", "ordini.json"),
            "ordini_docx": os.path.join(base, "ordini_docx"),
            "riepilogo_cucina": os.path.join(base, "riepilogo_cucina"),
            "archivio_ordini": os.path.join(base, "archivio_ordini"),
            "backup": os.path.join(base, "backup"),
            "logo": os.path.join(base, "neri.png"),
        }
    
    def crea_cartelle_necessarie(self):
        """Crea le cartelle necessarie se non esistono"""
        cartelle_da_creare = [
            self.percorsi_file["ordini_docx"],
            self.percorsi_file["riepilogo_cucina"], 
            self.percorsi_file["archivio_ordini"],
            self.percorsi_file["backup"],
        ]
        
        for cartella in cartelle_da_creare:
            if not os.path.exists(cartella):
                try:
                    os.makedirs(cartella, exist_ok=True)
                except:
                    pass
    
    def get_percorso(self, nome_file):
        """Restituisce il percorso completo per un file specifico"""
        return self.percorsi_file.get(nome_file, os.path.join(self.percorso_base, nome_file))

# Inizializza il gestore percorsi
gestore_percorsi = GestorePercorsi()

# ========== AGGIUNGI QUESTA CLASSE DOPO LE ALTRE CLASSI CUSTOM ==========
class SearchBar(QFrame):
    """Barra di ricerca elegante per filtrare le referenze"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setup_ui()
        self.is_visible = False
        
    def setup_ui(self):
        self.setStyleSheet("""
            QFrame {
                background-color: #f8f9fa;
                border: 2px solid #8b5cf6;
                border-radius: 8px;
                padding: 8px;
                margin: 5px;
            }
        """)
        
        layout = QHBoxLayout(self)
        layout.setContentsMargins(10, 8, 10, 8)
        layout.setSpacing(10)
        
        # Icona ricerca
        search_icon = QLabel("🔍")
        search_icon.setStyleSheet("font-size: 16px; color: #8b5cf6;")
        layout.addWidget(search_icon)
        
        # Campo di ricerca
        self.search_field = QLineEdit()
        self.search_field.setPlaceholderText("Cerca referenze... (Ctrl+F per aprire, Esc per chiudere)")
        self.search_field.setStyleSheet("""
            QLineEdit {
                border: 1px solid #e5e7eb;
                border-radius: 4px;
                padding: 8px 12px;
                font-size: 13px;
                background-color: white;
            }
            QLineEdit:focus {
                border-color: #8b5cf6;
                outline: none;
            }
        """)
        layout.addWidget(self.search_field)
        
        # Label risultati
        self.results_label = QLabel("")
        self.results_label.setStyleSheet("color: #6b7280; font-size: 12px;")
        layout.addWidget(self.results_label)
        
        # Inizialmente nascosta
        self.hide()

class LoadingWorker(QThread):
    finished = pyqtSignal()
    
    def run(self):
        time.sleep(1.0)  # 1 secondo
        self.finished.emit()

class CleanLogoSplash(QSplashScreen):
    def __init__(self):
        pixmap = self.load_logo()
        super().__init__(pixmap)
        
        self.setWindowFlags(Qt.WindowType.WindowStaysOnTopHint | 
                           Qt.WindowType.FramelessWindowHint)
        
        self.worker = LoadingWorker()
        self.worker.finished.connect(self.loading_finished)
        
        QTimer.singleShot(100, self.start_loading)
    
    
    
    def load_logo(self):
        """Carica icona.png con debug completo"""
        
        # DEBUG: Controlla se il file esiste
        icon_path = gestore_percorsi.get_percorso("logo")
        if not os.path.exists(icon_path):
            print(f"❌ ERRORE: {icon_path} non trovato nella cartella corrente")
            print(f"📁 Cartella corrente: {os.getcwd()}")
            print(f"📋 File nella cartella: {os.listdir('.')}")
            return self.create_fallback_logo()
        
        try:
            # Carica l'icona
            original_pixmap = QPixmap(icon_path)
            
            if original_pixmap.isNull():
                print(f"❌ ERRORE: {icon_path} esiste ma non è un'immagine valida")
                return self.create_fallback_logo()
            
            print(f"✅ Icona caricata correttamente: {original_pixmap.width()}x{original_pixmap.height()}")
            
            # Ridimensiona mantenendo proporzioni
            target_size = 200
            scaled_pixmap = original_pixmap.scaled(
                target_size, target_size,
                Qt.AspectRatioMode.KeepAspectRatio,
                Qt.TransformationMode.SmoothTransformation
            )
            
            print(f"✅ Icona ridimensionata: {scaled_pixmap.width()}x{scaled_pixmap.height()}")
            return scaled_pixmap
            
        except Exception as e:
            print(f"❌ ERRORE caricamento icona: {e}")
            return self.create_fallback_logo()
    
    
    def create_fallback_logo(self):
        """Logo di fallback se icona.png non funziona"""
        print("🔄 Creazione logo di fallback...")
        
        size = 150
        pixmap = QPixmap(size, size)
        pixmap.fill(QColor(255, 255, 255, 0))  # Trasparente
        
        painter = QPainter(pixmap)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        # Cerchio viola
        painter.setPen(QPen(QColor(139, 92, 246), 6))
        painter.setBrush(QColor(255, 255, 255))
        painter.drawEllipse(15, 15, 120, 120)
        
        # Lettera "N" al centro
        painter.setPen(QColor(139, 92, 246))
        font = QFont("Arial", 36, QFont.Weight.Bold)
        painter.setFont(font)
        painter.drawText(60, 90, "N")
        
        painter.end()
        print("✅ Logo di fallback creato")
        return pixmap
    
    def start_loading(self):
        self.worker.start()
    
    def loading_finished(self):
        QTimer.singleShot(100, self.close)
        
        
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure

# Impostazioni locali
try:
    locale.setlocale(locale.LC_TIME, 'italian')
except:
    try:
        locale.setlocale(locale.LC_TIME, 'it_IT.UTF-8')
    except:
        pass

ORDINI_DIR = gestore_percorsi.get_percorso("archivio_ordini")
ORDINI_FILE = os.path.join(ORDINI_DIR, "ordini.json")

# Stile professionale business-like
# Sostituisci tutto il blocco PROFESSIONAL_STYLE con questo:
PROFESSIONAL_STYLE = """
QMainWindow {
    background-color: #f8f9fa;
    color: #212529;
}

QWidget {
    background-color: #f8f9fa;
    color: #212529;
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 13px;
}

/* Header principale */
QLabel#appTitle {
    background-color: #8b5cf6;
    color: white;
    font-family: 'Acquatype', 'Arial Black', 'Impact', sans-serif;
    font-size: 36px;
    font-weight: bold;
    padding: 30px;
    border-bottom: 4px solid #7c3aed;
    letter-spacing: 2px;
    text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
}

QLabel#sectionHeader {
    background-color: #3b82f6;  /* BLU ACCESO per nomi servizi */
    color: white;
    font-size: 16px;
    font-weight: bold;
    padding: 12px 15px;
    border: 1px solid #3b82f6;
    border-radius: 6px;
    margin: 10px 0;
}

/* Campi di input */
QLineEdit {
    background-color: white;
    border: 2px solid #ddd;
    border-radius: 6px;
    padding: 8px 12px;
    font-size: 13px;
    color: #333;
}

QLineEdit:focus {
    border-color: #8b5cf6;
    outline: none;
}

QDateEdit {
    background-color: white;
    border: 2px solid #ddd;
    border-radius: 6px;
    padding: 8px 12px;
    font-size: 13px;
    color: #333;
}

QDateEdit:focus {
    border-color: #8b5cf6;
}

/* Pulsanti - Default viola */
QPushButton {
    background-color: #8b5cf6;  /* Viola acceso default */
    color: white;
    border: none;
    border-radius: 6px;
    padding: 10px 20px;
    font-size: 13px;
    font-weight: 600;
    min-width: 120px;
}

QPushButton:hover {
    background-color: #7c3aed;  /* Viola più scuro al hover */
}

QPushButton:pressed {
    background-color: #6d28d9;  /* Viola ancora più scuro quando premuto */
}

/* Pulsante GENERA ORDINE - VERDE */
QPushButton#primaryBtn {
    background-color: #10b981;  /* Verde acceso */
    color: white;
}

QPushButton#primaryBtn:hover {
    background-color: #059669;  /* Verde più scuro al hover */
}

/* Pulsante DESELEZIONA TUTTO - ROSSO */
QPushButton#secondaryBtn {
    background-color: #ef4444;  /* Rosso acceso */
    color: white;
}

QPushButton#secondaryBtn:hover {
    background-color: #dc2626;  /* Rosso più scuro al hover */
}

/* Pulsante SELEZIONA TUTTO - BLU */
QPushButton#successBtn {
    background-color: #3b82f6;  /* Blu acceso */
    color: white;
}

QPushButton#successBtn:hover {
    background-color: #2563eb;  /* Blu più scuro al hover */
}

QPushButton#dangerBtn {
    background-color: #e74c3c;
}

QPushButton#dangerBtn:hover {
    background-color: #c0392b;
}

/* Tabs */
QTabWidget::pane {
    border: 1px solid #8b5cf6;
    background-color: white;
    border-radius: 6px;
}

QTabBar::tab {
    background-color: #f3f4f6;
    color: #8b5cf6;  /* Testo viola */
    padding: 12px 20px;
    margin-right: 2px;
    border: 1px solid #e5e7eb;
    border-bottom: none;
    border-top-left-radius: 6px;
    border-top-right-radius: 6px;
    font-weight: 600;
}

QTabBar::tab:selected {
    background-color: #8b5cf6;  /* Viola acceso per tab selezionato */
    color: white;
    border-bottom: 1px solid white;
}

QTabBar::tab:hover:!selected {
    background-color: #ede9fe;  /* Viola molto chiaro al hover */
    color: #8b5cf6;
}

/* GroupBox */
QGroupBox {
    background-color: white;
    border: 2px solid #e5e7eb;
    border-radius: 8px;
    margin: 10px 5px;
    padding-top: 25px;
    font-weight: bold;
    font-size: 14px;
    color: #4a5568;
}

QGroupBox::title {
    subcontrol-origin: margin;
    left: 15px;
    top: 8px;
    padding: 6px 12px;
    background-color: #8b5cf6;  /* Viola acceso */
    color: white;
    border-radius: 4px;
    font-weight: bold;
}

/* Checkbox */
QCheckBox {
    spacing: 8px;
    font-size: 13px;
    color: #4a5568;
    padding: 6px;
    margin: 2px;
}

QCheckBox::indicator {
    width: 18px;
    height: 18px;
    border: 2px solid #e5e7eb;
    border-radius: 4px;
    background-color: white;
}

QCheckBox::indicator:checked {
    background-color: #8b5cf6;  /* Viola acceso */
    border-color: #8b5cf6;
}

QCheckBox::indicator:hover {
    border-color: #8b5cf6;
    background-color: #f7fafc;
}

/* RadioButton */
QRadioButton {
    spacing: 8px;
    font-size: 14px;
    color: #4a5568;
    font-weight: 600;
}

QRadioButton::indicator {
    width: 16px;
    height: 16px;
    border: 2px solid #e5e7eb;
    border-radius: 8px;
    background-color: white;
}

QRadioButton::indicator:checked {
    background-color: #8b5cf6;  /* Viola acceso */
    border-color: #8b5cf6;
}

/* ScrollArea */
QScrollArea {
    border: 1px solid #e5e7eb;
    background-color: white;
    border-radius: 6px;
}

QScrollBar:vertical {
    background-color: #f3f4f6;
    width: 15px;
    border-radius: 7px;
}

QScrollBar::handle:vertical {
    background-color: #8b5cf6;  /* Viola acceso */
    border-radius: 7px;
    min-height: 20px;
}

QScrollBar::handle:vertical:hover {
    background-color: #7c3aed;  /* Viola più scuro */
}

/* ComboBox */
QComboBox {
    background-color: white;
    border: 2px solid #ddd;
    border-radius: 6px;
    padding: 8px 12px;
    font-size: 13px;
    color: #333;
    min-width: 150px;
}

QComboBox:focus {
    border-color: #8b5cf6;
}

QComboBox::drop-down {
    border: none;
    width: 25px;
}

QComboBox::down-arrow {
    image: none;
    border-left: 5px solid transparent;
    border-right: 5px solid transparent;
    border-top: 5px solid #8b5cf6;
    margin-right: 5px;
}

/* TreeWidget */
QTreeWidget {
    background-color: white;
    border: 1px solid #e5e7eb;
    border-radius: 6px;
    font-size: 13px;
    alternate-background-color: #f8fafc;
}

QTreeWidget::item {
    padding: 8px;
    border-bottom: 1px solid #e2e8f0;
}

QTreeWidget::item:selected {
    background-color: #8b5cf6;  /* Viola acceso */
    color: white;
}

QTreeWidget::item:hover {
    background-color: #edf2f7;
}

/* Menu Bar */
QMenuBar {
    background-color: #7c3aed;  /* Viola scuro */
    color: white;
    font-weight: 600;
    padding: 5px;
    border-bottom: 1px solid #6d28d9;
}

QMenuBar::item {
    background-color: transparent;
    padding: 8px 12px;
    border-radius: 4px;
    margin: 2px;
}

QMenuBar::item:selected {
    background-color: #8b5cf6;  /* Viola acceso */
}

QMenu {
    background-color: white;
    border: 1px solid #e5e7eb;
    border-radius: 6px;
    padding: 5px;
}

QMenu::item {
    padding: 8px 20px;
    border-radius: 4px;
    margin: 1px;
}

QMenu::item:selected {
    background-color: #8b5cf6;  /* Viola acceso */
    color: white;
}

/* Progress Bar */
QProgressBar {
    background-color: #f3f4f6;
    border: 1px solid #e5e7eb;
    border-radius: 6px;
    text-align: center;
    font-weight: bold;
    color: #4a5568;
}

QProgressBar::chunk {
    background-color: #8b5cf6;  /* Viola acceso */
    border-radius: 4px;
}

/* Cards/Frames */
QFrame#card {
    background-color: white;
    border: 1px solid #e2e8f0;
    border-radius: 8px;
    padding: 15px;
    margin: 10px;
}

QFrame#headerCard {
    background-color: #8b5cf6;  /* Viola acceso */
    border: none;
    border-radius: 10px;
    padding: 20px;
    margin: 10px;
}

/* Toolbar */
QToolBar {
    background-color: #f3f4f6;
    border: 1px solid #e5e7eb;
    border-radius: 6px;
    padding: 5px;
    spacing: 10px;
}

QToolBar::separator {
    background-color: #e5e7eb;
    width: 1px;
    margin: 5px;
}
/* Header cliccabile per informazioni cliente */
QPushButton#sectionHeader:flat {
    background-color: #3b82f6;  /* BLU ACCESO */
    color: white;
    font-size: 16px;
    font-weight: bold;
    padding: 12px 15px;
    border: 1px solid #3b82f6;
    border-radius: 6px;
    margin: 10px 0;
    text-align: left;
}

QPushButton#sectionHeader:flat:hover {
    background-color: #2563eb;
} 
SCROLL_IMPROVEMENTS = 
/* Miglioramenti scroll */
QScrollBar:vertical {
    width: 0px !important;
    background: transparent !important;
    border: none !important;
    margin: 0px !important;
}

QScrollBar::handle:vertical {
    background: transparent !important;
    width: 0px !important;
    border: none !important;
    margin: 0px !important;
}

QScrollBar::add-line:vertical,
QScrollBar::sub-line:vertical {
    width: 0px !important;
    height: 0px !important;
    background: none !important;
    border: none !important;
}

QScrollBar::add-page:vertical,
QScrollBar::sub-page:vertical {
    background: none !important;
    width: 0px !important;
} 
PROFESSIONAL_CALENDAR_STYLE = 
/* Base */
QWidget {
    font-family: 'Inter', 'Segoe UI', Arial, sans-serif;
    font-size: 13px;
}

/* Calendario principale */
QMainWindow {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
        stop:0 #f1f5f9, stop:1 #e2e8f0);
}

/* Celle giorni premium */
QFrame[objectName="dayCell"] {
    background: white;
    border: 1px solid #f1f5f9;
    border-radius: 8px;
    margin: 2px;
}

QFrame[objectName="dayCell"]:hover {
    border: 2px solid #3b82f6;
    background: #fefeff;
    transform: scale(1.02);
}

QFrame[objectName="dayCell"][isToday="true"] {
    border: 2px solid #ef4444;
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 #fef2f2, stop:1 #fecaca);
}

/* Badge eventi eleganti */
QFrame[objectName="eventBadge"] {
    border-radius: 4px;
    margin: 1px;
    padding: 2px 6px;
    font-size: 10px;
    font-weight: 600;
}

/* Dialog premium */
QDialog {
    background: white;
    border: 1px solid #e2e8f0;
    border-radius: 12px;
}

/* Buttons premium */
QPushButton {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 #f8fafc, stop:1 #e2e8f0);
    border: 1px solid #d1d5db;
    border-radius: 8px;
    color: #374151;
    font-weight: 600;
    padding: 10px 20px;
    min-height: 20px;
}

QPushButton:hover {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 #3b82f6, stop:1 #2563eb);
    border-color: #3b82f6;
    color: white;
    transform: translateY(-1px);
}

QPushButton:pressed {
    transform: translateY(0px);
}

/* Shadows and effects */
QFrame[objectName="card"] {
    background: white;
    border: 1px solid #e2e8f0;
    border-radius: 10px;
    /* box-shadow: 0 1px 3px rgba(0,0,0,0.1); */
}
"""


class BusinessCard(QFrame):
    """Card con stile business (se non esiste già nel tuo codice)"""
    def __init__(self):
        super().__init__()
        self.setObjectName("card")
        self.setFrameStyle(QFrame.Shape.StyledPanel)
        self.setStyleSheet("""
            QFrame#card {
                background-color: white;
                border: 1px solid #e2e8f0;
                border-radius: 8px;
                padding: 15px;
                margin: 5px;
            }
        """)

class HeaderCard(QFrame):
    def __init__(self):
        super().__init__()
        self.setObjectName("headerCard")

# ===================================================================
# CORREZIONE COMPLETA DELLA CLASSE StatisticsWindow
# ===================================================================

# ========================================
# SOSTITUISCI LA CLASSE StatisticsWindow CON QUESTA VERSIONE MIGLIORATA
# ========================================

class StatisticsWindow(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Statistiche Servizi Catering")
        
        # ========== GEOMETRIA AUTOMATICA E RESPONSIVE ==========
        # Ottieni dimensioni schermo
        if parent:
            parent_geo = parent.geometry()
            screen = parent.screen()
        else:
            screen = QApplication.primaryScreen()
        
        screen_geo = screen.geometry()
        
        # Finestra grande (85% dello schermo)
        window_width = int(screen_geo.width() * 0.85)
        window_height = int(screen_geo.height() * 0.85)
        
        # Centra la finestra
        x = (screen_geo.width() - window_width) // 2
        y = (screen_geo.height() - window_height) // 2
        
        self.setGeometry(x, y, window_width, window_height)
        self.setMinimumSize(1000, 700)  # Dimensione minima più grande
        # =====================================================
        
        self.setStyleSheet(PROFESSIONAL_STYLE + """
    /* Miglioramenti specifici per statistiche */
    QScrollBar:vertical {
        width: 12px;  /* Cambia da 0px a 12px per renderlo visibile */
        background: #f8f9fa;
        border-radius: 6px;
    }
    QScrollBar::handle:vertical {
        background: #8b5cf6;
        border-radius: 6px;
        min-height: 20px;
    }
    QScrollBar::handle:vertical:hover {
        background: #7c3aed;
    }
    QDialog {
        background-color: #f8f9fa;
    }
""")
        
        self.ordini = []
        self.mesi_mappatura = {}
        self.setup_ui()
        
    def setup_ui(self):
        # ========== LAYOUT PRINCIPALE CON SCROLL ==========
        main_layout = QVBoxLayout()
        main_layout.setSpacing(20)
        main_layout.setContentsMargins(20, 20, 20, 20)
        
        # Header
        header = QLabel("STATISTICHE SERVIZI CATERING")
        header.setObjectName("sectionHeader")
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(header)
        
        # ========== CARD FILTRI MIGLIORATA ==========
        filter_card = BusinessCard()
        filter_layout = QHBoxLayout(filter_card)
        filter_layout.setSpacing(15)
        
        filter_layout.addWidget(QLabel("Filtra per mese:"))
        self.mese_combo = QComboBox()
        self.mese_combo.setMinimumWidth(250)  # Più largo
        self.mese_combo.setMaximumHeight(40)   # Più alto
        filter_layout.addWidget(self.mese_combo)
        
        filter_layout.addStretch()
        
        # Pulsante refresh più grande
        refresh_btn = QPushButton("🔄 Aggiorna Statistiche")
        refresh_btn.setMinimumSize(150, 40)
        refresh_btn.clicked.connect(self.load_statistics)
        filter_layout.addWidget(refresh_btn)
        
        main_layout.addWidget(filter_card)
        
        # ========== AREA CONTENUTO CON SCROLL MIGLIORATA ==========
        # ========== AREA CONTENUTO CON SCROLL MIGLIORATA ==========
        # ========== AREA CONTENUTO CON SCROLL MIGLIORATA ==========
        # ========== AREA CONTENUTO CON SCROLL MIGLIORATA ==========
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        # Ottimizza lo scroll per fluidità
        self.scroll_area.verticalScrollBar().setSingleStep(20)  # Step più fluido
        self.scroll_area.verticalScrollBar().setPageStep(100)   # Page step ottimizzato
        self.scroll_area.setStyleSheet("""
            QScrollArea {
                border: 1px solid #e5e7eb;
                background-color: white;
                border-radius: 8px;
            }
        """)
        
        # Widget contenuto
        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setSpacing(25)
        self.content_layout.setContentsMargins(20, 20, 20, 20)
        
        self.scroll_area.setWidget(self.content_widget)
        main_layout.addWidget(self.scroll_area)
        
        # ========== PULSANTI AZIONE ==========
        buttons_layout = QHBoxLayout()
        buttons_layout.addStretch()
        
        export_btn = QPushButton("📊 Esporta Dati")
        export_btn.setMinimumSize(130, 40)
        export_btn.setObjectName("successBtn")
        buttons_layout.addWidget(export_btn)
        export_btn.clicked.connect(self.export_statistics_to_excel)

        
        close_btn = QPushButton("Chiudi")
        close_btn.setMinimumSize(100, 40)
        close_btn.setObjectName("secondaryBtn")
        close_btn.clicked.connect(self.close)
        buttons_layout.addWidget(close_btn)
        
        main_layout.addLayout(buttons_layout)
        self.setLayout(main_layout)
        
        # Carica dati iniziali
        self.load_statistics()

    def export_statistics_to_excel(self):
        """Esporta statistiche filtrate in Excel"""
        try:
            # Controlla se pandas è disponibile
            try:
                import pandas as pd # type: ignore
                from datetime import datetime
            except ImportError:
                QMessageBox.warning(self, "Libreria Mancante", 
                                "Per esportare in Excel installa pandas:\n\n"
                                "pip install pandas openpyxl\n\n"
                                "Riavvia il programma dopo l'installazione.")
                return
            
            if not self.ordini:
                QMessageBox.warning(self, "Nessun Dato", "Non ci sono dati da esportare.")
                return
            
            # Ottieni filtro corrente
            testo_selezionato = self.mese_combo.currentText()
            mese_filtro = self.mesi_mappatura.get(testo_selezionato)
            ordini_filtrati = self.filtra_ordini_per_mese(mese_filtro)
            
            if not ordini_filtrati:
                QMessageBox.warning(self, "Nessun Dato", 
                                f"Nessun ordine trovato per '{testo_selezionato}'")
                return
            
            # Progress bar
            progress = QProgressBar(self)
            progress.setWindowTitle("Generazione Excel...")
            progress.setGeometry(self.width()//2 - 150, self.height()//2, 300, 30)
            progress.setRange(0, 100)
            progress.show()
            QApplication.processEvents()
            
            progress.setValue(20)
            QApplication.processEvents()
            
            # ========== PREPARA DATI PER EXCEL ==========
            
            # 1. DATI ORDINI DETTAGLIATI
            ordini_data = []
            servizi_count = {}
            persone_per_servizio = {}
            
            for ordine in ordini_filtrati:
                campi = ordine["campi"]
                
                # Determina tipo servizio
                servizio = self.determina_tipo_servizio(ordine)
                servizi_count[servizio] = servizi_count.get(servizio, 0) + 1
                
                # Conta persone
                try:
                    persone = int(campi.get("Numero persone", 0))
                    persone_per_servizio[servizio] = persone_per_servizio.get(servizio, 0) + persone
                except (ValueError, TypeError):
                    persone = 0
                
                # Servizi selezionati
                servizi_selezionati = []
                for k, v in ordine.get("scelte", {}).items():
                    if v and "|" in k:
                        voce = k.split("|")[-1]  # Ultima parte dopo |
                        servizi_selezionati.append(voce)
                
                ordini_data.append({
                    "Data Evento": campi.get("Data evento", ""),
                    "Cliente": campi.get("Nome", ""),
                    "Cellulare": campi.get("Cellulare", ""),
                    "Indirizzo": campi.get("Via", ""),
                    "Persone": persone,
                    "Tipo Servizio": servizio,
                    "Tipo Aperitivo": ordine.get("tipo_ap", "") if "aperitivo" in servizio.lower() else "",
                    "Orario Allestimento": campi.get("Orario allestimento", ""),
                    "Orario Pronti": campi.get("Orario Pronti", ""),
                    "In Fiera": "Sì" if ordine.get("evento_in_fiera", False) else "No",
                    "Numero Servizi": len(servizi_selezionati),
                    "Servizi Dettaglio": "; ".join(servizi_selezionati[:10]) + ("..." if len(servizi_selezionati) > 10 else "")
                })
            
            progress.setValue(50)
            QApplication.processEvents()
            
            # 2. RIEPILOGO STATISTICHE
            total_ordini = len(ordini_filtrati)
            total_persone = sum(persone_per_servizio.values())
            media_persone = total_persone / total_ordini if total_ordini > 0 else 0
            
            riepilogo_data = [
                {"Metrica": "Periodo Analizzato", "Valore": testo_selezionato},
                {"Metrica": "Numero Ordini Totali", "Valore": total_ordini},
                {"Metrica": "Persone Servite Totali", "Valore": total_persone},
                {"Metrica": "Media Persone per Ordine", "Valore": f"{media_persone:.1f}"},
                {"Metrica": "Data Generazione Report", "Valore": datetime.now().strftime("%d/%m/%Y %H:%M")}
            ]
            
            # 3. DISTRIBUZIONE SERVIZI
            distribuzione_data = []
            for servizio, count in servizi_count.items():
                percentuale = (count / total_ordini) * 100
                persone_totali = persone_per_servizio.get(servizio, 0)
                media_servizio = persone_totali / count if count > 0 else 0
                
                distribuzione_data.append({
                    "Tipo Servizio": servizio,
                    "Numero Ordini": count,
                    "Percentuale": f"{percentuale:.1f}%",
                    "Persone Totali": persone_totali,
                    "Media Persone": f"{media_servizio:.1f}"
                })
            
            progress.setValue(80)
            QApplication.processEvents()
            
            # ========== CREA FILE EXCEL ==========
            
            # Nome file con timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            periodo_safe = testo_selezionato.replace("/", "-").replace(" ", "_")
            filename = f"Statistiche_Catering_{periodo_safe}_{timestamp}.xlsx"
            
            # Crea Excel con pandas
            with pd.ExcelWriter(filename, engine='openpyxl') as writer:
                
                # Foglio 1: Riepilogo
                df_riepilogo = pd.DataFrame(riepilogo_data)
                df_riepilogo.to_excel(writer, sheet_name='Riepilogo', index=False, startrow=0)
                
                # Foglio 2: Distribuzione Servizi
                df_distribuzione = pd.DataFrame(distribuzione_data)
                df_distribuzione.to_excel(writer, sheet_name='Distribuzione Servizi', index=False, startrow=0)
                
                # Foglio 3: Ordini Dettagliati
                df_ordini = pd.DataFrame(ordini_data)
                df_ordini.to_excel(writer, sheet_name='Ordini Dettagliati', index=False, startrow=0)
                
                # ========== FORMATTAZIONE EXCEL ==========
                
                from openpyxl.styles import Font, PatternFill, Alignment # type: ignore
                from openpyxl.utils.dataframe import dataframe_to_rows # type: ignore
                
                workbook = writer.book
                
                # Stile header
                header_font = Font(bold=True, color="FFFFFF")
                header_fill = PatternFill(start_color="8B5CF6", end_color="8B5CF6", fill_type="solid")
                center_alignment = Alignment(horizontal="center", vertical="center")
                
                # Applica stile a tutti i fogli
                for sheet_name in workbook.sheetnames:
                    ws = workbook[sheet_name]
                    
                    # Header row styling
                    for cell in ws[1]:
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.alignment = center_alignment
                    
                    # Auto-width columns
                    for column in ws.columns:
                        max_length = 0
                        column_letter = column[0].column_letter
                        
                        for cell in column:
                            try:
                                if len(str(cell.value)) > max_length:
                                    max_length = len(str(cell.value))
                            except:
                                pass
                        
                        adjusted_width = min(max_length + 2, 50)  # Max 50 caratteri
                        ws.column_dimensions[column_letter].width = adjusted_width
            
            progress.setValue(100)
            QApplication.processEvents()
            progress.hide()
            
            # ========== MESSAGGIO DI SUCCESSO ==========
            
            success_msg = f"""📊 EXPORT EXCEL COMPLETATO!

    📁 File: {filename}
    📋 Fogli creati:
    • Riepilogo (5 metriche principali)
    • Distribuzione Servizi ({len(distribuzione_data)} tipi)
    • Ordini Dettagliati ({len(ordini_data)} ordini)

    📈 Dati esportati:
    • Periodo: {testo_selezionato}
    • Ordini: {total_ordini}
    • Persone: {total_persone}

    💡 Puoi aprire il file con Excel, LibreOffice o Google Sheets"""
            
            QMessageBox.information(self, "Export Completato", success_msg)
            
            # Opzione per aprire il file
            reply = QMessageBox.question(self, "Aprire File?", 
                                    f"Vuoi aprire il file Excel?\n\n{filename}",
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            
            if reply == QMessageBox.StandardButton.Yes:
                self.open_excel_file(filename)
            
        except Exception as e:
            if 'progress' in locals():
                progress.hide()
            
            error_msg = f"❌ Errore durante l'export:\n\n{str(e)}\n\n"
            
            if "pandas" in str(e) or "openpyxl" in str(e):
                error_msg += "💡 Soluzione: Installa le librerie necessarie:\n"
                error_msg += "pip install pandas openpyxl"
            
            QMessageBox.critical(self, "Errore Export", error_msg)

        
    def load_statistics(self):
        """Carica i dati delle statistiche"""
        if not os.path.exists(ORDINI_FILE):
            QMessageBox.warning(self, "Nessun dato", "Non ci sono ordini salvati per le statistiche.")
            return
            
        with open(ORDINI_FILE, "r", encoding="utf-8") as f:
            self.ordini = json.load(f)
        
        if not self.ordini:
            QMessageBox.warning(self, "Nessun dato", "Nessun ordine disponibile per le statistiche.")
            return
            
        # Popola combo mesi
        self.popola_combo_mesi()
        
        # Aggiorna statistiche
        self.update_statistics()

    def popola_combo_mesi(self):
        """Popola la combo box con i mesi disponibili"""
        self.mese_combo.blockSignals(True)
        
        self.mese_combo.clear()
        self.mesi_mappatura.clear()
        
        # Raccogli tutti i mesi disponibili
        mesi_disponibili = set()
        for ordine in self.ordini:
            data_evento = ordine["campi"].get("Data evento", "")
            if data_evento:
                try:
                    data_obj = datetime.strptime(data_evento, "%d/%m/%Y")
                    mese_anno = data_obj.strftime("%Y-%m")
                    mesi_disponibili.add(mese_anno)
                except ValueError:
                    continue
        
        mesi_ordinati = sorted(mesi_disponibili)
        
        # Aggiungi opzione "Tutti i mesi"
        self.mese_combo.addItem("Tutti i mesi")
        self.mesi_mappatura["Tutti i mesi"] = None
        
        # Aggiungi i mesi formattati
        for mese_anno in mesi_ordinati:
            try:
                data_obj = datetime.strptime(mese_anno, "%Y-%m")
                testo_mese = data_obj.strftime("%B %Y")
                self.mese_combo.addItem(testo_mese)
                self.mesi_mappatura[testo_mese] = mese_anno
            except ValueError:
                continue
        
        self.mese_combo.blockSignals(False)
        self.mese_combo.currentTextChanged.connect(self.update_statistics)

    def update_statistics(self):
        """Aggiorna le statistiche in base al filtro selezionato"""
        # Pulisce layout esistente
        while self.content_layout.count():
            child = self.content_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
        
        # Ottieni il filtro selezionato
        testo_selezionato = self.mese_combo.currentText()
        mese_filtro = self.mesi_mappatura.get(testo_selezionato)
        
        # Filtra ordini
        ordini_filtrati = self.filtra_ordini_per_mese(mese_filtro)
        
        if not ordini_filtrati:
            # ========== MESSAGGIO NESSUN DATO MIGLIORATO ==========
            no_data_card = BusinessCard()
            no_data_layout = QVBoxLayout(no_data_card)
            
            icon_label = QLabel("📊")
            icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            icon_label.setStyleSheet("font-size: 48px; margin: 20px;")
            no_data_layout.addWidget(icon_label)
            
            message_label = QLabel("Nessun ordine trovato per il periodo selezionato")
            message_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            message_label.setStyleSheet("font-size: 18px; color: #7f8c8d; padding: 20px;")
            no_data_layout.addWidget(message_label)
            
            self.content_layout.addWidget(no_data_card)
            return
        
        # Analizza dati filtrati
        servizi_count = Counter()
        persone_per_servizio = defaultdict(int)
        
        for ordine in ordini_filtrati:
            servizio = self.determina_tipo_servizio(ordine)
            servizi_count[servizio] += 1
            
            try:
                persone = int(ordine["campi"].get("Numero persone", 0))
                persone_per_servizio[servizio] += persone
            except (ValueError, TypeError):
                pass
        
        # Mostra riepilogo generale
        self.crea_riepilogo_generale_migliorato(ordini_filtrati, persone_per_servizio)
        
        # Crea grafico migliorato
        if servizi_count:
            self.create_chart_improved(servizi_count)
    
    def open_excel_file(self, filepath_completo):
        """Apre il file Excel generato"""
        try:
            import subprocess
            import platform
            
            if platform.system() == "Windows":
                subprocess.run(["start", filepath_completo], shell=True, check=True)
            elif platform.system() == "Darwin":  # macOS
                subprocess.run(["open", filepath_completo], check=True)
            else:  # Linux
                subprocess.run(["xdg-open", filepath_completo], check=True)
                
            print(f"📄 Excel aperto: {filepath_completo}")
            
        except Exception as e:
            print(f"❌ Errore apertura Excel: {e}")
            QMessageBox.information(self, "File Creato", 
                                f"File Excel creato:\n{filepath_completo}\n\n"
                                f"Aprilo manualmente per visualizzare i dati.")

    def crea_riepilogo_generale_migliorato(self, ordini_filtrati, persone_per_servizio):
        """Crea la sezione riepilogo generale migliorata"""
        summary_card = BusinessCard()
        summary_layout = QVBoxLayout(summary_card)
        summary_layout.setSpacing(20)
        
        # Header più grande
        summary_header = QLabel("📈 RIEPILOGO GENERALE")
        summary_header.setStyleSheet("""
            font-size: 20px; 
            font-weight: bold; 
            color: #2c3e50; 
            margin-bottom: 20px;
            padding: 15px;
            background-color: #f8f9fa;
            border-radius: 8px;
            border-left: 5px solid #3498db;
        """)
        summary_layout.addWidget(summary_header)
        
        total_ordini = len(ordini_filtrati)
        total_persone = sum(persone_per_servizio.values())
        media_persone = total_persone / total_ordini if total_ordini > 0 else 0
        
        # Info periodo
        periodo_info = self.get_periodo_info(ordini_filtrati)
        
        # ========== GRIGLIA STATISTICHE MIGLIORATA ==========
        stats_widget = QWidget()
        stats_grid = QGridLayout(stats_widget)
        stats_grid.setSpacing(20)
        
        stats_data = [
            ("📅 Periodo:", periodo_info, "#3498db"),
            ("📋 Ordini Totali:", str(total_ordini), "#2ecc71"),
            ("👥 Persone Servite:", str(total_persone), "#e74c3c"),
            ("📊 Media Persone/Ordine:", f"{media_persone:.1f}", "#f39c12")
        ]
        
        for i, (label, value, color) in enumerate(stats_data):
            # Card per ogni statistica
            stat_card = QFrame()
            stat_card.setStyleSheet(f"""
                QFrame {{
                    background-color: white;
                    border: 2px solid {color};
                    border-radius: 10px;
                    padding: 20px;
                    margin: 5px;
                }}
            """)
            
            stat_layout = QVBoxLayout(stat_card)
            
            label_widget = QLabel(label)
            label_widget.setStyleSheet("font-weight: bold; color: #34495e; font-size: 14px;")
            label_widget.setAlignment(Qt.AlignmentFlag.AlignCenter)
            
            value_widget = QLabel(value)
            value_widget.setStyleSheet(f"font-size: 24px; font-weight: bold; color: {color};")
            value_widget.setAlignment(Qt.AlignmentFlag.AlignCenter)
            
            stat_layout.addWidget(label_widget)
            stat_layout.addWidget(value_widget)
            
            stats_grid.addWidget(stat_card, i // 2, i % 2)
        
        summary_layout.addWidget(stats_widget)
        self.content_layout.addWidget(summary_card)

    def create_chart_improved(self, servizi_count):
        """Crea il grafico delle statistiche migliorato"""
        chart_card = BusinessCard()
        chart_layout = QVBoxLayout(chart_card)
        chart_layout.setSpacing(20)
        
        # Header
        chart_header = QLabel("📊 DISTRIBUZIONE SERVIZI")
        chart_header.setStyleSheet("""
            font-size: 20px; 
            font-weight: bold; 
            color: #2c3e50; 
            margin-bottom: 15px;
            padding: 15px;
            background-color: #f8f9fa;
            border-radius: 8px;
            border-left: 5px solid #9b59b6;
        """)
        chart_layout.addWidget(chart_header)
        
        # ========== GRAFICO MATPLOTLIB MIGLIORATO ==========
        # Dimensioni più grandi per il grafico
        # ========== GRAFICO MATPLOTLIB MIGLIORATO ==========
        # Calcola dimensioni dinamiche basate sulla finestra
        # Dimensioni fisse ottimizzate per le statistiche
        fig = Figure(figsize=(14, 8), facecolor='white', dpi=100)
        canvas = FigureCanvas(fig)
        canvas.setMinimumHeight(600)
        canvas.setFixedHeight(650)  # Altezza fissa per evitare lag

        # CORREZIONE: Disabilita la cattura degli eventi wheel su matplotlib
        canvas.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        # Funzione per passare gli eventi wheel al parent scroll area
        def wheelEvent(event):
            # Passa l'evento wheel al scroll area parent
            if hasattr(self, 'scroll_area'):
                self.scroll_area.wheelEvent(event)
            else:
                event.ignore()

        # Sovrascrivi l'evento wheel del canvas
        canvas.wheelEvent = wheelEvent
        
        
        ax = fig.add_subplot(111)
        
        # Colori moderni
        colors = ['#3498db', '#e74c3c', '#2ecc71', '#f39c12', '#9b59b6', '#1abc9c', '#34495e', '#e67e22']
        
        # Prepara dati per il grafico
        servizi = list(servizi_count.keys())
        conteggi = list(servizi_count.values())
        
        # Crea grafico a barre con stile moderno
        bars = ax.bar(servizi, conteggi, color=colors[:len(servizi)], alpha=0.8, edgecolor='white', linewidth=2)
        
        # Titolo e etichette
        ax.set_title("Numero di Ordini per Servizio", fontsize=18, fontweight='bold', color='#2c3e50', pad=20)
        ax.set_ylabel("Numero Ordini", fontweight='600', color='#34495e', fontsize=14)
        ax.set_xlabel("Tipo Servizio", fontweight='600', color='#34495e', fontsize=14)
        
        # Stile moderno
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('#bdc3c7')
        ax.spines['bottom'].set_color('#bdc3c7')
        ax.tick_params(colors='#34495e', labelsize=12)
        ax.grid(True, alpha=0.3, linestyle='-', linewidth=0.5)
        
        # Valori sopra le barre con stile
        for bar in bars:
            height = bar.get_height()
            ax.annotate(f'{int(height)}',
                       xy=(bar.get_x() + bar.get_width() / 2, height),
                       xytext=(0, 8),
                       textcoords="offset points",
                       ha='center', va='bottom',
                       fontweight='bold', fontsize=14, color='#2c3e50',
                       bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.8))
        
        # Ruota etichette se necessario
        if max(len(s) for s in servizi) > 12:
            plt.setp(ax.get_xticklabels(), rotation=45, ha='right')
        
        # Layout ottimizzato
        fig.tight_layout(pad=3.0)
        chart_layout.addWidget(canvas)
        self.content_layout.addWidget(chart_card)

    # ========== MANTIENI TUTTE LE ALTRE FUNZIONI ORIGINALI ==========
    def filtra_ordini_per_mese(self, mese_filtro):
        """Filtra gli ordini per il mese specificato"""
        if mese_filtro is None:
            return self.ordini
        
        ordini_filtrati = []
        for ordine in self.ordini:
            data_evento = ordine["campi"].get("Data evento", "")
            if data_evento:
                try:
                    data_obj = datetime.strptime(data_evento, "%d/%m/%Y")
                    mese_anno_ordine = data_obj.strftime("%Y-%m")
                    
                    if mese_anno_ordine == mese_filtro:
                        ordini_filtrati.append(ordine)
                        
                except ValueError:
                    continue
        
        return ordini_filtrati
    
    def determina_tipo_servizio(self, ordine):
        """Determina il tipo di servizio da un ordine - VERSIONE COMBINATA"""
        scelte = ordine.get("scelte", {})
        print(f"🔍 DEBUG determina_tipo_servizio - Scelte: {[k for k, v in scelte.items() if v]}")

        
        # Lista per raccogliere tutti i servizi trovati
        servizi_trovati = []
        
        # Verifica Coffee Break
        if any(k.startswith("coffee break|") and v for k, v in scelte.items()):
            servizi_trovati.append("Coffee Break")
        
        # Verifica Tea Break
        if any(k.startswith("tea break|") and v for k, v in scelte.items()):
            servizi_trovati.append("Tea Break")
        if any(k.startswith("lunch box|") and v for k, v in scelte.items()):
            servizi_trovati.append("Lunch box")
            
        if any(k.startswith("servito|") and v for k, v in scelte.items()):
            servizi_trovati.append("Servito")
        
        # Verifica Lunch Buffet
        if any(k.startswith("lunch buffet|") and v for k, v in scelte.items()):
            servizi_trovati.append("Lunch Buffet")
        
        # Verifica Aperitivo
        if any(k.startswith("aperitivo a buffet|") and v for k, v in scelte.items()):
            tipo_ap = ordine.get("tipo_ap", "Leggero")
            servizi_trovati.append(f"Aperitivo {tipo_ap}")
        
        # Restituisci tutti i servizi combinati
        if len(servizi_trovati) == 0:
            return "Altro"
        elif len(servizi_trovati) == 1:
            return servizi_trovati[0]
        else:
            # Ordini combinati - mostra tutti i servizi
            return " + ".join(servizi_trovati)
    
    def get_periodo_info(self, ordini_filtrati):
        """Ottiene info sul periodo degli ordini filtrati"""
        if not ordini_filtrati:
            return "Nessun dato"
        
        testo_combo = self.mese_combo.currentText()
        if testo_combo == "Tutti i mesi":
            return f"Tutti i mesi ({len(ordini_filtrati)} ordini)"
        else:
            return testo_combo

    def resizeEvent(self, event):
        """Gestisce il ridimensionamento della finestra"""
        super().resizeEvent(event)
        # Forza il ricalcolo del layout quando ridimensioni
        if hasattr(self, 'content_widget'):
            self.content_widget.updateGeometry()
            self.scroll_area.updateGeometry()

# SOSTITUISCI LA CLASSE OrderArchiveDialog CON QUESTA VERSIONE CORRETTA

class OrderArchiveDialog(QDialog):
    order_selected = pyqtSignal(dict)
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Archivio Ordini Salvati")
        self.setGeometry(100, 100, 1000, 700)
        self.setStyleSheet(PROFESSIONAL_STYLE)
        self.setup_ui()
        self.load_orders()
        
        # AGGIUNGI QUESTE RIGHE ALLA FINE:
        global archivio_dialogs_aperti
        archivio_dialogs_aperti.append(self)
        print(f"📋 Archivio registrato. Totale archivi aperti: {len(archivio_dialogs_aperti)}")
    def ricarica_dati(self):
        """Ricarica i dati dell'archivio dal file JSON - VERSIONE DEBUG"""
        print("🔄 === INIZIO RICARICAMENTO ARCHIVIO ===")
        
        # Salva l'elemento attualmente selezionato
        current_item = self.tree.currentItem()
        selected_cliente = None
        selected_data = None
        
        if current_item and current_item.parent() is not None:
            selected_cliente = current_item.text(1)
            selected_data = current_item.text(0)
            print(f"💾 Elemento selezionato: {selected_cliente} - {selected_data}")
        
        # Pulisci completamente l'albero
        print("🧹 Pulendo albero esistente...")
        self.tree.clear()
        
        # Ricarica dal file
        print("📂 Ricaricando dal file JSON...")
        self.load_orders()
        
        # Prova a ristabilire la selezione
        if selected_cliente and selected_data:
            print(f"🎯 Cercando di ristabilire selezione: {selected_cliente}")
            root = self.tree.invisibleRootItem()
            for i in range(root.childCount()):
                date_group = root.child(i)
                for j in range(date_group.childCount()):
                    order_item = date_group.child(j)
                    if (order_item.text(1) == selected_cliente and 
                        order_item.text(0) == selected_data):
                        self.tree.setCurrentItem(order_item)
                        self.tree.scrollToItem(order_item)
                        print(f"✅ Selezione ristabilita")
                        break
        
        print("✅ === FINE RICARICAMENTO ARCHIVIO ===")
        
        print("✅ Dati archivio ricaricati")
    def add_selected_services(self):
        """Aggiunge servizi a un ordine esistente - VERSIONE DEFINITIVA CORRETTA"""
        current = self.tree.currentItem()
        if not current or current.parent() is None:
            QMessageBox.information(self, "Selezione", "Seleziona un ordine specifico dalla lista.")
            return
            
        idx = current.data(0, Qt.ItemDataRole.UserRole)
        if idx is not None and 0 <= idx < len(self.ordini):
            ordine = self.ordini[idx]
            
            # Conferma aggiunta servizi
            cliente = ordine["campi"].get("Nome", "Cliente")
            data = ordine["campi"].get("Data evento", "")
            
            reply = QMessageBox.question(
                self, "Conferma Aggiunta Servizi", 
                f"Vuoi aggiungere servizi all'ordine di:\n\n"
                f"Cliente: {cliente}\n"
                f"Data: {data}\n\n"
                f"I nuovi servizi verranno aggiunti al documento Word esistente.\n"
                f"💡 L'archivio rimarrà aperto per vedere le modifiche",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                # ========== IMPOSTA MODALITÀ AGGIUNGI SERVIZI ==========
                print(f"🔧 Attivando modalità aggiungi servizi per ordine in posizione {idx}")
                
                if hasattr(self.parent(), 'modalita_aggiungi_servizi'):
                    self.parent().modalita_aggiungi_servizi = True
                    self.parent().ordine_da_modificare = ordine
                    # IMPORTANTE: Memorizza l'indice per il salvataggio
                    self.parent().indice_ordine_originale = idx
                    print(f"💡 Memorizzato indice ordine per aggiungi servizi: {idx}")
                    
                    # Debug
                    print(f"🔍 Debug ordine per aggiungi servizi:")
                    print(f"   - Nome: {cliente}")
                    print(f"   - Data: {data}")
                    print(f"   - Path: {ordine.get('path_docx', 'N/A')}")
                    print(f"   - Indice: {idx}")
                
                # CARICA I DATI PER AGGIUNGI SERVIZI (azzera scelte)
                self.parent().carica_dati_per_aggiungi_servizi(ordine)
                
                # AVVISO
                QMessageBox.information(self, "Modalità Aggiungi Servizi Attivata", 
                                    f"✅ Modalità aggiungi servizi attivata!\n\n"
                                    f"👤 Cliente: {cliente}\n"
                                    f"📅 Data: {data}\n\n"
                                    f"➕ Seleziona i nuovi servizi da aggiungere\n"
                                    f"🔄 Premi 'Aggiungi a Ordine Esistente' per salvare\n"
                                    f"👀 L'archivio si aggiornerà automaticamente")
                
                # ========== NON CHIUDERE L'ARCHIVIO! ==========
                # RIMUOVI ENTRAMBE LE RIGHE self.accept()
                print("✅ Modalità aggiungi servizi configurata, archivio rimane aperto")
            
    
    def setup_ui(self):
        layout = QVBoxLayout()
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)
        
        # Header
        header = QLabel("ARCHIVIO ORDINI SALVATI")
        header.setObjectName("sectionHeader")
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(header)
        
        # Search bar
        search_card = BusinessCard()
        search_layout = QHBoxLayout(search_card)
        
        search_label = QLabel("Cerca Cliente/Cellulare:")
        search_label.setStyleSheet("font-weight: bold;")
        search_layout.addWidget(search_label)
        
        self.search_entry = QLineEdit()
        self.search_entry.setPlaceholderText("Inserisci nome cliente o numero di telefono...")
        self.search_entry.textChanged.connect(self.filter_orders)
        search_layout.addWidget(self.search_entry)
        
        layout.addWidget(search_card)
        
        # ========== INFO SELEZIONE ==========
        # Card per mostrare info sulla selezione corrente
        self.selection_card = BusinessCard()
        selection_layout = QHBoxLayout(self.selection_card)
        
        self.selection_info = QLabel("Nessun ordine selezionato")
        self.selection_info.setStyleSheet("color: #6b7280; font-style: italic; font-size: 14px;")
        selection_layout.addWidget(self.selection_info)
        
        selection_layout.addStretch()
        
        # Istruzioni
        help_label = QLabel("💡 Usa Ctrl+Click per selezionare più ordini")
        help_label.setStyleSheet("color: #9ca3af; font-size: 12px;")
        selection_layout.addWidget(help_label)
        
        layout.addWidget(self.selection_card)
        
        # Tree widget con selezione multipla
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["Data", "Nome", "Cellulare", "Persone", "Servizio"])
        self.tree.setAlternatingRowColors(True)
        
        # ========== SELEZIONE MULTIPLA ABILITATA ==========
        self.tree.setSelectionMode(QTreeWidget.SelectionMode.ExtendedSelection)
        self.tree.setRootIsDecorated(True)
        
        # Connetti segnale per aggiornare info selezione
        self.tree.itemSelectionChanged.connect(self.update_selection_info)
        
        header_view = self.tree.header()
        header_view.setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        layout.addWidget(self.tree)
        
        # ========== PULSANTI AZIONE ==========
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        # Pulsante Carica (solo lettura)
        # Pulsante Carica (VERDE)
        load_btn = QPushButton("📖 Carica Ordine")
        load_btn.setObjectName("primaryBtn")  # Verde
        load_btn.setToolTip("Carica l'ordine nel form (solo visualizzazione)")
        load_btn.clicked.connect(self.load_selected_order)
        button_layout.addWidget(load_btn)

        # Pulsante Modifica (VIOLA)
        modify_btn = QPushButton("✏️ Modifica Ordine")
        modify_btn.setStyleSheet("""
            QPushButton {
                background-color: #8b5cf6;
                color: white;
                border: none;
                border-radius: 6px;
                padding: 10px 20px;
                font-size: 13px;
                font-weight: 600;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #7c3aed;
            }
            QPushButton:pressed {
                background-color: #6d28d9;
            }
        """)
        modify_btn.setToolTip("Carica l'ordine per modificarlo e rigenerarlo")
        modify_btn.clicked.connect(self.modify_selected_order)
        button_layout.addWidget(modify_btn)
        
        add_services_btn = QPushButton("➕ Aggiungi Servizi")
        add_services_btn.setStyleSheet("""
            QPushButton {
                background-color: #10b981;
                color: white;
                font-weight: bold;
                padding: 8px 16px;
                border-radius: 6px;
            }
            QPushButton:hover {
                background-color: #059669;
            }
        """)
        add_services_btn.clicked.connect(self.add_selected_services)  # <-- COLLEGAMENTO
        button_layout.addWidget(add_services_btn)
        # Pulsante Elimina (dinamico)
        self.delete_btn = QPushButton("🗑️ Elimina Ordine")
        self.delete_btn.setObjectName("dangerBtn")  # Rosso
        self.delete_btn.setToolTip("Elimina definitivamente l'ordine selezionato")
        self.delete_btn.setEnabled(False)  # Inizialmente disabilitato
        self.delete_btn.clicked.connect(self.delete_selected_orders)
        button_layout.addWidget(self.delete_btn)
        
        # Pulsante Chiudi
        # Pulsante Chiudi (BLU)
        close_btn = QPushButton("Chiudi")
        close_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #3b82f6, stop:1 #2563eb);
                border: 1px solid #2563eb;
                color: white;
                font-weight: 600;
                border-radius: 6px;
                padding: 10px 20px;
                min-width: 100px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #2563eb, stop:1 #1d4ed8);
                transform: translateY(-1px);
            }
            QPushButton:pressed {
                transform: translateY(0px);
            }
        """)
        close_btn.clicked.connect(self.close)
        button_layout.addWidget(close_btn)
        
        layout.addLayout(button_layout)
        self.setLayout(layout)
    
    def update_selection_info(self):
        """Aggiorna info sulla selezione corrente"""
        selected_items = self.tree.selectedItems()
        
        # Filtra solo gli ordini (non i gruppi per data)
        selected_orders = [item for item in selected_items if item.parent() is not None]
        count = len(selected_orders)
        
        if count == 0:
            self.selection_info.setText("Nessun ordine selezionato")
            self.selection_info.setStyleSheet("color: #6b7280; font-style: italic; font-size: 14px;")
            self.delete_btn.setEnabled(False)
            self.delete_btn.setText("🗑️ Elimina Ordine")
        elif count == 1:
            # Singola selezione
            order_item = selected_orders[0]
            cliente = order_item.text(1)  # Nome cliente
            self.selection_info.setText(f"Selezionato: {cliente}")
            self.selection_info.setStyleSheet("color: #3b82f6; font-weight: 500; font-size: 14px;")
            self.delete_btn.setEnabled(True)
            self.delete_btn.setText("🗑️ Elimina Ordine")
        else:
            # Selezione multipla
            self.selection_info.setText(f"🔢 Selezionati {count} ordini")
            self.selection_info.setStyleSheet("color: #dc2626; font-weight: bold; font-size: 14px;")
            self.delete_btn.setEnabled(True)
            self.delete_btn.setText(f"🗑️ Elimina {count} Ordini")
    
    def load_orders(self):
        """Carica gli ordini dall'archivio - VERSIONE CHE CONTROLLA MODIFICHE"""
        
        # Controlla se ci sono state modifiche recenti
        if hasattr(self.parent(), '_ordine_appena_modificato') and self.parent()._ordine_appena_modificato:
            print("🔄 Rilevate modifiche recenti, ricaricando archivio...")
            self.parent()._ordine_appena_modificato = False
        
        if not os.path.exists(ORDINI_FILE):
            QMessageBox.warning(self, "Nessun archivio", "Non ci sono ordini salvati.")
            return
            
        try:
            with open(ORDINI_FILE, "r", encoding="utf-8") as f:
                self.ordini = json.load(f)
        except Exception as e:
            QMessageBox.critical(self, "Errore", f"Errore caricamento archivio:\n{str(e)}")
            self.ordini = []
            
        self.populate_tree()
        
    def populate_tree(self, filter_text=""):
        """Popola l'albero con gli ordini - TUTTO CHIUSO E ORDINATO CORRETTAMENTE"""
        self.tree.clear()
        
        if not hasattr(self, 'ordini'):
            return
        
        # Import locale per sicurezza
        from datetime import datetime
        from collections import defaultdict
        
        # ========== STEP 1: RAGGRUPPA ORDINI PER DATA ==========
        ordini_per_data = defaultdict(list)
        
        for i, ordine in enumerate(self.ordini):
            campi = ordine["campi"]
            nome = campi.get("Nome", "")
            cellulare = campi.get("Cellulare", "")
            
            # Filtro ricerca
            if filter_text and (filter_text.lower() not in nome.lower() and 
                filter_text not in cellulare):
                continue
                
            data = campi.get("Data evento", "")
            n_pers = campi.get("Numero persone", "")
            
            # Determina servizio
            servizio = self.determina_tipo_servizio(ordine)
            
            # Aggiungi all'array della data
            ordini_per_data[data].append((i, nome, cellulare, n_pers, servizio))
        
        # ========== STEP 2: ORDINAMENTO CORRETTO ==========
        def get_datetime(data_str):
            """Converte data in datetime per ordinamento"""
            try:
                return datetime.strptime(data_str, "%d/%m/%Y")
            except (ValueError, TypeError):
                return datetime.min  # Date malformate vanno in fondo

        # ========== ORDINAMENTO: DAL PIÙ LONTANO IN ALTO ==========
        # reverse=False mette le date più LONTANE PRIMA (in alto)
        date_ordinate = sorted(ordini_per_data.keys(), 
                            key=get_datetime, 
                            reverse=False)  # ✅ CAMBIATO: False = ordine cronologico normale

        print(f"🔍 DEBUG ORDINAMENTO CRONOLOGICO:")
        for data in date_ordinate[:5]:  # Mostra prime 5 date
            try:
                data_obj = datetime.strptime(data, "%d/%m/%Y")
                print(f"   📅 {data} ({data_obj.strftime('%A')})")
            except:
                print(f"   📅 {data}")
        
        # ========== STEP 3: CREA L'ALBERO ==========
        for data in date_ordinate:
            # Header senza conteggio
            try:
                data_obj = datetime.strptime(data, "%d/%m/%Y")
                data_formattata = data_obj.strftime("%A %d %B %Y").title()
                header_text = f"📅 {data_formattata}"
            except (ValueError, TypeError):
                header_text = f"📅 {data}"
            
            # Crea parent item per la data
            parent = QTreeWidgetItem(self.tree, [header_text])
            parent.setFont(0, QFont("Arial", 11, QFont.Weight.Bold))
            
            # ========== SEMPRE CHIUSO ==========
            parent.setExpanded(False)  # ✅ SEMPRE CHIUSO
            
            # Stile header
            parent.setBackground(0, QColor("#f8fafc"))
            parent.setForeground(0, QColor("#374151"))
            
            # Ordina ordini dentro la data per nome
            ordini_data = ordini_per_data[data]
            ordini_data.sort(key=lambda x: x[1].lower())
            
            # Crea item figli per ogni ordine
            for i, nome, cellulare, n_pers, servizio in ordini_data:
                item = QTreeWidgetItem(parent, [
                    "",
                    f"👤 {nome}",
                    f"📱 {cellulare}",
                    f"👥 {n_pers}",
                    f"🍽️ {servizio}"
                ])
                
                item.setData(0, Qt.ItemDataRole.UserRole, i)
                
                # Colori per tipo servizio
                if "Coffee" in servizio:
                    item.setBackground(4, QColor("#f0fdf4"))
                    item.setForeground(4, QColor("#166534"))
                if "Lunch Box" in servizio:
                    item.setBackground(4, QColor("#f0fdf4"))
                    item.setForeground(4, QColor("#166534"))
                elif "Tea" in servizio:
                    item.setBackground(4, QColor("#fdf4ff"))
                    item.setForeground(4, QColor("#7c2d12"))
                elif "Lunch" in servizio:
                    item.setBackground(4, QColor("#eff6ff"))
                    item.setForeground(4, QColor("#1e40af"))
                elif "Aperitivo" in servizio:
                    item.setBackground(4, QColor("#fff7ed"))
                    item.setForeground(4, QColor("#c2410c"))
        
    
    def determina_tipo_servizio(self, ordine):
        """Determina il tipo di servizio da un ordine - VERSIONE COMBINATA"""
        scelte = ordine.get("scelte", {})
        
        # Lista per raccogliere tutti i servizi trovati
        servizi_trovati = []
        
        # Verifica Coffee Break
        if any(k.startswith("coffee break|") and v for k, v in scelte.items()):
            servizi_trovati.append("Coffee Break")
        
        # Verifica Tea Break
        if any(k.startswith("tea break|") and v for k, v in scelte.items()):
            servizi_trovati.append("Tea Break")
            
        if any(k.startswith("lunch box|") and v for k, v in scelte.items()):
            servizi_trovati.append("Lunch Box")
        
        if any(k.startswith("servito|") and v for k, v in scelte.items()):
            servizi_trovati.append("Servito")
        
        # Verifica Lunch Buffet
        if any(k.startswith("lunch buffet|") and v for k, v in scelte.items()):
            servizi_trovati.append("Lunch Buffet")
        
        # Verifica Aperitivo
        if any(k.startswith("aperitivo a buffet|") and v for k, v in scelte.items()):
            tipo_ap = ordine.get("tipo_ap", "Leggero")
            servizi_trovati.append(f"Aperitivo {tipo_ap}")
        
        # Restituisci tutti i servizi combinati
        if len(servizi_trovati) == 0:
            return "Altro"
        elif len(servizi_trovati) == 1:
            return servizi_trovati[0]
        else:
            # Ordini combinati - mostra tutti i servizi
            return " + ".join(servizi_trovati)
    
    def filter_orders(self, text):
        """Filtra gli ordini in base al testo"""
        self.populate_tree(text)
        
    def load_selected_order(self):
        """Carica un singolo ordine per visualizzazione"""
        current = self.tree.currentItem()
        if not current or current.parent() is None:
            QMessageBox.information(self, "Selezione", "Seleziona un ordine specifico dalla lista.")
            return
            
        idx = current.data(0, Qt.ItemDataRole.UserRole)
        if idx is not None and 0 <= idx < len(self.ordini):
            ordine = self.ordini[idx]
            self.order_selected.emit(ordine)
            self.accept()
    
    def modify_selected_order(self):
        """Carica ordine per modifica - VERSIONE ARCHIVE DIALOG CORRETTA"""
        current = self.tree.currentItem()
        if not current or current.parent() is None:
            QMessageBox.information(self, "Selezione", "Seleziona un ordine specifico dalla lista.")
            return
            
        idx = current.data(0, Qt.ItemDataRole.UserRole)
        if idx is not None and 0 <= idx < len(self.ordini):
            ordine = self.ordini[idx]
            
            # Conferma modifica
            cliente = ordine["campi"].get("Nome", "Cliente")
            data = ordine["campi"].get("Data evento", "")
            
            reply = QMessageBox.question(
                self, "Conferma Modifica", 
                f"Vuoi modificare l'ordine di:\n\n"
                f"Cliente: {cliente}\n"
                f"Data: {data}\n\n"
                f"⚠️ Il documento Word esistente verrà SOVRASCRITTO\n"
                f"💡 L'archivio rimarrà aperto per vedere le modifiche",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                # ========== IMPOSTA MODALITÀ MODIFICA E MEMORIZZA INDICE ==========
                if hasattr(self.parent(), 'modalita_modifica'):
                    self.parent().modalita_modifica = True
                    self.parent().ordine_da_modificare = ordine
                    # CRUCIALE: Memorizza l'indice dell'ordine nell'archivio
                    self.parent().indice_ordine_originale = idx
                    print(f"💡 Memorizzato indice ordine originale: {idx}")
                    
                    # Debug aggiuntivo
                    print(f"🔍 Debug ordine da modificare:")
                    print(f"   - Nome: {cliente}")
                    print(f"   - Data: {data}")
                    print(f"   - Path: {ordine.get('path_docx', 'N/A')}")
                    print(f"   - Indice: {idx}")
                
                # CARICA I DATI SENZA CHIUDERE L'ARCHIVIO
                self.parent().load_order_data(ordine)
                
                # AVVISO
                QMessageBox.information(self, "Modalità Modifica Attivata", 
                                    f"✅ Ordine caricato per modifica!\n\n"
                                    f"👤 Cliente: {cliente}\n"
                                    f"📅 Data: {data}\n\n"
                                    f"🔧 Fai le tue modifiche nella finestra principale\n"
                                    f"📄 Premi 'Genera ordine' per salvare\n"
                                    f"👀 L'archivio si aggiornerà automaticamente")
    
    
    def delete_selected_orders(self):
        """Elimina ordini selezionati (singoli o multipli) - VERSIONE CORRETTA"""
        selected_items = self.tree.selectedItems()
        
        # Filtra solo gli ordini (non i gruppi per data)
        selected_orders = [item for item in selected_items if item.parent() is not None]
        
        if not selected_orders:
            QMessageBox.information(self, "Selezione", "Seleziona uno o più ordini da eliminare.")
            return
        
        count = len(selected_orders)
        
        # Prepara info per conferma
        if count == 1:
            # Singola eliminazione
            item = selected_orders[0]
            idx = item.data(0, Qt.ItemDataRole.UserRole)
            if idx is None or idx >= len(self.ordini):
                QMessageBox.warning(self, "Errore", "Ordine non valido.")
                return
                
            ordine = self.ordini[idx]
            cliente = ordine["campi"].get("Nome", "Cliente sconosciuto")
            data = ordine["campi"].get("Data evento", "Data sconosciuta")
            
            messaggio = f"Vuoi eliminare questo ordine?\n\n" \
                    f"👤 Cliente: {cliente}\n" \
                    f"📅 Data: {data}"
        else:
            # Eliminazione multipla
            clienti_list = []
            for item in selected_orders[:5]:  # Mostra max 5 nomi
                idx = item.data(0, Qt.ItemDataRole.UserRole)
                if idx is not None and idx < len(self.ordini):
                    ordine = self.ordini[idx]
                    cliente = ordine["campi"].get("Nome", "Sconosciuto")
                    data = ordine["campi"].get("Data evento", "")
                    clienti_list.append(f"• {cliente} ({data})")
            
            clienti_text = "\n".join(clienti_list)
            if count > 5:
                clienti_text += f"\n... e altri {count - 5} ordini"
            
            messaggio = f"Vuoi eliminare {count} ordini?\n\n{clienti_text}"
        
        # Conferma eliminazione
        reply = QMessageBox.question(
            self, f"⚠️ CONFERMA ELIMINAZIONE {count} ORDINI" if count > 1 else "⚠️ CONFERMA ELIMINAZIONE", 
            f"{messaggio}\n\n"
            f"⚠️ ATTENZIONE:\n"
            f"• Gli ordini verranno rimossi dall'archivio\n"
            f"• I file Word verranno spostati nel cestino\n"
            f"• Questa operazione NON può essere annullata\n\n"
            f"Vuoi procedere?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self.esegui_eliminazione_multipla(selected_orders)

    def esegui_eliminazione_multipla(self, selected_orders):
        """Esegue eliminazione multipla con progress"""
        import shutil
        
        count = len(selected_orders)
        
        # Progress dialog per eliminazioni multiple
        if count > 1:
            progress = QProgressBar(self)
            progress.setWindowTitle("Eliminazione in corso...")
            progress.setGeometry(self.width()//2 - 150, self.height()//2, 300, 30)
            progress.setRange(0, count)
            progress.show()
            QApplication.processEvents()
        
        eliminati = 0
        errori = []
        files_spostati = 0
        
        # Raccogli indici da eliminare (dal più alto al più basso per evitare shift)
        indici_da_eliminare = []
        for item in selected_orders:
            idx = item.data(0, Qt.ItemDataRole.UserRole)
            if idx is not None and 0 <= idx < len(self.ordini):
                indici_da_eliminare.append(idx)
        
        indici_da_eliminare.sort(reverse=True)  # Dal più alto al più basso
        
        try:
            # Crea cartella cestino
            cestino_dir = "ordini_eliminati"
            os.makedirs(cestino_dir, exist_ok=True)
            
            for i, idx in enumerate(indici_da_eliminare):
                try:
                    ordine = self.ordini[idx]
                    cliente = ordine["campi"].get("Nome", f"Ordine_{idx}")
                    
                    # Sposta file Word se esiste
                    path_docx = ordine.get("path_docx")
                    if path_docx and os.path.exists(path_docx):
                        nome_file = os.path.basename(path_docx)
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        nuovo_nome = f"{timestamp}_{nome_file}"
                        cestino_path = os.path.join(cestino_dir, nuovo_nome)
                        shutil.move(path_docx, cestino_path)
                        files_spostati += 1
                    
                    # Rimuovi dall'archivio
                    del self.ordini[idx]
                    eliminati += 1
                    
                    if count > 1:
                        progress.setValue(i + 1)
                        QApplication.processEvents()
                    
                except Exception as e:
                    errori.append(f"• {cliente}: {str(e)}")
            
            # Salva archivio aggiornato
            with open(ORDINI_FILE, "w", encoding="utf-8") as f:
                json.dump(self.ordini, f, indent=2, ensure_ascii=False)
            
            # Aggiorna interfaccia
            self.populate_tree()
            
            if count > 1:
                progress.hide()
            
            # Messaggio risultato
            messaggio = f"✅ Eliminazione completata!\n\n" \
                    f"📋 Ordini eliminati: {eliminati}/{count}\n" \
                    f"📄 File spostati nel cestino: {files_spostati}"
            
            if errori:
                messaggio += f"\n\n⚠️ Errori durante l'eliminazione:\n" + "\n".join(errori)
            
            QMessageBox.information(self, "Eliminazione Completata", messaggio)
            
        except Exception as e:
            if count > 1 and 'progress' in locals():
                progress.hide()
            
            QMessageBox.critical(
                self, "Errore Eliminazione", 
                f"❌ Errore durante l'eliminazione multipla:\n\n{str(e)}"
            )

    def closeEvent(self, event):
        """Rimuovi dalla lista quando si chiude"""
        global archivio_dialogs_aperti
        if self in archivio_dialogs_aperti:
            archivio_dialogs_aperti.remove(self)
            print(f"📋 Archivio rimosso. Totale archivi aperti: {len(archivio_dialogs_aperti)}")
        event.accept()
        
# ===================================================================
# AGGIUNGI QUESTA CLASSE SUBITO DOPO GLI IMPORT
# ===================================================================

from PyQt6.QtGui import QPainter, QPen, QColor
from PyQt6.QtCore import QRect

class CustomCheckBox(QCheckBox):
    """Checkbox personalizzato con X blu invece del check standard"""
    
    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setStyleSheet("""
            QCheckBox {
                spacing: 8px;
                font-size: 13px;
                color: #4a5568;
                padding: 6px;
                margin: 2px;
            }
            
            QCheckBox::indicator {
                width: 18px;
                height: 18px;
                border: 2px solid #d1d5db;
                border-radius: 4px;
                background-color: white;
            }
            
            QCheckBox::indicator:hover {
                border-color: #3b82f6;
                background-color: #f8fafc;
            }
            
            QCheckBox::indicator:checked {
                background-color: #3b82f6;
                border-color: #3b82f6;
            }
            
            QCheckBox::indicator:checked:hover {
                background-color: #2563eb;
                border-color: #2563eb;
            }
            
            QCheckBox:checked {
                font-weight: 600;
                color: #1f2937;
            }
        """)
    
    def paintEvent(self, event):
        """Override per disegnare la X personalizzata"""
        super().paintEvent(event)
        
        if self.isChecked():
            painter = QPainter(self)
            painter.setRenderHint(QPainter.RenderHint.Antialiasing)
            
            # Calcola la posizione dell'indicatore con il metodo corretto per PyQt6
            from PyQt6.QtWidgets import QStyleOptionButton
            
            option = QStyleOptionButton()
            self.initStyleOption(option)
            
            style = self.style()
            indicator_rect = style.subElementRect(
                style.SubElement.SE_CheckBoxIndicator, 
                option, 
                self
            )
            
            # Se il rect non è valido, usa calcolo manuale migliorato
            if indicator_rect.isEmpty():
                # Spacing standard di Qt per checkbox
                spacing = 4
                indicator_size = 18
                
                indicator_rect = QRect(
                    spacing,  # Margine sinistro standard
                    (self.height() - indicator_size) // 2,
                    indicator_size,
                    indicator_size
                )
            
            # Disegna X bianca al centro dell'indicatore
            pen = QPen(QColor(255, 255, 255), 2.5)
            pen.setCapStyle(Qt.PenCapStyle.RoundCap)
            painter.setPen(pen)
            
            # Calcola le coordinate per la X centrata nell'indicatore
            center_x = indicator_rect.center().x()
            center_y = indicator_rect.center().y()
            offset = 5  # Aumento l'offset per una X più grande
            
            # Linea 1 della X (da top-left a bottom-right)
            painter.drawLine(
                center_x - offset, center_y - offset,
                center_x + offset, center_y + offset
            )
            
            # Linea 2 della X (da top-right a bottom-left)
            painter.drawLine(
                center_x + offset, center_y - offset,
                center_x - offset, center_y + offset
            )

            
class EventBadge(QFrame):
    """Badge per rappresentare un singolo evento nel calendario"""
    clicked = pyqtSignal(dict)
    
    def __init__(self, ordine_data, parent=None):
        super().__init__(parent)
        self.ordine_data = ordine_data
        self.setup_ui()
        self.setup_style()
        
    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(2, 1, 2, 1)
        layout.setSpacing(0)
        
        # Nome cliente (troncato)
        nome = self.ordine_data["campi"].get("Nome", "")
        nome_short = nome[:12] + "..." if len(nome) > 12 else nome
        
        nome_label = QLabel(nome_short)
        nome_label.setFont(QFont("Arial", 8, QFont.Weight.Bold))
        nome_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        # Persone
        persone = self.ordine_data["campi"].get("Numero persone", "0")
        persone_label = QLabel(f"👥{persone}")
        persone_label.setFont(QFont("Arial", 7))
        persone_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        layout.addWidget(nome_label)
        layout.addWidget(persone_label)
        
        # Tooltip dettagliato
        self.setup_tooltip()
        
    def setup_style(self):
        """Applica colore basato sul tipo di servizio"""
        servizio = self.determina_servizio()
        persone = int(self.ordine_data["campi"].get("Numero persone", "0"))
        in_fiera = self.ordine_data.get("evento_in_fiera", False)
        
        # Colori base per servizio
        colori = {
            "Coffee Break": "#10b981",      # Verde
            "Tea Break": "#8b5cf6",         # Viola  
            "Lunch Buffet": "#3b82f6",      # Blu
            "Lunch Box": "#f59e0b", 
            "Aperitivo Leggero": "#f59e0b", # Giallo/Arancio
            "Aperitivo Rinforzato": "#ef4444", # Rosso/Arancio
            "Altro": "#6b7280"              # Grigio
        }
        
        colore_base = colori.get(servizio, "#6b7280")
        
        # Intensità basata su numero persone (più persone = più scuro)
        if persone <= 20:
            opacity = "0.7"
        elif persone <= 50:
            opacity = "0.85"
        else:
            opacity = "1.0"
        
        # Bordo speciale per eventi in fiera
        border_style = "3px solid #ff6b35" if in_fiera else "1px solid #e5e7eb"
        
        self.setStyleSheet(f"""
            EventBadge {{
                background-color: {colore_base};
                opacity: {opacity};
                border: {border_style};
                border-radius: 4px;
                margin: 1px;
                max-height: 35px;
                min-height: 35px;
            }}
            EventBadge:hover {{
                opacity: 1.0;
                border: 2px solid #1f2937;
                transform: scale(1.05);
            }}
            QLabel {{
                color: white; 
                background: transparent;
                border: none;
            }}
        """)
        
    def setup_tooltip(self):
        """Crea tooltip dettagliato"""
        campi = self.ordine_data["campi"]
        servizio = self.determina_servizio()
        
        tooltip_text = f"""
<b>🏷️ {campi.get('Nome', 'Cliente')}</b><br>
📅 {campi.get('Data evento', '')}<br>
👥 {campi.get('Numero persone', '0')} persone<br>
🍽️ {servizio}<br>
📞 {campi.get('Cellulare', 'N/A')}<br>
📍 {campi.get('Via', 'N/A')}<br>
⏰ {campi.get('Orario allestimento', 'N/A')} - {campi.get('Orario Pronti', 'N/A')}
        """.strip()
        
        if self.ordine_data.get("evento_in_fiera", False):
            tooltip_text += "<br>🏢 <b>EVENTO IN FIERA</b>"
            
        self.setToolTip(tooltip_text)
        
    def determina_servizio(self):
        """Determina il tipo di servizio - AGGIORNATO CON SERVITO"""
        scelte = self.ordine_data.get("scelte", {})
        
        if any(k.startswith("coffee break|") and v for k, v in scelte.items()):
            return "Coffee Break"
        elif any(k.startswith("tea break|") and v for k, v in scelte.items()):
            return "Tea Break"
        elif any(k.startswith("lunch buffet|") and v for k, v in scelte.items()):
            return "Lunch Buffet"
        elif any(k.startswith("lunch box|") and v for k, v in scelte.items()):
            return "Lunch box"
        elif any(k.startswith("servito|") and v for k, v in scelte.items()):  # ← NUOVO
            return "Servito"
        elif any(k.startswith("aperitivo a buffet|") and v for k, v in scelte.items()):
            tipo_ap = self.ordine_data.get("tipo_ap", "Leggero")
            return f"Aperitivo {tipo_ap}"
        else:
            return "Altro"
    
    def mousePressEvent(self, event):
        """Gestisce click sul badge"""
        if event.button() == Qt.MouseButton.LeftButton:
            self.clicked.emit(self.ordine_data)
        super().mousePressEvent(event)

class DayCell(QFrame):
    """Cella per un singolo giorno del calendario"""
    day_clicked = pyqtSignal(int, list)
    
    def __init__(self, day_number, ordini_giorno=None, is_current_month=True):
        super().__init__()
        self.day_number = day_number
        self.ordini_giorno = ordini_giorno or []
        self.is_current_month = is_current_month
        
        # IMPORTANTE: Imposta object name per lo stile CSS
        self.setObjectName("dayCell")
        
        # Imposta proprietà per giorno corrente
        oggi = datetime.now().day
        if self.day_number == oggi and self.is_current_month:
            self.setProperty("isToday", "true")
        else:
            self.setProperty("isToday", "false")
        
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(4, 4, 4, 4)
        layout.setSpacing(2)
        
        # Header giorno con stile premium
        day_header = QLabel(str(self.day_number))
        day_header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        day_header.setFont(QFont("Inter", 11, QFont.Weight.Bold))
        
        # Stile per il numero del giorno
        oggi = datetime.now().day
        if self.day_number == oggi and self.is_current_month:
            day_header.setStyleSheet("""
                QLabel {
                    background-color: #ef4444;
                    color: white;
                    border-radius: 12px;
                    padding: 4px 8px;
                    font-weight: bold;
                }
            """)
        elif not self.is_current_month:
            day_header.setStyleSheet("""
                QLabel {
                    color: #9ca3af;
                    font-weight: normal;
                }
            """)
        else:
            day_header.setStyleSheet("""
                QLabel {
                    color: #374151;
                    font-weight: 600;
                }
            """)
        
        layout.addWidget(day_header)
        
        # Container eventi con scroll automatico
        events_widget = QWidget()
        events_layout = QVBoxLayout(events_widget)
        events_layout.setContentsMargins(0, 0, 0, 0)
        events_layout.setSpacing(1)
        
        # Badge eventi premium (max 3 visibili)
        for i, ordine in enumerate(self.ordini_giorno[:3]):
            badge = EventBadge(ordine)
            badge.setObjectName("eventBadge")  # Per lo stile CSS
            events_layout.addWidget(badge)
        
        # Indicatore eventi extra
        if len(self.ordini_giorno) > 3:
            more_label = QLabel(f"+{len(self.ordini_giorno) - 3}")
            more_label.setStyleSheet("""
                QLabel {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                        stop:0 #6b7280, stop:1 #4b5563);
                    color: white;
                    border-radius: 4px;
                    padding: 2px 6px;
                    font-size: 9px;
                    font-weight: 600;
                    text-align: center;
                }
            """)
            more_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            events_layout.addWidget(more_label)
        
        events_layout.addStretch()
        layout.addWidget(events_widget)
        
        # Dimensioni ottimizzate
        self.setMinimumHeight(80)
        self.setMaximumHeight(110)
        
        # Forza aggiornamento dello stile
        self.style().polish(self)
        
    def mousePressEvent(self, event):
        """Gestisce click sulla cella del giorno"""
        if event.button() == Qt.MouseButton.LeftButton:
            self.day_clicked.emit(self.day_number, self.ordini_giorno)
        super().mousePressEvent(event)
class CalendarioEventiProfessionale(QWidget):
    """Versione professionale del calendario con design premium"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.current_date = datetime.now()
        self.ordini_data = []
        self.view_mode = "month"  # month, week
        self.filter_servizio = "tutti"
        self.setup_professional_ui()
        self.load_ordini()
        self.refresh_calendar()
        
    def setup_professional_ui(self):
        """Setup interfaccia professionale premium"""
        main_layout = QHBoxLayout(self)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)
        
        # ========== SIDEBAR SINISTRA ==========
        self.create_sidebar(main_layout)
        
        # ========== AREA CALENDARIO PRINCIPALE ==========
        calendar_container = QWidget()
        calendar_container.setStyleSheet("""
            QWidget {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #f8fafc, stop:1 #e2e8f0);
                border-radius: 12px;
            }
        """)
        
        calendar_layout = QVBoxLayout(calendar_container)
        calendar_layout.setContentsMargins(25, 25, 25, 25)
        calendar_layout.setSpacing(20)
        
        # Header professionale
        self.create_professional_header(calendar_layout)
        
        # Controlli vista
        self.create_view_controls(calendar_layout)
        
        # Calendario principale
        self.create_main_calendar(calendar_layout)
        
        main_layout.addWidget(calendar_container, stretch=4)
        
    def create_sidebar(self, parent_layout):
        """Crea sidebar con info e controlli"""
        sidebar = QWidget()
        sidebar.setMaximumWidth(280)
        sidebar.setStyleSheet("""
            QWidget {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #1e293b, stop:1 #334155);
                border-top-left-radius: 12px;
                border-bottom-left-radius: 12px;
            }
        """)
        
        sidebar_layout = QVBoxLayout(sidebar)
        sidebar_layout.setContentsMargins(20, 25, 20, 25)
        sidebar_layout.setSpacing(25)
        
        # Logo/Titolo
        title = QLabel("📅 CALENDARIO\nORDINI")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setStyleSheet("""
            QLabel {
                color: white;
                font-size: 18px;
                font-weight: bold;
                letter-spacing: 1px;
                padding: 15px;
                background: rgba(255,255,255,0.1);
                border-radius: 8px;
            }
        """)
        sidebar_layout.addWidget(title)
        
        # Mini calendario navigator
        self.create_mini_calendar(sidebar_layout)
        
        # Filtri
        self.create_filters(sidebar_layout)
        
        # Statistiche rapide
        self.create_quick_stats(sidebar_layout)
        
        # Actions
        self.create_quick_actions(sidebar_layout)
        
        sidebar_layout.addStretch()
        parent_layout.addWidget(sidebar)
    
    def create_professional_header(self, parent_layout):
        """Header con design professionale"""
        header_container = QWidget()
        header_container.setStyleSheet("""
            QWidget {
                background: white;
                border: 1px solid #e2e8f0;
                border-radius: 12px;
                padding: 10px;
            }
        """)
        
        header_layout = QHBoxLayout(header_container)
        header_layout.setContentsMargins(20, 15, 20, 15)
        
        # Navigazione mese
        nav_container = QWidget()
        nav_layout = QHBoxLayout(nav_container)
        nav_layout.setContentsMargins(0, 0, 0, 0)
        nav_layout.setSpacing(15)
        
        # Pulsanti eleganti
        self.prev_btn = self.create_nav_button("❮", self.prev_month)
        self.next_btn = self.create_nav_button("❯", self.next_month)
        
        # Titolo mese elegante
        self.month_label = QLabel()
        self.month_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.month_label.setStyleSheet("""
            QLabel {
                font-size: 24px;
                font-weight: 700;
                color: #1e293b;
                letter-spacing: 1px;
                padding: 10px 20px;
            }
        """)
        
        nav_layout.addWidget(self.prev_btn)
        nav_layout.addWidget(self.month_label, stretch=1)
        nav_layout.addWidget(self.next_btn)
        
        header_layout.addWidget(nav_container)
        
        # Info rapide a destra
        info_container = QWidget()
        info_layout = QVBoxLayout(info_container)
        info_layout.setContentsMargins(0, 0, 0, 0)
        
        self.month_stats = QLabel("Caricamento...")
        self.month_stats.setStyleSheet("""
            QLabel {
                color: #64748b;
                font-size: 12px;
                font-weight: 500;
            }
        """)
        info_layout.addWidget(self.month_stats)
        
        header_layout.addWidget(info_container)
        parent_layout.addWidget(header_container)
    
    def create_nav_button(self, text, callback):
        """Crea pulsanti di navigazione eleganti"""
        btn = QPushButton(text)
        btn.clicked.connect(callback)
        btn.setStyleSheet("""
            QPushButton {
                background: #f1f5f9;
                border: 2px solid #e2e8f0;
                border-radius: 8px;
                color: #475569;
                font-size: 16px;
                font-weight: bold;
                padding: 8px 12px;
                min-width: 40px;
            }
            QPushButton:hover {
                background: #3b82f6;
                border-color: #3b82f6;
                color: white;
                transform: translateY(-1px);
            }
            QPushButton:pressed {
                transform: translateY(0px);
            }
        """)
        return btn
    
    def create_view_controls(self, parent_layout):
        """Controlli per cambiare vista"""
        controls_container = QWidget()
        controls_container.setStyleSheet("""
            QWidget {
                background: white;
                border: 1px solid #e2e8f0;
                border-radius: 10px;
                padding: 5px;
            }
        """)
        
        controls_layout = QHBoxLayout(controls_container)
        controls_layout.setContentsMargins(15, 10, 15, 10)
        
        # Toggle vista
        view_group = QWidget()
        view_layout = QHBoxLayout(view_group)
        view_layout.setContentsMargins(0, 0, 0, 0)
        view_layout.setSpacing(5)
        
        month_btn = QPushButton("📅 Mese")
        week_btn = QPushButton("📋 Settimana")
        
        for btn in [month_btn, week_btn]:
            btn.setStyleSheet("""
                QPushButton {
                    background: #f8fafc;
                    border: 1px solid #e2e8f0;
                    border-radius: 6px;
                    padding: 8px 16px;
                    font-weight: 600;
                    color: #475569;
                }
                QPushButton:hover {
                    background: #e2e8f0;
                }
                QPushButton:checked {
                    background: #3b82f6;
                    color: white;
                    border-color: #3b82f6;
                }
            """)
            btn.setCheckable(True)
        
        month_btn.setChecked(True)
        view_layout.addWidget(month_btn)
        view_layout.addWidget(week_btn)
        
        controls_layout.addWidget(view_group)
        controls_layout.addStretch()
        
        # Oggi button elegante
        today_btn = QPushButton("Oggi")
        today_btn.clicked.connect(self.go_to_today)
        today_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #ef4444, stop:1 #dc2626);
                border: none;
                border-radius: 8px;
                color: white;
                font-weight: 600;
                padding: 8px 16px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #dc2626, stop:1 #b91c1c);
            }
        """)
        
        controls_layout.addWidget(today_btn)
        parent_layout.addWidget(controls_container)
    
    def create_main_calendar(self, parent_layout):
        """Calendario principale con design premium"""
        calendar_container = QWidget()
        calendar_container.setStyleSheet("""
            QWidget {
                background: white;
                border: 1px solid #e2e8f0;
                border-radius: 12px;
            }
        """)
        
        calendar_layout = QVBoxLayout(calendar_container)
        calendar_layout.setContentsMargins(20, 20, 20, 20)
        calendar_layout.setSpacing(15)
        
        # Header giorni settimana elegante
        days_header = QHBoxLayout()
        giorni = ["LUN", "MAR", "MER", "GIO", "VEN", "SAB", "DOM"]
        
        for i, giorno in enumerate(giorni):
            day_label = QLabel(giorno)
            day_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            day_label.setStyleSheet(f"""
                QLabel {{
                    background: {'#fef3f2' if i >= 5 else '#f8fafc'};
                    color: {'#dc2626' if i >= 5 else '#475569'};
                    font-size: 12px;
                    font-weight: 700;
                    padding: 12px 8px;
                    border-radius: 8px;
                    letter-spacing: 0.5px;
                }}
            """)
            days_header.addWidget(day_label)
        
        calendar_layout.addLayout(days_header)
        
        # Griglia calendario professionale
        self.calendar_grid = QGridLayout()
        self.calendar_grid.setSpacing(8)
        calendar_layout.addLayout(self.calendar_grid)
        
        parent_layout.addWidget(calendar_container, stretch=1)

# Dialog finestra principale per integrare il calendario
# =============================================================================
# AGGIUNGI QUESTE CLASSI AL TUO FILE PRINCIPALE (cateringFunzionante.py)
# SUBITO DOPO LE IMPORT E PRIMA DELLA CLASSE CateringApp
# =============================================================================

from datetime import datetime, timedelta
from collections import defaultdict

# ===================================================================
# AGGIUNGI QUESTO STILE SUBITO DOPO PROFESSIONAL_STYLE
# ===================================================================

PROFESSIONAL_CALENDAR_STYLE = """
/* Base */
QWidget {
    font-family: 'Inter', 'Segoe UI', Arial, sans-serif;
    font-size: 13px;
}

/* Calendario principale */
QMainWindow {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
        stop:0 #f1f5f9, stop:1 #e2e8f0);
}

/* Celle giorni premium */
QFrame[objectName="dayCell"] {
    background: white;
    border: 1px solid #f1f5f9;
    border-radius: 8px;
    margin: 2px;
}

QFrame[objectName="dayCell"]:hover {
    border: 2px solid #3b82f6;
    background: #fefeff;
    transform: scale(1.02);
}

QFrame[objectName="dayCell"][isToday="true"] {
    border: 2px solid #ef4444;
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 #fef2f2, stop:1 #fecaca);
}

/* Badge eventi eleganti */
QFrame[objectName="eventBadge"] {
    border-radius: 4px;
    margin: 1px;
    padding: 2px 6px;
    font-size: 10px;
    font-weight: 600;
}

/* Dialog premium */
QDialog {
    background: white;
    border: 1px solid #e2e8f0;
    border-radius: 12px;
}

/* Buttons premium */
QPushButton {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 #f8fafc, stop:1 #e2e8f0);
    border: 1px solid #d1d5db;
    border-radius: 8px;
    color: #374151;
    font-weight: 600;
    padding: 10px 20px;
    min-height: 20px;
}

QPushButton:hover {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 #3b82f6, stop:1 #2563eb);
    border-color: #3b82f6;
    color: white;
    transform: translateY(-1px);
}

QPushButton:pressed {
    transform: translateY(0px);
}

/* Shadows and effects */
QFrame[objectName="card"] {
    background: white;
    border: 1px solid #e2e8f0;
    border-radius: 10px;
    /* box-shadow: 0 1px 3px rgba(0,0,0,0.1); */
}
"""

# ===================================================================
# AGGIORNA LA CLASSE DayCell PER USARE L'OBJECT NAME
# ===================================================================

class DayCell(QFrame):
    """Cella per un singolo giorno del calendario"""
    day_clicked = pyqtSignal(int, list)
    
    def __init__(self, day_number, ordini_giorno=None, is_current_month=True):
        super().__init__()
        self.day_number = day_number
        self.ordini_giorno = ordini_giorno or []
        self.is_current_month = is_current_month
        
        # IMPORTANTE: Imposta object name per lo stile CSS
        self.setObjectName("dayCell")
        
        # Imposta proprietà per giorno corrente
        oggi = datetime.now().day
        if self.day_number == oggi and self.is_current_month:
            self.setProperty("isToday", "true")
        else:
            self.setProperty("isToday", "false")
        
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(4, 4, 4, 4)
        layout.setSpacing(2)
        
        # Header giorno con stile premium
        day_header = QLabel(str(self.day_number))
        day_header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        day_header.setFont(QFont("Inter", 11, QFont.Weight.Bold))
        
        # Stile per il numero del giorno
        oggi = datetime.now().day
        if self.day_number == oggi and self.is_current_month:
            day_header.setStyleSheet("""
                QLabel {
                    background-color: #ef4444;
                    color: white;
                    border-radius: 12px;
                    padding: 4px 8px;
                    font-weight: bold;
                }
            """)
        elif not self.is_current_month:
            day_header.setStyleSheet("""
                QLabel {
                    color: #9ca3af;
                    font-weight: normal;
                }
            """)
        else:
            day_header.setStyleSheet("""
                QLabel {
                    color: #374151;
                    font-weight: 600;
                }
            """)
        
        layout.addWidget(day_header)
        
        # Container eventi con scroll automatico
        events_widget = QWidget()
        events_layout = QVBoxLayout(events_widget)
        events_layout.setContentsMargins(0, 0, 0, 0)
        events_layout.setSpacing(1)
        
        # Badge eventi premium (max 3 visibili)
        for i, ordine in enumerate(self.ordini_giorno[:3]):
            badge = EventBadge(ordine)
            badge.setObjectName("eventBadge")  # Per lo stile CSS
            events_layout.addWidget(badge)
        
        # Indicatore eventi extra
        if len(self.ordini_giorno) > 3:
            more_label = QLabel(f"+{len(self.ordini_giorno) - 3}")
            more_label.setStyleSheet("""
                QLabel {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                        stop:0 #6b7280, stop:1 #4b5563);
                    color: white;
                    border-radius: 4px;
                    padding: 2px 6px;
                    font-size: 9px;
                    font-weight: 600;
                    text-align: center;
                }
            """)
            more_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            events_layout.addWidget(more_label)
        
        events_layout.addStretch()
        layout.addWidget(events_widget)
        
        # Dimensioni ottimizzate
        self.setMinimumHeight(80)
        self.setMaximumHeight(110)
        
        # Forza aggiornamento dello stile
        self.style().polish(self)
        
    def mousePressEvent(self, event):
        """Gestisce click sulla cella del giorno"""
        if event.button() == Qt.MouseButton.LeftButton:
            self.day_clicked.emit(self.day_number, self.ordini_giorno)
        super().mousePressEvent(event)

# ===================================================================
# AGGIORNA LA CLASSE EventBadge PER STILE PREMIUM
# ===================================================================

class EventBadge(QFrame):
    """Badge per rappresentare un singolo evento nel calendario"""
    clicked = pyqtSignal(dict)
    
    def __init__(self, ordine_data, parent=None):
        super().__init__(parent)
        self.ordine_data = ordine_data
        
        # IMPORTANTE: Imposta object name per lo stile CSS
        self.setObjectName("eventBadge")
        
        self.setup_ui()
        self.setup_premium_style()
        
    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(3, 2, 3, 2)
        layout.setSpacing(0)
        
        # Nome cliente con font premium
        nome = self.ordine_data["campi"].get("Nome", "")
        nome_short = nome[:10] + "..." if len(nome) > 10 else nome
        
        nome_label = QLabel(nome_short)
        nome_label.setFont(QFont("Inter", 8, QFont.Weight.Bold))
        nome_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        # Numero persone con icona
        persone = self.ordine_data["campi"].get("Numero persone", "0")
        persone_label = QLabel(f"👥{persone}")
        persone_label.setFont(QFont("Inter", 7, QFont.Weight.Medium))
        persone_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        layout.addWidget(nome_label)
        layout.addWidget(persone_label)
        
        # Tooltip migliorato
        self.setup_premium_tooltip()
        
    def setup_premium_style(self):
        """Applica stile premium basato sul servizio"""
        servizio = self.determina_servizio()
        persone = int(self.ordine_data["campi"].get("Numero persone", "0"))
        in_fiera = self.ordine_data.get("evento_in_fiera", False)
        
        # Palette colori premium
        colori_premium = {
            "Coffee Break": {
                "bg": "qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #10b981, stop:1 #059669)",
                "border": "#047857"
            },
            "Tea Break": {
                "bg": "qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #8b5cf6, stop:1 #7c3aed)",
                "border": "#6d28d9"
            },
            "Lunch Buffet": {
                "bg": "qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3b82f6, stop:1 #2563eb)",
                "border": "#1d4ed8"
            },
            "Aperitivo Leggero": {
                "bg": "qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #f59e0b, stop:1 #d97706)",
                "border": "#b45309"
            },
            "Aperitivo Rinforzato": {
                "bg": "qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #ef4444, stop:1 #dc2626)",
                "border": "#b91c1c"
            },
            "Altro": {
                "bg": "qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #6b7280, stop:1 #4b5563)",
                "border": "#374151"
            }
        }
        
        style_info = colori_premium.get(servizio, colori_premium["Altro"])
        
        # Opacità basata su numero persone
        if persone <= 20:
            opacity = "0.85"
        elif persone <= 50:
            opacity = "0.95"
        else:
            opacity = "1.0"
        
        # Bordo speciale per eventi in fiera
        border_extra = "border: 2px solid #ff6b35;" if in_fiera else ""
        
        # Applica stile premium
        self.setStyleSheet(f"""
            EventBadge {{
                background: {style_info['bg']};
                opacity: {opacity};
                border: 1px solid {style_info['border']};
                border-radius: 6px;
                margin: 1px;
                max-height: 32px;
                min-height: 32px;
                {border_extra}
            }}
            EventBadge:hover {{
                opacity: 1.0;
                border: 2px solid #1f2937;
                transform: scale(1.05);
                background: {style_info['bg']};
            }}
            QLabel {{
                color: white;
                background: transparent;
                border: none;
                text-shadow: 0px 1px 2px rgba(0,0,0,0.3);
            }}
        """)
        
    def setup_premium_tooltip(self):
        """Tooltip con design premium"""
        campi = self.ordine_data["campi"]
        servizio = self.determina_servizio()
        
        tooltip_html = f"""
        <div style="
            background: white;
            border: 1px solid #e2e8f0;
            border-radius: 8px;
            padding: 12px;
            font-family: 'Inter', sans-serif;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        ">
            <h3 style="
                margin: 0 0 8px 0;
                color: #1f2937;
                font-size: 14px;
                font-weight: 600;
            ">🏷️ {campi.get('Nome', 'Cliente')}</h3>
            
            <div style="color: #6b7280; font-size: 12px; line-height: 1.4;">
                📅 <strong>{campi.get('Data evento', '')}</strong><br>
                👥 {campi.get('Numero persone', '0')} persone<br>
                🍽️ {servizio}<br>
                📞 {campi.get('Cellulare', 'N/A')}<br>
                📍 {campi.get('Via', 'N/A')}<br>
                ⏰ {campi.get('Orario allestimento', 'N/A')} - {campi.get('Orario Pronti', 'N/A')}
            </div>
        """.strip()
        
        if self.ordine_data.get("evento_in_fiera", False):
            tooltip_html += """
            <div style="
                margin-top: 8px;
                padding: 4px 8px;
                background: #fef3f2;
                border: 1px solid #fecaca;
                border-radius: 4px;
                color: #dc2626;
                font-weight: 600;
                font-size: 11px;
            ">
                🏢 EVENTO IN FIERA
            </div>
            """
        
        tooltip_html += "</div>"
        self.setToolTip(tooltip_html)
        
    def determina_servizio(self):
        """Determina il tipo di servizio - AGGIORNATO CON SERVITO"""
        scelte = self.ordine_data.get("scelte", {})
        
        if any(k.startswith("coffee break|") and v for k, v in scelte.items()):
            return "Coffee Break"
        elif any(k.startswith("tea break|") and v for k, v in scelte.items()):
            return "Tea Break"
        elif any(k.startswith("lunch buffet|") and v for k, v in scelte.items()):
            return "Lunch Buffet"
        elif any(k.startswith("lunch box|") and v for k, v in scelte.items()):
            return "Lunch box"
        elif any(k.startswith("servito|") and v for k, v in scelte.items()):  # ← NUOVO
            return "Servito"
        elif any(k.startswith("aperitivo a buffet|") and v for k, v in scelte.items()):
            tipo_ap = self.ordine_data.get("tipo_ap", "Leggero")
            return f"Aperitivo {tipo_ap}"
        else:
            return "Altro"
    
    def mousePressEvent(self, event):
        """Gestisce click sul badge"""
        if event.button() == Qt.MouseButton.LeftButton:
            self.clicked.emit(self.ordine_data)
        super().mousePressEvent(event)

class DayCell(QFrame):
    """Cella per un singolo giorno del calendario"""
    day_clicked = pyqtSignal(int, list)
    
    def __init__(self, day_number, ordini_giorno=None, is_current_month=True):
        super().__init__()
        self.day_number = day_number
        self.ordini_giorno = ordini_giorno or []
        self.is_current_month = is_current_month
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(2, 2, 2, 2)
        layout.setSpacing(1)
        
        # Header giorno
        day_header = QLabel(str(self.day_number))
        day_header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        day_header.setFont(QFont("Arial", 10, QFont.Weight.Bold))
        
        # Stile per giorno corrente
        oggi = datetime.now().day
        if self.day_number == oggi and self.is_current_month:
            day_header.setStyleSheet("""
                QLabel {
                    background-color: #3b82f6;
                    color: white;
                    border-radius: 10px;
                    padding: 2px;
                }
            """)
        elif not self.is_current_month:
            day_header.setStyleSheet("color: #9ca3af;")
        
        layout.addWidget(day_header)
        
        # Container per eventi
        events_widget = QWidget()
        events_layout = QVBoxLayout(events_widget)
        events_layout.setContentsMargins(0, 0, 0, 0)
        events_layout.setSpacing(1)
        
        # Aggiungi badge per ogni ordine (max 3 visibili)
        for i, ordine in enumerate(self.ordini_giorno[:3]):
            badge = EventBadge(ordine)
            events_layout.addWidget(badge)
        
        # Indicatore "altri eventi"
        if len(self.ordini_giorno) > 3:
            more_label = QLabel(f"+{len(self.ordini_giorno) - 3}")
            more_label.setStyleSheet("""
                QLabel {
                    background-color: #6b7280;
                    color: white;
                    border-radius: 3px;
                    padding: 2px;
                    font-size: 8px;
                    text-align: center;
                }
            """)
            more_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            events_layout.addWidget(more_label)
        
        events_layout.addStretch()
        layout.addWidget(events_widget)
        
        # Stile cella
        self.setStyleSheet("""
            DayCell {
                background-color: white;
                border: 1px solid #e5e7eb;
                border-radius: 4px;
            }
            DayCell:hover {
                background-color: #f8fafc;
                border: 2px solid #3b82f6;
            }
        """)
        
        self.setMinimumHeight(70)
        self.setMaximumHeight(100)
        
    def mousePressEvent(self, event):
        """Gestisce click sulla cella del giorno"""
        if event.button() == Qt.MouseButton.LeftButton:
            self.day_clicked.emit(self.day_number, self.ordini_giorno)
        super().mousePressEvent(event)

class CalendarioEventi(QWidget):
    """Widget principale del calendario mensile"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.current_date = datetime.now()
        self.ordini_data = []
        self.setup_premium_ui()
        self.load_ordini()
        self.refresh_calendar()
        
    def setup_premium_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(15, 15, 15, 15)
        
        # Header con navigazione premium
        header_card = QFrame()
        header_card.setObjectName("card")
        header_layout = QHBoxLayout(header_card)
        header_layout.setContentsMargins(20, 15, 20, 15)
        
        # Pulsanti navigazione premium
        # Pulsanti navigazione premium (BLU)
        self.prev_btn = QPushButton("◀ Prec")
        self.prev_btn.setMinimumSize(80, 35)
        self.prev_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #3b82f6, stop:1 #2563eb);
                border: 1px solid #2563eb;
                color: white;
                font-weight: 600;
                border-radius: 8px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #2563eb, stop:1 #1d4ed8);
                transform: translateY(-1px);
            }
            QPushButton:pressed {
                transform: translateY(0px);
            }
        """)
        self.prev_btn.clicked.connect(self.prev_month)
        
        self.next_btn = QPushButton("Succ ▶")
        self.next_btn.setMinimumSize(80, 35)
        self.next_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #3b82f6, stop:1 #2563eb);
                border: 1px solid #2563eb;
                color: white;
                font-weight: 600;
                border-radius: 8px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #2563eb, stop:1 #1d4ed8);
                transform: translateY(-1px);
            }
            QPushButton:pressed {
                transform: translateY(0px);
            }
        """)
        self.next_btn.clicked.connect(self.next_month)
        
        # Titolo mese premium
        self.month_label = QLabel()
        self.month_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.month_label.setStyleSheet("""
            QLabel {
                font-size: 22px;
                font-weight: 700;
                color: #1f2937;
                letter-spacing: 1px;
                padding: 8px 20px;
            }
        """)
        
        # Pulsante oggi premium
        today_btn = QPushButton("🔴 Oggi")
        today_btn.setMinimumSize(80, 35)
        today_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #ef4444, stop:1 #dc2626);
                border: 1px solid #dc2626;
                color: white;
                font-weight: 600;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #dc2626, stop:1 #b91c1c);
            }
        """)
        today_btn.clicked.connect(self.go_to_today)
        
        header_layout.addWidget(self.prev_btn)
        header_layout.addStretch()
        header_layout.addWidget(self.month_label)
        header_layout.addStretch()
        header_layout.addWidget(today_btn)
        header_layout.addWidget(self.next_btn)
        
        layout.addWidget(header_card)
        
        # Header giorni settimana premium
        days_card = QFrame()
        days_card.setObjectName("card")
        days_layout = QHBoxLayout(days_card)
        days_layout.setContentsMargins(10, 10, 10, 10)
        
        giorni = ["Lunedì", "Martedì", "Mercoledì", "Giovedì", "Venerdì", "Sabato", "Domenica"]
        
        for i, giorno in enumerate(giorni):
            day_label = QLabel(giorno[:3].upper())  # Prime 3 lettere
            day_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            
            # Stile weekend vs giorni feriali
            if i >= 5:  # Weekend
                day_label.setStyleSheet("""
                    QLabel {
                        background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                            stop:0 #fef2f2, stop:1 #fecaca);
                        color: #dc2626;
                        font-size: 12px;
                        font-weight: 700;
                        padding: 10px 8px;
                        border-radius: 6px;
                        letter-spacing: 0.5px;
                    }
                """)
            else:  # Giorni feriali
                day_label.setStyleSheet("""
                    QLabel {
                        background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                            stop:0 #f8fafc, stop:1 #f1f5f9);
                        color: #374151;
                        font-size: 12px;
                        font-weight: 700;
                        padding: 10px 8px;
                        border-radius: 6px;
                        letter-spacing: 0.5px;
                    }
                """)
            
            days_layout.addWidget(day_label)
        
        layout.addWidget(days_card)
        
        # Container calendario premium
        calendar_card = QFrame()
        calendar_card.setObjectName("card")
        calendar_layout = QVBoxLayout(calendar_card)
        calendar_layout.setContentsMargins(15, 15, 15, 15)
        
        # Griglia calendario
        self.calendar_grid = QGridLayout()
        self.calendar_grid.setSpacing(6)
        calendar_layout.addLayout(self.calendar_grid)
        
        layout.addWidget(calendar_card)
        
    def load_ordini(self):
        """Carica ordini dal file JSON"""
        try:
            if os.path.exists(ORDINI_FILE):
                with open(ORDINI_FILE, "r", encoding="utf-8") as f:
                    self.ordini_data = json.load(f)
            else:
                self.ordini_data = []
        except Exception as e:
            print(f"❌ Errore caricamento ordini: {e}")
            self.ordini_data = []
    
    def refresh_calendar(self):
        """Aggiorna la visualizzazione del calendario"""
        # Pulisce griglia esistente
        while self.calendar_grid.count():
            child = self.calendar_grid.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
        
        # Aggiorna titolo
        mese_nome = self.current_date.strftime("%B %Y").title()
        self.month_label.setText(mese_nome)
        
        # Raggruppa ordini per giorno
        ordini_per_giorno = self.raggruppa_ordini_per_giorno()
        
        # Calcola primo giorno del mese
        primo_giorno = self.current_date.replace(day=1)
        
        # Calcola numero giorni nel mese
        if primo_giorno.month == 12:
            ultimo_giorno = primo_giorno.replace(year=primo_giorno.year + 1, month=1) - timedelta(days=1)
        else:
            ultimo_giorno = primo_giorno.replace(month=primo_giorno.month + 1) - timedelta(days=1)
        
        giorni_nel_mese = ultimo_giorno.day
        
        # Giorno della settimana del primo giorno (0=lunedì)
        primo_giorno_settimana = primo_giorno.weekday()
        
        # Riempi calendario
        row, col = 0, 0
        
        # Giorni del mese precedente
        if primo_giorno_settimana > 0:
            mese_prec = primo_giorno - timedelta(days=1)
            giorni_mese_prec = mese_prec.day
            
            for i in range(primo_giorno_settimana):
                day_num = giorni_mese_prec - primo_giorno_settimana + i + 1
                cell = DayCell(day_num, [], is_current_month=False)
                cell.day_clicked.connect(self.on_day_clicked)
                self.calendar_grid.addWidget(cell, row, col)
                col += 1
        
        # Giorni del mese corrente
        for day in range(1, giorni_nel_mese + 1):
            ordini_giorno = ordini_per_giorno.get(day, [])
            cell = DayCell(day, ordini_giorno, is_current_month=True)
            cell.day_clicked.connect(self.on_day_clicked)
            self.calendar_grid.addWidget(cell, row, col)
            
            col += 1
            if col >= 7:
                col = 0
                row += 1
        
        # Giorni del mese successivo
        next_day = 1
        while col < 7 and row < 6:
            cell = DayCell(next_day, [], is_current_month=False)
            cell.day_clicked.connect(self.on_day_clicked)
            self.calendar_grid.addWidget(cell, row, col)
            col += 1
            next_day += 1
            
            if col >= 7:
                col = 0
                row += 1
    
    def raggruppa_ordini_per_giorno(self):
        """Raggruppa ordini per giorno del mese corrente"""
        ordini_per_giorno = defaultdict(list)
        
        target_month = self.current_date.month
        target_year = self.current_date.year
        
        for ordine in self.ordini_data:
            data_evento = ordine["campi"].get("Data evento", "")
            if data_evento:
                try:
                    data_obj = datetime.strptime(data_evento, "%d/%m/%Y")
                    if data_obj.month == target_month and data_obj.year == target_year:
                        ordini_per_giorno[data_obj.day].append(ordine)
                except ValueError:
                    continue
        
        return ordini_per_giorno
    
    def on_day_clicked(self, day_number, ordini_giorno):
        """Gestisce click su giorno specifico - VERSIONE COMPLETA"""
        if not ordini_giorno:
            QMessageBox.information(self, "Nessun Evento", 
                                f"Nessun ordine per il {day_number}")
            return
        
        # Crea dialog con opzioni
        self.show_day_options(day_number, ordini_giorno)

    def show_day_options(self, day_number, ordini_giorno):
        """Mostra opzioni per il giorno selezionato"""
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Opzioni per il {day_number} {self.current_date.strftime('%B %Y')}")
        dialog.setGeometry(300, 300, 500, 450)
        
        # Applica lo stesso stile dell'app principale
        dialog.setStyleSheet("""
            QDialog {
                background-color: #f8f9fa;
                color: #212529;
                font-family: 'Segoe UI', Arial, sans-serif;
            }
            QPushButton {
                background-color: #8b5cf6;
                color: white;
                border: none;
                border-radius: 6px;
                padding: 10px 20px;
                font-size: 13px;
                font-weight: 600;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #7c3aed;
            }
            QPushButton#primaryBtn {
                background-color: #10b981;
                color: white;
            }
            QPushButton#primaryBtn:hover {
                background-color: #059669;
            }
            QPushButton#successBtn {
                background-color: #3b82f6;
                color: white;
            }
            QPushButton#successBtn:hover {
                background-color: #2563eb;
            }
            QPushButton#secondaryBtn {
                background-color: #ef4444;
                color: white;
            }
            QPushButton#secondaryBtn:hover {
                background-color: #dc2626;
            }
            QFrame {
                background-color: white;
                border: 1px solid #e2e8f0;
                border-radius: 8px;
                padding: 15px;
                margin: 5px;
            }
        """)
        
        layout = QVBoxLayout(dialog)
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Header
        header = QLabel(f"📅 {day_number} {self.current_date.strftime('%B %Y')}")
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        header.setFont(QFont("Arial", 16, QFont.Weight.Bold))
        header.setStyleSheet("""
            QLabel {
                background-color: #3b82f6;
                color: white;
                padding: 12px;
                border-radius: 6px;
                margin-bottom: 10px;
            }
        """)
        layout.addWidget(header)
        
        # Info riepilogo
        info_card = QFrame()
        info_layout = QVBoxLayout(info_card)
        
        total_persone = sum(int(ordine["campi"].get("Numero persone", "0")) for ordine in ordini_giorno)
        
        info_text = f"""
        📋 <b>{len(ordini_giorno)} ordini programmati</b><br>
        👥 <b>{total_persone} persone totali</b><br><br>
        
        <b>Servizi del giorno:</b><br>
        """
        
        # Conta servizi
        servizi_count = {}
        for ordine in ordini_giorno:
            servizio = self.determina_servizio_ordine(ordine)
            servizi_count[servizio] = servizi_count.get(servizio, 0) + 1
        
        for servizio, count in servizi_count.items():
            info_text += f"• {servizio}: {count} ordini<br>"
        
        info_label = QLabel(info_text)
        info_label.setWordWrap(True)
        info_label.setStyleSheet("padding: 15px; font-size: 13px;")
        info_layout.addWidget(info_label)
        
        layout.addWidget(info_card)
        
        # Pulsanti azione
        buttons_layout = QVBoxLayout()
        buttons_layout.setSpacing(10)
        
        # Pulsante Visualizza Dettagli
        view_btn = QPushButton("👁️ Visualizza Dettagli Ordini")
        view_btn.setObjectName("successBtn")
        view_btn.clicked.connect(lambda: self.show_day_details_table(day_number, ordini_giorno, dialog))
        buttons_layout.addWidget(view_btn)
        
        # Pulsante Export Tabella Operativa - QUESTO È IL PULSANTE PRINCIPALE!
        export_btn = QPushButton("📋 Genera Tabella Operativa")
        export_btn.setObjectName("primaryBtn")
        export_btn.setToolTip("Genera tabella Excel per gestione operativa giornaliera")
        export_btn.clicked.connect(lambda: self.export_day_to_excel(day_number, ordini_giorno))
        buttons_layout.addWidget(export_btn)
        
        
        # Pulsante Chiudi
        close_btn = QPushButton("Chiudi")
        close_btn.setObjectName("secondaryBtn")
        close_btn.clicked.connect(dialog.close)
        buttons_layout.addWidget(close_btn)
        
        layout.addLayout(buttons_layout)
        dialog.exec()

    def show_day_details_table(self, day_number, ordini_giorno, parent_dialog):
        """Mostra tabella dettagliata ordini"""
        parent_dialog.close()  # Chiudi dialog opzioni
        
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Dettagli Ordini - {day_number} {self.current_date.strftime('%B %Y')}")
        dialog.setGeometry(150, 150, 900, 600)
        
        layout = QVBoxLayout(dialog)
        
        # Header
        header = QLabel(f"📅 {len(ordini_giorno)} ordini programmati")
        header.setFont(QFont("Arial", 16, QFont.Weight.Bold))
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        header.setStyleSheet("padding: 15px; background-color: #3b82f6; color: white; border-radius: 8px;")
        layout.addWidget(header)
        
        # Tabella ordini
        table = QTableWidget()
        table.setColumnCount(6)
        table.setHorizontalHeaderLabels(["Cliente", "Cellulare", "Persone", "Servizio", "Orario", "Indirizzo"])
        table.setRowCount(len(ordini_giorno))
        
        for i, ordine in enumerate(ordini_giorno):
            campi = ordine["campi"]
            servizio = self.determina_servizio_ordine(ordine)
            
            table.setItem(i, 0, QTableWidgetItem(campi.get("Nome", "")))
            table.setItem(i, 1, QTableWidgetItem(campi.get("Cellulare", "")))
            table.setItem(i, 2, QTableWidgetItem(campi.get("Numero persone", "")))
            table.setItem(i, 3, QTableWidgetItem(servizio))
            table.setItem(i, 4, QTableWidgetItem(f"{campi.get('Orario allestimento', '')} - {campi.get('Orario Pronti', '')}"))
            table.setItem(i, 5, QTableWidgetItem(campi.get("Via", "")))
        
        table.resizeColumnsToContents()
        layout.addWidget(table)
        
        # Pulsante chiudi
        close_btn = QPushButton("Chiudi")
        close_btn.clicked.connect(dialog.close)
        layout.addWidget(close_btn)
        
        dialog.exec()

    def export_day_to_excel(self, day_number, ordini_giorno):
        """Esporta ordini del giorno in Excel - CON CARTELLA ORGANIZZATA"""
        try:
            import pandas as pd # type: ignore
            from datetime import datetime
            import os
            
            # Progress bar
            progress = QProgressBar(self)
            progress.setWindowTitle("Generazione Tabella Operativa...")
            progress.setGeometry(self.width()//2 - 150, self.height()//2, 300, 30)
            progress.setRange(0, 100)
            progress.show()
            QApplication.processEvents()
            
            progress.setValue(20)
            QApplication.processEvents()
            
            # Data del giorno
            data_giorno = self.current_date.replace(day=day_number)
            data_formattata = data_giorno.strftime("%d/%m/%Y")
            giorno_settimana = data_giorno.strftime("%A").upper()
            
            # ========== CREA STRUTTURA CARTELLE ORGANIZZATA ==========
            # Cartella principale
            cartella_principale = "Tabelle Ordini"
            
            # Sottocartella per anno/mese
            anno_mese = data_giorno.strftime("%Y-%m_%B")  # Es: "2025-01_Gennaio"
            cartella_mese = os.path.join(cartella_principale, anno_mese)
            
            # Crea le cartelle se non esistono
            os.makedirs(cartella_mese, exist_ok=True)
            
            progress.setValue(30)
            QApplication.processEvents()
            
            # ========== PREPARA DATI TABELLA OPERATIVA ==========
            tabella_operativa = []
            
            for ordine in ordini_giorno:
                campi = ordine["campi"]
                
                # Combina Nome Cliente e Indirizzo
                nome_cliente = campi.get("Nome", "")
                indirizzo = campi.get("Via", "")
                cliente_indirizzo = f"{nome_cliente}\n{indirizzo}" if indirizzo else nome_cliente
                
                # Orario Pronti (corrisponde a Orario Allestimento)
                orario_pronti = campi.get("Orario allestimento", "")
                
                tabella_operativa.append({
                    "N.Prog": "",  # VUOTO - da compilare a mano
                    "Nome Cliente/Indirizzo": cliente_indirizzo,
                    "Orario Pronti": orario_pronti,
                    "Cameriere": "",  # VUOTO - da compilare a mano
                    "Uscita": "",     # VUOTO - da compilare a mano
                    "Consegna": "",   # VUOTO - da compilare a mano
                    "Ritiro": ""      # VUOTO - da compilare a mano
                })
            
            progress.setValue(60)
            QApplication.processEvents()
            
            # ========== CREA FILE EXCEL NELLA CARTELLA ORGANIZZATA ==========
            data_file = data_formattata.replace("/", "-")
            nome_file = f"Tabella_Operativa_{data_file}.xlsx"
            
            # Path completo del file
            filepath_completo = os.path.join(cartella_mese, nome_file)
            
            # Crea DataFrame
            df = pd.DataFrame(tabella_operativa)
            
            with pd.ExcelWriter(filepath_completo, engine='openpyxl') as writer:
                # Scrivi tabella con spazio per il titolo
                df.to_excel(writer, sheet_name='Tabella Operativa', index=False, startrow=3)
                
                # ========== FORMATTAZIONE BIANCO E NERO ==========
                from openpyxl.styles import Font, Alignment, Border, Side # type: ignore
                
                workbook = writer.book
                ws = workbook['Tabella Operativa']
                
                # ========== STILI BIANCO E NERO ==========
                title_font = Font(bold=True, size=18, color="000000")
                header_font = Font(bold=True, color="000000", size=12)
                
                thick_border = Border(
                    left=Side(style='thick', color='000000'),
                    right=Side(style='thick', color='000000'),
                    top=Side(style='thick', color='000000'),
                    bottom=Side(style='thick', color='000000')
                )
                
                thin_border = Border(
                    left=Side(style='thin', color='000000'),
                    right=Side(style='thin', color='000000'),
                    top=Side(style='thin', color='000000'),
                    bottom=Side(style='thin', color='000000')
                )
                
                center_alignment = Alignment(horizontal="center", vertical="center")
                left_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                
                # ========== TITOLO PRINCIPALE ==========
                ws['A1'] = f"TABELLA OPERATIVA CATERING"
                ws['A1'].font = title_font
                ws['A1'].alignment = center_alignment
                ws.merge_cells('A1:G1')
                
                # ========== SOTTOTITOLO ==========
                ws['A2'] = f"{giorno_settimana} {data_formattata} - {len(ordini_giorno)} ORDINI"
                ws['A2'].font = Font(bold=True, size=14, color="000000")
                ws['A2'].alignment = center_alignment
                ws.merge_cells('A2:G2')
                
                progress.setValue(80)
                QApplication.processEvents()
                
                # ========== FORMATTA HEADER TABELLA (riga 4) ==========
                header_row = 4
                for col_num, column_title in enumerate(df.columns, 1):
                    cell = ws.cell(row=header_row, column=col_num)
                    cell.value = column_title
                    cell.font = header_font
                    cell.alignment = center_alignment
                    cell.border = thick_border
                
                # ========== FORMATTA DATI ==========
                start_row = 5
                end_row = start_row + len(df) - 1
                
                for row in range(start_row, end_row + 1):
                    for col in range(1, len(df.columns) + 1):
                        cell = ws.cell(row=row, column=col)
                        cell.border = thin_border
                        
                        if col == 2:  # Colonna "Nome Cliente/Indirizzo"
                            cell.alignment = left_alignment
                        else:
                            cell.alignment = center_alignment
                
                # ========== DIMENSIONI COLONNE OTTIMIZZATE ==========
                column_widths = {
                    'A': 8,   # N.Prog
                    'B': 35,  # Nome Cliente/Indirizzo 
                    'C': 12,  # Orario Pronti
                    'D': 15,  # Cameriere
                    'E': 10,  # Uscita
                    'F': 12,  # Consegna
                    'G': 10   # Ritiro
                }
                
                for col_letter, width in column_widths.items():
                    ws.column_dimensions[col_letter].width = width
                
                # ========== ALTEZZA RIGHE ==========
                ws.row_dimensions[header_row].height = 25
                
                for row in range(start_row, end_row + 1):
                    ws.row_dimensions[row].height = 40
                
                # ========== RIGHE EXTRA VUOTE ==========
                extra_rows = 3
                for i in range(extra_rows):
                    row_num = end_row + 1 + i
                    for col in range(1, len(df.columns) + 1):
                        cell = ws.cell(row=row_num, column=col)
                        cell.border = thin_border
                        if col == 2:
                            cell.alignment = left_alignment
                        else:
                            cell.alignment = center_alignment
                    ws.row_dimensions[row_num].height = 40
                
                # ========== IMPOSTAZIONI STAMPA ==========
                ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
                ws.page_setup.paperSize = ws.PAPERSIZE_A4
                ws.page_setup.fitToPage = True
                ws.page_setup.fitToHeight = 1
                ws.page_setup.fitToWidth = 1
                
                ws.page_margins.left = 0.5
                ws.page_margins.right = 0.5
                ws.page_margins.top = 0.75
                ws.page_margins.bottom = 0.75
                
                ws.oddHeader.center.text = f"&16&B{giorno_settimana} {data_formattata}"
                ws.oddFooter.left.text = f"Generato: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
                ws.oddFooter.right.text = "Pag. &P di &N"
            
            progress.setValue(100)
            QApplication.processEvents()
            progress.hide()
            
            # ========== MESSAGGIO DI SUCCESSO CON INFO CARTELLE ==========
            success_msg = f"""📋 TABELLA OPERATIVA CREATA!

  """
            
            QMessageBox.information(self, "Tabella Operativa Generata", success_msg)
            
            # Opzione per aprire il file
            
        except ImportError:
            QMessageBox.warning(self, "Libreria Mancante", 
                            "Per esportare in Excel installa:\n\n"
                            "pip install pandas openpyxl\n\n"
                            "Riavvia il programma dopo l'installazione.")
        except Exception as e:
            if 'progress' in locals():
                progress.hide()
            
            QMessageBox.critical(self, "Errore Export", 
                            f"❌ Errore durante l'export:\n\n{str(e)}")

    # =============================================================================
    # AGGIORNA ANCHE IL METODO open_excel_file (se necessario)
    # =============================================================================

    def open_excel_file(self, filepath_completo):
        """Apre il file Excel generato - VERSIONE AGGIORNATA"""
        try:
            import subprocess
            import platform
            
            if platform.system() == "Windows":
                subprocess.run(["start", filepath_completo], shell=True, check=True)
            elif platform.system() == "Darwin":  # macOS
                subprocess.run(["open", filepath_completo], check=True)
            else:  # Linux
                subprocess.run(["xdg-open", filepath_completo], check=True)
                
            print(f"📄 Excel aperto: {filepath_completo}")
            
        except Exception as e:
            print(f"❌ Errore apertura Excel: {e}")
            QMessageBox.information(self, "File Creato", 
                                f"File Excel creato:\n{filepath_completo}\n\n"
                                f"Aprilo manualmente per visualizzare i dati.")

    def generate_kitchen_summary(self, day_number, ordini_giorno):
        """Genera riepilogo cucina per il giorno specifico"""
        try:
            data_target = self.current_date.replace(day=day_number).strftime("%d/%m/%Y")
            
            # Chiama il metodo esistente della tua app principale
            if hasattr(self.parent(), 'crea_riepilogo_cucina'):
                self.parent().crea_riepilogo_cucina(data_target)
            else:
                QMessageBox.information(self, "Funzionalità", 
                                    "Riepilogo cucina non disponibile dal calendario.\n"
                                    "Usa il menu principale 'Statistiche > Riepilogo Cucina'")
        except Exception as e:
            QMessageBox.warning(self, "Errore", f"Errore generazione riepilogo cucina:\n{str(e)}")
    
    def determina_servizio_ordine(self, ordine):
        """Determina tipo servizio da ordine - AGGIORNATO CON SERVITO"""
        scelte = ordine.get("scelte", {})
        
        # Lista per raccogliere tutti i servizi trovati
        servizi_trovati = []
        
        # Verifica Coffee Break
        if any(k.startswith("coffee break|") and v for k, v in scelte.items()):
            servizi_trovati.append("Coffee Break")
        
        # Verifica Tea Break
        if any(k.startswith("tea break|") and v for k, v in scelte.items()):
            servizi_trovati.append("Tea Break")
        
        # Verifica Lunch Buffet
        if any(k.startswith("lunch buffet|") and v for k, v in scelte.items()):
            servizi_trovati.append("Lunch Buffet")
            
        if any(k.startswith("lunch box|") and v for k, v in scelte.items()):
            servizi_trovati.append("Lunch box")
        
        # ========== NUOVO: Verifica Servito ==========
        if any(k.startswith("servito|") and v for k, v in scelte.items()):
            servizi_trovati.append("Servito")
        
        # Verifica Aperitivo
        if any(k.startswith("aperitivo a buffet|") and v for k, v in scelte.items()):
            tipo_ap = ordine.get("tipo_ap", "Leggero")
            servizi_trovati.append(f"Aperitivo {tipo_ap}")
        
        # Restituisci servizio principale o combinato
        if len(servizi_trovati) == 1:
            return servizi_trovati[0]
        elif len(servizi_trovati) > 1:
            return " + ".join(servizi_trovati)
        else:
            return "Altro" 
    def prev_month(self):
        """Va al mese precedente"""
        if self.current_date.month == 1:
            self.current_date = self.current_date.replace(year=self.current_date.year - 1, month=12)
        else:
            self.current_date = self.current_date.replace(month=self.current_date.month - 1)
        self.refresh_calendar()

    def next_month(self):
        """Va al mese successivo"""
        if self.current_date.month == 12:
            self.current_date = self.current_date.replace(year=self.current_date.year + 1, month=1)
        else:
            self.current_date = self.current_date.replace(month=self.current_date.month + 1)
        self.refresh_calendar()

    def go_to_today(self):
        """Torna al mese corrente"""
        self.current_date = datetime.now()
        self.refresh_calendar()

class CalendarioEventiDialog(QDialog):
    """Dialog principale che contiene il calendario con stile premium"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("📅 Calendario Eventi Catering")
        self.setGeometry(100, 100, 1300, 850)
        
        # APPLICA STILE COMBINATO PROFESSIONAL + CALENDAR
        combined_style = PROFESSIONAL_STYLE + PROFESSIONAL_CALENDAR_STYLE
        self.setStyleSheet(combined_style)
        
        # Imposta object name per il dialog
        self.setObjectName("calendarDialog")
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        
        # Header premium
        header_card = QFrame()
        header_card.setObjectName("card")
        header_layout = QVBoxLayout(header_card)
        
        header = QLabel("📅 CALENDARIO EVENTI")
        header.setStyleSheet("""
            QLabel {
                font-size: 24px;
                font-weight: 700;
                color: #1f2937;
                text-align: center;
                padding: 20px;
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #3b82f6, stop:1 #2563eb);
                color: white;
                border-radius: 8px;
                letter-spacing: 1px;
            }
        """)
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        header_layout.addWidget(header)
        
        layout.addWidget(header_card)
        
        # Widget calendario con stile premium
        self.calendario = CalendarioEventi()
        layout.addWidget(self.calendario)
        
        # Footer con pulsanti premium
        footer_layout = QHBoxLayout()
        footer_layout.addStretch()
        
        # Info badge
        info_label = QLabel("💡 Clicca su un giorno per gestire gli ordini")
        info_label.setStyleSheet("""
            QLabel {
                color: #6b7280;
                font-style: italic;
                font-size: 12px;
                padding: 8px;
            }
        """)
        footer_layout.addWidget(info_label)
        
        footer_layout.addStretch()
        
        # Pulsante chiudi premium
        # Pulsante chiudi premium (BLU)
        close_btn = QPushButton("Chiudi")
        close_btn.setMinimumSize(120, 40)
        close_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #3b82f6, stop:1 #2563eb);
                border: 1px solid #2563eb;
                color: white;
                font-weight: 600;
                border-radius: 8px;
                padding: 10px 20px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #2563eb, stop:1 #1d4ed8);
                transform: translateY(-1px);
            }
            QPushButton:pressed {
                transform: translateY(0px);
            }
        """)
        close_btn.clicked.connect(self.close)
        footer_layout.addWidget(close_btn)
        
        layout.addLayout(footer_layout)
        
        # Forza aggiornamento stili
        self.style().polish(self)

    

class CateringApp(QMainWindow):
    REFERENZE_IN_KG = [
            # FINGER FOOD E DOLCI
            "biscotteria", "spicchietti crostate", "trancetti tenerina crema e amerene",
            "mini tramezzini vegetariani", "rustici vegetariani", "rustici mix", 
            "parmigianini", "pasticceria", "cubetti mortadella", "cubetti mortadella scottati all'aceto",
            "dadini di mortadella", "rombetti di frittata", "mini tramezzini mix",
            "mozzarelline in carrozza", "mini cotolettine", "olive all'ascolana", "olive ascolane",
            "crocchette", "tagliata di frutta", "mozzarelle", "polpettine prosciutto e ricotta",
            "minisw farciti mix", "trancetti tenerina cioccolato", "crescenta a dadini",
            
            # INSALATE / SECONDI FREDDI
            "insalata di pollo", "insalatona mista con ananas", "insalata di pollo alla nizzarda", 
            "caprese", "insalata di pollo marinato", "tacchino con asparagi", "straccetti con scorza",
            "spezzatino con topinambur", "bocconcini di tacchino", "straccetti di pollo",
            "spezzatino di maiale", "bocconcini di maiale", "rombetti di tacchino", 
            "piccatine al balsamico", "filettino di maiale", "bocconcini di pollo", 
            "rombi di tacchino", "straccetti di maiale", "filetto di tacchino",
            
            # CONTORNI CHE VANNO IN KG (ALCUNI, NON TUTTI)
            "macedonia di verdure", "patate al rosmarino", "spinaci alla tirolese",
            "melanzane grigliate", "zucchine grigliate", "peperoni grigliati",
            "verdure miste grigliate", "caponata", "ratatouille"
            
            # NOTA: "verdure alla brace" NON è in questo array, quindi andrà in pezzi
        ]
    def _va_in_kg(self, voce):
        """Controlla se una voce deve essere mostrata in kg usando l'array centralizzato"""
        voce_lower = voce.lower()
        
        for ref in self.REFERENZE_IN_KG:
            ref_lower = ref.lower()
            if ref_lower in voce_lower:
                return True
        
        return False
    def __init__(self):
        super().__init__()
        self.menu_extra_widgets = {}

        self.setWindowTitle("Gestione Ordini Catering - Versione Professionale")
        # Sostituiscila con:
        self.setGeometry(100, 100, 1400, 1300)  # larghezza=1400, altezza=1200 (più lungo!)  # Finestra più grande
        self.setMinimumSize(400, 400)  # Dimensione minima
        self.setStyleSheet(PROFESSIONAL_STYLE)
        self.modalita_modifica = False
        self.ordine_da_modificare = None
        
        self.modalita_aggiungi_servizi = False  # <-- NUOVA

        # ============ PULSANTE RINOMINA REFERENZA ============
        rinomina_btn = QPushButton("🔄 Rinomina Referenza")
        rinomina_btn.setMinimumSize(180, 40)  # Dimensione appropriata
        rinomina_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #f59e0b, stop:1 #d97706);
                border: 1px solid #d97706;
                color: white;
                font-weight: 600;
                border-radius: 8px;
                padding: 10px 15px;
                margin: 2px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #d97706, stop:1 #b45309);
                transform: translateY(-1px);
            }
            QPushButton:pressed {
                transform: translateY(0px);
            }
        """)
        rinomina_btn.clicked.connect(self.dialog_rinomina_referenza)

        
        # Variabili di stato
        self.voce_categoria = {}
        self.selected = {}
        self.voce_originale = {}
        self.ap_buffet_tipo = ""
        
        # Carica dati referenze
        try:
            with open(gestore_percorsi.get_percorso("referenze_cibi.json"), "r", encoding="utf-8") as f:
                raw_data = json.load(f)
                self.dati_referenze = {
                    k.strip().lower(): v for k, v in raw_data.items()
                }
                self.referenze_cibi = raw_data  # Versione originale per menu extra
    
                # DEBUG: Stampa cosa viene caricato per "verdi"
                print("🔍 DEBUG CARICAMENTO JSON:")
                for k, v in raw_data.items():
                    if "verdi" in k.lower():
                        print(f"   JSON raw: '{k}' → {v}")
                        print(f"   Convertito: '{k.strip().lower()}' → {v}")
        except FileNotFoundError:
            self.dati_referenze = {}
            print("Warning: referenze_cibi.json non trovato. Verrà utilizzato il calcolo base.")
        
        os.makedirs(ORDINI_DIR, exist_ok=True)
        
        self.setup_ui()
        self.create_menu()
        
    def setup_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # ========== AGGIUNGI SCROLL PRINCIPALE ==========
        # Layout principale con scroll
        main_scroll = QScrollArea()
        main_scroll.setStyleSheet("QScrollBar:vertical { width: 0px; }")  # AGGIUNGI QUESTA
        main_scroll.setWidgetResizable(True)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        # Widget contenitore
        main_widget = QWidget()
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(15, 15, 15, 15)
        main_layout.setSpacing(20)
        # ===============================================
        
        # Header (STESSO CODICE)
        header_card = HeaderCard()
        header_layout = QVBoxLayout(header_card)
        
        title = QLabel("𝓝𝓮𝓻𝓲 𝓖𝓮𝓼𝓽𝓲𝓸𝓷𝓮 𝓒𝓪𝓽𝓮𝓻𝓲𝓷𝓰")
        title.setObjectName("appTitle")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        header_layout.addWidget(title)
        
        
        
        main_layout.addWidget(header_card)
        
        # Sezione informazioni ordine (STESSO CODICE)
        self.create_order_info_section(main_layout)
        
        # Tabs servizi (STESSO CODICE)
        self.create_services_tabs(main_layout)

        # ========== AGGIUNGI QUESTE RIGHE ALLA FINE ==========
        # Setup scroll
        main_scroll.setWidget(main_widget)
        
        # Layout per central widget  
        central_layout = QVBoxLayout(central_widget)
        central_layout.setContentsMargins(0, 0, 0, 0)
        central_layout.addWidget(main_scroll)
    
    
        
    def create_menu(self):
        menubar = self.menuBar()

        # Menu File (uguale a prima)
        file_menu = menubar.addMenu("File")
        
        new_action = QAction("🆕Nuovo Ordine", self)
        new_action.triggered.connect(self.nuovo_ordine)
        file_menu.addAction(new_action)
        
        file_menu.addSeparator()
        
        exit_action = QAction("❎Esci", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # Menu Archivio (uguale a prima)
        archivio_menu = menubar.addMenu("Archivio")
        
        carica_action = QAction("🗂️Carica Ordine", self)
        carica_action.triggered.connect(self.carica_ordine)
        archivio_menu.addAction(carica_action)
        
        archivio_menu.addSeparator()
        
        svuota_action = QAction("🗑️Svuota Archivio", self)
        svuota_action.triggered.connect(self.svuota_archivio)
        archivio_menu.addAction(svuota_action)
        
        # Menu Calendario
        calendario_menu = menubar.addMenu("Calendario")

        calendario_action = QAction("📅 Visualizza Calendario", self)
        calendario_action.triggered.connect(self.mostra_calendario)
        calendario_menu.addAction(calendario_action)
        
        # Menu Statistiche (uguale a prima)
        stats_menu = menubar.addMenu("Statistiche")
        
        stats_action = QAction("📊Visualizza Statistiche", self)
        stats_action.triggered.connect(self.mostra_statistiche)
        stats_menu.addAction(stats_action)
        
        riepilogo_action = QAction("🧑🏻‍🍳Riepilogo Cucina", self)
        riepilogo_action.triggered.connect(self.genera_riepilogo_cucina)
        stats_menu.addAction(riepilogo_action)
        
        # ============= NUOVO MENU GESTIONE ============= 
        gestione_menu = menubar.addMenu("Gestione")

        rinomina_action = QAction("🔄 Rinomina Referenza", self)
        rinomina_action.triggered.connect(self.dialog_rinomina_referenza)
        gestione_menu.addAction(rinomina_action)
        # ===============================================
        
        # Menu Aiuto (uguale a prima)
        help_menu = menubar.addMenu("Aiuto")
        
        about_action = QAction("👨🏻‍💻Informazioni", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)


    def mostra_calendario(self):
        """AGGIUNGI questo metodo alla tua classe CateringApp"""
        calendario_dialog = CalendarioEventiDialog(self)
        calendario_dialog.exec()
    def setup_search_functionality(self):
        """Configura la funzionalità di ricerca"""
        # Crea la barra di ricerca
        self.search_bar = SearchBar(self)
        
        # Aggiungi la barra di ricerca sopra i tabs (nel layout dei servizi)
        # Trova il layout dei servizi e inserisci la search bar
        services_layout = self.services_card.layout()
        services_layout.insertWidget(1, self.search_bar)  # Dopo il header, prima dei tabs
        
        # Colleghi gli eventi
        self.search_bar.search_field.textChanged.connect(self.filter_referenze)
        self.search_bar.search_field.returnPressed.connect(self.focus_first_result)
        
        # Timer per ritardare la ricerca durante la digitazione
        self.search_timer = QTimer()
        self.search_timer.setSingleShot(True)
        self.search_timer.timeout.connect(self.perform_search)
        
        # Shortcut Ctrl+F per aprire la ricerca
        self.search_shortcut = QShortcut(QKeySequence("Ctrl+F"), self)
        self.search_shortcut.activated.connect(self.toggle_search)
        
        # Shortcut Esc per chiudere la ricerca
        self.escape_shortcut = QShortcut(QKeySequence("Escape"), self)
        self.escape_shortcut.activated.connect(self.close_search)
        
    def toggle_search(self):
        """Apre/chiude la barra di ricerca con Ctrl+F"""
        if self.search_bar.is_visible:
            self.close_search()
        else:
            self.open_search()

    def open_search(self):
        """Apre la barra di ricerca e le dà il focus"""
        self.search_bar.show()
        self.search_bar.is_visible = True
        self.search_bar.search_field.setFocus()
        self.search_bar.search_field.selectAll()

    def close_search(self):
        """Chiude la barra di ricerca e ripristina tutte le referenze"""
        self.search_bar.hide()
        self.search_bar.is_visible = False
        self.search_bar.search_field.clear()
        self.show_all_referenze()

    def filter_referenze(self, text):
        """Filtra le referenze in base al testo di ricerca (con delay)"""
        # Ferma il timer precedente se in corso
        self.search_timer.stop()
        
        # Se il testo è vuoto, mostra tutto subito
        if not text.strip():
            self.show_all_referenze()
            self.search_bar.results_label.setText("")
            return
        
        # Avvia il timer per la ricerca ritardata (300ms)
        self.search_timer.start(300)

    def perform_search(self):
        """Esegue la ricerca effettiva"""
        search_text = self.search_bar.search_field.text().strip().lower()
        
        if not search_text:
            self.show_all_referenze()
            return
        
        visible_count = 0
        total_count = 0
        
        # Scorri tutti i tabs
        for tab_index in range(self.tabs.count()):
            tab_widget = self.tabs.widget(tab_index)
            if not tab_widget:
                continue
                
            # Trova tutti i GroupBox (sezioni) in questo tab
            group_boxes = tab_widget.findChildren(QGroupBox)
            
            for group_box in group_boxes:
                section_visible_count = 0
                
                # Trova tutte le checkbox in questo GroupBox
                checkboxes = group_box.findChildren(CustomCheckBox)
                
                for checkbox in checkboxes:
                    total_count += 1
                    checkbox_text = checkbox.text().lower()
                    
                    # Verifica se il testo di ricerca è presente
                    if search_text in checkbox_text:
                        checkbox.show()
                        visible_count += 1
                        section_visible_count += 1
                    else:
                        checkbox.hide()
                
                # Nascondi la sezione se non ha risultati
                if section_visible_count == 0:
                    group_box.hide()
                else:
                    group_box.show()
        
        # Aggiorna il label dei risultati
        if visible_count == 0:
            self.search_bar.results_label.setText("🚫 Nessun risultato trovato")
            self.search_bar.results_label.setStyleSheet("color: #dc2626; font-size: 12px;")
        else:
            self.search_bar.results_label.setText(f"✅ {visible_count} di {total_count} referenze")
            self.search_bar.results_label.setStyleSheet("color: #059669; font-size: 12px;")

    def show_all_referenze(self):
        """Mostra tutte le referenze e sezioni"""
        # Scorri tutti i tabs
        for tab_index in range(self.tabs.count()):
            tab_widget = self.tabs.widget(tab_index)
            if not tab_widget:
                continue
                
            # Trova tutti i GroupBox e checkbox
            group_boxes = tab_widget.findChildren(QGroupBox)
            for group_box in group_boxes:
                group_box.show()
                
                checkboxes = group_box.findChildren(CustomCheckBox)
                for checkbox in checkboxes:
                    checkbox.show()

    def focus_first_result(self):
        """Focalizza il primo risultato quando si preme Invio"""
        # Trova la prima checkbox visibile
        for tab_index in range(self.tabs.count()):
            tab_widget = self.tabs.widget(tab_index)
            if not tab_widget:
                continue
                
            checkboxes = tab_widget.findChildren(CustomCheckBox)
            for checkbox in checkboxes:
                if checkbox.isVisible():
                    # Scorri fino alla checkbox e dagli il focus
                    self.tabs.setCurrentIndex(tab_index)
                    checkbox.setFocus()
                    return
        
    def create_order_info_section(self, parent_layout):
        # Card per le informazioni ordine con tendina
        info_card = BusinessCard()
        info_layout = QVBoxLayout(info_card)
        
        # Header cliccabile per espandere/comprimere
        self.info_header = QPushButton("🔼 INFORMAZIONI CLIENTE 👥 ")
        self.info_header.setObjectName("sectionHeader")
        self.info_header.setFlat(True)
        self.info_header.clicked.connect(self.toggle_info_section)
        info_layout.addWidget(self.info_header)
        
        # Container per il contenuto (inizialmente visibile)
        self.info_content = QWidget()
        content_layout = QVBoxLayout(self.info_content)
        
        # Grid per i campi (come prima)
        fields_grid = QGridLayout()
        fields_grid.setSpacing(15)
        fields_grid.setVerticalSpacing(20)
        
        # Campi di input (stesso codice di prima)
        fields = [
            ("Nome Cliente:", "Nome"),
            ("Numero Cellulare:", "Cellulare"),
            ("Indirizzo:", "Via"),
            ("Numero Persone:", "Numero persone"),
            ("Orario Allestimento:", "Orario allestimento"),
            ("Orario Servizio:", "Orario Pronti")
        ]
        
        self.entries = {}
        
        for i, (label_text, field_name) in enumerate(fields):
            label = QLabel(label_text)
            label.setStyleSheet("font-weight: bold; color: #2c3e50;")
            fields_grid.addWidget(label, i // 2, (i % 2) * 2)
            
            entry = QLineEdit()
            entry.setPlaceholderText(f"Inserisci {field_name.lower()}")
            fields_grid.addWidget(entry, i // 2, (i % 2) * 2 + 1)
            self.entries[field_name] = entry
        
        # Data evento
        date_label = QLabel("Data Evento:")
        date_label.setStyleSheet("font-weight: bold; color: #2c3e50;")
        fields_grid.addWidget(date_label, 3, 0)
        
        self.date_entry = QDateEdit()
        self.date_entry.setDate(QDate.currentDate())
        self.date_entry.setDisplayFormat("dd/MM/yyyy")
        self.date_entry.setCalendarPopup(True)

        # DISABILITA SCROLL CHE CAMBIA ANNO
        def disable_wheel_event(event):
            event.ignore()
        self.date_entry.wheelEvent = disable_wheel_event
        fields_grid.addWidget(self.date_entry, 3, 1)
        self.entries["Data evento"] = self.date_entry
        
        # ========== AGGIUNGI SOLO QUESTO: CHECKBOX FIERA ==========
        # Checkbox "In Fiera" - piccolo e discreto
       # Checkbox "In Fiera" - usa CustomCheckBox per X blu
        self.fiera_checkbox = CustomCheckBox("🏢 Evento in Fiera")
        self.fiera_checkbox.toggled.connect(self.set_evento_fiera)
        
        # Aggiungi il checkbox nella griglia (riga 4, span su 2 colonne)
        fields_grid.addWidget(self.fiera_checkbox, 4, 0, 1, 2)
        # ========== FINE AGGIUNTA ==========
        
        content_layout.addLayout(fields_grid)
        
        # Pulsanti azione
        buttons_layout = QHBoxLayout()
        buttons_layout.addStretch()
        
        generate_btn = QPushButton("Genera Ordine")
        generate_btn.setObjectName("primaryBtn")
        generate_btn.clicked.connect(self.generate_doc)
        buttons_layout.addWidget(generate_btn)

        self.genera_btn = generate_btn
        self.aggiungi_ordine_btn = QPushButton("➕ Aggiorna l'ordine")
        self.aggiungi_ordine_btn.setStyleSheet("""
            QPushButton {
                background-color: #10b981;
                color: white;
                font-weight: bold;
                padding: 12px 24px;
                border-radius: 8px;
                font-size: 14px;
                min-width: 160px;
            }
            QPushButton:hover {
                background-color: #059669;
            }
        """)
        self.aggiungi_ordine_btn.clicked.connect(self.aggiungi_al_documento_esistente)
        self.aggiungi_ordine_btn.hide()  # Inizialmente nascosto
        buttons_layout.addWidget(self.aggiungi_ordine_btn)
        
        
        
        clear_btn = QPushButton("Deseleziona Tutto")
        clear_btn.setObjectName("dangerBtn")
        clear_btn.clicked.connect(self.deseleziona_tutto)
        buttons_layout.addWidget(clear_btn)
        
        content_layout.addLayout(buttons_layout)
        info_layout.addWidget(self.info_content)
        
        # Stato iniziale: espanso
        self.info_expanded = True
        
        parent_layout.addWidget(info_card)
    

    
    def rinomina_referenza_globale(self, vecchio_nome, nuovo_nome):
        """Rinomina una referenza in entrambi i file automaticamente - VERSIONE CORRETTA"""
        try:
            import shutil
            from datetime import datetime
            
            # 1. BACKUP automatico
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            shutil.copy2("catering_data.json", f"catering_data_backup_{timestamp}.json")
            shutil.copy2(gestore_percorsi.get_percorso("referenze_cibi"), f"referenze_cibi_backup_{timestamp}.json")
            
            # 2. Carica catering_data.json
            with open(gestore_percorsi.get_percorso("catering_data"), "r", encoding="utf-8") as f:
                catering_data = json.load(f)
            
            # 3. Carica referenze_cibi.json  
            with open(gestore_percorsi.get_percorso("referenze_cibi.json"), "r", encoding="utf-8") as f:
                referenze_data = json.load(f)
            
            # 4. Sostituisci in catering_data (ricorsivamente)
            def sostituisci_ricorsivo(obj):
                modificato = False
                if isinstance(obj, list):
                    for i, item in enumerate(obj):
                        if isinstance(item, str) and item.strip() == vecchio_nome:
                            obj[i] = nuovo_nome
                            modificato = True
                            print(f"🔄 Sostituito in lista: '{vecchio_nome}' → '{nuovo_nome}'")
                        elif isinstance(item, (dict, list)):
                            if sostituisci_ricorsivo(item):
                                modificato = True
                elif isinstance(obj, dict):
                    for value in obj.values():
                        if isinstance(value, (dict, list)):
                            if sostituisci_ricorsivo(value):
                                modificato = True
                return modificato
            
            modificato_catering = sostituisci_ricorsivo(catering_data)
            
            # 5. Gestisci referenze_cibi.json - CORREZIONE QUI
            modificato_referenze = False
            
            if vecchio_nome in referenze_data:
                # Salva i dati della vecchia referenza
                dati_vecchia_referenza = referenze_data[vecchio_nome]
                print(f"📋 Dati vecchia referenza: {dati_vecchia_referenza}")
                
                # PRIMA: Aggiungi la nuova referenza con gli stessi dati
                referenze_data[nuovo_nome] = dati_vecchia_referenza
                print(f"✅ Aggiunta nuova referenza: '{nuovo_nome}'")
                
                # POI: Rimuovi la vecchia referenza
                del referenze_data[vecchio_nome]
                print(f"🗑️ Rimossa vecchia referenza: '{vecchio_nome}'")
                
                modificato_referenze = True
            else:
                print(f"⚠️ Referenza '{vecchio_nome}' non trovata in referenze_cibi.json")
                # Aggiungi comunque la nuova con valori di default
                referenze_data[nuovo_nome] = {
                    "difficolta": 1.0,
                    "peso_unitario": 0.03333
                }
                modificato_referenze = True
                print(f"➕ Aggiunta nuova referenza con valori default: '{nuovo_nome}'")
            
            # 6. Salva entrambi i file SOLO se modificati
            messaggi_log = []
            
            if modificato_catering:
                with open("catering_data.json", "w", encoding="utf-8") as f:
                    json.dump(catering_data, f, indent=2, ensure_ascii=False)
                messaggi_log.append("✅ catering_data.json aggiornato")
            else:
                messaggi_log.append("⚠️ Nessuna modifica in catering_data.json")
            
            if modificato_referenze:
                with open(gestore_percorsi.get_percorso("referenze_cibi.json"), "w", encoding="utf-8") as f:
                    json.dump(referenze_data, f, indent=2, ensure_ascii=False)
                messaggi_log.append("✅ referenze_cibi.json aggiornato")
            else:
                messaggi_log.append("⚠️ Nessuna modifica in referenze_cibi.json")
            
            # 7. Aggiorna anche i dati in memoria
            self.ricarica_dati_referenze()
            messaggi_log.append("🔄 Dati in memoria ricaricati")
            
            # 8. Verifica finale
            print("\n" + "="*50)
            print("📊 VERIFICA FINALE:")
            print("="*50)
            
            # Verifica che la vecchia referenza sia stata rimossa
            with open(gestore_percorsi.get_percorso("referenze_cibi.json"), "r", encoding="utf-8") as f:
                verifica_referenze = json.load(f)
            
            vecchia_presente = vecchio_nome in verifica_referenze
            nuova_presente = nuovo_nome in verifica_referenze
            
            print(f"🔍 Vecchia referenza '{vecchio_nome}': {'❌ ANCORA PRESENTE!' if vecchia_presente else '✅ Rimossa'}")
            print(f"🔍 Nuova referenza '{nuovo_nome}': {'✅ Presente' if nuova_presente else '❌ MANCANTE!'}")
            
            # Messaggio finale
            messaggio_successo = (
                f"🎉 RINOMINAZIONE COMPLETATA!\n\n"
                f"Da: {vecchio_nome}\n"
                f"A: {nuovo_nome}\n\n"
                f"📋 DETTAGLI:\n" + "\n".join(messaggi_log) + "\n\n"
                f"💾 BACKUP CREATI:\n"
                f"- catering_data_backup_{timestamp}.json\n"
                f"- referenze_cibi_backup_{timestamp}.json\n\n"
                f"🔍 VERIFICA:\n"
                f"- Vecchia referenza: {'❌ Ancora presente!' if vecchia_presente else '✅ Rimossa'}\n"
                f"- Nuova referenza: {'✅ Presente' if nuova_presente else '❌ Mancante!'}"
            )
            
            QMessageBox.information(self, "Rinominazione Completata", messaggio_successo)
            
            return True
            
        except Exception as e:
            print(f"❌ ERRORE: {str(e)}")
            QMessageBox.critical(self, "Errore", f"Errore durante la rinominazione: {str(e)}")
            return False

    def dialog_rinomina_referenza(self):
        """Dialog per rinominare una referenza - VERSIONE SMART tollerante"""
        dialog = QDialog(self)
        dialog.setWindowTitle("🔄 Rinomina Referenza Globale")
        dialog.setFixedSize(700, 400)
        
        layout = QVBoxLayout(dialog)
        
        
        
        # Campo nome attuale
        layout.addWidget(QLabel("Nome attuale della referenza:"))
        vecchio_nome_input = QLineEdit()
        vecchio_nome_input.setPlaceholderText('es: "Mezze maniche con pomodorini ciliegini, basilico e bruciatini di pancetta",')
        layout.addWidget(vecchio_nome_input)
        
        # Label che mostra cosa verrà cercato
        vecchio_pulito_label = QLabel("")
        vecchio_pulito_label.setStyleSheet("background: #f8fafc; padding: 8px; border-radius: 3px; font-style: italic; color: #64748b;")
        layout.addWidget(vecchio_pulito_label)
        
        # Campo nuovo nome
        layout.addWidget(QLabel("Nuovo nome:"))
        nuovo_nome_input = QLineEdit()
        nuovo_nome_input.setPlaceholderText("es: Mezze maniche pomodorini e pancetta")
        layout.addWidget(nuovo_nome_input)
        
        # Label che mostra il nuovo nome pulito
        nuovo_pulito_label = QLabel("")
        nuovo_pulito_label.setStyleSheet("background: #f8fafc; padding: 8px; border-radius: 3px; font-style: italic; color: #64748b;")
        layout.addWidget(nuovo_pulito_label)
        
        # Anteprima
        anteprima_label = QLabel("")
        anteprima_label.setStyleSheet("background: #fef3f2; padding: 10px; border-radius: 5px; font-weight: bold;")
        layout.addWidget(anteprima_label)

        def pulisci_nome_input(testo_raw):
            """Pulisce l'input dall'utente rimuovendo caratteri JSON"""
            if not testo_raw:
                return ""
            
            # Rimuovi spazi iniziali e finali
            testo = testo_raw.strip()
            
            # Rimuovi virgolette all'inizio e alla fine
            if testo.startswith('"') and testo.endswith('"'):
                testo = testo[1:-1]
            elif testo.startswith("'") and testo.endswith("'"):
                testo = testo[1:-1]
            
            # Rimuovi virgolette solo all'inizio
            if testo.startswith('"') or testo.startswith("'"):
                testo = testo[1:]
            
            # Rimuovi virgola finale
            if testo.endswith(','):
                testo = testo[:-1]
            
            # Rimuovi virgolette finali dopo la virgola
            if testo.endswith('"') or testo.endswith("'"):
                testo = testo[:-1]
            
            # Rimuovi spazi extra
            testo = testo.strip()
            
            return testo
        
        def aggiorna_anteprima():
            vecchio_raw = vecchio_nome_input.text()
            nuovo_raw = nuovo_nome_input.text()
            
            vecchio_pulito = pulisci_nome_input(vecchio_raw)
            nuovo_pulito = pulisci_nome_input(nuovo_raw)
            
            # Mostra i nomi puliti
            if vecchio_pulito:
                vecchio_pulito_label.setText(f"🔍 Cercherò: {vecchio_pulito}")
            else:
                vecchio_pulito_label.setText("")
                
            if nuovo_pulito:
                nuovo_pulito_label.setText(f"✨ Diventerà: {nuovo_pulito}")
            else:
                nuovo_pulito_label.setText("")
            
            # Anteprima finale
            if vecchio_pulito and nuovo_pulito:
                anteprima_label.setText(f"🔄 OPERAZIONE: '{vecchio_pulito}' → '{nuovo_pulito}'")
                anteprima_label.setStyleSheet("background: #dcfce7; padding: 10px; border-radius: 5px; font-weight: bold; color: #166534;")
            else:
                anteprima_label.setText("")
                anteprima_label.setStyleSheet("background: #fef3f2; padding: 10px; border-radius: 5px; font-weight: bold;")
        
        vecchio_nome_input.textChanged.connect(aggiorna_anteprima)
        nuovo_nome_input.textChanged.connect(aggiorna_anteprima)
        
        # Pulsanti
        buttons_layout = QHBoxLayout()
        
        # Pulsante per testare se esiste
        test_btn = QPushButton("🔍 Testa se Esiste")
        test_btn.setStyleSheet("""
            QPushButton {
                background: #3b82f6;
                color: white;
                font-weight: bold;
                padding: 10px 15px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background: #2563eb;
            }
        """)
        
        def testa_esistenza():
            vecchio_pulito = pulisci_nome_input(vecchio_nome_input.text())
            if not vecchio_pulito:
                QMessageBox.information(dialog, "Test", "Inserisci prima un nome da testare!")
                return
            
            # Testa in catering_data.json
            try:
                with open(gestore_percorsi.get_percorso("catering_data"), "r", encoding="utf-8") as f:
                    catering_data = json.load(f)
                
                def cerca_ricorsivo(obj, target):
                    if isinstance(obj, list):
                        for item in obj:
                            if isinstance(item, str) and item.strip() == target:
                                return True
                            elif isinstance(item, (dict, list)):
                                if cerca_ricorsivo(item, target):
                                    return True
                    elif isinstance(obj, dict):
                        for value in obj.values():
                            if isinstance(value, (dict, list)):
                                if cerca_ricorsivo(value, target):
                                    return True
                    return False
                
                trovato_catering = cerca_ricorsivo(catering_data, vecchio_pulito)
                
                # Testa in referenze_cibi.json
                with open(gestore_percorsi.get_percorso("referenze_cibi.json"), "r", encoding="utf-8") as f:
                    referenze_data = json.load(f)
                
                trovato_referenze = vecchio_pulito in referenze_data
                
                # Risultato
                msg = f"🔍 RISULTATO TEST per: '{vecchio_pulito}'\n\n"
                msg += f"📄 catering_data.json: {'✅ TROVATO' if trovato_catering else '❌ NON trovato'}\n"
                msg += f"📊 referenze_cibi.json: {'✅ TROVATO' if trovato_referenze else '❌ NON trovato'}\n\n"
                
                if trovato_catering or trovato_referenze:
                    msg += "✅ Puoi procedere con la rinominazione!"
                else:
                    msg += "⚠️ Nome non trovato in nessun file. Controlla l'ortografia."
                
                QMessageBox.information(dialog, "Test Esistenza", msg)
                
            except Exception as e:
                QMessageBox.critical(dialog, "Errore Test", f"Errore durante il test: {str(e)}")
        
        test_btn.clicked.connect(testa_esistenza)
        buttons_layout.addWidget(test_btn)
        
        buttons_layout.addStretch()
        
        rinomina_btn = QPushButton("🔄 RINOMINA")
        rinomina_btn.setStyleSheet("""
            QPushButton {
                background: #dc2626;
                color: white;
                font-weight: bold;
                padding: 10px 20px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background: #b91c1c;
            }
        """)
        
        def esegui_rinomina():
            vecchio_raw = vecchio_nome_input.text()
            nuovo_raw = nuovo_nome_input.text()
            
            # Pulisci entrambi gli input
            vecchio_pulito = pulisci_nome_input(vecchio_raw)
            nuovo_pulito = pulisci_nome_input(nuovo_raw)
            
            if not vecchio_pulito or not nuovo_pulito:
                QMessageBox.warning(dialog, "Errore", "Inserisci sia il nome attuale che quello nuovo!")
                return
            
            # ✅ CORREZIONE QUI - Sintassi PyQt6
            reply = QMessageBox.question(dialog, "Conferma Rinominazione", 
                f"🔄 CONFERMI LA RINOMINAZIONE?\n\n"
                f"Da: {vecchio_pulito}\n"
                f"A:  {nuovo_pulito}\n\n"
                f"⚠️ Questa operazione modificherà entrambi i file JSON!\n"
                f"💾 Verranno creati backup automatici.",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            
            if reply == QMessageBox.StandardButton.Yes:
                # Usa i nomi puliti per la rinominazione
                if self.rinomina_referenza_globale(vecchio_pulito, nuovo_pulito):
                    dialog.accept()
        
        rinomina_btn.clicked.connect(esegui_rinomina)
        buttons_layout.addWidget(rinomina_btn)
        
        annulla_btn = QPushButton("Annulla")
        annulla_btn.clicked.connect(dialog.reject)
        buttons_layout.addWidget(annulla_btn)
        
        layout.addLayout(buttons_layout)
        
        dialog.exec()

    def ricarica_dati(self):
        """Ricarica i dati dell'archivio dal file JSON"""
        print("🔄 Ricaricando dati archivio...")
        self.load_orders()
        print("✅ Dati archivio ricaricati")
    
    
    
    
    def toggle_compact_mode(self):
            """Alterna tra modalità normale e compatta"""
            if hasattr(self, 'compact_mode') and self.compact_mode:
                # Torna alla modalità normale
                self.info_card.setMaximumHeight(16777215)  # Rimuovi limite altezza
                self.tabs.setMinimumHeight(500)
                self.compact_mode = False
                self.compact_btn.setText("📋 Modalità Compatta")
            else:
                # Attiva modalità compatta
                self.info_card.setMaximumHeight(200)  # Comprimi info ordine
                self.tabs.setMinimumHeight(700)  # Più spazio per i menu
                self.compact_mode = True
                self.compact_btn.setText("📋 Modalità Normale")
    
    def set_evento_fiera(self, checked):
        """Imposta se l'evento è in fiera"""
        self.evento_in_fiera = checked
        if checked:
            print("✅ Evento IN FIERA - 200g per persona")
        else:
            print("📍 Evento STANDARD - calcolo normale")
        
    def create_services_tabs(self, parent_layout):
        # Card per i servizi
        self.services_card = BusinessCard()
        services_layout = QVBoxLayout(self.services_card)
        
        # Header sezione
        section_header = QLabel("MENU E SERVIZI 📄​")
        section_header.setObjectName("sectionHeader")
        services_layout.addWidget(section_header)
        
        # ========== AGGIUNGI LA BARRA DI RICERCA QUI ==========
        # La barra di ricerca sarà aggiunta nel setup_search_functionality
        # =====================================================
        
        # Tabs con altezza dinamica
        self.tabs = QTabWidget()
        
        screen_height = QApplication.primaryScreen().geometry().height()
        self.tabs.setMinimumHeight(int(screen_height * 0.55))
        
        # Carica dati tabs
        try:
            with open(gestore_percorsi.get_percorso("catering_data"), "r", encoding="utf-8") as f:
                data = json.load(f)
                
            for tab_name, sections in data["tabs"].items():
                self.add_service_tab(tab_name, sections)
        except FileNotFoundError:
            print("Warning: catering_data.json non trovato.")
            self.add_demo_tab()
        
        services_layout.addWidget(self.tabs)
        
        # ========== CONFIGURA LA RICERCA DOPO AVER CREATO I TABS ==========
        self.setup_search_functionality()
        # ==================================================================
        
        parent_layout.addWidget(self.services_card)


    
    def toggle_info_section(self):
        """Espande/comprime la sezione Informazioni Cliente"""
        if self.info_expanded:
            # Comprimi informazioni cliente
            self.info_content.hide()
            self.info_header.setText("🔽 INFORMAZIONI CLIENTE 👥")
            self.info_expanded = False
            
            # ========== AGGIUNGI QUESTE RIGHE ==========
            # Più spazio per i tabs quando compressi
            screen_height = QApplication.primaryScreen().geometry().height()
            self.tabs.setMinimumHeight(int(screen_height * 0.7))  # 70% dello schermo
            # =========================================
            
        else:
            # Espandi informazioni cliente
            self.info_content.show()
            self.info_header.setText("🔼 INFORMAZIONI CLIENTE 👥")
            self.info_expanded = True
            
            # ========== AGGIUNGI QUESTE RIGHE ==========
            # Dimensione normale per i tabs
            screen_height = QApplication.primaryScreen().geometry().height()
            self.tabs.setMinimumHeight(int(screen_height * 0.55))  # 55% dello schermo
            # =========================================
    def resizeEvent(self, event):
        """Gestisce il ridimensionamento della finestra"""
        super().resizeEvent(event)
        
        # Aggiorna le dimensioni dei tabs quando ridimensioni la finestra
        if hasattr(self, 'tabs'):
            if hasattr(self, 'info_expanded') and not self.info_expanded:
                # Se info è compressa, usa più spazio
                new_height = int(self.height() * 0.65)
            else:
                # Se info è espansa, usa spazio normale
                new_height = int(self.height() * 0.5)
            
            self.tabs.setMinimumHeight(max(400, new_height))  # Minimo 400px

    
    def resize_scroll_areas(self, new_height):
        """Ridimensiona tutte le scroll area nei tabs"""
        for i in range(self.tabs.count()):
            tab_widget = self.tabs.widget(i)
            if tab_widget:
                # Trova tutte le scroll area in questo tab
                scroll_areas = tab_widget.findChildren(QScrollArea)
                for scroll_area in scroll_areas:
                    scroll_area.setMinimumHeight(new_height - 100)  # Lascia spazio per header
                    scroll_area.updateGeometry()
                    
                # Aggiorna anche i GroupBox
                group_boxes = tab_widget.findChildren(QGroupBox)
                for group_box in group_boxes:
                    group_box.updateGeometry()
    
    def add_demo_tab(self):
        """Aggiunge un tab dimostrativo se il file JSON non esiste"""
        tab_widget = QWidget()
        tab_layout = QVBoxLayout(tab_widget)
        
        info_label = QLabel("⚠️ File catering_data.json non trovato.\n\nCreare il file di configurazione per utilizzare l'applicazione.")
        info_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        info_label.setStyleSheet("font-size: 16px; color: #e74c3c; padding: 50px;")
        tab_layout.addWidget(info_label)
        
        self.tabs.addTab(tab_widget, "Demo")
    
    def add_service_tab(self, name, sections):
        # Widget principale del tab
        tab_widget = QWidget()
        main_layout = QVBoxLayout(tab_widget)
        main_layout.setContentsMargins(15, 15, 15, 15)
        main_layout.setSpacing(15)
        
        # Se è aperitivo a buffet, aggiungi controlli tipo
        if name.lower() == "aperitivo a buffet":
            tipo_card = BusinessCard()
            tipo_layout = QHBoxLayout(tipo_card)
            
            tipo_label = QLabel("Tipo Aperitivo:")
            tipo_label.setStyleSheet("font-weight: bold; color: #2c3e50;")
            tipo_layout.addWidget(tipo_label)
            
            self.ap_group = QButtonGroup()
            leggero_radio = CustomCheckBox("Leggero")
            rinforzato_radio = CustomCheckBox("Rinforzato")
            
            # CORREZIONE: Non impostare nessuna checkbox come checked all'inizio
            # Lascia entrambe deselezionate per evitare confusione
            leggero_radio.setChecked(False)
            rinforzato_radio.setChecked(False)
            
            def on_leggero_clicked():
                if leggero_radio.isChecked():
                    rinforzato_radio.setChecked(False)
                    self.set_ap_tipo("Leggero")
                    print("✅ Selezionato: Aperitivo Leggero")
                else:
                    # Se si deseleziona leggero, azzera il tipo
                    self.set_ap_tipo("")
                    print("❌ Deselezionato: Aperitivo Leggero")

            def on_rinforzato_clicked():
                if rinforzato_radio.isChecked():
                    leggero_radio.setChecked(False)
                    self.set_ap_tipo("Rinforzato")
                    print("✅ Selezionato: Aperitivo Rinforzato")
                else:
                    # Se si deseleziona rinforzato, azzera il tipo
                    self.set_ap_tipo("")
                    print("❌ Deselezionato: Aperitivo Rinforzato")
            
            leggero_radio.toggled.connect(on_leggero_clicked)
            rinforzato_radio.toggled.connect(on_rinforzato_clicked)
            
            tipo_layout.addWidget(leggero_radio)
            tipo_layout.addWidget(rinforzato_radio)
            tipo_layout.addStretch()
            
            main_layout.addWidget(tipo_card)

        
        # Contenitore per tutte le sezioni
        sections_widget = QWidget()
        sections_layout = QVBoxLayout(sections_widget)
        sections_layout.setSpacing(15)
        
        # Aggiungi sezioni direttamente
        for section_name, items in sections.items():
            self.add_service_section(sections_layout, name, section_name, items)
        
        # Scroll area semplificata
        scroll = QScrollArea()
        scroll.setStyleSheet("QScrollBar:vertical { width: 0px; }")  # AGGIUNGI QUESTA
        scroll.setWidget(sections_widget)
        scroll.setWidgetResizable(True)
        
        # Aggiungi lo scroll al layout principale
        main_layout.addWidget(scroll)
        
        # Aggiungi il tab
        self.tabs.addTab(tab_widget, name)
    
    
    def contains_food_items(self, items):
        """Controlla se una lista di items contiene chiaramente cibo"""
        food_keywords = [
            # Cibi base
            "indivia", "rustici", "sandwich", "calzoncini", "pizza", "focaccia", 
            "tramezzino", "bruschett", "crostini", "antipast", "salum", "formagg",
            "prosciutto", "mortadella", "mozzarell", "parmigiano", "gorgonzola",
            "pasta", "riso", "lasagne", "gnocchi", "risotto", "dolc", "tiramisù",
            "torta", "crostat", "biscott", "macaron", "cannol", "sfogliat",
            "tarall", "olive", "pomodor", "basilico", "rucola", "spinaci","crostini","rombetti","bocconcini","torta","frittata","mousse"
            
            # Aperitivi specifici
            "mousse", "quinoa", "frittata", "cous cous", "quiche", "piroghe","fioriere quadrate + decori",
            "bicchier", "coppette", "rotolini", "cubetti", "spiedini", "bocconcini",
            "pizzette", "dadini", "pagoda", "ananas", "speck", "caprino",
            "melanzane", "ratatouille", "guanciale", "pistacchio", "mandorle",
            "fagiolini", "pesto", "orzo", "venere", "balsamico", "confit",
            "crescenti", "edamame", "bufala", "ciliegini", "margherita",
            "vegetariani", "parma", "verdurine", "lime", "menta","crostini","rombetti","bocconcini","torta","frittata","mousse"
        ]
        
        for item in items:
            item_lower = item.lower()
            if any(keyword in item_lower for keyword in food_keywords):
                return True
        
        return False

    def are_real_accessories(self, items):
        """Controlla se gli items sono VERAMENTE accessori (non cibo) - VERSIONE MIGLIORATA"""
        accessory_keywords = [
            "tavoli", "tovagliato", "servizio", "camerieri", "cameriere", 
            "accessori monouso", "cestini", "smaltimento", "consegna", 
            "allestimento", "runner", "divisa", "biodegradabili", 
            "differenziata", "rifiuti", "attrezzature", "personale",
            "pinze", "guanti", "alzate", "cucchiai", "forchette", 
            "acc.bio", "nr.", "cam", "coprimacche", "posate", "piatti",
            "bicchieri", "flutes", "calici", "bowl", "sgabelli", "tavolini"
        ]
        
        # Debug degli items
        print(f"🔍 Analizzando {len(items)} items per accessori:")
        
        accessory_count = 0
        for item in items:
            item_lower = item.lower()
            is_accessory = any(keyword in item_lower for keyword in accessory_keywords)
            # AGGIUNGI QUESTO DEBUG:
        if "crostini" in item_lower or "rombetti" in item_lower or "bocconcini" in item_lower:
            print(f"🔍 ANALISI SPECIFICA: '{item}'")
            print(f"  → Parole trovate: {[kw for kw in accessory_keywords if kw in item_lower]}")
            print(f"  → Categorizzato come: {'ACCESSORIO' if is_accessory else 'CIBO'}")
            
            if is_accessory:
                accessory_count += 1
                print(f"  ✅ ACCESSORIO: {item}")
            else:
                print(f"  🍕 CIBO: {item}")
        
        percentage = (accessory_count / len(items)) * 100 if len(items) > 0 else 0
        print(f"📊 Accessori: {accessory_count}/{len(items)} ({percentage:.1f}%)")
        
        # Se più del 30% sono accessori, considera la sezione come accessori
        is_accessory_section = percentage > 30
        print(f"🏷️ Risultato: {'ACCESSORI' if is_accessory_section else 'CIBO'}")
        
        return is_accessory_section

    def add_service_section(self, parent_layout, tab_name, section_name, items):
        """Crea una sezione nel tab con layout a griglia per le referenze"""
        
        # Determina il tipo di sezione (stesso codice esistente)
        section_name_lower = section_name.lower()
        
        if any(keyword in section_name_lower for keyword in ["bevande", "bibite", "acqua", "prosecco", "aperol", "caffè", "té"]):
            section_type = "bevande"
            print(f"🥤 Categorizzata come: BEVANDE")
        elif any(keyword in section_name_lower for keyword in ["accessori", "sala", "attrezzature", "servizio"]):
            section_type = "accessori"
            print(f"🛠️ Categorizzata come: ACCESSORI")
        elif self.are_real_accessories(items):
            section_type = "accessori"
            print(f"🛠️ Categorizzata come: ACCESSORI (by content)")
        else:
            section_type = "cibo"
            # WHITELIST: Sezioni che sono sempre cibo
        food_sections = ["finger food", "fingerfood", "salati", "dolci", "antipasti", "crostini", "bocconcini"]

        if any(keyword in section_name_lower for keyword in food_sections):
            section_type = "cibo"
            print(f"🍕 FORZATO COME CIBO (whitelist): {section_name}")
        elif any(keyword in section_name_lower for keyword in ["bevande", "bibite", "acqua"]):
            section_type = "bevande"
        elif any(keyword in section_name_lower for keyword in ["accessori", "sala", "attrezzature"]):
            section_type = "accessori"
        elif self.are_real_accessories(items):
            section_type = "accessori"
        else:
            section_type = "cibo"
            print(f"🍕 Categorizzata come: CIBO")
        
        menu_extra_types = ["vegano", "no lattosio", "no uova", "celiaci"]
        is_menu_extra = any(menu_type in section_name.lower() for menu_type in menu_extra_types)
        
        if is_menu_extra:
            # Crea sezione speciale per menu extra
            self.create_menu_extra_section_special(parent_layout, section_name, items)
            return  # Esci dal metodo normale
        # GroupBox per la sezione
        section_group = QGroupBox(section_name.upper())
        section_layout = QVBoxLayout(section_group)
        section_layout.setSpacing(10)
        
        # Pulsanti di controllo (stesso codice esistente)
        controls_layout = QHBoxLayout()
        
        select_all_btn = QPushButton("Seleziona Tutto")
        select_all_btn.setObjectName("successBtn")
        select_all_btn.setMaximumWidth(120)
        
        deselect_all_btn = QPushButton("Deseleziona Tutto")
        deselect_all_btn.setObjectName("dangerBtn")
        deselect_all_btn.setMaximumWidth(120)
        
        controls_layout.addWidget(select_all_btn)
        controls_layout.addWidget(deselect_all_btn)
        controls_layout.addStretch()
        
        section_layout.addLayout(controls_layout)
        
        # ========== MODIFICA PRINCIPALE: LAYOUT A GRIGLIA PER LE CHECKBOX ==========
        # Crea un layout a griglia per le checkbox (2 colonne)
        checkboxes_widget = QWidget()
        checkboxes_grid = QGridLayout(checkboxes_widget)
        checkboxes_grid.setSpacing(8)
        checkboxes_grid.setContentsMargins(10, 10, 10, 10)
        
        # Aggiungi le checkbox alla griglia (2 per riga)
        for i, item in enumerate(items):
            checkbox = CustomCheckBox(item)
            
            # Crea la chiave univoca per la voce
            key = f"{tab_name.lower()}|{section_name.lower()}|{item.lower()}"
            
            # REGISTRA CATEGORIA E VOCE ORIGINALE
            self.voce_categoria[key] = section_type
            self.voce_originale[key] = item
            self.selected[key] = checkbox
            
            # NUOVO: Se è un primo, collega la funzione di auto-selezione
            if section_type == "cibo" and ("primi piatti" in section_name.lower()):
                checkbox.toggled.connect(lambda checked, k=key: self.gestisci_selezione_primo(self.selected[k], k))

            # NUOVO: Se è una bevanda con CONS, collega la funzione di auto-selezione
            if section_type == "bevande" and "+cons" in item.lower():
                checkbox.toggled.connect(lambda checked, k=key: self.gestisci_selezione_bevanda(self.selected[k], k))

            # CORREZIONE: Sistema di riconoscimento insalate con esclusione condimenti
            keywords_insalata = ["insalat","caprese"]
            keywords_da_escludere = ["condimento", "condimenti"]

            # È un'insalata se contiene le parole chiave MA NON contiene "condimento"
            is_insalata = (any(keyword in item.lower() for keyword in keywords_insalata) and
                        not any(esclusione in item.lower() for esclusione in keywords_da_escludere))

            if is_insalata:
                checkbox.toggled.connect(lambda checked, k=key: self.gestisci_selezione_insalata(self.selected[k], k))
                print(f"   🥗 Collegata insalata: {item}")
            # NUOVO: Se è un secondo piatto o contorno, collega la funzione di auto-selezione
                
            if section_type == "cibo" and ("secondi piatti" in section_name.lower() or "contorni" in section_name.lower()):
                checkbox.toggled.connect(lambda checked, k=key: self.gestisci_selezione_secondi_contorni(self.selected[k], k))


            
            # Debug delle voci accessori
            if section_type == "accessori":
                print(f"  ✅ Registrato accessorio: {item} → {key[:50]}...")
            
            # Calcola la posizione nella griglia (2 colonne)
            row = i // 2  # Riga
            col = i % 2   # Colonna (0 o 1)
            
            # Aggiungi la checkbox alla griglia
            checkboxes_grid.addWidget(checkbox, row, col)
        
        # Aggiungi il widget con la griglia al layout della sezione
        section_layout.addWidget(checkboxes_widget)
        # ========== FINE MODIFICA ==========
        
        # Connetti i pulsanti di controllo (stesso codice esistente)
        def select_all():
            for item in items:
                key = f"{tab_name.lower()}|{section_name.lower()}|{item.lower()}"
                if key in self.selected:
                    self.selected[key].setChecked(True)
        
        def deselect_all():
            for item in items:
                key = f"{tab_name.lower()}|{section_name.lower()}|{item.lower()}"
                if key in self.selected:
                    self.selected[key].setChecked(False)
        
        select_all_btn.clicked.connect(select_all)
        deselect_all_btn.clicked.connect(deselect_all)
        
        parent_layout.addWidget(section_group)
    
    
    def select_all_in_section(self, checkboxes, value):
        for checkbox in checkboxes:
            checkbox.setChecked(value)
    
    def set_ap_tipo(self, tipo):
        """Imposta il tipo di aperitivo"""
        self.ap_buffet_tipo = tipo
        print(f"🔧 Tipo aperitivo impostato: '{tipo}'")
    def reset_aperitivo_selection(self):
        """Reset della selezione aperitivo - da chiamare quando serve"""
        self.ap_buffet_tipo = ""
        
        # Se esistono i radio button, deselezionali
        if hasattr(self, 'ap_group'):
            for button in self.ap_group.buttons():
                button.setChecked(False)
    def reset_fields(self):
        """Reset tutti i campi di input"""
        for field_name, entry in self.entries.items():
            if field_name == "Data evento":
                entry.setDate(QDate.currentDate())
            else:
                entry.clear()
        
        QMessageBox.information(self, "Reset", "Tutti i campi sono stati resettati.")
    def create_menu_extra_section_special(self, parent_layout, section_name, items):
        """Crea sezione speciale per menu extra con stile colorato"""
        
        # Determina tipo menu e colore
        menu_type = None
        color = "#2c3e50"  # Default
        
        if "vegano" in section_name.lower():
            menu_type = "Vegano"
            color = "#27ae60"  # Verde
        elif "no lattosio" in section_name.lower():
            menu_type = "No Lattosio"
            color = "#3498db"  # Blu
        elif "no uova" in section_name.lower():
            menu_type = "No Uova"
            color = "#f39c12"  # Arancione
        elif "celiaci" in section_name.lower():
            menu_type = "Celiaci"
            color = "#e74c3c"  # Rosso
        
        # Frame principale
        menu_frame = QGroupBox()
        menu_frame.setStyleSheet(f"""
            QGroupBox {{
                border: 3px solid {color};
                border-radius: 12px;
                margin-top: 20px;
                background-color: rgba({int(color[1:3], 16)}, {int(color[3:5], 16)}, {int(color[5:7], 16)}, 0.05);
            }}
        """)
        
        menu_layout = QVBoxLayout(menu_frame)
        menu_layout.setSpacing(15)
        
        # Header con titolo e controlli
        header_layout = QHBoxLayout()
        
        # Titolo
        title_label = QLabel(f"🌱 {section_name.upper()}")
        title_label.setStyleSheet(f"""
            color: {color}; 
            font-size: 16px; 
            font-weight: bold;
            margin: 10px;
        """)
        header_layout.addWidget(title_label)
        
        header_layout.addStretch()
        
        # SpinBox per numero persone
        label_persone = QLabel("Persone:")
        label_persone.setStyleSheet(f"color: {color}; font-weight: bold;")
        header_layout.addWidget(label_persone)
        
        spinbox = QSpinBox()
        spinbox.setMinimum(0)
        spinbox.setMaximum(999)
        spinbox.setStyleSheet(f"""
            QSpinBox {{
                padding: 8px;
                border: 2px solid {color};
                border-radius: 8px;
                font-size: 14px;
                font-weight: bold;
                background-color: white;
                min-width: 80px;
            }}
        """)
        header_layout.addWidget(spinbox)
        
        # Salva il widget per uso successivo
        if menu_type:
            self.menu_extra_widgets[menu_type] = spinbox
        
        menu_layout.addLayout(header_layout)
        
        # ========== NUOVA SEZIONE: PULSANTI DI CONTROLLO ==========
        controls_layout = QHBoxLayout()
        
        # Pulsante Seleziona Tutto
        select_all_btn = QPushButton("Seleziona Tutto")
        select_all_btn.setObjectName("successBtn")
        select_all_btn.setMaximumWidth(120)
        
        # Pulsante Deseleziona Tutto
        deselect_all_btn = QPushButton("Deseleziona Tutto")
        deselect_all_btn.setObjectName("dangerBtn")
        deselect_all_btn.setMaximumWidth(120)
        
        # Funzioni per i pulsanti
        def seleziona_tutto_menu():
            """Seleziona tutte le checkbox di questo menu intolleranti"""
            for key, checkbox in self.selected.items():
                if (hasattr(checkbox, 'isChecked') and 
                    menu_type and menu_type.lower() in key.lower()):
                    checkbox.setChecked(True)
            print(f"✅ Selezionate tutte le voci del menu {menu_type}")
        
        def deseleziona_tutto_menu():
            """Deseleziona tutte le checkbox di questo menu intolleranti E azzera persone"""
            for key, checkbox in self.selected.items():
                if (hasattr(checkbox, 'isChecked') and 
                    menu_type and menu_type.lower() in key.lower()):
                    checkbox.setChecked(False)
            
            # NUOVO: Azzera anche il numero di persone
            spinbox.setValue(0)
            print(f"❌ Deselezionate tutte le voci del menu {menu_type} e azzerato numero persone")
        
        # Collega i pulsanti alle funzioni
        select_all_btn.clicked.connect(seleziona_tutto_menu)
        deselect_all_btn.clicked.connect(deseleziona_tutto_menu)
        
        controls_layout.addWidget(select_all_btn)
        controls_layout.addWidget(deselect_all_btn)
        controls_layout.addStretch()
        
        menu_layout.addLayout(controls_layout)
        
        # Lista delle voci (come sezione normale ma più compatta)
        items_layout = QVBoxLayout()
        items_layout.setSpacing(5)
        
        for item in items:
            key = f"{section_name}|{item}"
            checkbox = CustomCheckBox(item)
            
            # Collegamento normale            
            # Salva nei dizionari esistenti
            self.selected[key] = checkbox
            self.voce_originale[key] = item
            self.voce_categoria[key] = "intolleranti"  # I menu extra sono sempre cibo
            
            items_layout.addWidget(checkbox)
        
        menu_layout.addLayout(items_layout)
        
        # Nota informativa
        info_label = QLabel("💡 Le quantità verranno calcolate automaticamente per il numero di persone specificato")
        info_label.setStyleSheet(f"""
            color: {color}; 
            font-style: italic; 
            font-size: 11px;
            margin-top: 10px;
        """)
        menu_layout.addWidget(info_label)
        
        parent_layout.addWidget(menu_frame)
    def aggiungi_riquadro_note(self, doc):
        """Aggiunge un piccolo quadratino per le note sotto l'elenco del cibo"""
        
        # Aggiungi un po' di spazio prima
        doc.add_paragraph()
        
        # Crea una piccola tabella per le note
        note_table = doc.add_table(rows=1, cols=1)
        note_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        note_table.autofit = False
        note_table.columns[0].width = Inches(1.0)  # Larghezza ridotta
        
        # Cella per le note (piccola)
        note_cell = note_table.cell(0, 0)
        
        # Aggiungi titolo "NOTE:"
        p_title = note_cell.paragraphs[0]
        run_title = p_title.add_run("NOTE:")
        run_title.bold = True
        run_title.font.size = Pt(10)  # Font più piccolo
        run_title.font.name = "Calibri"
        
        # Aggiungi solo 2 righe vuote (piccolo spazio)
        for _ in range(1):
            note_cell.add_paragraph()
        
        # Applica bordi al quadratino
        self.apply_borders(note_cell)
    def deseleziona_tutto(self):
        """Deseleziona tutte le checkbox E azzera tutti i widget persone"""
        
        # Deseleziona tutte le checkbox (logica esistente)
        for checkbox in self.selected.values():
            if hasattr(checkbox, 'setChecked'):
                checkbox.setChecked(False)
        
        # NUOVO: Azzera tutti i widget persone dei menu intolleranti
        if hasattr(self, 'menu_extra_widgets'):
            for menu_type, widget in self.menu_extra_widgets.items():
                widget.setValue(0)
                print(f"   🔄 Azzerato widget persone per {menu_type}")
        
        print("❌ Tutte le selezioni sono state annullate e tutti i contatori persone azzerati")

    
    def nuovo_ordine(self):
        """Crea un nuovo ordine pulito"""
        
        # ========== RESET MODALITÀ AGGIUNGI SERVIZI ==========
        self.reset_modalita_aggiungi_servizi()
        
        # Reset normale
        self.reset_fields()
        self.deseleziona_tutto()
        
        print("✅ Nuovo ordine - tutte le modalità resettate")
        
    def show_about(self):
        """Mostra informazioni sull'applicazione"""
        QMessageBox.about(self, "Informazioni", 
                         "Gestione Ordini Catering v1.0\n\n"
                         "Sistema professionale per la gestione\n"
                         "degli ordini di catering.\n\n"
                         "© 2025 - Versione PyQt6")
        
    def mostra_statistiche(self):
        stats_window = StatisticsWindow(self)
        stats_window.exec()
    
    def genera_riepilogo_cucina(self):
        # Dialog per selezionare data
        dialog = QDialog(self)
        dialog.setWindowTitle("Riepilogo Cucina per Data")
        dialog.setGeometry(200, 200, 400, 200)
        dialog.setStyleSheet(PROFESSIONAL_STYLE)
        
        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)
        
        # Header
        header = QLabel("RIEPILOGO CUCINA")
        header.setObjectName("sectionHeader")
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(header)
        
        # Descrizione
        desc = QLabel("Seleziona la data per generare il riepilogo dettagliato delle quantità necessarie in cucina.")
        desc.setWordWrap(True)
        desc.setStyleSheet("color: #7f8c8d; margin: 10px 0;")
        layout.addWidget(desc)
        
        # Data picker
        date_layout = QHBoxLayout()
        date_layout.addWidget(QLabel("Data evento:"))
        
        date_edit = QDateEdit()
        date_edit.setDate(QDate.currentDate())
        date_edit.setDisplayFormat("dd/MM/yyyy")
        date_edit.setCalendarPopup(True)
        date_layout.addWidget(date_edit)
        
        layout.addLayout(date_layout)
        
        # Pulsanti
        buttons_layout = QHBoxLayout()
        buttons_layout.addStretch()
        
        cancel_btn = QPushButton("Annulla")
        cancel_btn.setObjectName("secondaryBtn")
        cancel_btn.clicked.connect(dialog.reject)
        buttons_layout.addWidget(cancel_btn)
        
        generate_btn = QPushButton("Genera Riepilogo")
        generate_btn.clicked.connect(lambda: self.crea_riepilogo_cucina(date_edit.date().toString("dd/MM/yyyy")))
        generate_btn.clicked.connect(dialog.accept)
        buttons_layout.addWidget(generate_btn)
        
        layout.addLayout(buttons_layout)
        dialog.exec()
    
    def crea_riepilogo_cucina(self, data_target):
        try:
            with open(ORDINI_FILE, "r", encoding="utf-8") as f:
                ordini = json.load(f)
        except FileNotFoundError:
            QMessageBox.warning(self, "Nessun dato", "Non ci sono ordini salvati.")
            return
        
        # Filtra ordini per la data selezionata
        ordini_giorno = []
        for ordine in ordini:
            if ordine["campi"].get("Data evento") == data_target:
                ordini_giorno.append(ordine)
        
        if not ordini_giorno:
            QMessageBox.warning(self, "Nessun ordine", f"Nessun ordine trovato per il {data_target}")
            return

        # Crea documento Word
        doc = Document()
        
        # Formatta la data
        data_obj = datetime.strptime(data_target, "%d/%m/%Y")
        giorno_str = data_obj.strftime("%A %d %B %Y").upper()
        
        # Titolo
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = 1  # centra
        title_run = title_paragraph.add_run(f"RIEPILOGO CUCINA\n{giorno_str}")
        title_run.bold = True
        title_run.font.size = Pt(18)
        
        doc.add_paragraph()  # spazio
        
        # Raggruppa tutte le voci di cibo con le relative quantità
        riepilogo_cibo = defaultdict(lambda: {"quantita": 0, "unita": ""})
        
        for i, ordine in enumerate(ordini_giorno, 1):
            persone = int(ordine["campi"].get("Numero persone", 0))
            nome_cliente = ordine["campi"].get("Nome", f"Cliente {i}")
            
            print(f"\n📋 === ORDINE {i}: {nome_cliente} ({persone} persone) ===")
            
            # ========== PRIMA: PROVA A LEGGERE DAL WORD ==========
            quantita_da_word = {}
            path_docx = ordine.get("path_docx")
            if path_docx and os.path.exists(path_docx):
                quantita_da_word = self.leggi_quantita_da_word(path_docx)
                if quantita_da_word:
                    print(f"✅ Usato quantità modificate dal Word: {len(quantita_da_word)} voci")
                else:
                    print("⚠️  Nessuna modifica trovata nel Word, uso calcolo automatico")
            else:
                print("❌ File Word non trovato, uso calcolo automatico")
            
            # Determina tipo servizio
            servizio_lower = ""
            scelte = ordine.get("scelte", {})
            tipo_ap = ordine.get("tipo_ap", "leggero").lower()
            
            if any(k.startswith("coffee break|") and v for k, v in scelte.items()):
                servizio_lower = "coffee break"
            elif any(k.startswith("tea break|") and v for k, v in scelte.items()):
                servizio_lower = "tea break"
            elif any(k.startswith("servito|") and v for k, v in scelte.items()):
                servizio_lower = "servito"
            elif any(k.startswith("lunch buffet|") and v for k, v in scelte.items()):
                servizio_lower = "lunch buffet"
            elif any(k.startswith("aperitivo a buffet|") and v for k, v in scelte.items()):
                servizio_lower = "aperitivo a buffet"
            
            # ========== GESTIONE EVENTO_IN_FIERA PER RIEPILOGO ==========
            stato_fiera_originale = getattr(self, 'evento_in_fiera', False)
            ordine_in_fiera = ordine.get("evento_in_fiera", False)
            self.evento_in_fiera = ordine_in_fiera
            print(f"🏢 Fiera: {ordine_in_fiera}")
            
            # Calcola referenze come nel metodo originale
            self.referenze_primi = []
            self.referenze_dolci = []
            self.referenze_salato = []
            
            for k, v in scelte.items():
                if v and self.voce_categoria.get(k) == "cibo":
                    voce_lower = self.voce_originale.get(k, "").strip().lower()
                    if "bis di primi piatti" in k or any(x in voce_lower for x in ["lasagne", "crespelle", "cannelloni", "rosette", "nidi", "gnocchi alla parigina", "intrighi", "tagliatelle", "strettine"]):
                        self.referenze_primi.append(k)
                    elif "dolci" in k:
                        self.referenze_dolci.append(k)
                    else:
                        self.referenze_salato.append(k)
            
            referenze_cibo = self.referenze_primi + self.referenze_dolci + self.referenze_salato
            
            # ========== CALCOLA O USA QUANTITÀ DAL WORD (VERSIONE CORRETTA) ==========
            if quantita_da_word:
                # USA TUTTE LE QUANTITÀ DAL WORD
                print("📄 Usando TUTTE le quantità dal Word:")
                for key, info in quantita_da_word.items():
                    voce = info['nome']
                    quantita = info['quantita']
                    print(f"   📝 {voce}: {quantita} (dal Word)")
                    self.aggrega_quantita_riepilogo(riepilogo_cibo, voce, quantita)
            else:
                # CALCOLO AUTOMATICO NORMALE
                print("🧮 Calcolando automaticamente:")
                for key in referenze_cibo:
                    voce = self.voce_originale.get(key, key)
                    quantita = self.calcola_quantita(key, referenze_cibo, persone, servizio_lower, tipo_ap)
                    print(f"   🧮 {voce}: {quantita} (calcolato)")
                    self.aggrega_quantita_riepilogo(riepilogo_cibo, voce, quantita)

            # ========== RIPRISTINA STATO ORIGINALE ==========
            self.evento_in_fiera = stato_fiera_originale
        
        # Crea tabella riepilogo (FUORI dal loop degli ordini)
        if riepilogo_cibo:
            header_p = doc.add_paragraph()
            header_run = header_p.add_run("RIEPILOGO QUANTITÀ TOTALI")
            header_run.bold = True
            header_run.underline = True
            header_run.font.size = Pt(14)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            header_cells = table.rows[0].cells
            header_cells[0].text = "PRODOTTO"
            header_cells[1].text = "QUANTITÀ"
            
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)
            
            def ordina_cibo(item):
                voce = item[0].lower()
                if any(x in voce for x in ["lasagne", "crespelle", "cannelloni", "rosette", "nidi", "gnocchi"]):
                    return (1, voce)
                elif "dolc" in voce or "torta" in voce or "crostat" in voce or "tiramisù" in voce:
                    return (3, voce)
                else:
                    return (2, voce)
            
            for prodotto, info in sorted(riepilogo_cibo.items(), key=ordina_cibo):
                row_cells = table.add_row().cells
                row_cells[0].text = prodotto
                
                quantita = info["quantita"]
                unita = info["unita"]
                
                row_cells[1].text = self.formatta_quantita_riepilogo(quantita, unita, prodotto)
        
        # Salva documento
        data_file = data_target.replace("/", "-")
        cartella_cucina = gestore_percorsi.get_percorso("riepilogo_cucina")
        os.makedirs(cartella_cucina, exist_ok=True)
        
        filename = os.path.join(cartella_cucina, f"Riepilogo_Cucina_{data_file}.docx")
        doc.save(filename)
        
        # Mostra messaggio di successo
        QMessageBox.information(self, "Riepilogo Creato", 
                            f"Riepilogo cucina salvato con successo!\n\n"
                            f"Percorso: {filename}\n"
                            f"Ordini elaborati: {len(ordini_giorno)}\n"
                            f"Data: {data_target}")
    

    # PROBLEMA: Il riepilogo legge solo il primo servizio e ignora il secondo
# SOLUZIONE: Migliorare il metodo leggi_quantita_da_word per leggere TUTTO
    def aggrega_quantita_riepilogo(self, riepilogo_cibo, voce, quantita):
        """Aggrega quantità nel riepilogo cucina RISPETTANDO l'unità originale"""
        
        # ========== RISPETTA L'UNITÀ ORIGINALE DAL WORD ==========
        if isinstance(quantita, str):
            if "kg" in quantita.lower():
                # MANTIENI KG se la quantità originale è in kg
                try:
                    peso = float(quantita.split()[0])
                    riepilogo_cibo[voce]["quantita"] += peso
                    riepilogo_cibo[voce]["unita"] = "kg"
                    print(f"   ✅ Aggregato {voce}: +{peso} kg (totale: {riepilogo_cibo[voce]['quantita']} kg)")
                except (ValueError, IndexError):
                    riepilogo_cibo[voce]["quantita"] += 1
                    riepilogo_cibo[voce]["unita"] = "kg"
            
            elif any(x in quantita.lower() for x in ["lasagne", "crespelle", "cannelloni", "rosette", "nidi"]):
                # Primi elaborati: mantieni unità specifica
                try:
                    numero = int(quantita.split()[0])
                    tipo_porzione = " ".join(quantita.split()[1:])
                    riepilogo_cibo[voce]["quantita"] += numero
                    riepilogo_cibo[voce]["unita"] = tipo_porzione
                    print(f"   ✅ Aggregato {voce}: +{numero} {tipo_porzione}")
                except (ValueError, IndexError):
                    riepilogo_cibo[voce]["quantita"] += 1
                    riepilogo_cibo[voce]["unita"] = "porzioni"
            
            elif "pz" in quantita.lower():
                # MANTIENI PZ se la quantità originale è in pezzi
                try:
                    numero = int(quantita.split()[0])
                    riepilogo_cibo[voce]["quantita"] += numero
                    riepilogo_cibo[voce]["unita"] = "pz"
                    print(f"   ✅ Aggregato {voce}: +{numero} pz (totale: {riepilogo_cibo[voce]['quantita']} pz)")
                except (ValueError, IndexError):
                    riepilogo_cibo[voce]["quantita"] += 1
                    riepilogo_cibo[voce]["unita"] = "pz"
            
            else:
                # Numero senza unità: controlla se è lasagne
                try:
                    numero = int(quantita)
                    
                    # ========== GESTIONE SPECIALE PER LASAGNE ==========
                    if "lasagne" in voce.lower():
                        riepilogo_cibo[voce]["quantita"] += numero
                        riepilogo_cibo[voce]["unita"] = "teglie"
                        print(f"   ✅ Aggregato {voce}: +{numero} teglie (totale: {riepilogo_cibo[voce]['quantita']} teglie)")
                    else:
                        riepilogo_cibo[voce]["quantita"] += numero
                        riepilogo_cibo[voce]["unita"] = "pz"
                        print(f"   ✅ Aggregato {voce}: +{numero} pz (totale: {riepilogo_cibo[voce]['quantita']} pz)")
                except (ValueError, TypeError):
                    riepilogo_cibo[voce]["quantita"] += 1
                    riepilogo_cibo[voce]["unita"] = "pz"
        
        else:
            # Quantità numerica: controlla se è lasagne
            try:
                numero = int(quantita)
                
                # ========== GESTIONE SPECIALE PER LASAGNE ==========
                if "lasagne" in voce.lower():
                    riepilogo_cibo[voce]["quantita"] += numero
                    riepilogo_cibo[voce]["unita"] = "teglie"
                    print(f"   ✅ Aggregato {voce}: +{numero} teglie (totale: {riepilogo_cibo[voce]['quantita']} teglie)")
                else:
                    riepilogo_cibo[voce]["quantita"] += numero
                    riepilogo_cibo[voce]["unita"] = "pz"
                    print(f"   ✅ Aggregato {voce}: +{numero} pz (totale: {riepilogo_cibo[voce]['quantita']} pz)")
            except (ValueError, TypeError):
                riepilogo_cibo[voce]["quantita"] += 1
                riepilogo_cibo[voce]["unita"] = "pz"
    def formatta_quantita_riepilogo(self, quantita, unita, voce):
        """Formatta la quantità per il riepilogo cucina SENZA duplicazioni"""
        
        if unita == "kg":
            if quantita == int(quantita):
                return f"{int(quantita)} kg"
            else:
                return f"{quantita:.1f} kg"
        elif unita == "pz" or unita == "pezzi":
            return f"{int(quantita)} pz"
        else:
            # Per lasagne, teglie, crespelle, etc.
            return f"{int(quantita)} {unita}"

    def leggi_quantita_da_word(self, path_docx):
        """Legge quantità da Word - VERSIONE MIGLIORATA per tutti i formati"""
        if not os.path.exists(path_docx):
            print(f"❌ File Word non trovato: {path_docx}")
            return {}
        
        try:
            from docx import Document
            doc = Document(path_docx)
            
            prodotti = {}
            
            print(f"\n📄 === LETTURA COMPLETA DA: {os.path.basename(path_docx)} ===")
            
            # Leggi TUTTO il testo del documento
            all_text = ""
            
            # 1. LEGGI DA PARAGRAFI
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text += para.text + "\n"
            
            # 2. LEGGI DA TABELLE  
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            all_text += cell.text + "\n"
            
            # 3. PROCESSA TUTTO IL TESTO RIGA PER RIGA
            righe = [r.strip() for r in all_text.split('\n') if r.strip()]
            
            print(f"📋 Analizzando {len(righe)} righe totali...")
            
            for riga in righe:
                prodotto_trovato = False
                
                # ========== PATTERN 1: PESO IN KG (PRIORITÀ ALTA) ==========
                import re
                
                # Pattern per peso in kg (normale e grassetto)
                patterns_kg = [
                    r'^\*\*(\d+\.?\d*)\s*kg\s+(.+?)\*\*$',  # **3.00 kg Cous cous con verdure**
                    r'^(\d+\.?\d*)\s*kg\s+(.+)$',           # 3.00 kg Cous cous con verdure
                ]
                
                for pattern in patterns_kg:
                    match = re.match(pattern, riga, re.IGNORECASE)
                    if match:
                        peso = float(match.group(1))
                        nome = match.group(2).strip()
                        
                        # Pulisci nome da markdown
                        nome = nome.replace("**", "").replace("*", "").strip()
                        
                        if nome in prodotti:
                            # Accumula peso esistente
                            if "kg" in str(prodotti[nome]):
                                peso_esistente = float(str(prodotti[nome]).split()[0])
                                prodotti[nome] = f"{peso_esistente + peso} kg"
                                print(f"   🔄 {nome}: {peso_esistente} + {peso} = {peso_esistente + peso} kg")
                            else:
                                prodotti[nome] = f"{peso} kg"
                                print(f"   ✅ {nome}: {peso} kg (nuovo)")
                        else:
                            prodotti[nome] = f"{peso} kg"
                            print(f"   ✅ {nome}: {peso} kg (nuovo)")
                        
                        prodotto_trovato = True
                        break
                
                if prodotto_trovato:
                    continue
                
                # ========== PATTERN 2: NUMERI NORMALI ==========
                # ========== PATTERN 2: NUMERI NORMALI E INTOLLERANTI ==========
                patterns_num = [
                    r'^\*\*(\d+)\s+(.+?)\*\*$',              # **3 Lasagne verdi al ragù**
                    r'^\((\d+)\)\d+\s+pz\s+(.+)$',           # (2)1 pz Pizzette (con "pz")
                    r'^\((\d+)\)\d+\s+(.+)$',                 # (2)1 Prodotto (senza "pz")
                    r'^(\d+)\s+pz\s+(.+)$',                   # 15 pz Prodotto
                    r'^(\d+)\s+(.+)$',                        # 15 Prodotto
                ]

                for i, pattern in enumerate(patterns_num):
                    match = re.match(pattern, riga, re.IGNORECASE)
                    if match:
                        print(f"   ✅ MATCH con pattern {i+1}: {pattern}")
                        
                        if i in [1, 2]:  # Pattern per intolleranti
                            quantita = int(match.group(1))  # Numero nelle parentesi
                            nome = match.group(2).strip()
                            quantita_str = f"{quantita} pz"
                            print(f"   📋 INTOLLERANTI: quantità={quantita}, nome='{nome}'")
                        else:  # Pattern normali
                            quantita = int(match.group(1))
                            nome = match.group(2).strip()
                            
                            # Pulisci nome da markdown
                            nome = nome.replace("**", "").replace("*", "").strip()
                            nome = " ".join(nome.split())
                            
                            if nome.lower().startswith("pz "):
                                nome = nome[3:].strip()
                                quantita_str = f"{quantita} pz"
                            else:
                                quantita_str = str(quantita)
                        
                        # ========== FILTRI PER ESCLUDERE RIGHE NON UTILI ==========
                        nome_lower = nome.lower()
                        skip_keywords = [
                            "nr progr", "numero progressivo", "all'attenzione", "presso", 
                            "persone", "allestimento", "pronti", "aperitivo", "coffee", 
                            "tea", "lunch", "buffet", "break", "entro", "ore",
                            "stand", "lunedì", "martedì", "mercoledì", "giovedì", 
                            "venerdì", "sabato", "domenica", "gennaio", "febbraio", 
                            "marzo", "aprile", "maggio", "giugno", "luglio", "agosto",
                            "settembre", "ottobre", "novembre", "dicembre"
                        ]
                        
                        if any(keyword == nome_lower or nome_lower.startswith(keyword + " ") for keyword in skip_keywords):
                            print(f"   ⏭️ SALTATO (keyword): {nome}")
                            continue
                        
                        # Salta se è solo un numero o troppo corto
                        if len(nome.split()) == 0 or nome.isdigit() or len(nome) < 3:
                            print(f"   ⏭️ SALTATO (troppo corto): {nome}")
                            continue
                        
                        # ========== ACCUMULA QUANTITÀ ==========
                        if nome in prodotti:
                            # Prodotto già esistente, somma le quantità
                            try:
                                if "kg" not in str(prodotti[nome]):
                                    quantita_esistente = int(str(prodotti[nome]).split()[0]) if " " in str(prodotti[nome]) else int(prodotti[nome])
                                    nuova_quantita = quantita_esistente + quantita
                                    prodotti[nome] = f"{nuova_quantita} pz" if "pz" in quantita_str else str(nuova_quantita)
                                    print(f"   🔄 {nome}: {quantita_esistente} + {quantita} = {nuova_quantita}")
                                else:
                                    print(f"   ⚠️ Saltato accumulo kg/pezzi per: {nome}")
                            except:
                                prodotti[nome] = quantita_str
                                print(f"   ⚠️ Errore accumulo, impostato: {nome} = {quantita}")
                        else:
                            # Nuovo prodotto
                            prodotti[nome] = quantita_str
                            print(f"   ✅ {nome}: {quantita_str} (nuovo)")
                        
                        prodotto_trovato = True
                        break
                
                if not prodotto_trovato:
                    # Debug per righe non riconosciute
                    if len(riga) > 5 and not any(skip in riga.lower() for skip in ["bevande", "sala", "nr progr"]):
                        print(f"   ❓ NON RICONOSCIUTO: '{riga}'")
            
            # ========== CONVERTI IN FORMATO FINALE ==========
            risultato = {}
            for i, (nome, quantita) in enumerate(prodotti.items()):
                risultato[f"extracted_{i}"] = {
                    'nome': nome,
                    'quantita': quantita
                }
            
            print(f"\n📊 === RIEPILOGO FINALE MIGLIORATO ===")
            print(f"✅ Trovati {len(risultato)} prodotti:")
            for key, info in risultato.items():
                print(f"   📝 {info['nome']}: {info['quantita']}")
            
            return risultato
            
        except Exception as e:
            print(f"❌ Errore lettura Word: {e}")
            import traceback
            traceback.print_exc()
            return {}

    
    
    def trova_chiave_corrispondente_migliorata(self, nome_prodotto):
        """Trova la chiave corrispondente - VERSIONE MIGLIORATA"""
        nome_lower = nome_prodotto.lower()
        
        print(f"   🔍 Cercando chiave per: '{nome_prodotto}' (lower: '{nome_lower}')")
        
        # 1. MATCH SPECIFICI PRIORITARI per prodotti problematici
        match_specifici = {
            'mini pan au chocolat': ['mini', 'pan', 'chocolat'],
            'mini pan chocolat': ['mini', 'pan', 'chocolat'],
            'pan au chocolat': ['pan', 'chocolat'],
            'pan chocolat': ['pan', 'chocolat'],
        }
        
        for pattern_nome, parole_chiave in match_specifici.items():
            if all(parola in nome_lower for parola in parole_chiave):
                print(f"   🎯 Pattern specifico riconosciuto: {pattern_nome}")
                # Cerca nelle voci originali
                for key, voce_originale in self.voce_originale.items():
                    voce_lower = voce_originale.lower()
                    if all(parola in voce_lower for parola in parole_chiave):
                        print(f"   ✅ MATCH SPECIFICO: {key} → '{voce_originale}'")
                        return key
        
        # 2. Match diretto
        for key, voce_originale in self.voce_originale.items():
            voce_lower = voce_originale.lower()
            if nome_lower == voce_lower:
                print(f"   ✅ MATCH DIRETTO: {key}")
                return key
        
        # 3. Match parziale (più permissivo)
        for key, voce_originale in self.voce_originale.items():
            voce_lower = voce_originale.lower()
            
            # Verifica se almeno il 70% delle parole coincide
            parole_nome = set(nome_lower.split())
            parole_voce = set(voce_lower.split())
            
            if parole_nome and parole_voce:
                overlap = len(parole_nome & parole_voce)
                percentuale = overlap / max(len(parole_nome), len(parole_voce))
                
                if percentuale >= 0.7:  # 70% di sovrapposizione
                    print(f"   ✅ MATCH PARZIALE ({percentuale:.1%}): {key} → '{voce_originale}'")
                    return key
        
        # 4. Match per sottostringa (fallback)
        for key, voce_originale in self.voce_originale.items():
            voce_lower = voce_originale.lower()
            
            if nome_lower in voce_lower or voce_lower in nome_lower:
                if len(nome_lower) >= 4:  # Evita match troppo corti
                    print(f"   ✅ MATCH SOTTOSTRINGA: {key} → '{voce_originale}'")
                    return key
        
        print(f"   ❌ Nessuna corrispondenza trovata per: '{nome_prodotto}'")
        
        # DEBUG: Stampa tutte le voci disponibili
        print(f"   📋 Voci disponibili nel menu:")
        for key, voce in self.voce_originale.items():
            if "pan" in voce.lower() or "chocolat" in voce.lower():
                print(f"      {key} → '{voce}'")
        
        return None

    def pulisci_nome_prodotto(self, nome):
        """Pulisce il nome del prodotto da testo extra"""
        # Rimuovi parole comuni che potrebbero essere alla fine
        parole_da_rimuovere = [
            'verdi', 'al ragù', 'antica bologna', 'con', 'di', 'e', 'al', 'alla', 'alle',
            'melanzane', 'mozzarella', 'vegetariani', 'mix', 'farcite', 'in spiedo'
        ]
        
        nome_pulito = nome
        for parola in parole_da_rimuovere:
            # Rimuovi la parola se è alla fine
            if nome_pulito.lower().endswith(' ' + parola.lower()):
                nome_pulito = nome_pulito[:-len(' ' + parola)]
        
        return nome_pulito.strip()

    def trova_chiave_corrispondente(self, nome_prodotto):
        """Trova la chiave corrispondente nelle voci originali"""
        nome_lower = nome_prodotto.lower()
        
        # Cerca corrispondenza diretta o parziale
        for key, voce_originale in self.voce_originale.items():
            voce_lower = voce_originale.lower()
            
            # Corrispondenza diretta
            if nome_lower == voce_lower:
                return key
            
            # Corrispondenza parziale (nome contenuto nella voce o viceversa)
            if nome_lower in voce_lower or voce_lower in nome_lower:
                # Controllo aggiuntivo per evitare false corrispondenze
                if len(nome_lower) >= 4:  # Evita match troppo corti
                    return key
            
            # Corrispondenze specifiche per casi particolari
            parole_chiave_mapping = {
                'pizzette': 'pizzette',
                'calzoncini': 'calzoncini',
                'tramezzini': 'tramezzini',
                'focaccine': 'focaccine',
                'polpettine': 'polpettine',
                'lasagne': 'lasagne',
                'crespelle': 'crespelle',
                'cannelloni': 'cannelloni',
                'cous cous': 'cous',
                'verdi': 'verdi'
            }
            
            for parola_chiave, pattern in parole_chiave_mapping.items():
                if parola_chiave in nome_lower and pattern in voce_lower:
                    return key
        
        return None
    
    def svuota_archivio(self):
        reply = QMessageBox.question(self, "Conferma Eliminazione", 
                                   "Sei sicuro di voler eliminare tutti gli ordini?\n\n"
                                   "Questa operazione non può essere annullata.",
                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                   QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            if os.path.exists(ORDINI_FILE):
                os.remove(ORDINI_FILE)
            QMessageBox.information(self, "Archivio", "Archivio svuotato con successo!")
    
    def salva_ordine(self, ordine):
        """Salva ordine nel file JSON - VERSIONE CORRETTA PER MODALITÀ MODIFICA"""
        import json
        
        ORDINI_FILE = gestore_percorsi.get_percorso("ordini_json")
        
        print(f"\n🔍 === DEBUG SALVA_ORDINE MODALITÀ MODIFICA ===")
        print(f"🔧 modalita_modifica: {getattr(self, 'modalita_modifica', 'NON DEFINITA')}")
        print(f"📋 ordine_da_modificare presente: {hasattr(self, 'ordine_da_modificare') and self.ordine_da_modificare is not None}")
        print(f"🔥 FIREBASE DEBUG - Ordine da salvare:")
        print(f"   - Intolleranze: {ordine.get('intolleranze', 'MANCANTI')}")
        print(f"   - Keys ordine: {list(ordine.keys())}")
        if not os.path.exists(ORDINI_FILE):
            ordini = []
            print("📁 File ordini.json non esiste, creo lista vuota")
        else:
            with open(ORDINI_FILE, "r", encoding="utf-8") as f:
                ordini = json.load(f)
            print(f"📁 Caricati {len(ordini)} ordini dal file")

        # ========== MODALITÀ MODIFICA: ALGORITMO MIGLIORATO ==========
        if hasattr(self, 'modalita_modifica') and self.modalita_modifica and hasattr(self, 'ordine_da_modificare') and self.ordine_da_modificare:
            print("🔄 MODALITÀ MODIFICA CONFERMATA: Cercando ordine da sostituire...")
            
            ordine_originale = self.ordine_da_modificare
            
            # Debug ordine originale
            print(f"🎯 ORDINE ORIGINALE:")
            print(f"   Nome: {ordine_originale.get('campi', {}).get('Nome', 'N/A')}")
            print(f"   Data: {ordine_originale.get('campi', {}).get('Data evento', 'N/A')}")
            print(f"   Cellulare: {ordine_originale.get('campi', {}).get('Cellulare', 'N/A')}")
            print(f"   Path: {ordine_originale.get('path_docx', 'N/A')}")
            
            # Debug nuovo ordine
            print(f"🆕 NUOVO ORDINE:")
            print(f"   Nome: {ordine.get('campi', {}).get('Nome', 'N/A')}")
            print(f"   Data: {ordine.get('campi', {}).get('Data evento', 'N/A')}")
            print(f"   Cellulare: {ordine.get('campi', {}).get('Cellulare', 'N/A')}")
            print(f"   Path: {ordine.get('path_docx', 'N/A')}")
            
            ordine_trovato = False
            indice_trovato = -1
            
            # ========== STRATEGIA MULTI-LIVELLO PER TROVARE L'ORDINE ==========
            
            # STRATEGIA 1: Match per indice nell'array (se disponibile)
            if hasattr(self, 'indice_ordine_originale'):
                try:
                    if 0 <= self.indice_ordine_originale < len(ordini):
                        ordini[self.indice_ordine_originale] = ordine
                        ordine_trovato = True
                        indice_trovato = self.indice_ordine_originale
                        print(f"✅ TROVATO PER INDICE: posizione {self.indice_ordine_originale}")
                except:
                    pass
            
            # STRATEGIA 2: Match per combinazione Nome + Data + Cellulare (più affidabile)
            if not ordine_trovato:
                nome_orig = ordine_originale.get("campi", {}).get("Nome", "").strip()
                data_orig = ordine_originale.get("campi", {}).get("Data evento", "").strip()
                cell_orig = ordine_originale.get("campi", {}).get("Cellulare", "").strip()
                
                print(f"🔍 STRATEGIA 2: Cercando Nome='{nome_orig}' + Data='{data_orig}' + Cell='{cell_orig}'")
                
                for i, ordine_esistente in enumerate(ordini):
                    nome_esist = ordine_esistente.get("campi", {}).get("Nome", "").strip()
                    data_esist = ordine_esistente.get("campi", {}).get("Data evento", "").strip()
                    cell_esist = ordine_esistente.get("campi", {}).get("Cellulare", "").strip()
                    
                    print(f"   Ordine {i}: Nome='{nome_esist}' Data='{data_esist}' Cell='{cell_esist}'")
                    
                    if (nome_esist == nome_orig and 
                        data_esist == data_orig and 
                        cell_esist == cell_orig):
                        
                        print(f"✅ TROVATO PER DATI CLIENTE! Sostituendo ordine in posizione {i}")
                        ordini[i] = ordine
                        ordine_trovato = True
                        indice_trovato = i
                        break
            
            # STRATEGIA 3: Match per path del file (fallback)
            if not ordine_trovato:
                path_originale = ordine_originale.get("path_docx", "")
                print(f"🔍 STRATEGIA 3: Cercando path: '{path_originale}'")
                
                for i, ordine_esistente in enumerate(ordini):
                    path_esistente = ordine_esistente.get("path_docx", "")
                    print(f"   Ordine {i}: Path = '{path_esistente}'")
                    
                    if path_esistente == path_originale and path_originale != "":
                        print(f"✅ TROVATO PER PATH! Sostituendo ordine in posizione {i}")
                        ordini[i] = ordine
                        ordine_trovato = True
                        indice_trovato = i
                        break
            
            # STRATEGIA 4: Match per Nome + Data (senza cellulare, meno rigoroso)
            if not ordine_trovato:
                nome_orig = ordine_originale.get("campi", {}).get("Nome", "").strip()
                data_orig = ordine_originale.get("campi", {}).get("Data evento", "").strip()
                
                print(f"🔍 STRATEGIA 4: Cercando solo Nome='{nome_orig}' + Data='{data_orig}'")
                
                for i, ordine_esistente in enumerate(ordini):
                    nome_esist = ordine_esistente.get("campi", {}).get("Nome", "").strip()
                    data_esist = ordine_esistente.get("campi", {}).get("Data evento", "").strip()
                    
                    if nome_esist == nome_orig and data_esist == data_orig:
                        print(f"✅ TROVATO PER NOME+DATA! Sostituendo ordine in posizione {i}")
                        ordini[i] = ordine
                        ordine_trovato = True
                        indice_trovato = i
                        break
            
            if not ordine_trovato:
                print("❌ ORDINE ORIGINALE NON TROVATO CON NESSUNA STRATEGIA!")
                print("   Aggiungendo come nuovo ordine...")
                ordini.append(ordine)
                indice_trovato = len(ordini) - 1
            else:
                print(f"🎉 MODIFICA COMPLETATA! Ordine aggiornato in posizione {indice_trovato}")
        
        else:
            # MODALITÀ NORMALE: Aggiungi nuovo ordine
            print("➕ MODALITÀ NORMALE: Aggiungendo nuovo ordine")
            ordini.append(ordine)

        # ========== SALVA IL FILE ==========
        try:
            # Crea la cartella se non esiste
            os.makedirs(os.path.dirname(ORDINI_FILE), exist_ok=True)
            
            with open(ORDINI_FILE, "w", encoding="utf-8") as f:
                json.dump(ordini, f, indent=2, ensure_ascii=False)
            
            print(f"💾 Archivio aggiornato - Totale ordini: {len(ordini)}")
            
        except Exception as e:
            print(f"❌ ERRORE SALVATAGGIO: {e}")
            # Messaggio di errore all'utente
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Errore Salvataggio", f"Errore salvando l'archivio:\n{str(e)}")
            return
        
        print(f"=== FINE DEBUG SALVA_ORDINE ===\n")
        
        # ========== AGGIORNA ARCHIVI APERTI ==========
        global archivio_dialogs_aperti
        print(f"📊 Archivi aperti da aggiornare: {len(archivio_dialogs_aperti)}")
        
        for dialog in archivio_dialogs_aperti:
            try:
                dialog.ricarica_dati()
                print(f"🔄 Archivio dialog aggiornato")
            except Exception as e:
                print(f"⚠️ Errore aggiornamento archivio dialog: {e}")
        
        # MEMORIZZA che c'è stata una modifica per archivi futuri
        if hasattr(self, 'modalita_modifica') and self.modalita_modifica:
            self._ordine_appena_modificato = True
            print("💡 Flag modifica impostato per prossimi archivi")




    def ricarica_dati_debug(self):
        """Ricarica i dati dell'archivio dal file JSON - VERSIONE DEBUG MIGLIORATA"""
        print("🔄 === INIZIO RICARICAMENTO ARCHIVIO DEBUG ===")
        
        # Salva l'elemento attualmente selezionato
        current_item = self.tree.currentItem()
        selected_cliente = None
        selected_data = None
        selected_index = None
        
        if current_item and current_item.parent() is not None:
            selected_cliente = current_item.text(1).replace("👤 ", "")
            selected_data = current_item.text(0)
            selected_index = current_item.data(0, Qt.ItemDataRole.UserRole)
            print(f"💾 Elemento selezionato: {selected_cliente} - {selected_data} (idx: {selected_index})")
        
        # Pulisci completamente l'albero
        print("🧹 Pulendo albero esistente...")
        self.tree.clear()
        
        # Ricarica dal file con controllo errori
        print("📂 Ricaricando dal file JSON...")
        try:
            self.load_orders()
            print(f"✅ Ricaricati {len(getattr(self, 'ordini', []))} ordini")
        except Exception as e:
            print(f"❌ Errore ricaricamento: {e}")
            return
        
        # Prova a ristabilire la selezione
        if selected_cliente and selected_data:
            print(f"🎯 Cercando di ristabilire selezione: {selected_cliente}")
            root = self.tree.invisibleRootItem()
            for i in range(root.childCount()):
                date_group = root.child(i)
                for j in range(date_group.childCount()):
                    order_item = date_group.child(j)
                    item_cliente = order_item.text(1).replace("👤 ", "")
                    item_data = order_item.text(0)
                    
                    if item_cliente == selected_cliente and item_data == selected_data:
                        self.tree.setCurrentItem(order_item)
                        self.tree.scrollToItem(order_item)
                        print(f"✅ Selezione ristabilita per: {selected_cliente}")
                        break
        
        print("✅ === FINE RICARICAMENTO ARCHIVIO DEBUG ===")
    
    def carica_ordine(self):
        dialog = OrderArchiveDialog(self)
        dialog.order_selected.connect(self.load_order_data)
        dialog.exec()
    
    def load_order_data(self, ordine):
        """Carica i dati dell'ordine nei campi - VERSIONE CORRETTA CON INTOLLERANZE"""
        print(f"📥 Caricando ordine: {ordine['campi'].get('Nome', 'Cliente')}")
        
        # Carica i dati dell'ordine nei campi
        for k, v in ordine["campi"].items():
            if k in self.entries:
                if k == "Data evento":
                    try:
                        date_obj = datetime.strptime(v, "%d/%m/%Y")
                        qdate = QDate(date_obj.year, date_obj.month, date_obj.day)
                        self.entries[k].setDate(qdate)
                    except:
                        pass
                else:
                    self.entries[k].setText(v)
        
        # Carica le scelte
        for k, checkbox in self.selected.items():
            if hasattr(checkbox, 'setChecked'):
                checkbox.setChecked(ordine["scelte"].get(k, False))
        
        # ========== NUOVO: RIPRISTINA LE INTOLLERANZE ==========
        intolleranze_salvate = ordine.get("intolleranze", {})
        print(f"🌱 Ripristinando intolleranze: {list(intolleranze_salvate.keys())}")
        
        # Prima azzera tutti i widget menu extra
        if hasattr(self, 'menu_extra_widgets'):
            for menu_type, widget in self.menu_extra_widgets.items():
                widget.setValue(0)  # Azzera tutti
                print(f"   🔄 Azzerato {menu_type}")
        
        # Poi ripristina i valori salvati
        for menu_type, dati_intolleranza in intolleranze_salvate.items():
            persone_salvate = dati_intolleranza.get('persone', 0)
            
            # Trova il widget corrispondente
            if hasattr(self, 'menu_extra_widgets') and menu_type in self.menu_extra_widgets:
                widget = self.menu_extra_widgets[menu_type]
                widget.setValue(persone_salvate)
                print(f"   ✅ Ripristinato {menu_type}: {persone_salvate} persone")
            else:
                print(f"   ⚠️ Widget non trovato per {menu_type}")
        
        # Carica stato fiera
        self.evento_in_fiera = ordine.get("evento_in_fiera", False)
        if hasattr(self, 'fiera_checkbox'):
            self.fiera_checkbox.setChecked(self.evento_in_fiera)
        
        # Imposta tipo aperitivo
        self.ap_buffet_tipo = ordine.get("tipo_ap", "Leggero")
        
        # Aggiorna i radio button se necessario
        if hasattr(self, 'ap_group'):
            for button in self.ap_group.buttons():
                if button.text() == self.ap_buffet_tipo:
                    button.setChecked(True)
        
        # SE SIAMO IN MODALITÀ MODIFICA, NON MOSTRARE IL MESSAGGIO
        if not (hasattr(self, 'modalita_modifica') and self.modalita_modifica):
            QMessageBox.information(self, "Ordine Caricato", 
                                f"Ordine di {ordine['campi'].get('Nome', 'Cliente')} caricato con successo!")
        else:
            print(f"✅ Ordine caricato per modifica: {ordine['campi'].get('Nome', 'Cliente')}")
    # --- Helper: riconoscimento torta (unica verità) ---
    def _is_torta(self, label: str) -> bool:
        return "kg torta" in (label or "").lower()
    def calcola_bilanciamento_dinamico(self, servizio_lower, tipo_ap, num_pers):
        """
        Calcola le riduzioni di primi e finger food basate su secondi e contorni selezionati
        Logica: se aggiungi 300g di secondi/contorni, togli 150g dai primi e 150g dai finger food
        """
        # Controlla se è un servizio che richiede bilanciamento
        is_lunch = "lunch" in servizio_lower or "pranzo" in servizio_lower
        is_aperitivo_rinforzato = "aperitivo" in servizio_lower and tipo_ap == "rinforzato"
        
        if not (is_lunch or is_aperitivo_rinforzato):
            return {"riduzione_primi": 0, "riduzione_finger": 0, "peso_secondi_aggiunto": 0, "peso_contorni_aggiunto": 0}
        
        print(f"\n🔄 CALCOLO BILANCIAMENTO DINAMICO:")
        print(f"   - Servizio: {servizio_lower}")
        print(f"   - Tipo AP: {tipo_ap}")
        print(f"   - Is lunch: {is_lunch}")
        print(f"   - Is aperitivo rinforzato: {is_aperitivo_rinforzato}")
        
        # Calcola peso aggiunto da secondi
        peso_secondi_aggiunto = 0
        if hasattr(self, 'referenze_secondi') and self.referenze_secondi:
            num_secondi = len(self.referenze_secondi)
            peso_secondi_per_persona = 0.100  # 100g per persona (quantità ridotta)
            peso_secondi_totale = peso_secondi_per_persona * num_pers
            peso_secondi_aggiunto = peso_secondi_totale
            print(f"   - Secondi selezionati: {num_secondi}")
            print(f"   - Peso secondi aggiunto: {peso_secondi_aggiunto:.3f} kg ({peso_secondi_aggiunto*1000:.0f}g)")
        
        # Calcola peso aggiunto da contorni  
        peso_contorni_aggiunto = 0
        if hasattr(self, 'referenze_contorni') and self.referenze_contorni:
            num_contorni = len(self.referenze_contorni)
            peso_contorni_per_persona = 0.050  # 50g per persona (quantità ridotta)
            peso_contorni_totale = peso_contorni_per_persona * num_pers
            peso_contorni_aggiunto = peso_contorni_totale
            print(f"   - Contorni selezionati: {num_contorni}")
            print(f"   - Peso contorni aggiunto: {peso_contorni_aggiunto:.3f} kg ({peso_contorni_aggiunto*1000:.0f}g)")
        
        # Calcola peso totale aggiunto
        peso_totale_aggiunto = peso_secondi_aggiunto + peso_contorni_aggiunto
        print(f"   - Peso totale aggiunto: {peso_totale_aggiunto:.3f} kg ({peso_totale_aggiunto*1000:.0f}g)")
        
        # Se non ci sono secondi/contorni, nessuna riduzione
        if peso_totale_aggiunto == 0:
            print(f"   - Nessun secondo/contorno selezionato: nessuna riduzione")
            return {"riduzione_primi": 0, "riduzione_finger": 0, "peso_secondi_aggiunto": 0, "peso_contorni_aggiunto": 0}
        
        # LOGICA DI BILANCIAMENTO: dividi la riduzione 50%-50% tra primi e finger food
        riduzione_totale = peso_totale_aggiunto
        riduzione_primi = riduzione_totale * 0.3
        riduzione_finger = riduzione_totale * 0.30
        
        print(f"   - Riduzione primi: {riduzione_primi:.3f} kg ({riduzione_primi*1000:.0f}g)")
        print(f"   - Riduzione finger food: {riduzione_finger:.3f} kg ({riduzione_finger*1000:.0f}g)")
        
        return {
            "riduzione_primi": riduzione_primi,
            "riduzione_finger": riduzione_finger, 
            "peso_secondi_aggiunto": peso_secondi_aggiunto,
            "peso_contorni_aggiunto": peso_contorni_aggiunto
        }
    def calcola_quantita(self, key, referenze_cibo, num_pers, servizio_lower, tipo_ap):
        """Calcola la quantità per una voce del menu con bilanciamento dinamico"""
        # DEBUG TEMPORANEO - per verificare categorizzazione contorni
        voce = self.voce_originale.get(key, key).strip()
        voce_lower = voce.lower()
        if "verdure" in voce.lower():
            print(f"🔍 DEBUG VERDURE ALLA BRACE:")
            print(f"   - key: {key}")
            print(f"   - voce: {voce}")
            print(f"   - in referenze_contorni: {key in self.referenze_contorni}")
            print(f"   - self.referenze_contorni contenuto: {self.referenze_contorni}")
            
            
        print(f"\n🔧 CALCOLA QUANTITÀ per: {voce}")
        print(f"   - Servizio: {servizio_lower}")
        print(f"   - Tipo AP: {tipo_ap}")
        print(f"   - Persone: {num_pers}")
        if "arrosto" in voce.lower():
            print(f"🔍 DEBUG ARROSTO:")
            print(f"   - key: {key}")
            print(f"   - voce: {voce}")
            print(f"   - In referenze_secondi: {key in self.referenze_secondi}")
            print(f"   - In referenze_salato: {key in self.referenze_salato}")
            print(f"   - In referenze_contorni: {key in self.referenze_contorni}")
        
        # GESTIONE SPECIALE PER TORTE (mantieni logica originale)
        if "kg torta" in voce_lower:
            print(f"🎂 TORTA: {num_pers} persone → {num_pers/10:.1f} kg")
            kg_torta = num_pers / 10

            if kg_torta <= int(kg_torta) + 0.25:
                kg_finali = int(kg_torta)
            elif kg_torta <= int(kg_torta) + 0.75:
                kg_finali = int(kg_torta) + 0.5
            else:
                kg_finali = int(kg_torta) + 1

            kg_finali = max(0.5, kg_finali)

            return f"{int(kg_finali)} kg" if kg_finali == int(kg_finali) else f"{kg_finali:.1f} kg"
        # GESTIONE SPECIALE PER TAGLIERI (indipendente dal bilanciamento)
        if "taglieri" in voce.lower():
            print(f"🧀 DEBUG TAGLIERI per {voce}:")
            print(f"   - Persone: {num_pers}")
            
            if "taglieri di salumi medi" in voce.lower():
                # 1 tagliere ogni 20 persone
                num_taglieri = max(1, int(num_pers / 20))
                print(f"   🥓 Taglieri salumi: {num_pers} persone ÷ 20 = {num_taglieri} taglieri")
                return f"{num_taglieri} pz"
            
            elif "taglieri formaggi piccoli" in voce.lower():
                # 1 tagliere ogni 20 persone  
                num_taglieri = max(1, int(num_pers / 20))
                print(f"   🧀 Taglieri formaggi: {num_pers} persone ÷ 20 = {num_taglieri} taglieri")
                return f"{num_taglieri} pz"
                # GESTIONE SPECIALE PER ACCOMPAGNAMENTI TAGLIERI (prima di qualsiasi altro calcolo)
            # GESTIONE SPECIALE PER ACCOMPAGNAMENTI TAGLIERI (blocco separato!)
            # GESTIONE SPECIALE PER ACCOMPAGNAMENTI TAGLIERI
        if any(item in voce.lower() for item in ["miele", "composta di frutta", "composte di frutta"]):
            taglieri_formaggi_selezionati = False
            for k, checkbox in self.selected.items():
                if hasattr(checkbox, 'isChecked') and checkbox.isChecked():
                    voce_k = self.voce_originale.get(k, "").lower()
                    if "taglieri formaggi piccoli" in voce_k:
                        taglieri_formaggi_selezionati = True
                        break
            
            if taglieri_formaggi_selezionati:
                # Miele e composta: stesso numero dei taglieri (3 per 60 persone)
                num_accompagnamenti = max(1, int(num_pers / 20))
                print(f"   🍯 {voce} per taglieri: {num_accompagnamenti} pz")
                return f"{num_accompagnamenti} pz"

        # GESTIONE SPECIALE PER GRISSINI CON TAGLIERI (logica separata)
        elif "grissini" in voce.lower():
            taglieri_formaggi_selezionati = False
            for k, checkbox in self.selected.items():
                if hasattr(checkbox, 'isChecked') and checkbox.isChecked():
                    voce_k = self.voce_originale.get(k, "").lower()
                    if "taglieri formaggi piccoli" in voce_k:
                        taglieri_formaggi_selezionati = True
                        break
            
            if taglieri_formaggi_selezionati:
                # Grissini: 1 per persona quando ci sono taglieri
                print(f"   🥖 Grissini per taglieri: {num_pers} pz (1 per persona)")
                return f"{num_pers} pz"
        # Calcola peso unitario dalla referenza
        # Ottieni dati dalla referenza (peso unitario e difficoltà)
        dati = self.trova_dati_referenza(voce, self.dati_referenze, {"difficolta": 1.0, "peso_unitario": 0.03333})
        coeff = dati["difficolta"]
        peso_unitario = dati["peso_unitario"]

        print(f"   - Peso unitario: {peso_unitario}")
        print(f"   - Coefficiente difficoltà: {coeff}")

        if not peso_unitario or peso_unitario <= 0:
            print("   ⚠️ peso_unitario nullo: ritorno 1 pezzo di sicurezza")
            return "1 pezzo"
        
        # CALCOLA BILANCIAMENTO DINAMICO per lunch e aperitivo rinforzato
        bilanciamento = self.calcola_bilanciamento_dinamico(servizio_lower, tipo_ap, num_pers)
        
        if "coffee" in servizio_lower or "tea" in servizio_lower:
            return self._calcola_finger_food_con_bilanciamento(key, referenze_cibo, num_pers, servizio_lower, tipo_ap, peso_unitario, bilanciamento)
        # Controlla se è lunch o aperitivo rinforzato
        is_lunch_service = "lunch" in servizio_lower or "pranzo" in servizio_lower
        is_aperitivo_rinforzato = "aperitivo" in servizio_lower and tipo_ap == "rinforzato"
        use_dynamic_balancing = is_lunch_service or is_aperitivo_rinforzato
        
        print(f"   - Use dynamic balancing: {use_dynamic_balancing}")
        if use_dynamic_balancing:
            print(f"   - Riduzione primi: {bilanciamento['riduzione_primi']*1000:.0f}g")
            print(f"   - Riduzione finger: {bilanciamento['riduzione_finger']*1000:.0f}g")
        
        # GESTIONE PRIMI CON BILANCIAMENTO
        if key in self.referenze_primi:
            return self._calcola_primi_con_bilanciamento(key, num_pers, servizio_lower, tipo_ap, bilanciamento, peso_unitario)
        
        if key in self.referenze_secondi:
            if use_dynamic_balancing:
                peso_secondi_per_persona = 0.100  # 100g invece di 150g per lunch/aperitivo rinforzato
                print(f"   🥩 Secondo ridotto (lunch/aperitivo rinforzato): 100g per persona")
            else:
                peso_secondi_per_persona = 0.150  # 150g per altri servizi
                print(f"   🥩 Secondo standard: 150g per persona")
                
            peso_secondi_totale = peso_secondi_per_persona * num_pers
            
            if len(self.referenze_secondi) > 0:
                peso_per_secondo = peso_secondi_totale / len(self.referenze_secondi)
            else:
                peso_per_secondo = peso_secondi_totale
            
            print(f"   🥩 Secondo: {peso_per_secondo:.3f} kg")
            
            # DETERMINA SE IL SECONDO VA IN KG O PEZZI
            # DETERMINA SE IL SECONDO VA IN KG O PEZZI USANDO ARRAY CENTRALIZZATO
            if self._va_in_kg(voce):
                print(f"   - Secondo va in kg: True")
                return f"{peso_per_secondo:.1f} kg"
            else:
                print(f"   - Secondo va in kg: False")
                pezzi = peso_per_secondo / peso_unitario
                pezzi_finali = max(1, int(round(pezzi)))
                print(f"   - Pezzi calcolati: {pezzi:.1f} → {pezzi_finali}")
                return f"{pezzi_finali} pz"
        if key in self.referenze_contorni:
            if use_dynamic_balancing:
                peso_contorni_per_persona = 0.070  # 50g invece di 80g per lunch/aperitivo rinforzato
                print(f"   🥗 Contorno ridotto (lunch/aperitivo rinforzato): 50g per persona")
            else:
                peso_contorni_per_persona = 0.080  # 80g per altri servizi
                print(f"   🥗 Contorno standard: 80g per persona")
                
            peso_contorni_totale = peso_contorni_per_persona * num_pers
            
            if len(self.referenze_contorni) > 0:
                peso_per_contorno = peso_contorni_totale / len(self.referenze_contorni)
            else:
                peso_per_contorno = peso_contorni_totale
            
            print(f"   🥗 Contorno: {peso_per_contorno:.3f} kg")
            
            # USA LA FUNZIONE HELPER INVECE DI peso_unitario <= 0.05
            if self._va_in_kg(voce):
                print(f"   - Contorno va in kg: True")
                return f"{peso_per_contorno:.1f} kg"
            else:
                print(f"   - Contorno va in kg: False")
                pezzi = peso_per_contorno / peso_unitario
                pezzi_finali = max(1, int(round(pezzi)))
                print(f"   - Pezzi calcolati: {pezzi:.1f} → {pezzi_finali}")
                return f"{pezzi_finali} pz"
        # GESTIONE ACCOMPAGNAMENTI
        if key in self.referenze_accompagnamenti:
            return self._calcola_accompagnamenti(key, num_pers, servizio_lower)

        # GESTIONE DOLCI
        if key in self.referenze_dolci:
            # Aperitivo rinforzato → usa la funzione speciale
            if "aperitivo" in servizio_lower and tipo_ap == "rinforzato":
                return self.calcola_quantita_menu_extra_dolci_aperitivo_rinforzato(
                    voce, num_pers, servizio_lower, tipo_ap
                )
            # Altri servizi → logica standard
            return self._calcola_dolci(key, num_pers, servizio_lower, tipo_ap, peso_unitario)

        

        # GESTIONE CONTORNI (quantità ridotte per lunch e aperitivo rinforzato)
        
        # GESTIONE SPECIALE PER APERITIVO
        if "aperitivo" in servizio_lower and key not in self.referenze_primi:
            
            # PRIMA: Controlla se è un accompagnamento (grissini, pane, bocconcini)
            accompagnamenti_keywords = ["grissini", "pane a fette", "bocconcini alle olive"]
            is_accompagnamento = any(keyword in voce_lower for keyword in accompagnamenti_keywords)
            
            if is_accompagnamento:
                print(f"🍞 ACCOMPAGNAMENTO APERITIVO: chiamando _calcola_accompagnamenti")
                return self._calcola_accompagnamenti(key, num_pers, servizio_lower)
            
            # SECONDO: Verifica se è un dolce
            is_dolce = any(parola in voce_lower for parola in ["dolc", "torta", "crostat", "tiramisu", "mousse", "spicchietti", "frutta", "biscotteria", "pasticceria"])
            
            if is_dolce and tipo_ap == "rinforzato":
                print(f"🍰 DOLCE APERITIVO RINFORZATO: chiamando funzione specifica")
                return self.calcola_quantita_menu_extra_dolci_aperitivo_rinforzato(voce, num_pers, servizio_lower, tipo_ap)
            else:
                print(f"🥂 FINGER FOOD APERITIVO: chiamando funzione originale")
                return self._calcola_aperitivo_originale(key, referenze_cibo, num_pers, servizio_lower, tipo_ap, bilanciamento)
        
        

        # GESTIONE FINGER FOOD CON BILANCIAMENTO
        return self._calcola_finger_food_con_bilanciamento(key, referenze_cibo, num_pers, servizio_lower, tipo_ap, peso_unitario, bilanciamento)
    # ==============================================================================
# PROBLEMA: gestisci_selezione_bevanda seleziona accessori di TUTTE le sezioni
# SOLUZIONE: Limitare l'auto-selezione solo alla sezione corrente
# ==============================================================================
    def trova_dati_referenza(self,voce, dizionario, default):
        import re
        voce_norm = re.sub(r"\s*\(.*?\)\*?", "", voce.strip().lower())
        voce_norm = voce_norm.rstrip('*')
        
        # DEBUG GENERICO
        if any(x in voce for x in ["mozzarelline", "ricotta", "millefoglie"]):
            print(f"🔍 DEBUG MATCHING per '{voce[:30]}...':")
            print(f"   - voce_norm (senza parentesi): '{voce_norm}'")
            print(f"   - Sto cercando in {len(dizionario)} chiavi")
        
        # PRIMO: Cerca match esatto CON parentesi
        voce_originale_lower = voce.strip().lower()
        for chiave in dizionario:
            chiave_lower = chiave.strip().lower()
            if chiave_lower == voce_originale_lower:
                if any(x in voce for x in ["mozzarelline", "ricotta", "millefoglie"]):
                    print(f"   ✅ MATCH ESATTO CON PARENTESI: trovato!")
                return dizionario[chiave]
        
        # SECONDO: Cerca match SENZA parentesi
        for chiave in dizionario:
            chiave_norm = re.sub(r"\s*\(.*?\)\*?", "", chiave.strip().lower())
            chiave_norm = chiave_norm.rstrip('*')
            if chiave_norm == voce_norm:
                if any(x in voce for x in ["mozzarelline", "ricotta", "millefoglie"]):
                    print(f"   ✅ MATCH SENZA PARENTESI: '{chiave_norm}' == '{voce_norm}'")
                return dizionario[chiave]
        
        # TERZO: Match parziale per crostini
        if "crostini" in voce_norm and "radicchio" in voce_norm:
            for chiave in dizionario:
                chiave_norm = chiave.strip().lower()
                if "crostini" in chiave_norm and "radicchio" in chiave_norm:
                    return dizionario[chiave]
        
        # QUARTO: FALLBACK PER SPICCHIETTI
        if "spicchietti" in voce_norm and "crostat" in voce_norm:
            print("   🔧 USANDO FALLBACK per spicchietti crostate")
            return {"difficolta": 1.0, "peso_unitario": 0.025}
        
        if any(x in voce for x in ["mozzarelline", "ricotta", "millefoglie"]):
            print(f"   ❌ NESSUN MATCH TROVATO - uso default")
        
        return default
    def _calcola_primi_con_bilanciamento(self, key, num_pers, servizio_lower, tipo_ap, bilanciamento, peso_unitario):
        """Calcola quantità primi con bilanciamento dinamico"""
        voce = self.voce_originale.get(key, key).strip()
        
        # Determina se è servizio servito
        is_servito_service = False
        for k, checkbox in self.selected.items():
            if hasattr(checkbox, 'isChecked') and checkbox.isChecked():
                if k.startswith("servito|"):
                    is_servito_service = True
                    break
        
        # Determina peso base primi per servizio
        if is_servito_service:
            peso_primi_base_per_persona = 0.250  # 220g per servito
            print(f"   🍝 SERVITO - peso base primi: 220g per persona")
        elif "lunch" in servizio_lower or "pranzo" in servizio_lower or ("aperitivo" in servizio_lower and tipo_ap == "rinforzato"):
            peso_primi_base_per_persona = 0.160  # 160g per lunch e aperitivo rinforzato
            print(f"   🍝 LUNCH/APERITIVO RINFORZATO - peso base primi: 160g per persona")
        else:
            # Altri servizi (non dovrebbero arrivare qui con bilanciamento)
            peso_primi_base_per_persona = 0.160
        
        peso_primi_totale_base = peso_primi_base_per_persona * num_pers
        
        # Applica riduzione da bilanciamento (solo per lunch/aperitivo rinforzato, non servito)
        if not is_servito_service:
            riduzione_primi = bilanciamento["riduzione_primi"]
            peso_primi_totale_finale = max(0.5, peso_primi_totale_base - riduzione_primi)
            print(f"   🍝 PRIMI CON BILANCIAMENTO:")
            print(f"   - Peso base: {peso_primi_totale_base:.3f} kg")
            print(f"   - Riduzione: {riduzione_primi:.3f} kg")
            print(f"   - Peso finale: {peso_primi_totale_finale:.3f} kg")
        else:
            # Servito non ha bilanciamento
            peso_primi_totale_finale = peso_primi_totale_base
            print(f"   🍝 PRIMI SERVITO (senza bilanciamento): {peso_primi_totale_finale:.3f} kg")
        
        # IDENTIFICA TIPOLOGIA PRIMI
        primi_freddi_keywords = ["cous cous", "orecchiette", "mezze penne", "mezze maniche", 
                                "insalata di riso", "insalatina di farro","gnocchetti sardi con carpaccio",
                                "penne", "quinoa", "pennette", "fioriere"]
        
        primi_elaborati_keywords = ["lasagne", "cannelloni", "rosette", "nidi", "crespelle"]
        
        # Classifica TUTTI i primi selezionati
        primi_freddi = []
        primi_elaborati = []
        primi_normali = []
        
        for k in self.referenze_primi:
            voce_primo = self.voce_originale.get(k, "").lower()
            
            if any(pf in voce_primo for pf in primi_freddi_keywords):
                primi_freddi.append(k)
            elif any(pe in voce_primo for pe in primi_elaborati_keywords):
                primi_elaborati.append(k) 
            else:
                primi_normali.append(k)
        
        print(f"🍝 CLASSIFICAZIONE PRIMI:")
        print(f"   - Freddi: {len(primi_freddi)}")
        print(f"   - Elaborati: {len(primi_elaborati)}")
        print(f"   - Normali: {len(primi_normali)}")
        
        # CALCOLA PESO CONSUMATO DA PRIMI ELABORATI
        peso_consumato_elaborati = 0
        import math

        for k in primi_elaborati:
            voce_elaborato = self.voce_originale.get(k, "").lower()
            if "lasagne" in voce_elaborato:
                if is_servito_service:
                    teglie = max(1, math.ceil(num_pers / 15.0))  # Servito: più teglie
                else:
                    teglie = max(1, math.ceil(num_pers / 20.0)) if len(self.referenze_primi) > 1 else math.ceil(num_pers / 10.0)
                peso_consumato_elaborati += teglie * 2.0  # 2 kg per teglia
            elif "crespelle" in voce_elaborato:
                if is_servito_service:
                    crespelle = max(12, math.ceil((num_pers * 0.9) / 8.0) * 8)  # Servito: più crespelle
                else:
                    crespelle = max(8, math.ceil((num_pers * 0.64) / 8.0) * 8) if len(self.referenze_primi) > 1 else math.ceil((num_pers * 2) / 8.0) * 8
                peso_consumato_elaborati += crespelle * 0.080  # 80g per crespella
            elif "cannelloni" in voce_elaborato:
                if is_servito_service:
                    cannelloni = max(12, math.ceil((num_pers * 0.9) / 12.0) * 12)  # Servito: più cannelloni
                else:
                    cannelloni = max(12, math.ceil((num_pers * 0.64) / 12.0) * 12) if len(self.referenze_primi) > 1 else math.ceil((num_pers * 0.64) / 6.0) * 12
                peso_consumato_elaborati += cannelloni * 0.040  # 40g per cannellone
            elif "nidi" in voce_elaborato or "rosette" in voce_elaborato:
                if is_servito_service:
                    pezzi = max(18, math.ceil((num_pers * 0.9) / 18.0) * 18)  # Servito: più pezzi
                else:
                    pezzi = max(18, math.ceil((num_pers * 0.64) / 18.0) * 18) if len(self.referenze_primi) > 1 else math.ceil((num_pers * 0.64) / 9.0) * 18
                peso_consumato_elaborati += pezzi * 0.045  # 45g per pezzo

        # CALCOLA PESO RIMANENTE PER FREDDI E NORMALI
        peso_rimanente = max(0.5, peso_primi_totale_finale - peso_consumato_elaborati)
        primi_non_elaborati = len(primi_freddi) + len(primi_normali)
        peso_per_primo_non_elaborato = peso_rimanente / max(1, primi_non_elaborati)

        print(f"   📊 CALCOLO BILANCIATO:")
        print(f"   - Peso totale primi: {peso_primi_totale_finale:.1f} kg")
        print(f"   - Peso consumato elaborati: {peso_consumato_elaborati:.1f} kg") 
        print(f"   - Peso rimanente: {peso_rimanente:.1f} kg")
        print(f"   - Primi non elaborati: {primi_non_elaborati}")
        print(f"   - Peso per primo non elaborato: {peso_per_primo_non_elaborato:.1f} kg")

        # GESTISCI LA TIPOLOGIA DEL PRIMO CORRENTE
        voce_lower = voce.lower()
        if any(pe in voce_lower for pe in primi_elaborati_keywords):
            # PRIMI ELABORATI: usa quantità già calcolate sopra
            import math
            
            if "lasagne" in voce_lower:
                if is_servito_service:
                    teglie = max(1, math.ceil(num_pers / 15.0))
                    print(f"   🍝 Lasagne SERVITO: {teglie} teglie")
                else:
                    teglie = max(1, math.ceil(num_pers / 20.0)) if len(self.referenze_primi) > 1 else math.ceil(num_pers / 10.0)
                    print(f"   🍝 Lasagne BILANCIATO: {teglie} teglie")
                return f"{int(teglie)} lasagne"
            
            elif "crespelle" in voce_lower:
                if is_servito_service:
                    crespelle = max(12, math.ceil((num_pers * 0.9) / 8.0) * 8)
                    print(f"   🥟 Crespelle SERVITO: {crespelle}")
                else:
                    crespelle = max(8, math.ceil((num_pers * 0.64) / 8.0) * 8) if len(self.referenze_primi) > 1 else math.ceil((num_pers * 2) / 8.0) * 8
                    print(f"   🥟 Crespelle BILANCIATO: {crespelle}")
                return f"{int(crespelle)} crespelle"
            
            elif "cannelloni" in voce_lower:
                if is_servito_service:
                    cannelloni = max(12, math.ceil((num_pers * 0.9) / 12.0) * 12)
                    print(f"   🥘 Cannelloni SERVITO: {cannelloni}")
                else:
                    cannelloni = max(12, math.ceil((num_pers * 0.64) / 12.0) * 12) if len(self.referenze_primi) > 1 else math.ceil((num_pers * 0.64) / 6.0) * 12
                    print(f"   🥘 Cannelloni BILANCIATO: {cannelloni}")
                return f"{int(cannelloni)} cannelloni"
            
            elif "nidi" in voce_lower or "rosette" in voce_lower:
                tipo = 'nidi' if 'nidi' in voce_lower else 'rosette'
                if is_servito_service:
                    pezzi = max(18, math.ceil((num_pers * 0.9) / 18.0) * 18)
                    print(f"   🍜 {tipo.title()} SERVITO: {pezzi}")
                else:
                    pezzi = max(18, math.ceil((num_pers * 0.64) / 18.0) * 18) if len(self.referenze_primi) > 1 else math.ceil((num_pers * 0.64) / 9.0) * 18
                    print(f"   🍜 {tipo.title()} BILANCIATO: {pezzi}")
                return f"{int(pezzi)} {tipo}"
        
        else:
            # PRIMI FREDDI E NORMALI: usano il peso rimanente
            peso_finale = peso_per_primo_non_elaborato
            
            # Arrotonda a 0.5 kg
            intero = int(peso_finale)
            decimale = peso_finale - intero
            if decimale < 0.25:
                peso_finale = float(intero)
            elif decimale < 0.75:
                peso_finale = intero + 0.5
            else:
                peso_finale = float(intero + 1)
            
            peso_finale = max(0.5, peso_finale)  # Minimo 500g
            
            # CONTROLLO PRIMI CRUDI (inox) - DEVONO ESSERE DIMEZZATI
            primi_crudi = [
                "maccheroncini pasticciati al ragù bolognese(inox)*",
                "paccheri con fiammiferi di speck croccante e zucchine marinate(inox)*", 
                "gnocchetti sardi con pesto di noci e radicchio di treviso(inox)*",
                "trofie al pesto fagiolini e noci(inox)*",
                "gramigna paglia e fieno alla boscaiola(inox)*",
                "paccheri con guanciale di ariccia croccante e radicchio trevigiano(inox)*",
                "pennette alla norma con melanzane al forno e ricotta salata (inox)*",
                "maccheroncini panna,prosciutto e piselli(inox)*",
                "garganelli ai porcini e ricotta affumicata(inox)*",
                "tortiglioni crudo di parma e pachino con carciofi croccanti(inox)*",
                "garganelli al pettine con crema di zucchine e fiammiferi di speck arrostito(inox)*"
            ]
            
            # Controlla se questo primo deve essere dimezzato
            voce_lower = voce.lower()
            is_primo_crudo = any(primo_crudo.lower() in voce_lower for primo_crudo in primi_crudi)
            
            if is_primo_crudo:
                # DIMEZZA il peso per i primi crudi
                peso_finale = peso_finale / 2
                print(f"   🍝 Primo crudo (inox) DIMEZZATO: {peso_finale:.1f} kg (da cuocere)")
                
                # Riapplica l'arrotondamento dopo il dimezzamento
                if peso_finale < 0.25:
                    peso_finale = 0.25  # Minimo 250g
                elif peso_finale < 0.75:
                    peso_finale = 0.5
                elif peso_finale < 1.25:
                    peso_finale = 1.0
                else:
                    # Arrotonda a 0.5 kg per valori superiori
                    peso_finale = round(peso_finale * 2) / 2
            
            tipo_primo = "freddo" if key in primi_freddi else "normale"
            print(f"   🍝 Primo {tipo_primo} BILANCIATO: {peso_finale:.1f} kg")
            
            return f"{peso_finale:.1f} kg"
    def _calcola_aperitivo_originale(self, key, referenze_cibo, num_pers, servizio_lower, tipo_ap, bilanciamento=None):
        """Logica per aperitivo - SOLO FINGER FOOD SALATI con bilanciamento dinamico"""
        voce = self.voce_originale.get(key, key).strip()
        # Ottieni dati dalla referenza
        dati = self.trova_dati_referenza(voce, self.dati_referenze, {"difficolta": 1.0, "peso_unitario": 1.0})
        coeff = dati["difficolta"]
        peso_unitario = dati["peso_unitario"]
        
        # PESO BASE per finger food salati (esclusi i dolci)
        if tipo_ap == "rinforzato":
            peso_salati_base_persona = 0.260  # 280g salati base
        else:  # leggero
            peso_salati_base_persona = 0.165  # 230g totali - 40g dolci = 190g salati
        
        peso_totale_base = num_pers * peso_salati_base_persona
        
        # APPLICA BILANCIAMENTO DINAMICO se fornito
        if bilanciamento and "riduzione_finger" in bilanciamento:
            riduzione_finger = bilanciamento["riduzione_finger"]
            peso_totale_finale = max(0.5, peso_totale_base - riduzione_finger)
            
            print(f"🥂 DEBUG APERITIVO SALATI CON BILANCIAMENTO per {voce}:")
            print(f"   - Tipo: {tipo_ap}")
            print(f"   - Peso base per persona: {peso_salati_base_persona*1000}g")
            print(f"   - Peso totale base: {peso_totale_base:.3f} kg")
            print(f"   - Riduzione bilanciamento: {riduzione_finger:.3f} kg")
            print(f"   - Peso finale: {peso_totale_finale:.3f} kg")
            print(f"   - Peso finale per persona: {(peso_totale_finale/num_pers)*1000:.1f}g")
        else:
            peso_totale_finale = peso_totale_base
            
            print(f"🥂 DEBUG APERITIVO SALATI per {voce}:")
            print(f"   - Tipo: {tipo_ap}")
            print(f"   - Peso SALATI per persona: {peso_salati_base_persona*1000}g")
            print(f"   - Peso totale salati: {peso_totale_finale:.3f} kg")
        
        # CALCOLA COEFFICIENTI SOLO PER FINGER FOOD SALATI (esclude dolci e contorni)
        pesi_difficalta_salati = []
        referenze_salati = []
        
        for k in referenze_cibo:
            voce_k = self.voce_originale.get(k, "").strip().lower()
            
            # ESCLUDI i dolci dal calcolo
            is_dolce = any(parola in voce_k for parola in ["dolc", "torta", "crostat", "tiramisu", "mousse", "spicchietti", "frutta", "biscotteria", "pasticceria"])
            
            # ESCLUDI anche i contorni (che hanno logica separata)
            is_contorno = any(parola in voce_k for parola in ["verdure", "patate", "spinaci", "macedonia", "ricotta"])
            
            # ESCLUDI i primi piatti
            is_primo = any(parola in voce_k for parola in ["pasta", "orecchiette", "pennette", "crespelle", "rosette", "lasagne"])
            
            if is_dolce or is_contorno or is_primo:
                continue  # Salta dolci, contorni e primi
            
            referenze_salati.append(k)
            coeff_k = self.dati_referenze.get(voce_k, {"difficolta": 1.0})["difficolta"]
            pesi_difficalta_salati.append(coeff_k)
        
        print(f"   - Referenze salati considerate: {len(referenze_salati)} (esclusi dolci/contorni/primi)")
        
        # Calcola quota basata solo sui finger food salati
        if pesi_difficalta_salati:
            somma_coeff = sum(pesi_difficalta_salati)
            quota = coeff / somma_coeff if somma_coeff > 0 else 1 / len(referenze_salati)
        else:
            quota = 1.0
            
        peso_target = peso_totale_finale * quota
        
        # Converti in pezzi usando peso unitario
        pezzi = peso_target / peso_unitario
        risultato = int(round(pezzi))
        
        print(f"   - Quota difficoltà: {quota:.3f}")
        print(f"   - Peso target: {peso_target:.3f} kg")
        print(f"   - Pezzi calcolati: {risultato}")
        
        return risultato
    def _calcola_accompagnamenti(self, key, num_pers, servizio_lower):
        """Logica per accompagnamenti"""
        voce = self.voce_originale.get(key, key).strip()
        print(f"🍞 DEBUG ACCOMPAGNAMENTI per {voce}:")
        print(f"   - Persone: {num_pers}")
        
        # Calcolo specifico per ogni accompagnamento
        if "grissini" in voce.lower():
            # Logica normale per grissini (solo se non gestiti altrove)
            pezzi = int(num_pers * 0.5)
            print(f"   🥖 Grissini: {num_pers} persone × 0.5 = {pezzi} pezzi")
            return f"{pezzi} pz"
        
        elif "bocconcini alle olive" in voce.lower():
            pezzi = int(num_pers * 0.5)
            print(f"   🫒 Bocconcini alle olive: {num_pers} persone × 0.5 = {pezzi} pezzi")
            return f"{pezzi} pz"
        
        elif "pane a fette" in voce.lower():
                    # Ogni 10 persone = 0.5 kg = 0.5 pz (1 kg = 1 pz)
                    pz = int(num_pers / 20)  # Stesso calcolo, ma ora interpretiamo come pezzi
                    
                    # Arrotondamento intelligente
                    if pz <= int(pz) + 0.25:
                        pz_finali = int(pz)
                    elif pz <= int(pz) + 0.75:
                        pz_finali = int(pz) + 0.5
                    else:
                        pz_finali = int(pz) + 1
                    
                    pz_finali = max(0.5, pz_finali)  # Minimo 0.5 pz
                    
                    print(f"   🍞 Pane a fette: {num_pers} persone × 0.05 = {pz:.2f} pz → {pz_finali} pz")
                    
                    if pz_finali == int(pz_finali):
                        return f"{int(pz_finali)} pz"
                    else:
                        return f"{pz_finali:.1f} pz"
        
        # Gestisce crescenta e crescentine per taglieri salumi
        elif "mini crescentine fritte" in voce.lower():
            # Controlla se è selezionato "Taglieri di salumi medi"
            taglieri_salumi_selezionati = False
            for k, checkbox in self.selected.items():
                if hasattr(checkbox, 'isChecked') and checkbox.isChecked():
                    voce_k = self.voce_originale.get(k, "").lower()
                    if "taglieri di salumi medi" in voce_k:
                        taglieri_salumi_selezionati = True
                        break
            
            if taglieri_salumi_selezionati:
                # Crescentine: 1 per persona
                print(f"   🥮 Crescentine: {num_pers} pz (1 per persona)")
                return num_pers
        
        # Gestisce miele, composta di frutta per taglieri formaggi
        elif any(voce_item in voce.lower() for voce_item in ["miele", "composta di frutta", "composte di frutta"]):
            # Controlla se è selezionato "Taglieri formaggi piccoli"
            taglieri_formaggi_selezionati = False
            for k, checkbox in self.selected.items():
                if hasattr(checkbox, 'isChecked') and checkbox.isChecked():
                    voce_k = self.voce_originale.get(k, "").lower()
                    if "taglieri formaggi piccoli" in voce_k:
                        taglieri_formaggi_selezionati = True
                        break
            
            if taglieri_formaggi_selezionati:
                # Stesso numero dei taglieri formaggi (1 ogni 20 persone)
                num_accompagnamenti = max(1, int(num_pers / 20))
                print(f"   🍯 {voce} per taglieri: {num_accompagnamenti} pz")
                return num_accompagnamenti
        
        # Fallback per altri accompagnamenti
        else:
            print(f"   ⚠️ Accompagnamento non riconosciuto: {voce}")
            return 1

    def _calcola_dolci(self, key, num_pers, servizio_lower, tipo_ap, peso_unitario):
        """Logica per dolci"""
        voce = self.voce_originale.get(key, key).strip()
        
        # Ottieni coefficiente difficoltà
        coeff = self.dati_referenze.get(voce.strip().lower(), {"difficolta": 1.0})["difficolta"]
        
        # Lista dei bicchierini mousse
        bicchierini_mousse = [
            "bicch carta mousse chantilly+cucch bio",
            "bicch carta mousse pistacchio+cucch bio", 
            "bicch. carta mousse cioccolato+cucch bio"
        ]
        
        # Conta quanti bicchierini mousse sono selezionati
        bicchierini_selezionati = []
        altri_dolci_selezionati = []
        
        for k in self.referenze_dolci:
            voce_dolce = self.voce_originale.get(k, "").lower()
            if any(bic.lower() in voce_dolce for bic in bicchierini_mousse):
                bicchierini_selezionati.append(k)
            else:
                altri_dolci_selezionati.append(k)
        
        print(f"🍮 DEBUG DOLCI SPECIALI:")
        print(f"   - Bicchierini mousse: {len(bicchierini_selezionati)}")
        print(f"   - Altri dolci: {len(altri_dolci_selezionati)}")
        
        # Determina se ha primi (per calcolare peso dolci)
        ha_primi = bool(self.referenze_primi) and len(self.referenze_primi) > 0
        
        # Caso 1: Solo bicchierini (2 o più)
        if len(bicchierini_selezionati) >= 2 and len(altri_dolci_selezionati) == 0:
            if key in bicchierini_selezionati:
                porzioni_per_persona = 1.2  # Media tra 1.2 e 1.5
                totale_porzioni = num_pers * porzioni_per_persona
                porzioni_per_bicchierino = totale_porzioni / len(bicchierini_selezionati)
                
                risultato = max(1, int(round(porzioni_per_bicchierino)))
                print(f"   🍮 Solo bicchierini: {risultato} porzioni")
                return risultato
        
        # Caso 2: Un bicchierino + altri dolci (rapporto 35%-65%)
        elif len(bicchierini_selezionati) == 1 and len(altri_dolci_selezionati) > 0:
            # Determina peso standard per il tipo di servizio
            if "aperitivo" in servizio_lower:
                peso_dolci_per_persona = 0.075 if tipo_ap == "rinforzato" else 0.040
            else:
                peso_dolci_per_persona = 0.080 if ha_primi else 0.070
            
            peso_totale_dolci = peso_dolci_per_persona * num_pers
            
            if key in bicchierini_selezionati:
                # Bicchierino: 35% del peso totale dolci
                peso_target = peso_totale_dolci * 0.35
                print(f"   🍮 Bicchierino (35%): peso_target = {peso_target:.3f} kg")
            else:
                # Altri dolci: 65% diviso tra loro
                peso_target = (peso_totale_dolci * 0.65) / len(altri_dolci_selezionati)
                print(f"   🍮 Altro dolce (65%): peso_target = {peso_target:.3f} kg")
            
            pezzi = peso_target / peso_unitario
            risultato = max(1, int(round(pezzi)))
            return risultato
        
        # Caso 3: Logica standard (nessun bicchierino o solo altri dolci)
        else:
            if "aperitivo" in servizio_lower:
                peso_dolci_per_persona = 0.060 if tipo_ap == "rinforzato" else 0.040
            else:
                peso_dolci_per_persona = 0.080 if ha_primi else 0.070
            
            pesi_difficalta = [
                self.dati_referenze.get(
                    self.voce_originale.get(k, "").strip().lower(), {"difficolta": 1.0}
                )["difficolta"] for k in self.referenze_dolci
            ]
            somma_coeff = sum(pesi_difficalta)
            quota = coeff / somma_coeff if somma_coeff > 0 else 1 / len(self.referenze_dolci)
            peso_target = peso_dolci_per_persona * num_pers * quota
            
            pezzi = peso_target / peso_unitario
            print(f"   🍮 Logica standard: {int(round(pezzi))} pezzi")
            return max(1, int(round(pezzi)))
        
    def _calcola_standard(self, key, num_pers, peso_unitario):
        """Logica standard per altri servizi"""
        voce = self.voce_originale.get(key, key).strip()
        
        print(f"🔧 CALCOLO STANDARD per {voce}:")
        print(f"   - Persone: {num_pers}")
        print(f"   - Peso unitario: {peso_unitario}")
        
        # Logica base: 1 pezzo per persona
        pezzi_base = num_pers
        
        # USA LA FUNZIONE HELPER CENTRALIZZATA
        if self._va_in_kg(voce):
            print(f"   - Va in kg: True")
            kg_totali = pezzi_base * peso_unitario
            return f"{kg_totali:.1f} kg"
        else:
            print(f"   - Va in kg: False")
            return f"{pezzi_base} pz"

    def _calcola_finger_food_con_bilanciamento(self, key, referenze_cibo, num_pers, servizio_lower, tipo_ap, peso_unitario, bilanciamento):
        """Calcola finger food con bilanciamento dinamico"""
        voce = self.voce_originale.get(key, key).strip()
        coeff = self.dati_referenze.get(voce.strip().lower(), {"difficolta": 1.0})["difficolta"]
        
        # Determina se applicare bilanciamento
        # Determina se applicare bilanciamento SOLO se ci sono effettivamente secondi/contorni
        is_lunch = "lunch" in servizio_lower or "pranzo" in servizio_lower
        is_aperitivo_rinforzato = "aperitivo" in servizio_lower and tipo_ap == "rinforzato"
        ha_secondi_contorni = bool(self.referenze_secondi) or bool(self.referenze_contorni)
        use_balancing = (is_lunch or is_aperitivo_rinforzato) and ha_secondi_contorni
        
        if use_balancing:
            print(f"   🍤 FINGER FOOD CON BILANCIAMENTO per {voce}:")
            
            # Calcola peso base finger food
            ha_primi = bool(self.referenze_primi) and len(self.referenze_primi) > 0
            
            if is_lunch:
                if ha_primi:
                    peso_finger_base_per_persona = 0.180  # 180g salato con primi (già ridotto)
                else:
                    peso_finger_base_per_persona = 0.350  # 350g salato senza primi
            else:  # aperitivo rinforzato
                peso_finger_base_per_persona = 0.280  # 280g salato aperitivo rinforzato
            
            peso_finger_totale_base = peso_finger_base_per_persona * num_pers
            
            # Applica riduzione da bilanciamento
            riduzione_finger = bilanciamento["riduzione_finger"]
            peso_finger_totale_finale = max(0.5, peso_finger_totale_base - riduzione_finger)
            
            print(f"   - Peso base finger food: {peso_finger_totale_base:.3f} kg")
            print(f"   - Riduzione: {riduzione_finger:.3f} kg")
            print(f"   - Peso finale finger food: {peso_finger_totale_finale:.3f} kg")
            
            # Calcola quota per questa voce
            # Calcola quota per questa voce (ESCLUDI taglieri e accompagnamenti)
            pesi_difficalta = []
            referenze_finger_reali = []

            for k in referenze_cibo:
                voce_k = self.voce_originale.get(k, "").strip().lower()
                
                # ESCLUDI taglieri e loro accompagnamenti dal calcolo finger food
                if ("taglieri" in voce_k or 
                    "miele" in voce_k or 
                    "composta" in voce_k):
                    continue
                
                referenze_finger_reali.append(k)
                dati_k = self.dati_referenze.get(voce_k, {"difficolta": 1.0})
                pesi_difficalta.append(dati_k["difficolta"])

            # Usa solo le referenze finger food reali per il calcolo
            if referenze_finger_reali:
                somma_coeff = sum(pesi_difficalta)
                quota = coeff / somma_coeff if somma_coeff > 0 else 1 / len(referenze_finger_reali)
            else:
                quota = 1.0
            peso_target = peso_finger_totale_finale * quota
            
            print(f"   - Quota: {quota:.3f}")
            print(f"   - Peso target questa voce: {peso_target:.3f} kg")
            
            # USA LA FUNZIONE HELPER CENTRALIZZATA
            if self._va_in_kg(voce):
                print(f"   - Va in kg: True")
                return f"{peso_target:.1f} kg"
            else:
                print(f"   - Va in kg: False")
                # Guard per evitare divisione per zero
                if not peso_unitario or peso_unitario <= 0:
                    print("   ⚠️ peso_unitario nullo/<=0: ritorno 1 pezzo di sicurezza")
                    return "1 pz"
                
                pezzi = peso_target / peso_unitario
                
                # PRODOTTI PESANTI vs ALTRI
                if peso_unitario >= 0.5:
                    risultato_pezzi = max(1, int(round(pezzi)))
                    print(f"   🥬 Prodotto pesante bilanciato: {risultato_pezzi} pezzi (no blocchi)")
                else:
                    def arrotonda_blocchi(x, base):
                        x = int(round(x))
                        resto = x % base
                        return x + (base - resto) if resto >= base * 0.6 else x - resto
                    risultato_pezzi = max(1, arrotonda_blocchi(pezzi, 5))
                    print(f"   🍤 Finger food bilanciato: {risultato_pezzi} pezzi")
                
                return f"{risultato_pezzi} pz"
        
        else:
            # Logica originale per servizi senza bilanciamento
            print(f"   🍤 FINGER FOOD SENZA BILANCIAMENTO per {voce}:")
            
            
             # PESO VARIABILE A SECONDA DEL SERVIZIO
            if "coffee" in servizio_lower or "tea" in servizio_lower:
                peso_salato_per_persona = 0.090  # 80g per coffee/tea break
                print(f"   ☕ Coffee/Tea break: 80g per persona")
            elif "aperitivo" in servizio_lower and tipo_ap == "leggero":
                peso_salato_per_persona = 0.150  # 150g per aperitivo leggero
                print(f"   🥂 Aperitivo leggero: 150g per persona")
                
            else:
                peso_salato_per_persona = 0.220  # 200g per servito e altri servizi
                print(f"   🍽️ Altri servizi: 250g per persona")
            
            # Calcola quota per questa voce (ESCLUDI taglieri e accompagnamenti)
            pesi_difficalta = []
            referenze_finger_reali = []

            for k in referenze_cibo:
                voce_k = self.voce_originale.get(k, "").strip().lower()
                
                # ESCLUDI taglieri e loro accompagnamenti dal calcolo finger food
                if ("taglieri" in voce_k or 
                    "miele" in voce_k or 
                    "composta" in voce_k):
                    continue
                
                referenze_finger_reali.append(k)
                dati_k = self.dati_referenze.get(voce_k, {"difficolta": 1.0})
                pesi_difficalta.append(dati_k["difficolta"])

            # Usa solo le referenze finger food reali per il calcolo
            if referenze_finger_reali:
                somma_coeff = sum(pesi_difficalta)
                quota = coeff / somma_coeff if somma_coeff > 0 else 1 / len(referenze_finger_reali)
            else:
                quota = 1.0
            peso_salato = peso_salato_per_persona * num_pers
            peso_target = peso_salato * quota
            
            print(f"   🍽️ Servito - Peso target: {peso_target:.3f} kg")
            
            # USA LA FUNZIONE HELPER CENTRALIZZATA
            if self._va_in_kg(voce):
                print(f"   - Va in kg: True")
                return f"{peso_target:.1f} kg"
            else:
                print(f"   - Va in kg: False")
                if peso_unitario <= 0:
                    return "1 pz"
                
                pezzi = peso_target / peso_unitario
                
                if peso_unitario >= 0.5:
                    risultato_pezzi = max(1, int(round(pezzi)))
                else:
                    def arrotonda_blocchi(x, base):
                        x = int(round(x))
                        resto = x % base
                        return x + (base - resto) if resto >= base * 0.6 else x - resto
                    risultato_pezzi = max(1, arrotonda_blocchi(pezzi, 5))
                
                print(f"   🍽️ Servito finger food: {risultato_pezzi} pezzi")
                return f"{risultato_pezzi} pz"
    def gestisci_selezione_insalata(self, checkbox, key_insalata):
        """Quando viene selezionata un'insalata, attiva automaticamente i condimenti per insalata"""
        
        voce_insalata = self.voce_originale.get(key_insalata, "").lower()
        
        # CORREZIONE: Usa keywords ma escludi i condimenti
        keywords_insalata = ["insalat", "caprese"]
        keywords_da_escludere = ["condimento", "condimenti"]  # NUOVO: Escludi condimenti
        
        # È un'insalata se contiene le parole chiave MA NON contiene "condimento"
        is_insalata_speciale = (any(keyword in voce_insalata for keyword in keywords_insalata) and
                            not any(esclusione in voce_insalata for esclusione in keywords_da_escludere))
        
        if is_insalata_speciale:
            print(f"🥗 Gestione insalata: {voce_insalata[:50]}...")
            print(f"   🔍 checkbox.isChecked() = {checkbox.isChecked()}")
            
            if checkbox.isChecked():
                print(f"   ➡️ Entrando nel ramo SELEZIONE")
                # SELEZIONE: attiva i condimenti per insalata
                for k, cb in self.selected.items():
                    if hasattr(cb, 'isChecked'):
                        voce_condimenti = self.voce_originale.get(k, "").lower()
                        
                        # Cerca "Condimento per insalata"
                        if "condimento" in voce_condimenti and "insalata" in voce_condimenti:
                            print(f"   🎯 Trovato condimento: chiave={k}")
                            if not cb.isChecked():
                                cb.setChecked(True)
                                print(f"✅ Auto-selezionato: {self.voce_originale.get(k, k)} (per insalata)")
                            else:
                                print(f"ℹ️ Condimento già selezionato")
                            break
            else:
                print(f"   ➡️ Entrando nel ramo DESELEZIONE")
                # DESELEZIONE: controlla se ci sono altre insalate selezionate
                altre_insalate_attive = False
                
                for k, cb in self.selected.items():
                    if k != key_insalata and hasattr(cb, 'isChecked') and cb.isChecked():
                        voce_altra = self.voce_originale.get(k, "").lower()
                        
                        # CORREZIONE: Controlla che sia un'insalata vera (non un condimento)
                        is_altra_insalata = (any(keyword in voce_altra for keyword in keywords_insalata) and
                                            not any(esclusione in voce_altra for esclusione in keywords_da_escludere))
                        
                        if is_altra_insalata:
                            altre_insalate_attive = True
                            print(f"   🔍 Trovata altra insalata attiva: {voce_altra[:30]}...")
                            break
                
                print(f"   📊 Altre insalate attive: {altre_insalate_attive}")
                
                # Se non ci sono altre insalate attive, disattiva i condimenti
                if not altre_insalate_attive:
                    for k, cb in self.selected.items():
                        if hasattr(cb, 'isChecked'):
                            voce_condimenti = self.voce_originale.get(k, "").lower()
                            
                            if "condimento" in voce_condimenti and "insalata" in voce_condimenti:
                                cb.setChecked(False)
                                print(f"❌ Auto-deselezionato: {self.voce_originale.get(k, k)} (nessuna insalata attiva)")
                                break
                else:
                    print(f"ℹ️ Condimento mantenuto (altre insalate ancora attive)")
        else:
            print(f"   ⚠️ Voce '{voce_insalata[:30]}...' non riconosciuta come insalata speciale")
    def gestisci_selezione_secondi_contorni(self, checkbox, key_secondo_contorno):
        """Quando viene selezionato un secondo o contorno, attiva automaticamente pane e accompagnamenti"""
        
        voce_corrente = self.voce_originale.get(key_secondo_contorno, "").lower()
        
        # Lista degli extra da auto-selezionare (dalla sezione "Optionals" > "Cibo Extra")
        extras_da_selezionare = [
            "pane a fette",
            "bocconcini alle olive", 
            "grissini"
        ]
        
        # Verifica se questa voce è un secondo o contorno
        is_secondo_o_contorno = ("secondi piatti" in key_secondo_contorno.lower() or 
                                "contorni" in key_secondo_contorno.lower())
        
        if is_secondo_o_contorno:
            print(f"🍖 Gestione secondo/contorno: {voce_corrente[:50]}...")
            print(f"   🔍 checkbox.isChecked() = {checkbox.isChecked()}")
            
            if checkbox.isChecked():
                print(f"   ➡️ Entrando nel ramo SELEZIONE")
                # SELEZIONE: attiva pane e accompagnamenti nella sezione "Optionals"
                for k, cb in self.selected.items():
                    if hasattr(cb, 'isChecked'):
                        voce_extra = self.voce_originale.get(k, "").lower()
                        
                        # Verifica che sia nella sezione corretta (Optionals)
                        if "optionals" in k.lower():
                            # Cerca gli extra nella lista
                            for extra in extras_da_selezionare:
                                if extra in voce_extra:
                                    print(f"   🎯 Trovato extra: {extra} (chiave={k})")
                                    if not cb.isChecked():
                                        cb.setChecked(True)
                                        print(f"✅ Auto-selezionato: {self.voce_originale.get(k, k)} (per secondo/contorno)")
                                    else:
                                        print(f"ℹ️ Extra già selezionato: {extra}")
                                    break
            else:
                print(f"   ➡️ Entrando nel ramo DESELEZIONE")
                # DESELEZIONE: controlla se ci sono altri secondi/contorni selezionati
                altri_secondi_contorni_attivi = False
                
                for k, cb in self.selected.items():
                    if k != key_secondo_contorno and hasattr(cb, 'isChecked') and cb.isChecked():
                        if ("secondi piatti" in k.lower() or "contorni" in k.lower()):
                            altri_secondi_contorni_attivi = True
                            print(f"   🔍 Trovato altro secondo/contorno attivo: {k[:30]}...")
                            break
                
                print(f"   📊 Altri secondi/contorni attivi: {altri_secondi_contorni_attivi}")
                
                # Se non ci sono altri secondi/contorni attivi, disattiva gli extra
                if not altri_secondi_contorni_attivi:
                    for k, cb in self.selected.items():
                        if hasattr(cb, 'isChecked'):
                            voce_extra = self.voce_originale.get(k, "").lower()
                            
                            # Verifica che sia nella sezione corretta (Optionals)
                            if "optionals" in k.lower():
                                for extra in extras_da_selezionare:
                                    if extra in voce_extra:
                                        cb.setChecked(False)
                                        print(f"❌ Auto-deselezionato: {self.voce_originale.get(k, k)} (nessun secondo/contorno attivo)")
                                        break
                else:
                    print(f"ℹ️ Extra mantenuti (altri secondi/contorni ancora attivi)")
    def gestisci_selezione_bevanda(self, checkbox, key_bevanda):
        """Quando viene selezionata/deselezionata una bevanda, gestisce gli accessori SOLO DELLA STESSA SEZIONE"""
        
        voce_bevanda = self.voce_originale.get(key_bevanda, "").lower()
        
        print(f"🍷 Gestione bevanda: {voce_bevanda[:50]}...")
        
        # ========== ESTRAI LA SEZIONE DELLA BEVANDA ==========
        # key_bevanda formato: "lunch buffet|bevande|nr.    prosecco+cons"
        parti_key = key_bevanda.split('|')
        if len(parti_key) >= 2:
            sezione_bevanda = parti_key[0].lower()  # es: "lunch buffet", "aperitivo a buffet"
            print(f"   🔍 Sezione bevanda: '{sezione_bevanda}'")
        else:
            print(f"   ⚠️ Formato key non riconosciuto: {key_bevanda}")
            return
        
        if checkbox.isChecked():
            # SELEZIONE: attiva gli accessori SOLO DELLA STESSA SEZIONE
            for k, cb in self.selected.items():
                if hasattr(cb, 'isChecked') and self.voce_categoria.get(k) == "accessori":
                    
                    # ========== FILTRO PER SEZIONE ==========
                    parti_k = k.split('|')
                    if len(parti_k) >= 2:
                        sezione_accessorio = parti_k[0].lower()
                        
                        # CONTROLLA SE L'ACCESSORIO È NELLA STESSA SEZIONE DELLA BEVANDA
                        if sezione_accessorio != sezione_bevanda:
                            continue  # Salta accessori di altre sezioni
                    else:
                        continue  # Salta se il formato non è riconosciuto
                    
                    voce_accessorio = self.voce_originale.get(k, "").lower()
                    
                    # Bowl+ice per tutte le bevande con CONS
                    if ("bowl" in voce_accessorio and "ice" in voce_accessorio) or "bowl+ice" in voce_accessorio:
                        if "+cons" in voce_bevanda:
                            cb.setChecked(True)
                            print(f"✅ Auto-selezionato: {self.voce_originale.get(k, k)} (bevanda CONS, sezione: {sezione_accessorio})")
                    
                    # Flutes per Prosecco
                    if "prosecco" in voce_bevanda and "flutes" in voce_accessorio and "mono" in voce_accessorio:
                        cb.setChecked(True)
                        print(f"✅ Auto-selezionato: {self.voce_originale.get(k, k)} (per Prosecco, sezione: {sezione_accessorio})")
                    
                    # Calici per Chardonnay e Cabernet
                    if ("chardonnay" in voce_bevanda or "cabernet" in voce_bevanda) and "calici" in voce_accessorio and "mono" in voce_accessorio:
                        cb.setChecked(True)
                        print(f"✅ Auto-selezionato: {self.voce_originale.get(k, k)} (per vino, sezione: {sezione_accessorio})")
        else:
            # DESELEZIONE: controlla se ci sono altre bevande selezionate NELLA STESSA SEZIONE che usano gli stessi accessori
            altre_bevande_cons = False
            prosecco_selezionato = False
            vino_selezionato = False
            
            for k, cb in self.selected.items():
                if k != key_bevanda and hasattr(cb, 'isChecked') and cb.isChecked():
                    if self.voce_categoria.get(k) == "bevande":
                        
                        # ========== FILTRO PER SEZIONE ANCHE QUI ==========
                        parti_k = k.split('|')
                        if len(parti_k) >= 2:
                            sezione_altra_bevanda = parti_k[0].lower()
                            # Controlla solo bevande della stessa sezione
                            if sezione_altra_bevanda != sezione_bevanda:
                                continue
                        else:
                            continue
                        
                        voce_altra = self.voce_originale.get(k, "").lower()
                        if "+cons" in voce_altra:
                            altre_bevande_cons = True
                        if "prosecco" in voce_altra:
                            prosecco_selezionato = True
                        if "chardonnay" in voce_altra or "cabernet" in voce_altra:
                            vino_selezionato = True
            
            # Deseleziona solo se non ci sono altre bevande NELLA STESSA SEZIONE che li richiedono
            for k, cb in self.selected.items():
                if hasattr(cb, 'isChecked') and self.voce_categoria.get(k) == "accessori":
                    
                    # ========== FILTRO PER SEZIONE ANCHE PER LA DESELEZIONE ==========
                    parti_k = k.split('|')
                    if len(parti_k) >= 2:
                        sezione_accessorio = parti_k[0].lower()
                        # Deseleziona solo accessori della stessa sezione
                        if sezione_accessorio != sezione_bevanda:
                            continue
                    else:
                        continue
                    
                    voce_accessorio = self.voce_originale.get(k, "").lower()
                    
                    # Bowl+ice solo se non ci sono altre bevande CONS NELLA STESSA SEZIONE
                    if not altre_bevande_cons and (("bowl" in voce_accessorio and "ice" in voce_accessorio) or "bowl+ice" in voce_accessorio):
                        cb.setChecked(False)
                        print(f"❌ Auto-deselezionato: {self.voce_originale.get(k, k)} (sezione: {sezione_accessorio})")
                    
                    # Flutes solo se non c'è prosecco NELLA STESSA SEZIONE
                    if not prosecco_selezionato and "flutes" in voce_accessorio:
                        cb.setChecked(False)
                        print(f"❌ Auto-deselezionato: {self.voce_originale.get(k, k)} (sezione: {sezione_accessorio})")
                    
                    # Calici solo se non ci sono vini NELLA STESSA SEZIONE
                    if not vino_selezionato and "calici" in voce_accessorio:
                        cb.setChecked(False)
                        print(f"❌ Auto-deselezionato: {self.voce_originale.get(k, k)} (sezione: {sezione_accessorio})")

    # ==============================================================================
    # STESSA LOGICA PER gestisci_selezione_primo (se serve)
    # ==============================================================================

    def gestisci_selezione_primo(self, checkbox, key_primo):
        """Quando viene selezionato/deselezionato un primo, gestisce gli accessori SOLO DELLA STESSA SEZIONE"""
        
        voce_primo = self.voce_originale.get(key_primo, "").lower()
        
        # Determina se è un primo caldo o freddo
        primi_caldi = ["lasagne", "crespelle", "cannelloni", "rosette", "nidi", 
                    "gnocchi", "intrighi", "tagliatelle", "gramigna", "fusilli",
                    "garganelli", "rotolini", "(inox)"]
        
        is_primo_caldo = any(x in voce_primo for x in primi_caldi)
        
        print(f"🍝 Gestione primo: {voce_primo[:50]}... | Caldo: {is_primo_caldo}")
        
        # ========== ESTRAI LA SEZIONE DEL PRIMO ==========
        parti_key = key_primo.split('|')
        if len(parti_key) >= 2:
            sezione_primo = parti_key[0].lower()
            print(f"   🔍 Sezione primo: '{sezione_primo}'")
        else:
            print(f"   ⚠️ Formato key non riconosciuto: {key_primo}")
            return
        
        if checkbox.isChecked():
            # SELEZIONE: attiva gli accessori SOLO DELLA STESSA SEZIONE
            for k, cb in self.selected.items():
                if hasattr(cb, 'isChecked') and self.voce_categoria.get(k) == "accessori":
                    
                    # ========== FILTRO PER SEZIONE ==========
                    parti_k = k.split('|')
                    if len(parti_k) >= 2:
                        sezione_accessorio = parti_k[0].lower()
                        if sezione_accessorio != sezione_primo:
                            continue  # Salta accessori di altre sezioni
                    else:
                        continue
                    
                    voce_accessorio = self.voce_originale.get(k, "").lower()
                    
                    # "Acc x porz" per tutti i primi
                    if "acc x porz" in voce_accessorio or "acc." in voce_accessorio:
                        cb.setChecked(True)
                        print(f"✅ Auto-selezionato: {self.voce_originale.get(k, k)} (sezione: {sezione_accessorio})")
                    
                    # "Scaldavivande" solo per primi caldi
                    if is_primo_caldo and "scaldavivande" in voce_accessorio:
                        cb.setChecked(True)
                        print(f"✅ Auto-selezionato: {self.voce_originale.get(k, k)} (primo caldo, sezione: {sezione_accessorio})")
        else:
            # DESELEZIONE: controlla se ci sono altri primi selezionati NELLA STESSA SEZIONE
            altri_primi_selezionati = False
            altri_primi_caldi_selezionati = False
            
            for k, cb in self.selected.items():
                if k != key_primo and hasattr(cb, 'isChecked') and cb.isChecked():
                    categoria = self.voce_categoria.get(k)
                    if categoria == "cibo":
                        
                        # ========== FILTRO PER SEZIONE ==========
                        parti_k = k.split('|')
                        if len(parti_k) >= 2:
                            sezione_altro_primo = parti_k[0].lower()
                            if sezione_altro_primo != sezione_primo:
                                continue  # Considera solo primi della stessa sezione
                        else:
                            continue
                        
                        voce_altro = self.voce_originale.get(k, "").lower()
                        # Controlla se è un primo piatto
                        if any(primo in voce_altro for primo in ["primi", "pasta", "risotto", "gnocchi", "lasagne", "crespelle"]):
                            altri_primi_selezionati = True
                            if any(x in voce_altro for x in primi_caldi):
                                altri_primi_caldi_selezionati = True
            
            # Deseleziona solo se non ci sono altri primi NELLA STESSA SEZIONE che li richiedono
            if not altri_primi_selezionati:
                for k, cb in self.selected.items():
                    if hasattr(cb, 'isChecked') and self.voce_categoria.get(k) == "accessori":
                        
                        # ========== FILTRO PER SEZIONE ==========
                        parti_k = k.split('|')
                        if len(parti_k) >= 2:
                            sezione_accessorio = parti_k[0].lower()
                            if sezione_accessorio != sezione_primo:
                                continue  # Deseleziona solo accessori della stessa sezione
                        else:
                            continue
                        
                        voce_accessorio = self.voce_originale.get(k, "").lower()
                        if "acc x porz" in voce_accessorio or "acc." in voce_accessorio:
                            cb.setChecked(False)
                            print(f"❌ Auto-deselezionato: {self.voce_originale.get(k, k)} (sezione: {sezione_accessorio})")
            
            if not altri_primi_caldi_selezionati:
                for k, cb in self.selected.items():
                    if hasattr(cb, 'isChecked') and self.voce_categoria.get(k) == "accessori":
                        
                        # ========== FILTRO PER SEZIONE ==========
                        parti_k = k.split('|')
                        if len(parti_k) >= 2:
                            sezione_accessorio = parti_k[0].lower()
                            if sezione_accessorio != sezione_primo:
                                continue
                        else:
                            continue
                        
                        voce_accessorio = self.voce_originale.get(k, "").lower()
                        if "scaldavivande" in voce_accessorio:
                            cb.setChecked(False)
                            print(f"❌ Auto-deselezionato: {self.voce_originale.get(k, k)} (sezione: {sezione_accessorio})")



    def determina_unita_misura(self, voce_originale, quantita):
        """Determina se mostrare kg o pz basandoci sulla lista definita"""
        
        voce_lower = voce_originale.lower().strip()
        try:
            quantita_numerica = float(quantita)
        except (ValueError, TypeError):
            quantita_numerica = 1
        
        if isinstance(quantita, str) and "nr" in quantita.lower() and "porz" in quantita.lower():
            return quantita 
        if isinstance(quantita, str) and quantita.startswith("(") and ")" in quantita:
            # Formato (50)1 → estrai i numeri
            import re
            match = re.match(r'\((\d+)\)(\d+)', quantita)
            if match:
                totale = match.group(1)
                per_persona = match.group(2)
                return f"({totale}){per_persona} pz"
        # Ottieni i dati della referenza
        dati_referenza = self.dati_referenze.get(voce_lower, {"peso_unitario": 0.03333})
        peso_unitario = dati_referenza["peso_unitario"]
        
        
        # ========== GESTIONE PRIMI ELABORATI (PRIORITÀ ALTA) ==========
        primi_elaborati = ["lasagne", "crespelle", "cannelloni", "nidi", "rosette"]
        for primo in primi_elaborati:
            if primo in voce_lower:
                return str(quantita)  # Mantieni formato originale
        
        # ========== GESTIONE STRINGHE CON UNITÀ GIÀ DEFINITE ==========
        if isinstance(quantita, str):
            if "kg" in quantita.lower():
                return quantita  # Già formattato correttamente
            elif any(x in quantita.lower() for x in ["lasagne", "crespelle", "cannelloni", "nidi", "rosette", "porzioni"]):
                return quantita  # Già formattato correttamente
            else:
                # Prova a estrarre il numero dalla stringa
                try:
                    import re
                    numeri = re.findall(r'\d+\.?\d*', quantita)
                    if numeri:
                        quantita_numerica = float(numeri[0])
                    else:
                        return f"{quantita} pz"  # Fallback
                except:
                    return f"{quantita} pz"  # Fallback
        else:
            quantita_numerica = quantita
        
        # ========== REFERENZE CHE VANNO MOSTRATE IN KG (AGGIORNATO) ==========
        referenze_in_kg = [
            # FINGER FOOD E DOLCI (mantieni quelli che erano già in kg)
            "biscotteria", "spicchietti crostate", "trancetti tenerina crema e amerene","mini tramezzini vegetariani","rustici vegetariani",
            "rustici mix", "parmigianini", "pasticceria", "crescenta a dadini",
            "cubetti mortadella","cubetti mortadella scottati all'aceto",
            "dadini di mortadella", "rombetti di frittata", "mini tramezzini mix",
            "mozzarelline in carrozza", "mini cotolettine",
            "olive all'ascolana", "olive ascolane", "crocchette", "tagliata di frutta",
            "mozzarelle", "polpettine prosciutto e ricotta","minisw farciti mix","trancetti tenerina cioccolato",
            
            "insalata di pollo",
            "insalatona mista con ananas", 
            "insalata di pollo alla nizzarda",
            "caprese",
            "insalata di pollo marinato",
            "tacchino con asparagi", "straccetti con scorza",
            "spezzatino con topinambur", "bocconcini di tacchino", "straccetti di pollo",
            "spezzatino di maiale", "bocconcini di maiale",
            "rombetti di tacchino", "piccatine al balsamico",
            "filettino di maiale", "bocconcini di pollo", "rombi di tacchino",
            "straccetti di maiale", "filetto di tacchino",
            
            # ========== CONTORNI (AGGIUNTI) ==========
            "macedonia di verdure", "patate al rosmarino",
             "spinaci alla tirolese",
            ]
        
        # Controlla se questa referenza va mostrata in kg
        for ref_kg in referenze_in_kg:
            if ref_kg in voce_lower:
                print(f"🔍 DEBUG determina_unita_misura per '{voce_originale}':")
                print(f"   - voce_lower: '{voce_lower}'")
                print(f"   - quantita ricevuta: {quantita}")
                print(f"   - quantita_numerica calcolata: {quantita_numerica}")
                print(f"   - dati_referenza trovati: {dati_referenza}")
                print(f"   - peso_unitario usato: {peso_unitario}")
                print(f"   - match trovato con: '{ref_kg}'")
                # QUANTITA_NUMERICA è il numero di pezzi calcolato
                # Dobbiamo convertire in kg moltiplicando per peso_unitario
                kg_effettivi = quantita_numerica * peso_unitario
                print(f"   - Conversione in kg: {quantita_numerica} pezzi × {peso_unitario} = {kg_effettivi:.3f} kg")

                # ========== ARROTONDAMENTO INTELLIGENTE A 0.5 KG ==========
                intero = int(kg_effettivi)
                decimale = kg_effettivi - intero
                
                if decimale < 0.2:
                    kg_finali = intero
                elif decimale < 0.7:
                    kg_finali = intero + 0.5
                else:
                    kg_finali = intero + 1
                print(f"   - kg_finali dopo arrotondamento: {kg_finali}")

                # Formatta output
                if kg_finali == int(kg_finali):
                    risultato = f"{int(kg_finali)} kg"
                else:
                    risultato = f"{kg_finali:.1f} kg"
                
                print(f"   ✅ Risultato finale: {risultato}")
                return risultato
        
        # ========== PRIMI PIATTI NORMALI ==========
        primi_parole = ["insalata di riso", "cous cous", "quinoa", "orecchiette", "mezze penne", "penne","fioriere quadrate + decori"]
        for parola in primi_parole:
            if parola in voce_lower:
                kg_effettivi = quantita_numerica * peso_unitario
                # Stesso arrotondamento per i primi
                intero = int(kg_effettivi)
                decimale = kg_effettivi - intero
                
                if decimale < 0.2:
                    kg_finali = intero
                elif decimale < 0.7:
                    kg_finali = intero + 0.5
                else:
                    kg_finali = intero + 1
                
                if kg_finali == int(kg_finali):
                    return f"{int(kg_finali)} kg"
                else:
                    return f"{kg_finali:.1f} kg"
        
        # ========== DEFAULT: PEZZI ==========
        return f"{int(quantita_numerica)} pz"
    def stampa_quantita_docx(self, quant, voce, cella, bold=False):
        
        """Stampa quantità nel documento Word con unità corretta"""
        print(f"📝 DEBUG stampa_quantita_docx per '{voce}':")
        print(f"   - quant ricevuto: {quant}")
        print(f"   - tipo quant: {type(quant)}")
        p = cella.add_paragraph()
        
        # Usa il nuovo metodo per determinare l'unità
        quantita_formattata = self.determina_unita_misura(voce, quant)
        
        # Testo finale
        # Gestione speciale per primi elaborati che già contengono il tipo
        primi_elaborati = ["lasagne", "crespelle", "cannelloni", "nidi", "rosette", "porzioni", "gnocchi"]
        if any(primo in quantita_formattata.lower() for primo in primi_elaborati):
            # Se la quantità già contiene il tipo, usa solo quella
            text = f"{quantita_formattata} {voce.split()[0] if len(voce.split()) > 1 else voce}"
            # Oppure ancora meglio, rimuovi la parola duplicata:
            parole_voce = voce.split()
            if len(parole_voce) > 1 and parole_voce[0].lower() in quantita_formattata.lower():
                text = f"{quantita_formattata} {' '.join(parole_voce[1:])}"
            else:
                text = f"{quantita_formattata} {voce}"
        else:
            # Normale
            text = f"{quantita_formattata} {voce}"        
        run = p.add_run(text)
        if bold:
            run.bold = True

    def apply_borders(self, cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ('top', 'start', 'bottom', 'end'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '6')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def generate_doc(self):
        """Genera il documento Word dell'ordine - VERSIONE CORRETTA"""
        
        # Import necessari all'inizio
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        # Validazione input
        try:
            persone = int(self.entries["Numero persone"].text())
            if persone <= 0:
                raise ValueError("Il numero di persone deve essere positivo")
        except ValueError:
            QMessageBox.warning(self, "Errore Input", "Inserire un numero valido per le persone.")
            return

        if not self.entries["Nome"].text().strip():
            QMessageBox.warning(self, "Errore Input", "Inserire il nome del cliente.")
            return

        # Progress dialog
        progress = QProgressBar(self)
        progress.setWindowTitle("Generazione Ordine")
        progress.setGeometry(self.width()//2 - 150, self.height()//2, 300, 30)
        progress.setRange(0, 100)
        progress.show()
        progress.setValue(10)
        QApplication.processEvents()

        # ========== DETERMINA MODALITÀ E CREA DOCUMENTO ==========
        doc = None
        final_path = None
        
        # Controlla se siamo in modalità aggiungi servizi
        if hasattr(self, 'modalita_aggiungi_servizi') and self.modalita_aggiungi_servizi and hasattr(self, 'ordine_da_modificare') and self.ordine_da_modificare:
            # MODALITÀ AGGIUNGI SERVIZI
            original_path = self.ordine_da_modificare.get("path_docx")
            if original_path and os.path.exists(original_path):
                print(f"➕ MODALITÀ AGGIUNGI SERVIZI: Aggiungendo al documento {original_path}")
                final_path = original_path
                
                try:
                    doc = Document(original_path)
                    print("✅ Documento esistente caricato per aggiungi servizi")
                    
                    # Aggiungi solo spazio bianco per separare i servizi
                    doc.add_paragraph()  # Spazio
                    doc.add_paragraph()  # Spazio aggiuntivo
                    
                    # Aggiungi tabella NR PROGR anche per servizi aggiuntivi
                    table_title_add = doc.add_table(rows=1, cols=1)
                    table_title_add.autofit = False
                    table_title_add.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table_title_add.columns[0].width = Inches(2.0)

                    cell_add = table_title_add.cell(0, 0)
                    cell_add.width = Inches(3.0)

                    run_add = cell_add.paragraphs[0].add_run("NR PROGR")
                    run_add.bold = True
                    run_add.font.name = "Calibri"
                    run_add.font.size = Pt(16)
                    cell_add.paragraphs[0].alignment = 1

                    self.apply_borders(cell_add)
                    p_leggenda = doc.add_paragraph()
                    p_leggenda.alignment = 2  # Destra
                    p_leggenda.paragraph_format.space_before = Pt(0)
                    p_leggenda.paragraph_format.space_after = Pt(6)

                    # Primo run: "deluxe+decori"
                    run1 = p_leggenda.add_run("vassoi deluxe+decori")
                    run1.bold = True
                    run1.font.size = Pt(11)
                    run1.font.name = "Calibri"

                    # Vai a capo
                    p_leggenda.add_run("\n")

                    # Secondo run: "* indica i caldi" 
                    run2 = p_leggenda.add_run("* indica i caldi")
                    run2.font.size = Pt(9)
                    run2.font.name = "Calibri"
                    run2.italic = True
                    
                    # Aggiungi anche intestazione completa per servizio aggiuntivo
                    # Data evento - CORREZIONE BUG
                    if hasattr(self.entries["Data evento"], 'date'):
                        qdate = self.entries["Data evento"].date()
                        data_evento = qdate.toString("dd/MM/yyyy")
                    else:
                        data_evento = self.entries["Data evento"].text()
                    
                    data_obj = datetime.strptime(data_evento, "%d/%m/%Y")
                    giorno_str = data_obj.strftime("%A %d %B %Y").upper()
                    
                    # Aggiungi intestazione completa per servizio aggiuntivo
                    par_data = doc.add_paragraph()
                    par_data.alignment = 1
                    run_data = par_data.add_run(giorno_str)
                    run_data.bold = True
                    run_data.underline = True

                    nome = self.entries["Nome"].text()
                    cellulare = self.entries["Cellulare"].text()
                    via = self.entries["Via"].text()
                    par_cliente = doc.add_paragraph()
                    par_cliente.alignment = 1
                    run_cliente = par_cliente.add_run(f"All'attenzione di {nome} {cellulare}\nPresso {via}")
                    run_cliente.bold = True
                    run_cliente.underline = True
                    run_data.bold = True
                    run_data.underline = True

                    
                except Exception as e:
                    print(f"❌ Errore caricamento documento per aggiungi servizi: {e}")
                    doc = Document()  # Crea nuovo se fallisce
                    final_path = None  # Verrà creato dopo
                    
        # Controlla se siamo in modalità modifica normale
        elif hasattr(self, 'modalita_modifica') and self.modalita_modifica and hasattr(self, 'ordine_da_modificare') and self.ordine_da_modificare:
            # MODALITÀ MODIFICA NORMALE (sovrascrive)
            original_path = self.ordine_da_modificare.get("path_docx")
            if original_path and os.path.exists(original_path):
                final_path = original_path
                print(f"🔄 MODALITÀ MODIFICA: Sovrascrivendo {original_path}")
            doc = Document()  # Sempre nuovo documento per modifica
            
        else:
            # MODALITÀ NORMALE: Nuovo ordine
            print("📄 MODALITÀ NORMALE: Nuovo ordine")
            doc = Document()

        # Se doc è ancora None, crealo
        if doc is None:
            print("⚠️ Creazione documento di emergenza")
            doc = Document()

        progress.setValue(20)
        QApplication.processEvents()

       # ========== GENERA INTESTAZIONE (solo se NON aggiungi servizi) ==========
        # ========== GENERA INTESTAZIONE (solo se NON aggiungi servizi) ==========
        # ========== GENERA INTESTAZIONE (solo se NON aggiungi servizi) ==========
        if not (hasattr(self, 'modalita_aggiungi_servizi') and self.modalita_aggiungi_servizi):
            
            # Crea tabella NR PROGR per prima (andrà in centro)
            table_title = doc.add_table(rows=1, cols=1)
            table_title.autofit = False
            table_title.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_title.columns[0].width = Inches(2.0)

            cell = table_title.cell(0, 0)
            cell.width = Inches(3.0)

            run = cell.paragraphs[0].add_run("NR PROGR")
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(16)
            cell.paragraphs[0].alignment = 0

            self.apply_borders(cell)
            
            # Aggiungi leggenda subito dopo (stesso livello, a destra)
            p_leggenda = doc.add_paragraph()
            p_leggenda.alignment = 2  # Destra
            p_leggenda.paragraph_format.space_before = Pt(0)   # Nessuno spazio prima
            p_leggenda.paragraph_format.space_after = Pt(0)    # Nessuno spazio dopo
            
            run1 = p_leggenda.add_run("vassoi deluxe+decori\n* indica i caldi")
            run1.font.size = Pt(11)
            run1.font.name = "Calibri"

        progress.setValue(30)
        QApplication.processEvents()

        # Data evento
        if hasattr(self.entries["Data evento"], 'date'):
            qdate = self.entries["Data evento"].date()
            data_evento = qdate.toString("dd/MM/yyyy")
        else:
            data_evento = self.entries["Data evento"].text()

        data_obj = datetime.strptime(data_evento, "%d/%m/%Y")
        giorno_str = data_obj.strftime("%A %d %B %Y").upper()

        # Intestazione completa (solo se NON aggiungi servizi)
        if not (hasattr(self, 'modalita_aggiungi_servizi') and self.modalita_aggiungi_servizi):
            par = doc.add_paragraph()
            par.alignment = 1
            run = par.add_run(giorno_str)
            run.bold = True
            run.underline = True

            nome = self.entries["Nome"].text()
            cellulare = self.entries["Cellulare"].text()
            via = self.entries["Via"].text()
            par = doc.add_paragraph()
            par.alignment = 1
            run = par.add_run(f"All'attenzione di {nome} {cellulare}\nPresso {via}")
            run.bold = True
            run.underline = True

        progress.setValue(50)
        QApplication.processEvents()

        # ========== DETERMINA SERVIZIO ==========
        servizi_possibili = ["Coffee Break", "Tea Break", "Lunch Buffet", "Aperitivo a Buffet","Lunch Box","Servito"]
        servizio_usato = "SERVIZIO NON DEFINITO"

        for servizio_nome in servizi_possibili:
            if any(checkbox.isChecked() and key.startswith(servizio_nome.lower() + '|') 
                for key, checkbox in self.selected.items() if hasattr(checkbox, 'isChecked')):
                servizio_usato = servizio_nome
                break

        if servizio_usato == "Aperitivo a Buffet":
            tipo = self.ap_buffet_tipo.upper()
            servizio_usato = f"APERITIVO A BUFFET {tipo}"
        else:
            servizio_usato = servizio_usato.upper()

        orario_allestimento = self.entries["Orario allestimento"].text()
        orario_pronti = self.entries["Orario Pronti"].text()

        # Logica intelligente per la stringa orario
        if orario_allestimento.strip() and orario_pronti.strip():
            # Entrambi i campi compilati: mostra entrambi
            allestimento_str = f"\nP.ALLESTIMENTO ENTRO LE ORE {orario_allestimento} - {orario_pronti} PRONTI"
        elif orario_allestimento.strip() and not orario_pronti.strip():
            # Solo allestimento compilato: mostra solo allestimento
            allestimento_str = f"\nP.ALLESTIMENTO ENTRO LE ORE {orario_allestimento}"
        elif not orario_allestimento.strip() and orario_pronti.strip():
            # Solo orario pronti compilato: mostra solo orario servizio
            allestimento_str = f"\n{orario_pronti} PRONTI"
        else:
            # Nessun campo compilato: nessuna stringa
            allestimento_str = ""

        # Intestazione servizio (sempre presente)
        par = doc.add_paragraph()
        par.alignment = 1
        run = par.add_run(f"{servizio_usato} CON {persone} PERSONE{allestimento_str}")
        run.bold = True
        run.underline = True

        progress.setValue(70)
        QApplication.processEvents()

        # ========== TABELLA PRINCIPALE ==========
        table = doc.add_table(rows=1, cols=2)
        table.allow_autofit = True
        table.autofit = True
        row = table.rows[0]
        cibo_cell = row.cells[0]
        right_cell = row.cells[1]

        # Calcolo intelligente delle quantità
        servizio_lower = servizio_usato.lower()
        tipo_ap = self.ap_buffet_tipo.lower()
        num_pers = persone

        # Calcola le referenze divise per tipo
        self.referenze_primi = []
        self.referenze_secondi = []
        self.referenze_contorni = []
        self.referenze_dolci = []
        self.referenze_salato = []  # Finger food
        self.referenze_accompagnamenti = []  # NUOVA LISTA per pane, grissini, bocconcini
        accompagnamenti_target = ["pane a fette", "bocconcini alle olive", "grissini"]



        for k, checkbox in self.selected.items():
            if hasattr(checkbox, 'isChecked') and checkbox.isChecked() and self.voce_categoria.get(k) == "cibo":
                if any(menu_type.lower() in k.lower() for menu_type in ["vegano", "no lattosio", "no uova", "celiaci"]):
                    continue  
                voce_lower = self.voce_originale.get(k, "").strip().lower()
                
                print(f"🔍 CATEGORIZZANDO: k='{k}', voce='{voce_lower}'")  # QUI va il debug
                
                # Primi piatti
                if "bis di primi piatti" in k or any(x in voce_lower for x in ["lasagne", "crespelle", "cannelloni", "rosette", "nidi", "gnocchi alla parigina", "intrighi", "rotolini di tagliatelle", "tagliatelle", "intrighi di strettine"]):
                    self.referenze_primi.append(k)
                    print(f"   ➡️ PRIMI")
                
                # Dolci
                elif "dolci" in k:
                    self.referenze_dolci.append(k)
                    print(f"   ➡️ DOLCI")
                
                # Secondi piatti
                elif "secondi piatti" in k:
                    self.referenze_secondi.append(k)
                    print(f"   ➡️ SECONDI")
                
                # Contorni E insalate specifiche
                elif "contorni" in k or any(insalata in voce_lower for insalata in [
                    "insalatona verde mista", "insalata waldorf", "caprese", 
                    "insalata di pollo", "insalatina di stagione"
                ]):
                    self.referenze_contorni.append(k)
                    print(f"   ➡️ CONTORNI (o insalata forzata)")
                
                # Finger food (tutto il resto)
                
                elif any(acc in voce_lower for acc in accompagnamenti_target):
                    self.referenze_accompagnamenti.append(k)
                    print(f"   ➡️ ACCOMPAGNAMENTI (pane/grissini)")
                else:
                    self.referenze_salato.append(k)
                    print(f"   ➡️ SALATO (finger food)")

        referenze_cibo = self.referenze_primi + self.referenze_secondi + self.referenze_contorni + self.referenze_dolci + self.referenze_salato

        # 1. FINGER FOOD (normale)
        # 1. FINGER FOOD (normale)
        for key in self.referenze_salato:
            voce = self.voce_originale.get(key, key)
            voce_lower = voce.lower()  # ✅ Definisci voce_lower qui!
            
            if ("taglieri" in voce_lower or 
                "miele" in voce_lower or 
                "composta" in voce_lower or
                "grissini" in voce_lower):  # ← AGGIUNGI QUESTA RIGA
                continue
                
            if "kg torta" in voce.strip().lower():
                continue
                
            quant = self.calcola_quantita(key, referenze_cibo, num_pers, servizio_lower, tipo_ap)
            print(f"🔧 VALORE RESTITUITO da calcola_quantità per {voce}: {quant}")
            self.stampa_quantita_docx(quant, voce, cibo_cell, bold=False)

        # 2. PRIMI (grassetto)
        for key in self.referenze_primi:
            voce = self.voce_originale.get(key, key)
            if "kg torta" in voce.strip().lower():
                continue
            quant = self.calcola_quantita(key, referenze_cibo, num_pers, servizio_lower, tipo_ap)
            print(f"🔧 VALORE RESTITUITO da calcola_quantità per {voce}: {quant}")  # ⬅️ AGGIUNGI QUESTO
            self.stampa_quantita_docx(quant, voce, cibo_cell, bold=True)
        # AGGIUNGI SUBITO DOPO I PRIMI: TAGLIERI IN GRASSETTO
        taglieri_keys = []
        for key in self.selected.keys():
            if hasattr(self.selected[key], 'isChecked') and self.selected[key].isChecked():
                voce = self.voce_originale.get(key, "").lower()
                if "taglieri" in voce:
                    taglieri_keys.append(key)

        for key in taglieri_keys:
            voce = self.voce_originale.get(key, key)
            quant = self.calcola_quantita(key, referenze_cibo, num_pers, servizio_lower, tipo_ap)
            self.stampa_quantita_docx(quant, voce, cibo_cell, bold=True)  # GRASSETTO
            
            # Stampa immediatamente gli accompagnamenti di questo tagliere
            voce_lower = voce.lower()
            if "taglieri formaggi piccoli" in voce_lower:
            # Cerca miele, composta E grissini
                for acc_key in self.selected.keys():
                    if hasattr(self.selected[acc_key], 'isChecked') and self.selected[acc_key].isChecked():
                        acc_voce = self.voce_originale.get(acc_key, "").lower()
                        if "miele" in acc_voce or "composta" in acc_voce or "grissini" in acc_voce:  # ← AGGIUNGI "grissini"
                            acc_voce_originale = self.voce_originale.get(acc_key, acc_key)
                            acc_quant = self.calcola_quantita(acc_key, referenze_cibo, num_pers, servizio_lower, tipo_ap)
                            self.stampa_quantita_docx(acc_quant, acc_voce_originale, cibo_cell, bold=False)
            
            elif "taglieri di salumi medi" in voce_lower:
                # Cerca crescentine
                for acc_key in self.selected.keys():
                    if hasattr(self.selected[acc_key], 'isChecked') and self.selected[acc_key].isChecked():
                        acc_voce = self.voce_originale.get(acc_key, "").lower()
                        if "crescentine" in acc_voce:
                            acc_voce_originale = self.voce_originale.get(acc_key, acc_key)
                            acc_quant = self.calcola_quantita(acc_key, referenze_cibo, num_pers, servizio_lower, tipo_ap)
                            self.stampa_quantita_docx(acc_quant, acc_voce_originale, cibo_cell, bold=False)  # NORMALE
        # 3. SECONDI (grassetto)
        for key in self.referenze_secondi:
            voce = self.voce_originale.get(key, key)
            if "kg torta" in voce.strip().lower():
                continue
            quant = self.calcola_quantita(key, referenze_cibo, num_pers, servizio_lower, tipo_ap)
            print(f"🔧 VALORE RESTITUITO da calcola_quantità per {voce}: {quant}")  # ⬅️ AGGIUNGI QUESTO
            self.stampa_quantita_docx(quant, voce, cibo_cell, bold=True)

        # 4. CONTORNI (grassetto)
        for key in self.referenze_contorni:
            voce = self.voce_originale.get(key, key)
            if "kg torta" in voce.strip().lower():
                continue
            quant = self.calcola_quantita(key, referenze_cibo, num_pers, servizio_lower, tipo_ap)
            print(f"🔧 VALORE RESTITUITO da calcola_quantità per {voce}: {quant}")  # ⬅️ AGGIUNGI QUESTO
            self.stampa_quantita_docx(quant, voce, cibo_cell, bold=True)

        # 5. ACCOMPAGNAMENTI (normale)
        for key in self.referenze_accompagnamenti:
            voce = self.voce_originale.get(key, key)
            voce_lower = voce.lower()

            # Escludi grissini se ci sono taglieri formaggi (già gestiti nella sezione taglieri)
            if "grissini" in voce_lower:
                taglieri_formaggi_selezionati = False
                for k, checkbox in self.selected.items():
                    if hasattr(checkbox, 'isChecked') and checkbox.isChecked():
                        voce_k = self.voce_originale.get(k, "").lower()
                        if "taglieri formaggi piccoli" in voce_k:
                            taglieri_formaggi_selezionati = True
                            break
                
                if taglieri_formaggi_selezionati:
                    continue  # Salta completamente i grissini
            if "kg torta" in voce.strip().lower():
                continue
            quant = self.calcola_quantita(key, referenze_cibo, num_pers, servizio_lower, tipo_ap)
            print(f"🔧 VALORE RESTITUITO da calcola_quantità per {voce}: {quant}")
            self.stampa_quantita_docx(quant, voce, cibo_cell, bold=False)

        # 6. DOLCI (normale)
        for key in self.referenze_dolci:
            voce = self.voce_originale.get(key, key)
            if "kg torta" in voce.strip().lower():
                continue
            quant = self.calcola_quantita(key, referenze_cibo, num_pers, servizio_lower, tipo_ap)
            print(f"🔧 VALORE RESTITUITO da calcola_quantità per {voce}: {quant}")  # ⬅️ AGGIUNGI QUESTO
            self.stampa_quantita_docx(quant, voce, cibo_cell, bold=False)

        # 7. TORTA (grassetto) — ULTIMA VOCE
        for key in referenze_cibo:
            voce = self.voce_originale.get(key, key)
            if "kg torta" in voce.strip().lower():
                quant = self.calcola_quantita(key, referenze_cibo, num_pers, servizio_lower, tipo_ap)
                print(f"🔧 VALORE RESTITUITO da calcola_quantità per {voce}: {quant}")  # ⬅️ AGGIUNGI QUESTO
                self.stampa_quantita_docx(quant, voce, cibo_cell, bold=True)
                break


        progress.setValue(85)
        QApplication.processEvents()
        # ========== BEVANDE E ACCESSORI CON FIX PER AGGIUNGI SERVIZI ==========
        def add_inner_table_to_cell(testo_titolo, categoria, parent_cell):
            if categoria != "bevande" and categoria != "accessori":
                return

            print(f"\n🔧 Generando sezione: {testo_titolo} (categoria: {categoria})")
            
            # Trova tutte le voci per questa categoria
            voci_categoria = []
            
            # ========== FIX PER MODALITÀ AGGIUNGI SERVIZI ==========
            # Se modalità aggiungi servizi, includi anche scelte esistenti dall'ordine precedente
            print(f"🔍 CERCANDO VOCI PER CATEGORIA: {categoria}")
            for k, checkbox in self.selected.items():
                if hasattr(checkbox, 'isChecked') and checkbox.isChecked():
                    voce_cat = self.voce_categoria.get(k)
                    voce_nome = self.voce_originale.get(k, k)
                    
                    if voce_cat == categoria:
                        if voce_nome not in voci_categoria:  # Evita duplicati
                            voci_categoria.append(voce_nome)
                            print(f"    ✅ Nuova: {voce_nome}")

            # ========== AGGIUNGI DEBUG ==========
            if hasattr(self, 'modalita_aggiungi_servizi') and self.modalita_aggiungi_servizi:
                print(f"    🎯 MODALITÀ AGGIUNGI SERVIZI: Mostrando SOLO nuove {categoria}")
                print(f"    📋 Trovate {len(voci_categoria)} nuove voci per {categoria}")
            
            # Crea la sezione nel documento
            inner_table = parent_cell.add_table(rows=1, cols=1)
            inner_cell = inner_table.cell(0, 0)
            p_title = inner_cell.paragraphs[0]
            run = p_title.add_run(testo_titolo.upper())
            run.bold = True
            run.underline = True
            
            # Aggiungi tutte le voci
            for voce in voci_categoria:
                inner_cell.add_paragraph(f"• {voce}")
            
            self.apply_borders(inner_cell)
            print(f"    ✅ Sezione {testo_titolo} creata con {len(voci_categoria)} voci")
            
        def add_inner_table_to_cell_menu_extra(testo_titolo, menu_type, parent_cell):
            """Crea riquadro separato per menu extra"""
            print(f"\n🌱 Generando riquadro: {testo_titolo}")
            
            # Trova le voci selezionate per questo menu extra
            voci_menu = []
            for k, checkbox in self.selected.items():
                if hasattr(checkbox, 'isChecked') and checkbox.isChecked():
                    print(f"   🔍 Checking key: {k} (checked: {checkbox.isChecked()})")
                    if menu_type.lower() in k.lower():
                        print(f"   ✅ MATCH trovato per {k}")
                        voce = self.voce_originale.get(k, k)
                        persone = self.menu_extra_widgets[menu_type].value()
                        quantita_box = self.calcola_quantita_menu_extra_simple(voce, persone)
                        # NUOVO FORMATO: (totale)pezzi_per_box pz Nome
                        voci_menu.append(f"{quantita_box} pz {voce}")
                        print(f"    ✅ {quantita_box} pz {voce}")
            
            print(f"🎯 Totale voci trovate: {len(voci_menu)}")
            
            if voci_menu:
                # Crea la tabella interna
                inner_table = parent_cell.add_table(rows=1, cols=1)
                inner_cell = inner_table.cell(0, 0)
                
                # Titolo del menu
                p_title = inner_cell.paragraphs[0]
                run = p_title.add_run(testo_titolo)
                run.bold = True
                run.underline = True
                
                # Aggiungi tutte le voci con quantità
                for voce_con_quantita in voci_menu:
                    inner_cell.add_paragraph(f"• {voce_con_quantita}")
                
                self.apply_borders(inner_cell)
                print(f"    ✅ Riquadro {testo_titolo} creato con {len(voci_menu)} voci")
        # ========== CONTROLLA SE CI SONO BEVANDE (esistenti + nuove) ==========
        has_bevande = False
        
        # In modalità aggiungi servizi, controlla anche bevande esistenti
        if hasattr(self, 'modalita_aggiungi_servizi') and self.modalita_aggiungi_servizi and hasattr(self, 'ordine_da_modificare') and self.ordine_da_modificare:
            scelte_esistenti = self.ordine_da_modificare.get("scelte", {})
            has_bevande = any(v and self.voce_categoria.get(k) == "bevande" for k, v in scelte_esistenti.items())
        
        # Controlla bevande nuove
        if not has_bevande:
            has_bevande = any(checkbox.isChecked() and self.voce_categoria.get(k) == "bevande" 
                            for k, checkbox in self.selected.items() if hasattr(checkbox, 'isChecked'))

        # ========== CONTROLLA SE CI SONO ACCESSORI (esistenti + nuovi) ==========
        has_accessori = False
        
        # In modalità aggiungi servizi, controlla anche accessori esistenti
        if hasattr(self, 'modalita_aggiungi_servizi') and self.modalita_aggiungi_servizi and hasattr(self, 'ordine_da_modificare') and self.ordine_da_modificare:
            scelte_esistenti = self.ordine_da_modificare.get("scelte", {})
            has_accessori = any(v and self.voce_categoria.get(k) == "accessori" for k, v in scelte_esistenti.items())
        
        # Controlla accessori nuovi
        if not has_accessori:
            has_accessori = any(checkbox.isChecked() and self.voce_categoria.get(k) == "accessori" 
                              for k, checkbox in self.selected.items() if hasattr(checkbox, 'isChecked'))

        # ========== CREA LE SEZIONI SE NECESSARIO ==========
        if has_bevande:
            add_inner_table_to_cell("Bevande", "bevande", right_cell)

        if has_accessori:
            add_inner_table_to_cell("Sala", "accessori", right_cell)
        # ========== SEZIONE INTOLLERANTI SEPARATA ==========
        menu_extra_attivi = {
            menu_type: widget.value() 
            for menu_type, widget in self.menu_extra_widgets.items() 
            if widget.value() > 0
        }

        if menu_extra_attivi:
            # Aggiungi interruzione di pagina o spazio
            doc.add_page_break()  # Oppure doc.add_paragraph() per stesso foglio
            
            # ========== NUOVA INTESTAZIONE COMPLETA ==========
            # Tabella NR PROGR
            table_title_intol = doc.add_table(rows=1, cols=1)
            table_title_intol.autofit = False
            table_title_intol.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_title_intol.columns[0].width = Inches(2.0)

            cell_intol = table_title_intol.cell(0, 0)
            cell_intol.width = Inches(3.0)

            run_intol = cell_intol.paragraphs[0].add_run("NR PROGR")
            run_intol.bold = True
            run_intol.font.name = "Calibri"
            run_intol.font.size = Pt(16)
            cell_intol.paragraphs[0].alignment = 1

            self.apply_borders(cell_intol)
            
            # Leggenda
            p_leggenda_intol = doc.add_paragraph()
            p_leggenda_intol.alignment = 2  # Destra
            p_leggenda_intol.paragraph_format.space_before = Pt(0)
            p_leggenda_intol.paragraph_format.space_after = Pt(0)
            
            run1_intol = p_leggenda_intol.add_run("vassoi deluxe+decori\n* indica i caldi")
            run1_intol.font.size = Pt(11)
            run1_intol.font.name = "Calibri"

            # Data
            par_data_intol = doc.add_paragraph()
            par_data_intol.alignment = 1
            run_data_intol = par_data_intol.add_run(giorno_str)
            run_data_intol.bold = True
            run_data_intol.underline = True

            # Cliente
            nome = self.entries["Nome"].text()
            cellulare = self.entries["Cellulare"].text()
            via = self.entries["Via"].text()
            par_cliente_intol = doc.add_paragraph()
            par_cliente_intol.alignment = 1
            run_cliente_intol = par_cliente_intol.add_run(f"All'attenzione di {nome} {cellulare}\nPresso {via}")
            run_cliente_intol.bold = True
            run_cliente_intol.underline = True

            # Intestazione servizio INTOLLERANTI
            par_servizio_intol = doc.add_paragraph()
            par_servizio_intol.alignment = 1
            run_servizio_intol = par_servizio_intol.add_run("INTOLLERANTI")
            run_servizio_intol.bold = True
            run_servizio_intol.underline = True

            # ========== TABELLA INTOLLERANTI ==========
            # ========== TABELLA INTOLLERANTI A GRIGLIA (2 COLONNE) ==========
            menu_types = list(menu_extra_attivi.keys())
            num_menus = len(menu_types)

            # Calcola righe necessarie (2 colonne per riga)
            righe_necessarie = (num_menus + 1) // 2  # Arrotonda per eccesso

            # Crea tabella con numero di righe necessario e 2 colonne
            table_intol = doc.add_table(rows=righe_necessarie, cols=2)
            table_intol.allow_autofit = True
            table_intol.autofit = True

            # Imposta larghezza colonne
            for col in table_intol.columns:
                col.width = Inches(3.5)  # Larghezza uguale per entrambe le colonne

            # Riempie la tabella 2x2
            for i, menu_type in enumerate(menu_types):
                row_index = i // 2  # Riga (0, 0, 1, 1, 2, 2...)
                col_index = i % 2   # Colonna (0, 1, 0, 1, 0, 1...)
                
                persone = menu_extra_attivi[menu_type]
                cell = table_intol.cell(row_index, col_index)
                
                # ========== TITOLO DEL RIQUADRO ==========
                p_title = cell.paragraphs[0]
                persone = menu_extra_attivi[menu_type]
                # Crea il titolo nel formato richiesto
                titolo_formattato = f"MENU X {persone} {menu_type.upper()} BOX"
                run_title = p_title.add_run(titolo_formattato)

                run_title.bold = True
                run_title.underline = True
                p_title.alignment = 1  # Centro
            
                # ========== VOCI DEL MENU ==========
                voci_trovate = []
                for key, checkbox in self.selected.items():
                    if (hasattr(checkbox, 'isChecked') and checkbox.isChecked() and 
                        menu_type.lower() in key.lower()):
                        voce = self.voce_originale.get(key, key.split('|')[-1])
                        servizio_principale = self.determina_servizio_principale()
                        quantita = self.calcola_quantita_menu_extra_simple(voce, persone, servizio_principale)
                        voci_trovate.append(f"{quantita} pz {voce}")
                
                # Aggiungi le voci al riquadro
                for voce_completa in voci_trovate:
                    p_voce = cell.add_paragraph()
                    run_voce = p_voce.add_run(voce_completa)
                    run_voce.font.size = Pt(11)
                
                # ========== BORDI PER IL RIQUADRO ==========
                self.apply_borders(cell)
                
                # Debug
                print(f"🌱 Riquadro {menu_type}: {len(voci_trovate)} voci in posizione [{row_index}, {col_index}]")

            # Se numero dispari, riempi l'ultima cella vuota
            if num_menus % 2 == 1:
                ultima_riga = righe_necessarie - 1
                cell_vuota = table_intol.cell(ultima_riga, 1)
                # Lascia vuota o aggiungi un messaggio
                p_vuota = cell_vuota.paragraphs[0]
                p_vuota.add_run("")  # Cella vuota
        progress.setValue(90)
        QApplication.processEvents()


        # ========== DETERMINA PATH E SALVA ==========
        if not final_path:
            # Genera nuovo path nella cartella corretta
            nome_file = self.entries["Nome"].text().strip().replace(" ", "_")
            data_file = data_evento.replace("/", "-")
            
            # ========== IMPORTANTE: USA ordini_docx CON SOTTOCARTELLE PER MESE ==========
            mese_cartella = data_obj.strftime("%Y-%m")  # Es: "2025-01"
            cartella_output = os.path.join(gestore_percorsi.get_percorso("ordini_docx"), mese_cartella)
            os.makedirs(cartella_output, exist_ok=True)
            
            final_path = os.path.join(cartella_output, f"{nome_file}_{data_file}.docx")
            
            # Evita duplicati solo per nuovi ordini
            if not (hasattr(self, 'modalita_modifica') and self.modalita_modifica):
                base, ext = os.path.splitext(final_path)
                counter = 1
                while os.path.exists(final_path):
                    final_path = os.path.join(cartella_output, f"{nome_file}_{data_file}_{counter}{ext}")
                    counter += 1

        # Salva il documento
        # ========== AGGIUNGI LEGGENDA E RIQUADRO NOTE ==========
        # (solo per nuovi documenti, non per aggiungi servizi)
        self.aggiungi_riquadro_note(doc)
        
        
        doc.save(final_path)
        print(f"📄 Documento salvato: {final_path}")

        # ========== SALVA ORDINE NEL JSON ==========
         # ========== SALVA ORDINE NEL JSON ==========
        # ========== SALVA ORDINE NEL JSON ==========
        # ========== AGGIUNGI INTOLLERANZE AL JSON ==========
        intolleranze = {}
        if hasattr(self, 'menu_extra_widgets'):
            for menu_type, widget in self.menu_extra_widgets.items():
                if widget.value() > 0:
                    persone = widget.value()
                    voci_intolleranza = []
                    
                    # Trova le voci selezionate per questa intolleranza
                    for key, checkbox in self.selected.items():
                        if (hasattr(checkbox, 'isChecked') and checkbox.isChecked() and 
                            menu_type.lower() in key.lower()):
                            voce = self.voce_originale.get(key, key.split('|')[-1])
                            quantita_raw = self.calcola_quantita_menu_extra_simple(voce, persone)

                            # Estrai solo il numero totale (quello tra parentesi) per il monitor
                            import re
                            if isinstance(quantita_raw, str) and '(' in quantita_raw:
                                # Formato (12)7 → estrai 12
                                match = re.search(r'\((\d+)\)', str(quantita_raw))
                                quantita_finale = int(match.group(1)) if match else quantita_raw
                            else:
                                quantita_finale = quantita_raw

                            voci_intolleranza.append({
                                'nome': voce,
                                'quantita': f"{quantita_finale} pz"  # Solo il totale
                            })

                            print(f"🌱 FIXED: {voce} → {quantita_raw} → {quantita_finale} pz")
                    
                    if voci_intolleranza:
                        intolleranze[menu_type] = {
                            'persone': persone,
                            'voci': voci_intolleranza
                        }

        print(f"🌱 Intolleranze da salvare: {intolleranze}")

        # Crea l'ordine CON le intolleranze
        ordine = {
            "campi": {k: (v.date().toString("dd/MM/yyyy") if hasattr(v, 'date') else v.text()) 
                    for k, v in self.entries.items()},
            "scelte": {k: checkbox.isChecked() for k, checkbox in self.selected.items() 
                    if hasattr(checkbox, 'isChecked')},
            "tipo_ap": self.ap_buffet_tipo,
            "evento_in_fiera": getattr(self, 'evento_in_fiera', False),
            "path_docx": final_path,
            "intolleranze": intolleranze  # ← AGGIUNGI QUESTA RIGA
        }

        # AGGIUNGI QUESTA RIGA DI DEBUG:
        print("🎯 === ORDINE CREATO, DEVO CHIAMARE SALVA_ORDINE ===")

        # POI AGGIUNGI QUI LA CHIAMATA:
        if hasattr(self, 'modalita_aggiungi_servizi') and self.modalita_aggiungi_servizi:
            print("💾 Aggiornando ordine esistente con nuovi servizi")
            self.aggiorna_ordine_esistente(ordine)
        else:
            print("💾 Salvando ordine...")
            self.salva_ordine(ordine)

        # ========== SALVA PRIMA DI RESETTARE LE MODALITÀ ==========
        nome = self.entries["Nome"].text()
        if hasattr(self, 'modalita_aggiungi_servizi') and self.modalita_aggiungi_servizi:
            QMessageBox.information(self, "Servizi Aggiunti", 
                                f"Servizi aggiunti con successo!\n\n"
                                f"Cliente: {nome}\n"
                                f"Nuovo servizio: {servizio_usato}\n"
                                f"File aggiornato: {os.path.basename(final_path)}")
        else:
            QMessageBox.information(self, "Ordine Generato", 
                                f"Ordine generato con successo!\n\n"
                                f"Cliente: {nome}\n"
                                f"File: {os.path.basename(final_path)}")

        # ========== RESET MODALITÀ ==========
        if hasattr(self, 'modalita_modifica'):
            print(f"🔄 Resettando modalita_modifica da {self.modalita_modifica} a False")
            self.modalita_modifica = False
        if hasattr(self, 'modalita_aggiungi_servizi'):
            self.modalita_aggiungi_servizi = False
        if hasattr(self, 'ordine_da_modificare'):
            self.ordine_da_modificare = None

        progress.setValue(100)
        QApplication.processEvents()
        progress.hide()

        print("✅ Generazione completata e modalità resettate")

    # ========== AGGIUNGI QUESTO METODO SEPARATAMENTE ALLA CLASSE (NON DENTRO generate_doc) ==========
    def debug_voce_categoria(self):
        """Debug per verificare la categorizzazione delle voci"""
        print("\n🔍 === DEBUG CATEGORIZZAZIONE VOCI ===")
        
        accessori_previsti = [
            "tavoli da buffet", "tovagliato cotone + coprimacche", "nr.    cam", 
            "acc.bio", "pinze", "guanti neri", "alzate", "cucch. e forch. bio x finger"
        ]
        
        print("📋 Accessori che dovrebbero essere in SALA:")
        for acc in accessori_previsti:
            found = False
            for key, categoria in self.voce_categoria.items():
                voce = self.voce_originale.get(key, key).lower()
                if acc.lower() in voce:
                    print(f"  • {acc} → Categoria: {categoria} (Key: {key[:50]}...)")
                    found = True
                    break
            if not found:
                print(f"  ❌ {acc} → NON TROVATO!")
        
        print("\n📊 Tutte le voci categorizzate come 'accessori':")
        for key, categoria in self.voce_categoria.items():
            if categoria == "accessori":
                voce = self.voce_originale.get(key, key)
                print(f"  ✅ {voce} (Key: {key[:50]}...)")
        
        print("\n" + "="*50)
    
    def aggiungi_al_documento_esistente(self):
        """Aggiungi servizio al documento esistente - FUNZIONE DEDICATA"""
        
        # Validazione input
        try:
            persone = int(self.entries["Numero persone"].text())
            if persone <= 0:
                raise ValueError("Il numero di persone deve essere positivo")
        except ValueError:
            QMessageBox.warning(self, "Errore Input", "Inserire un numero valido per le persone.")
            return

        if not any(checkbox.isChecked() for checkbox in self.selected.values() if hasattr(checkbox, 'isChecked')):
            QMessageBox.warning(self, "Errore Input", "Selezionare almeno un servizio da aggiungere.")
            return

        # Conferma aggiunta
        nome_cliente = self.entries["Nome"].text()
        reply = QMessageBox.question(
            self, "Conferma Aggiunta Servizio", 
            f"Vuoi aggiungere il nuovo servizio all'ordine di:\n\n"
            f"Cliente: {nome_cliente}\n\n"
            f"Il nuovo servizio verrà aggiunto al documento Word esistente.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply != QMessageBox.StandardButton.Yes:
            return

        # Usa la funzione generate_doc normale, che gestirà la modalità
        self.generate_doc()
        
        # Reset interfaccia
        self.reset_modalita_aggiungi_servizi()
    def aggiungi_leggenda_dopo_intestazione(self, doc):
        """Aggiunge la leggenda subito dopo l'intestazione NR PROGR"""
        
        # Trova l'ultima tabella (dovrebbe essere NR PROGR)
        tables = doc.tables
        if not tables:
            return
        
        # Aggiungi paragrafo per la leggenda dopo l'ultima tabella
        p = doc.add_paragraph()
        p.alignment = 2  # Allinea a destra
        
        # RIDUCI SPAZIO PRIMA E DOPO
        p.paragraph_format.space_before = Pt(0)  # Nessuno spazio prima
        p.paragraph_format.space_after = Pt(6)   # Poco spazio dopo
        
        # Primo run: "deluxe+decori"
        run1 = p.add_run("vassoi deluxe+decori")
        run1.bold = True
        run1.font.size = Pt(11)
        run1.font.name = "Calibri"
        
        # Vai a capo
        p.add_run("\n")
        
        # Secondo run: "* indica i caldi"
        run2 = p.add_run("* indica i caldi")
        run2.font.size = Pt(9)
        run2.font.name = "Calibri"
        run2.italic = True
        
        # NON aggiungere spazio extra dopo
    def reset_modalita_aggiungi_servizi(self):
        """Reset interfaccia dopo aggiunta servizio"""
        
        # Reset modalità
        if hasattr(self, 'modalita_aggiungi_servizi'):
            self.modalita_aggiungi_servizi = False
        if hasattr(self, 'ordine_da_modificare'):
            self.ordine_da_modificare = None
        
        # Ripristina pulsanti
        if hasattr(self, 'aggiungi_ordine_btn') and hasattr(self, 'genera_btn'):
            self.aggiungi_ordine_btn.hide()
            self.genera_btn.show()
        
        # Reset titolo finestra
        self.setWindowTitle("Gestione Ordini Catering - Versione Professionale")
        
        print("✅ Interfaccia resettata alla modalità normale")

    def set_ap_tipo(self, tipo):
        self.ap_buffet_tipo = tipo
        print(f"🎯 Tipo aperitivo impostato: '{tipo}'")  # Debug
    # ========== FUNZIONI HELPER (da aggiungere alla classe) ==========
    
    # SOLUZIONE PROBLEMA 2: Correzione quantità per Coffee/Tea Break e Aperitivo Rinforzato

    def calcola_quantita_menu_extra_simple(self, voce, persone, servizio_principale=None):
        """Calcolo quantità per menu extra - VERSIONE CORRETTA CON QUANTITÀ GIUSTE"""
        
        print(f"\n🌱 === CALCOLO MENU EXTRA SEMPLICE ===")
        print(f"   📝 Voce: {voce}")
        print(f"   👥 Persone: {persone}")
        
        voce_lower = voce.lower()
        
        # ========== GESTIONE PRIMI ==========
        is_primo = any(parola in voce_lower for parola in ["pasta corta", "pasta al forno"])
        if is_primo:
            print(f"   🍝 PRIMO: 1 per persona")
            return f"({persone})1"
        
        # ========== GESTIONE DOLCI ==========
        # ========== GESTIONE DOLCI ==========
        is_dolce = any(parola in voce_lower for parola in ["dolc", "torta", "crostat", "tiramisu", "spicchietti", "frutta","crema","tenerina","biscotteria","sacher","zuppa inglese","pistacchio+cucch bio",])
        if is_dolce:
            # USA IL SERVIZIO PASSATO COME PARAMETRO
            if servizio_principale:
                servizio_lower = servizio_principale.lower()
                print(f"   🍽️ Servizio per dolci: {servizio_principale}")
            else:
                # Fallback alla logica originale
                servizio = self.entries.get("Servizio", "").strip() if "Servizio" in self.entries else ""
                if not servizio:
                    tipo_ap = getattr(self, 'ap_buffet_tipo', '').lower()
                    if "leggero" in tipo_ap:
                        servizio = "Aperitivo leggero"
                    elif "rinforzato" in tipo_ap:
                        servizio = "Aperitivo rinforzato"
                    else:
                        servizio = "Lunch Buffet"
                servizio_lower = servizio.lower()
                print(f"   🍽️ Servizio fallback per dolci: {servizio}")
            
            # LOGICA SPECIALE PER APERITIVO RINFORZATO
            if "aperitivo" in servizio_lower and "rinforzato" in servizio_lower:
                dolci_per_persona = 3
                dolci_totali = dolci_per_persona * persone
                print(f"   🍰 APERITIVO RINFORZATO - DOLCI: {dolci_per_persona} per persona")
                return f"({dolci_totali}){dolci_per_persona}"
            else:
                print(f"   🍰 DOLCE STANDARD: 1 per persona")
                return f"({persone})1"

        
        # ========== DETERMINA SERVIZIO ==========
        # ========== DETERMINA SERVIZIO ==========
        if servizio_principale:
            servizio_lower = servizio_principale.lower()
            print(f"   🍽️ Servizio da parametro: {servizio_principale}")
        else:
            # Fallback alla logica originale
            servizio = self.entries.get("Servizio", "").strip() if "Servizio" in self.entries else ""
            print(f"   🔍 DEBUG Servizio field: '{servizio}'")
            
            if not servizio:
                tipo_ap = getattr(self, 'ap_buffet_tipo', '').lower()
                print(f"   🔍 DEBUG tipo_ap: '{tipo_ap}'")
                
                if "leggero" in tipo_ap:
                    servizio = "Aperitivo leggero"
                elif "rinforzato" in tipo_ap:
                    servizio = "Aperitivo rinforzato"
                else:
                    servizio = "Lunch Buffet"
            
            servizio_lower = servizio.lower()
            print(f"   🍽️ Servizio finale fallback: {servizio}")
        
        # ========== TROVA MENU TYPE PRIMA DI TUTTO ==========
        menu_type_trovato = None
        for tipo in ["vegano", "no lattosio", "no uova", "celiaci"]:
            for key in self.selected.keys():
                if (tipo.lower() in key.lower() and 
                    voce_lower in self.voce_originale.get(key, key.split('|')[-1]).lower()):
                    menu_type_trovato = tipo
                    break
            if menu_type_trovato:
                break
        
        if not menu_type_trovato:
            print(f"   ⚠️ Menu type non trovato")
            return f"({persone})1"
        
        print(f"   📋 Menu type: {menu_type_trovato}")
        
        # ========== VERIFICA PRIMI NEL MENU INTOLLERANTI (NON MENU PRINCIPALE) ==========
        ha_primi_menu_intolleranti = False

        # Controlla se ci sono primi SPECIFICAMENTE nel menu intolleranti
        for key, checkbox in self.selected.items():
            if (hasattr(checkbox, 'isChecked') and checkbox.isChecked() and 
                menu_type_trovato.lower() in key.lower()):
                voce_check = self.voce_originale.get(key, key.split('|')[-1])
                voce_check_lower = voce_check.lower()
                
                # Verifica se è un primo nel menu intolleranti
                if any(parola in voce_check_lower for parola in ["pasta corta", "pasta al forno"]):
                    ha_primi_menu_intolleranti = True
                    print(f"      ✅ PRIMO nel menu intolleranti: {voce_check}")
                    break

        print(f"   🍝 Ha primi nel menu INTOLLERANTI: {ha_primi_menu_intolleranti}")

        # FORZA LOGICA CORRETTA: se non ci sono primi nel menu intolleranti, usa 8 finger food
        ha_primi_menu = ha_primi_menu_intolleranti
        
        # ========== CONTA VOCI SALATE ==========
        voci_salate_count = 0
        for key, checkbox in self.selected.items():
            if (hasattr(checkbox, 'isChecked') and checkbox.isChecked() and 
                menu_type_trovato.lower() in key.lower()):
                voce_check = self.voce_originale.get(key, key.split('|')[-1])
                voce_check_lower = voce_check.lower()
                
                # Escludi dolci e primi
                is_dolce_check = any(parola in voce_check_lower for parola in ["dolc", "torta", "crostat", "tiramisu", "mousse", "spicchietti", "frutta"])
                is_primo_check = any(parola in voce_check_lower for parola in ["pasta corta", "pasta al forno"])
                
                if not is_dolce_check and not is_primo_check:
                    voci_salate_count += 1
                    print(f"      ✅ Voce salata: {voce_check}")
        
        print(f"   📊 Totale voci salate: {voci_salate_count}")
        
        if voci_salate_count == 0:
            return f"({persone})1"
        
        # ========== CALCOLA FINGER FOOD TOTALI PER PERSONA - CORREZIONE ==========
        finger_food_per_persona = 0
        
        if "coffee" in servizio_lower or "tea" in servizio_lower:
            # CORREZIONE: Coffee/Tea break = 2 finger food salati per persona (non 4)
            finger_food_per_persona = 2
            print(f"   ☕ COFFEE/TEA: 2 finger food salati per persona (CORRETTO)")
            
        elif "aperitivo" in servizio_lower and "leggero" in servizio_lower:
            finger_food_per_persona = 6
            print(f"   🥂 APERITIVO LEGGERO: 6 finger food per persona")
            
        elif "aperitivo" in servizio_lower and "rinforzato" in servizio_lower:
            # CORREZIONE: Aperitivo rinforzato = 8 finger food salati per persona
            finger_food_per_persona = 8
            print(f"   🥂 APERITIVO RINFORZATO: 8 finger food salati per persona (CORRETTO)")
            
        elif "lunch" in servizio_lower or "pranzo" in servizio_lower:
            if ha_primi_menu:
                finger_food_per_persona = 4
                print(f"   🍽️ LUNCH CON PRIMI: 4 finger food per persona")
            else:
                finger_food_per_persona = 8
                print(f"   🍽️ LUNCH SENZA PRIMI: 8 finger food per persona")
        else:
            # Default: considera come lunch
            if ha_primi_menu:
                finger_food_per_persona = 2
                print(f"   🍽️ DEFAULT CON PRIMI: 4 finger food per persona")
            else:
                finger_food_per_persona = 8
                print(f"   🍽️ DEFAULT SENZA PRIMI: 8 finger food per persona")
                
        # ========== OVERRIDE FINALE: Se non è né coffee né aperitivo, forza logica lunch ==========
        if not ("coffee" in servizio_lower or "tea" in servizio_lower or "aperitivo" in servizio_lower):
            if ha_primi_menu:
                finger_food_per_persona = 4
                print(f"   🔧 OVERRIDE FINALE: 4 finger food (con primi)")
            else:
                finger_food_per_persona = 8
                print(f"   🔧 OVERRIDE FINALE: 8 finger food (senza primi)")
        
        # ========== DISTRIBUZIONE SEMPLICE ==========
        if voci_salate_count == 1:
            pezzi_per_persona = finger_food_per_persona
            
        elif voci_salate_count == 2:
            # Per 2 voci: sempre divisione equa
            pezzi_per_persona = finger_food_per_persona // 2
            print(f"   📊 2 voci: {finger_food_per_persona} ÷ 2 = {pezzi_per_persona} per voce")
            
        elif voci_salate_count == 3:
            # Per 3 voci: distribuzione corretta per arrivare al totale esatto
            if finger_food_per_persona == 8:
                # 8 finger food: 3 + 3 + 2 = 8 (CORRETTO)
                voci_ordinate = []
                for key, checkbox in self.selected.items():
                    if (hasattr(checkbox, 'isChecked') and checkbox.isChecked() and 
                        menu_type_trovato.lower() in key.lower()):
                        voce_check = self.voce_originale.get(key, key.split('|')[-1])
                        voce_check_lower = voce_check.lower()
                        
                        is_dolce_check = any(parola in voce_check_lower for parola in ["dolc", "torta", "crostat", "tiramisu", "mousse", "spicchietti", "frutta"])
                        is_primo_check = any(parola in voce_check_lower for parola in ["pasta corta", "pasta al forno", "rosette", "gnocchi", "riso", "cous", "quinoa", "farro"])
                        
                        if not is_dolce_check and not is_primo_check:
                            voci_ordinate.append(voce_check)
                
                voci_ordinate.sort()
                print(f"   🔍 DEBUG 3 VOCI (8 finger): voci ordinate = {voci_ordinate}")
                print(f"   🔍 DEBUG: voce corrente = '{voce}'")
                
                if voce in voci_ordinate:
                    indice = voci_ordinate.index(voce)
                    print(f"   🔍 DEBUG: indice voce = {indice}")
                    if indice == 0:
                        pezzi_per_persona = 3  # Prima voce: 3
                        print(f"   ✅ Prima voce: 3 pezzi")
                    elif indice == 1:
                        pezzi_per_persona = 3  # Seconda voce: 3
                        print(f"   ✅ Seconda voce: 3 pezzi")
                    else:
                        pezzi_per_persona = 2  # Terza voce: 2
                        print(f"   ✅ Terza voce: 2 pezzi")
                else:
                    pezzi_per_persona = 3
                    print(f"   ⚠️ Voce non trovata in lista, default 3")
            elif finger_food_per_persona == 6:
                # 6 finger food: 2 + 2 + 2
                pezzi_per_persona = 2
            elif finger_food_per_persona == 4:
                # 4 finger food: 2 + 1 + 1
                voci_ordinate = []
                for key, checkbox in self.selected.items():
                    if (hasattr(checkbox, 'isChecked') and checkbox.isChecked() and 
                        menu_type_trovato.lower() in key.lower()):
                        voce_check = self.voce_originale.get(key, key.split('|')[-1])
                        voce_check_lower = voce_check.lower()
                        
                        is_dolce_check = any(parola in voce_check_lower for parola in ["dolc", "torta", "crostat", "tiramisu", "mousse", "spicchietti", "frutta"])
                        is_primo_check = any(parola in voce_check_lower for parola in ["pasta corta", "pasta al forno", "rosette", "gnocchi", "riso", "cous", "quinoa", "farro"])
                        
                        if not is_dolce_check and not is_primo_check:
                            voci_ordinate.append(voce_check)
                
                voci_ordinate.sort()
                
                if voce in voci_ordinate:
                    indice = voci_ordinate.index(voce)
                    if indice == 0:
                        pezzi_per_persona = 2  # Prima voce: 2
                    else:
                        pezzi_per_persona = 1  # Altre voci: 1
                else:
                    pezzi_per_persona = 1
            elif finger_food_per_persona == 2:
                # NUOVO: 2 finger food totali per 3 voci (coffee/tea con 3 voci salate) - NON POSSIBILE
                # In questo caso specifico, distribuisci come 1 + 1 + 0, ma almeno 1 per tutti
                voci_ordinate = []
                for key, checkbox in self.selected.items():
                    if (hasattr(checkbox, 'isChecked') and checkbox.isChecked() and 
                        menu_type_trovato.lower() in key.lower()):
                        voce_check = self.voce_originale.get(key, key.split('|')[-1])
                        voce_check_lower = voce_check.lower()
                        
                        is_dolce_check = any(parola in voce_check_lower for parola in ["dolc", "torta", "crostat", "tiramisu", "mousse", "spicchietti", "frutta"])
                        is_primo_check = any(parola in voce_check_lower for parola in ["pasta corta", "pasta al forno", "rosette", "gnocchi", "riso", "cous", "quinoa", "farro"])
                        
                        if not is_dolce_check and not is_primo_check:
                            voci_ordinate.append(voce_check)
                
                voci_ordinate.sort()
                
                if voce in voci_ordinate:
                    indice = voci_ordinate.index(voce)
                    if indice <= 1:
                        pezzi_per_persona = 1  # Prime 2 voci: 1 ciascuna
                    else:
                        pezzi_per_persona = 1  # Terza voce: 1 (minimo garantito)
                else:
                    pezzi_per_persona = 1
            else:
                # Altri casi: distribuzione standard
                base = finger_food_per_persona // 3
                resto = finger_food_per_persona % 3
                
                voci_ordinate = []
                for key, checkbox in self.selected.items():
                    if (hasattr(checkbox, 'isChecked') and checkbox.isChecked() and 
                        menu_type_trovato.lower() in key.lower()):
                        voce_check = self.voce_originale.get(key, key.split('|')[-1])
                        voce_check_lower = voce_check.lower()
                        
                        is_dolce_check = any(parola in voce_check_lower for parola in ["dolc", "torta", "crostat", "tiramisu", "mousse", "spicchietti", "frutta"])
                        is_primo_check = any(parola in voce_check_lower for parola in ["pasta corta", "pasta al forno", "rosette", "gnocchi", "riso", "cous", "quinoa", "farro"])
                        
                        if not is_dolce_check and not is_primo_check:
                            voci_ordinate.append(voce_check)
                
                voci_ordinate.sort()
                
                if voce in voci_ordinate:
                    indice = voci_ordinate.index(voce)
                    pezzi_per_persona = base + (1 if indice < resto else 0)
                else:
                    pezzi_per_persona = base
                
        elif voci_salate_count == 4:
            base = finger_food_per_persona // 4
            resto = finger_food_per_persona % 4
            
            voci_ordinate = []
            for key, checkbox in self.selected.items():
                if (hasattr(checkbox, 'isChecked') and checkbox.isChecked() and 
                    menu_type_trovato.lower() in key.lower()):
                    voce_check = self.voce_originale.get(key, key.split('|')[-1])
                    voce_check_lower = voce_check.lower()
                    
                    is_dolce_check = any(parola in voce_check_lower for parola in ["dolc", "torta", "crostat", "tiramisu", "mousse", "spicchietti", "frutta"])
                    is_primo_check = any(parola in voce_check_lower for parola in ["pasta corta", "pasta al forno", "rosette", "gnocchi", "riso", "cous", "quinoa", "farro"])
                    
                    if not is_dolce_check and not is_primo_check:
                        voci_ordinate.append(voce_check)
            
            voci_ordinate.sort()
            
            if voce in voci_ordinate:
                indice = voci_ordinate.index(voce)
                pezzi_per_persona = base + (1 if indice < resto else 0)
            else:
                pezzi_per_persona = base
        else:
            # 5+ voci: distribuzione uniforme
            pezzi_per_persona = max(1, finger_food_per_persona // voci_salate_count)
        
        # Assicura almeno 1 pezzo
        pezzi_per_persona = max(1, pezzi_per_persona)
        pezzi_totali = pezzi_per_persona * persone
        
        print(f"   🎯 Distribuzione {finger_food_per_persona} finger food tra {voci_salate_count} voci")
        print(f"   📊 Pezzi per persona calcolati: {pezzi_per_persona}")
        print(f"   📊 Pezzi totali: {pezzi_totali}")
        print(f"   🔍 DEBUG: voce corrente = '{voce}'")
        print(f"   🔍 DEBUG: finger_food_per_persona = {finger_food_per_persona}")
        print(f"   🔍 DEBUG: voci_salate_count = {voci_salate_count}")
        
        return f"({pezzi_totali}){pezzi_per_persona}"


    # CORREZIONE AGGIUNTIVA: Gestione dei dolci nell'Aperitivo Rinforzato
    # Nell'aperitivo rinforzato dovrebbe essere 8 finger food salati + 3 dolci
    # Questa logica va aggiunta dove si gestiscono i dolci nell'aperitivo rinforzato

    def calcola_quantita_menu_extra_dolci_aperitivo_rinforzato(self, voce, persone, servizio_lower, tipo_ap):
        """Calcolo speciale per dolci nell'aperitivo rinforzato: 70g totali per persona distribuiti tra tutti i dolci"""
        
        # DEBUG: Stampa i valori ricevuti come parametri
        print(f"   🔍 DEBUG CONDIZIONI DOLCI:")
        print(f"      - servizio_lower parametro: '{servizio_lower}'")
        print(f"      - tipo_ap parametro: '{tipo_ap}'")
        print(f"      - 'aperitivo' in servizio: {'aperitivo' in servizio_lower}")
        print(f"      - 'rinforzato' in tipo_ap: {'rinforzato' in tipo_ap}")
        
        # Se è aperitivo rinforzato, i dolci sono 70g totali per persona
        tipo_ap_lower = tipo_ap.lower()
        voce_lower = voce.lower()

        if "aperitivo" in servizio_lower and "rinforzato" in tipo_ap_lower:
            print("   ✅ ENTRANDO NELLA LOGICA APERITIVO RINFORZATO!")

            # 0.08 kg/persona = 80 g/persona (se vuoi 70g metti 0.07)
            peso_dolci_totale = 0.070 * persone

            dolci_selezionati = 0
            for key, checkbox in self.selected.items():
                if hasattr(checkbox, 'isChecked') and checkbox.isChecked():
                    voce_check = self.voce_originale.get(key, key.split('|')[-1]).lower()

                    # --- ESCLUSIONI PER SALATI CON MOUSSE SU CROSTINI/BRUSCHETTE ---
                    is_crostino_bruschetta = ("bruschettine" in voce_check) or ("crostini" in voce_check)
                    has_mousse = "mousse" in voce_check
                    # ingredienti tipici salati
                    mousse_salata = any(w in voce_check for w in [
                        "mortadella","formaggio","caprino","gorgonzola","prosciutto",
                        "speck","ricotta","tonno","salmone","guanciale","pancetta"
                    ])
                    if is_crostino_bruschetta and has_mousse and mousse_salata:
                        continue  # non contare come dolce

                    # --- RICONOSCIMENTO DOLCI VERI ---
                    is_dolce = any(k in voce_check for k in [
                        "dolc","torta","crostat","tiramis","cheesecake","brownie",
                        "biscott","pasticcer","macaron","cioccolat","crema pasticcera",
                        "sfogliat","frutta","cannoncin"
                    ])
                    # opzionale: consentire "mousse" dolce se con parole dolci
                    if (has_mousse and not mousse_salata) and any(s in voce_check for s in [
                        "cioccolat","fragola","frutti","limone","panna","vaniglia"
                    ]):
                        is_dolce = True

                    if is_dolce:
                        dolci_selezionati += 1
                        print(f"      - Dolce trovato: {voce_check}")

            if dolci_selezionati == 0:
                print("   🍰 APERITIVO RINFORZATO - Nessun dolce selezionato")
                return int(persone)  # 1 pz a persona, numerico, NON stringa

            # Distribuzione uniforme del peso dolci tra i dolci selezionati
            peso_per_dolce = peso_dolci_totale / dolci_selezionati

            # peso_unitario default 33.33g se manca
            peso_unitario = self.dati_referenze.get(voce_lower, {}).get("peso_unitario", 0.03333)

            pezzi_necessari = peso_per_dolce / peso_unitario
            import math
            pezzi_finali = max(1, math.ceil(pezzi_necessari))

            print("   🍰 APERITIVO RINFORZATO - DOLCI:")
            print(f"      - Peso totale dolci: {peso_dolci_totale:.3f} kg (80g/persona)")
            print(f"      - Dolci selezionati: {dolci_selezionati}")
            print(f"      - Peso per questo dolce: {peso_per_dolce:.3f} kg")
            print(f"      - Peso unitario {voce}: {peso_unitario:.5f} kg")
            print(f"      - Pezzi calcolati: {pezzi_finali}")

            return pezzi_finali

        # Fuori dalla condizione “aperitivo rinforzato”: niente stringhe speciali.
        print("   🍰 DOLCE STANDARD: 1 per persona (condizione aperitivo rinforzato non soddisfatta)")
        return int(persone)

    def determina_servizio_principale(self):
        """Determina il tipo di servizio principale dalle scelte selezionate"""
        
        for k, checkbox in self.selected.items():
            if hasattr(checkbox, 'isChecked') and checkbox.isChecked():
                if k.startswith("coffee break|"):
                    return "Coffee Break"
                elif k.startswith("tea break|"):
                    return "Tea Break"  
                elif k.startswith("lunch buffet|"):
                    return "Lunch Buffet"
                elif k.startswith("servito|"):
                    return "Servito"
                elif k.startswith("lunch box|"):  # CORREZIONE: Sintassi corretta
                    return "Lunch Box"
                elif k.startswith("aperitivo a buffet|"):
                    tipo_ap = getattr(self, 'ap_buffet_tipo', 'Leggero')
                    return f"Aperitivo {tipo_ap}"
        
        return "Lunch Buffet"  # Default
    def aggiorna_ordine_esistente(self, nuovo_ordine):
        """Aggiorna un ordine esistente nel file JSON - VERSIONE COMPLETA CON AGGIORNAMENTO ARCHIVI"""
        import json
        
        ORDINI_FILE = gestore_percorsi.get_percorso("ordini_json") # Path corretto
        
        print(f"\n🔧 === AGGIORNA_ORDINE_ESISTENTE ===")
        print(f"📂 File: {ORDINI_FILE}")
        
        if not os.path.exists(ORDINI_FILE):
            print("❌ File archivio non esiste, uso salva_ordine normale")
            self.salva_ordine(nuovo_ordine)
            return
        
        with open(ORDINI_FILE, "r", encoding="utf-8") as f:
            ordini = json.load(f)
        
        print(f"📋 Caricati {len(ordini)} ordini esistenti")
        
        ordine_trovato = False
        
        # STRATEGIA 1: Usa l'indice memorizzato se disponibile
        if hasattr(self, 'indice_ordine_originale') and self.indice_ordine_originale is not None:
            idx = self.indice_ordine_originale
            if 0 <= idx < len(ordini):
                print(f"🎯 Usando indice memorizzato: {idx}")
                
                # Combina le scelte esistenti con quelle nuove - VERSIONE CORRETTA
                scelte_esistenti = ordini[idx].get("scelte", {})
                scelte_nuove = nuovo_ordine["scelte"]

                print(f"🔍 SCELTE ESISTENTI: {sum(1 for v in scelte_esistenti.values() if v)} voci")
                print(f"🔍 SCELTE NUOVE: {sum(1 for v in scelte_nuove.values() if v)} voci")

                # COMBINA: mantieni esistenti E aggiungi nuove
                scelte_combinate = {}

                # Copia tutte le scelte esistenti
                for k, v in scelte_esistenti.items():
                    scelte_combinate[k] = v

                # Aggiungi solo le scelte nuove che sono TRUE
                for k, v in scelte_nuove.items():
                    if v:  # Solo se la nuova scelta è selezionata
                        scelte_combinate[k] = True

                print(f"🔍 SCELTE COMBINATE: {sum(1 for v in scelte_combinate.values() if v)} voci")
                
                # Aggiorna l'ordine mantenendo i dati esistenti
                ordini[idx]["scelte"] = scelte_combinate
                # INVECE DI sovrascrivere tutti i campi, mantieni le persone esistenti
                campi_esistenti = ordini[idx]["campi"]
                campi_nuovi = nuovo_ordine["campi"]

                # Somma le persone invece di sovrascrivere
                try:
                    persone_esistenti = int(campi_esistenti.get("Numero persone", 0))
                    persone_nuove = int(campi_nuovi.get("Numero persone", 0))
                    persone_totali = persone_esistenti + persone_nuove
                    
                    # Aggiorna i campi mantenendo le persone sommate
                    campi_aggiornati = campi_esistenti.copy()
                    for campo, valore in campi_nuovi.items():
                        if campo == "Numero persone":
                            campi_aggiornati[campo] = str(persone_totali)
                        elif valore:  # Solo se il nuovo valore non è vuoto
                            campi_aggiornati[campo] = valore
                    
                    ordini[idx]["campi"] = campi_aggiornati
                    print(f"✅ Persone aggiornate: {persone_esistenti} + {persone_nuove} = {persone_totali}")
                    
                except (ValueError, TypeError):
                    # Se c'è un errore, usa il comportamento originale
                    ordini[idx]["campi"] = nuovo_ordine["campi"]
                    print("⚠️ Errore nella somma, uso comportamento originale")
                ordini[idx]["tipo_ap"] = nuovo_ordine["tipo_ap"]
                ordini[idx]["evento_in_fiera"] = nuovo_ordine["evento_in_fiera"]
                ordini[idx]["path_docx"] = nuovo_ordine["path_docx"]  # Aggiorna path
                
                ordine_trovato = True
                print(f"✅ Ordine combinato e aggiornato in posizione {idx}")
                print(f"📊 Scelte totali dopo combinazione: {len(scelte_combinate)}")
        
        # STRATEGIA 2: Fallback per path
        if not ordine_trovato:
            path_da_cercare = nuovo_ordine["path_docx"]
            print(f"🔍 Cercando per path: {path_da_cercare}")
            
            for i, ordine in enumerate(ordini):
                if ordine.get("path_docx") == path_da_cercare:
                    # Combina scelte
                    scelte_esistenti = ordine.get("scelte", {})
                    scelte_nuove = nuovo_ordine["scelte"]
                    scelte_combinate = {**scelte_esistenti, **scelte_nuove}
                    
                    # Aggiorna ordine
                    ordini[i]["scelte"] = scelte_combinate
                    ordini[i]["campi"] = nuovo_ordine["campi"]
                    ordini[i]["tipo_ap"] = nuovo_ordine["tipo_ap"]
                    ordini[i]["evento_in_fiera"] = nuovo_ordine["evento_in_fiera"]
                    
                    ordine_trovato = True
                    print(f"✅ Ordine trovato per path e aggiornato in posizione {i}")
                    break
        
        if not ordine_trovato:
            print("❌ Ordine non trovato, aggiungendo come nuovo")
            ordini.append(nuovo_ordine)
            print(f"➕ Nuovo ordine aggiunto")
        
        # Salva il file
        try:
            # Crea la cartella se non esiste
            os.makedirs(os.path.dirname(ORDINI_FILE), exist_ok=True)
            
            with open(ORDINI_FILE, "w", encoding="utf-8") as f:
                json.dump(ordini, f, indent=2, ensure_ascii=False)
            
            print(f"💾 File salvato con {len(ordini)} ordini totali")
            
        except Exception as e:
            print(f"❌ ERRORE SALVATAGGIO: {e}")
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Errore Salvataggio", f"Errore salvando l'archivio:\n{str(e)}")
            return
        
        print(f"=== FINE AGGIORNA_ORDINE_ESISTENTE ===\n")
        
        # ========== QUESTA È LA PARTE MANCANTE - AGGIORNA ARCHIVI APERTI ==========
        global archivio_dialogs_aperti
        print(f"📊 Archivi aperti da aggiornare: {len(archivio_dialogs_aperti)}")
        
        for dialog in archivio_dialogs_aperti:
            try:
                dialog.ricarica_dati()
                print(f"🔄 Archivio dialog aggiornato")
            except Exception as e:
                print(f"⚠️ Errore aggiornamento archivio dialog: {e}")
        
        # MEMORIZZA che c'è stata una modifica per archivi futuri
        if hasattr(self, 'modalita_aggiungi_servizi') and self.modalita_aggiungi_servizi:
            self._ordine_appena_modificato = True
            print("💡 Flag modifica impostato per prossimi archivi")
        
        # Reset indice
        if hasattr(self, 'indice_ordine_originale'):
            self.indice_ordine_originale = None
            print("🧹 Indice ordine originale resettato")

    def carica_dati_per_aggiungi_servizi(self, ordine):
        """Carica dati cliente e azzera scelte - VERSIONE CORRETTA"""
        print(f"🔄 Modalità aggiungi servizi per: {ordine['campi'].get('Nome', 'Cliente')}")
        
        # Carica dati cliente (NON le scelte!)
        for campo, valore in ordine["campi"].items():
            if campo in self.entries:
                if campo == "Data evento" and hasattr(self.entries[campo], 'setDate'):
                    try:
                        data_obj = datetime.strptime(valore, "%d/%m/%Y")
                        qdate = QDate(data_obj.year, data_obj.month, data_obj.day)
                        self.entries[campo].setDate(qdate)
                    except:
                        pass
                else:
                    self.entries[campo].setText(valore)
        
        # AZZERA tutte le scelte (importante!)
        for key, checkbox in self.selected.items():
            if hasattr(checkbox, 'setChecked'):
                checkbox.setChecked(False)
        
        # Reset aperitivo
        self.ap_buffet_tipo = "Leggero"
        
        # Reset flag fiera
        if hasattr(self, 'evento_in_fiera'):
            self.evento_in_fiera = ordine.get("evento_in_fiera", False)
            if hasattr(self, 'fiera_checkbox'):
                self.fiera_checkbox.setChecked(self.evento_in_fiera)
        
        print("✅ Dati caricati, scelte azzerate")
        
        # MOSTRA IL PULSANTE AGGIUNGI E NASCONDI GENERA NORMALE
        if hasattr(self, 'aggiungi_ordine_btn') and hasattr(self, 'genera_btn'):
            self.aggiungi_ordine_btn.show()
            self.genera_btn.hide()
            
            # Aggiorna il titolo della finestra per indicare la modalità
            self.setWindowTitle("Neri Gestione Catering - AGGIUNGI SERVIZI")
        
        print("✅ Interfaccia aggiornata per modalità aggiungi servizi")

        
def apri_documento_word(self, file_path):
    """Apre il documento Word generato"""
    try:
        import subprocess
        import platform
        
        if platform.system() == "Windows":
            subprocess.run(["start", file_path], shell=True, check=True)
        elif platform.system() == "Darwin":  # macOS
            subprocess.run(["open", file_path], check=True)
        else:  # Linux
            subprocess.run(["xdg-open", file_path], check=True)
            
        print(f"📄 Documento aperto automaticamente: {file_path}")
        
    except Exception as e:
        print(f"❌ Errore apertura documento: {e}")
        QMessageBox.information(self, "Info", 
                               f"Documento salvato in:\n{file_path}\n\n"
                               f"Aprilo manualmente per modificare le quantità.")

# ===================================================================
# FUNZIONE INTELLIGENTE PER LEGGERE MODIFICHE DAL WORD
# ===================================================================


# MIGLIORA IL METODO leggi_quantita_da_word PER LEGGERE MEGLIO GLI ORDINI COMPLETI

def is_header_or_non_quantity(self, riga):
    """Controlla se una riga è un header o non contiene quantità"""
    riga_lower = riga.lower()
    
    # Salta header e righe non utili
    header_keywords = [
        "nr progr", "all'attenzione", "presso", "aperitivo", "coffee break", 
        "tea break", "lunch buffet", "persone", "allestimento", "pronti",
        "bevande", "sala", "lunedì", "martedì", "mercoledì", "giovedì", 
        "venerdì", "sabato", "domenica", "gennaio", "febbraio", "marzo",
        "aprile", "maggio", "giugno", "luglio", "agosto", "settembre",
        "ottobre", "novembre", "dicembre"
    ]
    
    # Se contiene parole header, salta
    if any(keyword in riga_lower for keyword in header_keywords):
        return True
    
    # Se è troppo corta (meno di 3 caratteri), salta
    if len(riga.strip()) < 3:
        return True
        
    # Se è solo punteggiatura, salta
    if riga.strip() in ["•", "-", "*", "○", "●"]:
        return True
    
    return False

def estrai_quantita_da_riga_migliorata(self, riga):
    """Estrae quantità da una singola riga - VERSIONE MIGLIORATA"""
    import re
    
    risultati = {}
    riga_originale = riga.strip()
    
    # Pattern migliorati per catturare tutti i formati possibili
    patterns = [
        # Pattern per peso in kg (più specifico)
        (r'^(\d+\.?\d*)\s*kg\s+(.+)$', "kg"),
        
        # Pattern per lasagne, crespelle, etc.
        (r'^(\d+)\s+(lasagne|crespelle|cannelloni|rosette|nidi)(?:\s+(.+))?$', "primi"),
        
        # Pattern generico numero + nome (più permissivo)
        (r'^(\d+)\s+(.+)$', "pezzi"),
        
        # Pattern per quantità con **grassetto** (primi piatti)
        (r'^\*\*(\d+\.?\d*)\s*(?:kg\s+)?(.+?)\*\*$', "primi_grassetto"),
    ]
    
    for pattern, tipo in patterns:
        match = re.match(pattern, riga_originale, re.IGNORECASE)
        if match:
            try:
                if tipo == "kg":
                    quantita_num = float(match.group(1))
                    nome_prodotto = match.group(2).strip()
                    quantita_str = f"{quantita_num} kg"
                    
                elif tipo == "primi":
                    quantita_num = int(match.group(1))
                    tipo_primo = match.group(2)
                    desc_extra = match.group(3) if match.group(3) else ""
                    nome_prodotto = f"{tipo_primo} {desc_extra}".strip()
                    quantita_str = f"{quantita_num} {tipo_primo}"
                    
                elif tipo == "primi_grassetto":
                    quantita_raw = match.group(1)
                    nome_prodotto = match.group(2).strip()
                    if "." in quantita_raw:
                        quantita_num = float(quantita_raw)
                        quantita_str = f"{quantita_num} kg"
                    else:
                        quantita_num = int(quantita_raw)
                        quantita_str = f"{quantita_num}"
                        
                else:  # pezzi
                    quantita_num = int(match.group(1))
                    nome_prodotto = match.group(2).strip()
                    quantita_str = str(quantita_num)
                
                # Pulisci il nome del prodotto
                nome_prodotto = self.pulisci_nome_prodotto(nome_prodotto)
                
                # Crea chiave univoca
                key = f"extracted_{len(risultati)}"
                
                risultati[key] = {
                    'nome': nome_prodotto,
                    'quantita': quantita_str
                }
                
                print(f"        ✅ Match {tipo}: '{riga_originale}' → {nome_prodotto}: {quantita_str}")
                break
                
            except (ValueError, AttributeError) as e:
                print(f"        ❌ Errore parsing: {e}")
                continue
    
    return risultati

def pulisci_nome_prodotto(self, nome):
    """Pulisce il nome del prodotto da caratteri indesiderati"""
    # Rimuovi caratteri markdown
    nome = nome.replace("**", "").replace("*", "")
    
    # Rimuovi bullet points
    nome = nome.replace("•", "").replace("○", "").replace("●", "").replace("-", "", 1)
    
    # Rimuovi spazi multipli
    nome = " ".join(nome.split())
    
    return nome.strip()

def trova_chiave_corrispondente_debug(self, nome_prodotto):
    """Trova la chiave corrispondente - VERSIONE DEBUG"""
    nome_lower = nome_prodotto.lower()
    print(f"            🔍 Ricerca chiave per: '{nome_prodotto}' (lower: '{nome_lower}')")
    
    print(f"            📋 Voci disponibili nel menu:")
    for key, voce_originale in self.voce_originale.items():
        print(f"               {key} → '{voce_originale}'")
    
    # 1. Corrispondenza diretta
    for key, voce_originale in self.voce_originale.items():
        voce_lower = voce_originale.lower()
        
        if nome_lower == voce_lower:
            print(f"            ✅ MATCH DIRETTO: '{nome_lower}' == '{voce_lower}'")
            return key
    
    # 2. Corrispondenza parziale
    for key, voce_originale in self.voce_originale.items():
        voce_lower = voce_originale.lower()
        
        # Nome contenuto nella voce o viceversa
        if nome_lower in voce_lower:
            print(f"            ✅ MATCH PARZIALE: '{nome_lower}' IN '{voce_lower}'")
            if len(nome_lower) >= 4:  # Evita match troppo corti
                return key
        
        if voce_lower in nome_lower:
            print(f"            ✅ MATCH PARZIALE: '{voce_lower}' IN '{nome_lower}'")
            if len(voce_lower) >= 4:
                return key
    
    # 3. Corrispondenze specifiche per "Mini pan chocolat"
    if "pan" in nome_lower and "chocolat" in nome_lower:
        for key, voce_originale in self.voce_originale.items():
            voce_lower = voce_originale.lower()
            if "pan" in voce_lower and "chocolat" in voce_lower:
                print(f"            ✅ MATCH SPECIFICO PAN CHOCOLAT: {key}")
                return key
    
    # 4. Altri mapping specifici
    mapping_specifici = {
        'mini pan chocolat': ['pan', 'chocolat'],
        'pizzette': ['pizzette'],
        'calzoncini': ['calzoncini'],
        'tramezzini': ['tramezzini'],
        'focaccine': ['focaccine'],
        'polpettine': ['polpettine'],
        'lasagne': ['lasagne'],
        'crespelle': ['crespelle'],
        'cannelloni': ['cannelloni'],
        'cous cous': ['cous'],
        'verdi': ['verdi']
    }
    
    for nome_mapping, parole_chiave in mapping_specifici.items():
        if all(parola in nome_lower for parola in parole_chiave):
            for key, voce_originale in self.voce_originale.items():
                voce_lower = voce_originale.lower()
                if all(parola in voce_lower for parola in parole_chiave):
                    print(f"            ✅ MATCH MAPPING: {nome_mapping} → {key}")
                    return key
    
    print(f"            ❌ Nessuna corrispondenza trovata per: '{nome_prodotto}'")
    return None

# ===================================================================
# TEST RAPIDO PER DEBUG
# ===================================================================

def test_parsing_mini_pan_chocolat(self):
    """Test specifico per Mini pan chocolat"""
    print("\n🧪 === TEST PARSING MINI PAN CHOCOLAT ===")
    
    test_righe = [
        "25 Mini pan chocolat",
        "25  Mini pan chocolat",
        "25 Mini pan chocolat al cioccolato",
        "30 Mini pan chocolat",
        "Mini pan chocolat 25",
        "25Mini pan chocolat"
    ]
    
    for riga in test_righe:
        print(f"\n🧪 Test riga: '{riga}'")
        risultato = self.estrai_quantita_da_riga_debug(riga)
        if risultato:
            print(f"   ✅ Successo: {risultato}")
        else:
            print(f"   ❌ Fallito")



def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    
    # DEBUG: Info sistema
    print("🚀 Avvio applicazione...")
    
    # Splash screen
    splash = CleanLogoSplash()
    splash.show()
    
    # Centra splash
    screen = app.primaryScreen().geometry()
    splash_rect = splash.geometry()
    splash.move(
        (screen.width() - splash_rect.width()) // 2,
        (screen.height() - splash_rect.height()) // 2
    )
    
    app.processEvents()
    
    main_window = None
    
    def start_main_application():
        nonlocal main_window
        
        app.setStyleSheet(PROFESSIONAL_STYLE)
        
        # Palette
        palette = QPalette()
        palette.setColor(QPalette.ColorRole.Window, QColor(248, 249, 250))
        palette.setColor(QPalette.ColorRole.WindowText, QColor(33, 37, 41))
        palette.setColor(QPalette.ColorRole.Base, QColor(255, 255, 255))
        palette.setColor(QPalette.ColorRole.AlternateBase, QColor(248, 249, 250))
        palette.setColor(QPalette.ColorRole.ToolTipBase, QColor(52, 73, 94))
        palette.setColor(QPalette.ColorRole.ToolTipText, QColor(255, 255, 255))
        palette.setColor(QPalette.ColorRole.Text, QColor(33, 37, 41))
        palette.setColor(QPalette.ColorRole.Button, QColor(236, 240, 241))
        palette.setColor(QPalette.ColorRole.ButtonText, QColor(44, 62, 80))
        palette.setColor(QPalette.ColorRole.BrightText, QColor(52, 152, 219))
        palette.setColor(QPalette.ColorRole.Link, QColor(52, 152, 219))
        palette.setColor(QPalette.ColorRole.Highlight, QColor(52, 152, 219))
        palette.setColor(QPalette.ColorRole.HighlightedText, QColor(255, 255, 255))
        app.setPalette(palette)
         
        main_window = CateringApp()
        main_window.show()
        splash.close()
        
        print("✅ Applicazione principale avviata")
    
    splash.worker.finished.connect(start_main_application)
    
    sys.exit(app.exec())

if __name__ == "__main__":
    main()