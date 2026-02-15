#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Server Centrale con Cloud Bridge - VERSIONE FIREBASE
Sincronizza ordini con Firebase per accesso da reti diverse
"""

import os
import json
import time
import threading
from datetime import datetime, timedelta
from collections import defaultdict
from pathlib import Path
import re
import hashlib
import shutil  # 🔥 AGGIUNGI QUESTA RIGA


from flask import Flask, jsonify, request
try:
    from flask_cors import CORS
    CORS_AVAILABLE = True
except ImportError:
    CORS_AVAILABLE = False
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

# ===== NUOVO: FIREBASE IMPORTS =====
try:
    import firebase_admin
    from firebase_admin import credentials, firestore
    FIREBASE_AVAILABLE = True
    print("✅ Firebase disponibile")
except ImportError:
    FIREBASE_AVAILABLE = False
    print("⚠️ Firebase non disponibile - installare con: pip install firebase-admin")

class FirebaseManager:
    """Gestisce la sincronizzazione con Firebase"""
    
    def __init__(self):
        self.db = None
        self.firebase_enabled = False
        self.init_firebase()
    
    def init_firebase(self):
        """Inizializza Firebase con credenziali"""
        if not FIREBASE_AVAILABLE:
            print("❌ Firebase non disponibile")
            return
        
        try:
            # Cerca il file delle credenziali
            cred_files = [
                "firebase-credentials.json",
                "serviceAccountKey.json", 
                "firebase-adminsdk.json"
            ]
            
            cred_path = None
            for file in cred_files:
                if os.path.exists(file):
                    cred_path = file
                    break
            
            if not cred_path:
                print("⚠️ File credenziali Firebase non trovato!")
                print("📋 Salva le credenziali come 'firebase-credentials.json'")
                return
            
            # Inizializza Firebase
            if not firebase_admin._apps:
                cred = credentials.Certificate(cred_path)
                firebase_admin.initialize_app(cred)
            
            self.db = firestore.client()
            self.firebase_enabled = True
            print(f"✅ Firebase inizializzato con {cred_path}")
            
            # Test connessione
            self.test_connection()
            
        except Exception as e:
            print(f"❌ Errore inizializzazione Firebase: {e}")
            self.firebase_enabled = False
    
    def test_connection(self):
        """Testa la connessione Firebase"""
        try:
            # Scrive un documento di test
            test_ref = self.db.collection('test').document('connection')
            test_ref.set({
                'timestamp': datetime.now().isoformat(),
                'status': 'ok'
            })
            print("🔥 Test Firebase OK - Connessione stabilita")
        except Exception as e:
            print(f"⚠️ Test Firebase fallito: {e}")
    
    def upload_ordini(self, ordini_singoli, ordini_aggregati):
        """Carica ordini su Firebase"""
        if not self.firebase_enabled:
            return False
        
        try:
            batch = self.db.batch()
            
            # ===== CARICA ORDINI SINGOLI =====
            # ===== CARICA ORDINI SINGOLI (CON CONTROLLO ELIMINAZIONI) =====
            for ordine_id, ordine in ordini_singoli.items():
                doc_ref = self.db.collection('ordini_singoli').document(ordine_id)
                
                # Prepara dati per Firebase (serializzabili)
                ordine_data = {
                    'id': ordine['id'],
                    'info': ordine['info'],
                    'cibo': ordine['cibo'],
                    'stato': ordine['stato'],
                    'ultimo_aggiornamento': ordine['ultimo_aggiornamento'],
                    'cloud_sync': datetime.now().isoformat()
                }
                
                batch.set(doc_ref, ordine_data)
                print(f"✅ Ordine {ordine_id} aggiornato su Firebase")
            
            # ===== CARICA AGGREGATI =====
            aggregati_doc = self.db.collection('sistema').document('ordini_aggregati')
            aggregati_data = {
                'ordini': dict(ordini_aggregati),
                'last_update': datetime.now().isoformat(),
                'status': 'ok'
            }
            batch.set(aggregati_doc, aggregati_data)
            
            # ===== METADATA SISTEMA =====
            meta_doc = self.db.collection('sistema').document('metadata')
            meta_data = {
                'totale_ordini': len(ordini_singoli),
                'date_attive': list(ordini_aggregati.keys()),
                'ultimo_sync': datetime.now().isoformat(),
                'versione': '2.0-CLOUD'
            }
            batch.set(meta_doc, meta_data)
            
            # Commit batch
            batch.commit()
            
            print(f"🔥 Firebase sync OK: {len(ordini_singoli)} ordini caricati")
            return True
            
        except Exception as e:
            print(f"❌ Errore upload Firebase: {e}")
            return False
    
    def cleanup_old_orders(self, max_days=30):
        """Pulisce ordini vecchi da Firebase"""
        if not self.firebase_enabled:
            return
        
        try:
            cutoff_date = datetime.now() - timedelta(days=max_days)
            
            # Query ordini vecchi
            old_orders = self.db.collection('ordini_singoli').where(
                'ultimo_aggiornamento', '<', cutoff_date.isoformat()
            ).get()
            
            deleted_count = 0
            for doc in old_orders:
                doc.reference.delete()
                deleted_count += 1
            
            if deleted_count > 0:
                print(f"🧹 Firebase cleanup: {deleted_count} ordini vecchi rimossi")
                
        except Exception as e:
            print(f"⚠️ Errore cleanup Firebase: {e}")


class WordFileHandler(FileSystemEventHandler):
    """Gestisce eventi di modifica sui file Word - VERSIONE CORRETTA"""
    
    def __init__(self, server_instance):
        self.server = server_instance
        
    def should_ignore_file(self, filepath):
        """Controlla se il file dovrebbe essere ignorato"""
        # 🔥 IGNORA file nella cartella _cancellati
        if "_cancellati" in filepath:
            print(f"⏭️ IGNORANDO file in cartella cancellati: {filepath}")
            return True
        
        # Ignora file temporanei Word
        filename = os.path.basename(filepath)
        if filename.startswith('~$'):
            print(f"⏭️ IGNORANDO file temporaneo: {filepath}")
            return True
            
        # Solo file .docx
        if not filepath.endswith('.docx'):
            return True
            
        return False
        
    def on_modified(self, event):
        if event.is_directory:
            return
            
        if self.should_ignore_file(event.src_path):
            return
            
        print(f"📄 File modificato (accettato): {event.src_path}")
        self.server.aggiorna_ordine_da_file(event.src_path)
    
    def on_created(self, event):
        if event.is_directory:
            return
            
        if self.should_ignore_file(event.src_path):
            return
            
        print(f"📄 Nuovo file (accettato): {event.src_path}")
        time.sleep(1)  # Aspetta che il file sia completamente scritto
        self.server.aggiorna_ordine_da_file(event.src_path)
    
    def on_moved(self, event):
        """🔥 NUOVO: Gestisce spostamenti di file"""
        if event.is_directory:
            return
        
        # Se un file viene spostato IN _cancellati, ignoralo
        if "_cancellati" in event.dest_path:
            print(f"📁 File spostato in _cancellati (ignorato): {event.src_path} → {event.dest_path}")
            return
            
        # Se un file viene spostato FUORI da _cancellati, trattalo come nuovo
        if "_cancellati" in event.src_path and not "_cancellati" in event.dest_path:
            print(f"📄 File ripristinato da _cancellati: {event.src_path} → {event.dest_path}")
            if not self.should_ignore_file(event.dest_path):
                time.sleep(1)
                self.server.aggiorna_ordine_da_file(event.dest_path)


class ServerCentraleCloud:
    """Server centrale con sincronizzazione Firebase"""
    
    def __init__(self):
        # Mantiene ordini singoli + aggregazione
        self.ordini_singoli = {}  # {ordine_id: {info_ordine + quantita_cibo}}
        self.ordini_aggregati = {}  # {data: {nome_cibo: quantita_totale}}
        
        self.cartella_ordini = "ordini_docx"
        self.cartella_drive = "ordini_drive"  # NUOVA cartella Drive

        self.last_update = datetime.now()
        
        # ===== NUOVO: FIREBASE MANAGER =====
        self.firebase = FirebaseManager()
        
        # Crea Flask app
        self.app = Flask(__name__)
        
        # Abilita CORS se disponibile
        if CORS_AVAILABLE:
            CORS(self.app)
            print("✅ CORS abilitato")
        else:
            print("⚠️ CORS non disponibile - installare con: pip install flask-cors")
        
        # Registra routes
        self.registra_routes()
        
        # Avvia monitoring
        self.avvia_file_watcher()
        
        # Carica ordini esistenti
        self.carica_ordini_esistenti()
        
        # ===== NUOVO: SYNC TIMER =====
        self.avvia_sync_automatico()
    
    def avvia_sync_automatico(self):
        """Sync automatico INTELLIGENTE - non sovrascrive cancellazioni"""
        def sync_loop():
            while True:
                time.sleep(30)  # Sync ogni 30 secondi
                if self.firebase.firebase_enabled:
                    # 🔥 SYNC INTELLIGENTE: Solo nuovi/modificati, non sovrascrive
                    self.smart_firebase_sync_with_cleanup()
        
        sync_thread = threading.Thread(target=sync_loop, daemon=True)
        sync_thread.start()
        print("🔄 Sync automatico Firebase INTELLIGENTE avviato (ogni 30s)")
    
    def smart_firebase_sync_with_cleanup(self):
        """Sync Firebase + pulizia locale ordini scaduti"""
        try:
            print("🧠 Sync Firebase intelligente con pulizia...")
            
            # 1. Prima pulisci ordini scaduti dal server locale
            from datetime import datetime, timedelta
            cutoff_date = datetime.now() - timedelta(hours=24)
            
            orders_to_remove = []
            for order_id, order_data in self.ordini_singoli.items():
                data_evento = order_data.get('info', {}).get('data_evento')
                if data_evento:
                    try:
                        data_obj = datetime.strptime(data_evento, '%d/%m/%Y')
                        if data_obj < cutoff_date:
                            orders_to_remove.append(order_id)
                            nome_cliente = order_data.get('info', {}).get('nome_cliente', 'Unknown')
                            print(f"🧹 Rimuovendo ordine scaduto dal server: {nome_cliente} ({data_evento})")
                    except (ValueError, TypeError):
                        pass
            
            # Rimuovi ordini scaduti dal server
            for order_id in orders_to_remove:
                del self.ordini_singoli[order_id]
            
            if orders_to_remove:
                print(f"🧹 Rimossi {len(orders_to_remove)} ordini scaduti dal server")
                self.rigenera_aggregazione()
            
            # 2. Poi fai il sync normale (che ora non includerà ordini scaduti)
            firebase_orders = {}
            docs = self.firebase.db.collection('ordini_singoli').stream()
            for doc in docs:
                firebase_orders[doc.id] = doc.to_dict()
            
            print(f"📊 Firebase ha {len(firebase_orders)} ordini, server ha {len(self.ordini_singoli)} ordini")
            
            # 3. Sync solo ordini validi
            to_add_or_update = {}
            for order_id, order_data in self.ordini_singoli.items():
                if order_id not in firebase_orders:
                    to_add_or_update[order_id] = order_data
                    print(f"➕ Da aggiungere: {order_id}")
                elif firebase_orders[order_id].get('ultimo_aggiornamento') != order_data.get('ultimo_aggiornamento'):
                    to_add_or_update[order_id] = order_data
                    print(f"🔄 Da aggiornare: {order_id}")
            
            if to_add_or_update:
                self.firebase.upload_ordini(to_add_or_update, self.ordini_aggregati)
                print(f"✅ Sync completato: {len(to_add_or_update)} ordini sincronizzati")
            else:
                print("✅ Nessun sync necessario")
                
        except Exception as e:
            print(f"❌ Errore smart sync con pulizia: {e}")
            import traceback
            traceback.print_exc()
    
    def genera_id_ordine(self, filepath):
        """Genera ID ordine basato solo sul nome del file"""
        import hashlib
        filename = os.path.basename(filepath)
        # Rimuovi estensione
        name_without_ext = os.path.splitext(filename)[0]
        ordine_id = hashlib.md5(name_without_ext.encode()).hexdigest()[:8]
        return ordine_id
    
    def estrai_info_ordine_da_filename(self, filename, filepath):
        """Estrae informazioni ordine dal nome file"""
        try:
            # Pattern: Nome_Cliente_DD-MM-YYYY.docx
            base_name = os.path.splitext(filename)[0]
            
            # Cerca pattern data
            date_match = re.search(r'(\d{2}-\d{2}-\d{4})', base_name)
            if date_match:
                data_str = date_match.group(1).replace('-', '/')
                # Rimuovi la data dal nome per ottenere il cliente
                nome_cliente = base_name.replace(f"_{date_match.group(1)}", "").replace(date_match.group(1), "")
                nome_cliente = nome_cliente.strip('_')
            else:
                data_str = None
                nome_cliente = base_name
            
            # Info file
            stat = os.stat(filepath)
            
            return {
                'nome_cliente': nome_cliente,
                'data_evento': data_str,
                'filename': filename,
                'filepath': filepath,
                'data_creazione': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'data_modifica': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'dimensione_file': stat.st_size
            }
            
        except Exception as e:
            print(f"❌ Errore estrazione info da {filename}: {e}")
            return {
                'nome_cliente': os.path.splitext(filename)[0],
                'data_evento': None,
                'filename': filename,
                'filepath': filepath,
                'data_creazione': datetime.now().isoformat(),
                'data_modifica': datetime.now().isoformat(),
                'dimensione_file': 0
            }
    
    def leggi_quantita_da_word(self, path_docx):
        """
        Legge quantità da Word - VERSIONE CORRETTA con supporto nuovo formato
        Riconosce: (2)1 pz Pasta corta fai tu, 1 pz Pasta corta fai tu, 3.5 kg Cous cous
        Ritorna: {chiave: {'nome': str, 'quantita': str}}
        """
        if not os.path.exists(path_docx):
            print(f"❌ File Word non trovato: {path_docx}")
            return {}
        
        try:
            from docx import Document
            import re
            doc = Document(path_docx)
            
            prodotti = {}
            
            print(f"\n📄 === LETTURA DA: {os.path.basename(path_docx)} ===")
            
            # Leggi TUTTO il testo del documento
            all_text = ""
            
            # 1. LEGGI DA PARAGRAFI
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text += para.text + "\n"
            
            # 2. LEGGI DA TABELLE  
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            all_text += cell.text + "\n"
            
            # 3. PROCESSA TUTTO IL TESTO RIGA PER RIGA
            righe = [r.strip() for r in all_text.split('\n') if r.strip()]
            
            print(f"📋 Analizzando {len(righe)} righe totali...")
            
            # Parole da saltare (headers, etc.)
            skip_keywords = [
                "nr progr", "all'attenzione", "presso", "aperitivo", "coffee break", 
                "tea break", "lunch buffet", "persone", "allestimento", "pronti",
                "bevande", "sala", "lunedì", "martedì", "mercoledì", "giovedì", 
                "venerdì", "sabato", "domenica", "gennaio", "febbraio", "marzo",
                "aprile", "maggio", "giugno", "luglio", "agosto", "settembre",
                "ottobre", "novembre", "dicembre", "totale", "pax", "euro", "€",
                "ore", "orario", "via", "telefono", "email", "fax", "intolleranti"
            ]
            
            for i, riga in enumerate(righe):
                riga_original = riga.strip()
                print(f"   DEBUG RIGA {i+1}: '{riga_original}'")

                riga_lower = riga_original.lower()
                
                # Salta righe vuote o troppo corte
                if len(riga_original) < 3:
                    continue
                
                # Salta righe con parole chiave da evitare
                if any(keyword in riga_lower for keyword in skip_keywords) and "kg" not in riga_lower:                    
                    continue
                
                # Salta se è solo punteggiatura
                if riga_original.strip() in ["•", "-", "*", "○", "●", "◦", "|", "+", "="]:
                    continue
                
                prodotto_trovato = False
                
                # ========== PATTERN 1: PESO IN KG ==========
                patterns_kg = [
                    r'^(\d+\.?\d*)\s*kg\s+(.+)$',
                    r'^\*\*(\d+\.?\d*)\s*kg\s+(.+?)\*\*$',
                    r'^(.+?)\s+(\d+\.?\d*)\s*kg$',
                    r'^\*\*(.+?)\s+(\d+\.?\d*)\s*kg\*\*$'
                ]
                
                for pattern in patterns_kg:
                    match = re.match(pattern, riga_original, re.IGNORECASE)
                    if match:
                        if pattern.endswith('kg$') or 'kg\\*\\*$' in pattern:
                            nome = match.group(1).strip()
                            peso = match.group(2)
                        else:
                            peso = match.group(1)
                            nome = match.group(2).strip()
                        
                        nome = nome.replace('*', '').strip()
                        peso_float = float(peso)
                        
                        key = nome.lower().replace(' ', '_')
                        
                        if key in prodotti and 'kg' in prodotti[key]['quantita']:
                            peso_esistente = float(prodotti[key]['quantita'].replace(' kg', ''))
                            peso_totale = peso_esistente + peso_float
                            prodotti[key]['quantita'] = f"{peso_totale} kg"
                            print(f"   🔄 KG ACCUMULATO: {nome} = {peso_esistente} + {peso_float} = {peso_totale} kg")
                        else:
                            prodotti[key] = {
                                'nome': nome,
                                'quantita': f"{peso_float} kg"
                            }
                            print(f"   ✅ KG: {nome} = {peso_float} kg")
                        
                        prodotto_trovato = True
                        break
                
                if prodotto_trovato:
                    continue
                
                # ========== PATTERN 2: NUOVO FORMATO CON PARENTESI ==========
                print(f"   TESTANDO PARENTESI per: '{riga_original}'")
                
                # Pattern per (X)Y pz Nome e (X)Y Nome
                pattern_pz = r'^\((\d+)\)(\d+)\s*pz\s+(.+)$'
                pattern_semplice = r'^\((\d+)\)(\d+)\s+(.+)$'
                pattern_kg = r'^\((\d+)\)(\d+\.?\d*)\s*kg\s+(.+)$'
                
                # Test pattern con pz
                match = re.match(pattern_pz, riga_original, re.IGNORECASE)
                if match:
                    print(f"      ✅ MATCH PARENTESI PZ TROVATO!")
                    numero_totale = int(match.group(1))         # 12
                    quantita_per_persona = int(match.group(2))  # 7 (non lo usiamo più)
                    nome = match.group(3).strip()
                    
                    nome = nome.replace('*', '').strip()
                    nome = " ".join(nome.split())
                    
                    key = nome.lower().replace(' ', '_')
                    
                    if key in prodotti:
                        try:
                            numero_esistente = int(prodotti[key]['quantita'].split()[0])
                            numero_totale_finale = numero_esistente + numero_totale  # ← CAMBIATO DA quantita_per_persona
                            prodotti[key]['quantita'] = f"{numero_totale_finale} pezzi"
                            print(f"   🔄 PARENTESI PEZZI ACCUMULATI: {nome} = {numero_esistente} + {numero_totale} = {numero_totale_finale}")  # ← CAMBIATO
                        except:
                            prodotti[key]['quantita'] = f"{numero_totale} pezzi"  # ← CAMBIATO DA quantita_per_persona
                    else:
                        prodotti[key] = {
                            'nome': nome,
                            'quantita': f"{numero_totale} pezzi"  # ← CAMBIATO DA quantita_per_persona
                        }
                        print(f"   ✅ PARENTESI PZ: {nome} = {numero_totale} pezzi (era parziale: {quantita_per_persona})")  # ← CAMBIATO
                    
                    prodotto_trovato = True
                # Test pattern semplice se non trovato con pz
                if not prodotto_trovato:
                    match = re.match(pattern_semplice, riga_original, re.IGNORECASE)
                    if match:
                        print(f"      ✅ MATCH PARENTESI SEMPLICE TROVATO!")
                        numero_totale = int(match.group(1))
                        quantita_per_persona = int(match.group(2))
                        nome = match.group(3).strip()
                        
                        nome = nome.replace('*', '').strip()
                        nome = " ".join(nome.split())
                        
                        key = nome.lower().replace(' ', '_')
                        
                        if key in prodotti:
                            try:
                                numero_esistente = int(prodotti[key]['quantita'].split()[0])
                                numero_totale_finale = numero_esistente + quantita_per_persona
                                prodotti[key]['quantita'] = f"{numero_totale_finale} pezzi"
                                print(f"   🔄 PARENTESI SEMPLICE ACCUMULATI: {nome} = {numero_esistente} + {quantita_per_persona} = {numero_totale_finale}")
                            except:
                                prodotti[key]['quantita'] = f"{quantita_per_persona} pezzi"
                        else:
                            prodotti[key] = {
                                'nome': nome,
                                'quantita': f"{quantita_per_persona} pezzi"
                            }
                            print(f"   ✅ PARENTESI SEMPLICE: {nome} = {quantita_per_persona} pezzi (totale persone: {numero_totale})")
                        
                        prodotto_trovato = True
                
                # Test pattern kg se non trovato
                if not prodotto_trovato:
                    match = re.match(pattern_kg, riga_original, re.IGNORECASE)
                    if match:
                        print(f"      ✅ MATCH PARENTESI KG TROVATO!")
                        numero_totale = int(match.group(1))
                        quantita_per_persona = float(match.group(2))
                        nome = match.group(3).strip()
                        
                        nome = nome.replace('*', '').strip()
                        
                        key = nome.lower().replace(' ', '_')
                        
                        if key in prodotti and 'kg' in prodotti[key]['quantita']:
                            quantita_esistente = float(prodotti[key]['quantita'].replace(' kg', ''))
                            quantita_totale = quantita_esistente + quantita_per_persona
                            prodotti[key]['quantita'] = f"{quantita_totale} kg"
                            print(f"   🔄 PARENTESI KG ACCUMULATO: {nome} = {quantita_esistente} + {quantita_per_persona} = {quantita_totale} kg")
                        else:
                            prodotti[key] = {
                                'nome': nome,
                                'quantita': f"{quantita_per_persona} kg"
                            }
                            print(f"   ✅ PARENTESI KG: {nome} = {quantita_per_persona} kg (totale persone: {numero_totale})")
                        
                        prodotto_trovato = True
                
                if prodotto_trovato:
                    continue
                
                # ========== PATTERN 3: FORMATO TRADIZIONALE PEZZI ==========
                patterns_pz = [
                    # NUOVO PATTERN per formato intolleranze (12)7 pz
                    r'^\((\d+)\)(\d+)\s*pz\s+(.+)$',           # (12)7 pz Pizzette per vegani  
                    r'^\*\*\((\d+)\)(\d+)\s*pz\s+(.+?)\*\*$',  # **(12)7 pz Pizzette per vegani**
                    r'^\*\*(\d+)\s*pz\s+(.+?)\*\*$',
                    r'^(\d+)\s*pz\s+(.+)$',
                    r'^\*\*(\d+)\s+(.+?)\*\*$',
                    r'^(\d+)\s+(.+)$'
                    
                ]
                
                for pattern in patterns_pz:
                    match = re.match(pattern, riga_original, re.IGNORECASE)
                    if match:
                        # ========== GESTIONE PATTERN INTOLLERANZE ==========
                        if pattern.startswith(r'^\('):  # Pattern con parentesi
                            totale = int(match.group(1))     # Il numero dentro le parentesi (12)
                            parziale = int(match.group(2))   # Il numero dopo le parentesi (7)
                            nome = match.group(3).strip()    # Nome del prodotto
                            
                            # USA IL TOTALE invece del parziale
                            numero = totale  # ← QUESTA È LA CORREZIONE PRINCIPALE
                            
                            print(f"   ✅ INTOLLERANZA: ({totale}){parziale} → uso totale {totale} per {nome}")
                            
                        else:  # Pattern normali
                            numero = int(match.group(1))
                            nome = match.group(2).strip()
                        
                        # Pulisci nome da markdown e spazi extra
                        nome = nome.replace('*', '').strip()
                        nome = " ".join(nome.split())  # Rimuovi spazi multipli
                        
                        # Filtra nomi troppo generici
                        if len(nome) < 3:
                            continue
                        
                        key = nome.lower().replace(' ', '_')
                        
                        if key in prodotti:
                            try:
                                numero_esistente = int(prodotti[key]['quantita'].split()[0])
                                numero_totale = numero_esistente + numero
                                prodotti[key]['quantita'] = f"{numero_totale} pezzi"
                                print(f"   🔄 PEZZI ACCUMULATI: {nome} = {numero_esistente} + {numero} = {numero_totale}")
                            except:
                                prodotti[key]['quantita'] = f"{numero} pezzi"
                        else:
                            prodotti[key] = {
                                'nome': nome,  # ✅ SOLO IL NOME PULITO
                                'quantita': f"{numero} pezzi"
                            }
                            print(f"   ✅ PEZZI: {nome} = {numero} pezzi")
                        
                        prodotto_trovato = True
                        break
                
                # Se non riconosciuto, stampa per debug
                if not prodotto_trovato and len(riga_original) > 5:
                    print(f"   ❓ NON RICONOSCIUTO: '{riga_original}'")
            
            print(f"🎯 === TOTALE TROVATI: {len(prodotti)} prodotti ===")
            for key, info in prodotti.items():
                print(f"   📝 {info['nome']}: {info['quantita']}")
            
            return prodotti
            
        except Exception as e:
            print(f"❌ Errore lettura Word {path_docx}: {e}")
            import traceback
            traceback.print_exc()
            return {}
    def aggiorna_ordine_da_file(self, filepath):
        """Aggiorna/crea ordine singolo da file Word - CON INTOLLERANZE"""
        try:
            filename = os.path.basename(filepath)
            ordine_id = self.genera_id_ordine(filepath)
            
            print(f"🔄 Elaborando ordine ID: {ordine_id}")
            
            # Estrai info ordine
            info_ordine = self.estrai_info_ordine_da_filename(filename, filepath)
            
            # Leggi quantità dal Word
            quantita_cibo = self.leggi_quantita_da_word(filepath)
            
            # 🔥 NUOVO: LEGGI INTOLLERANZE DAL JSON
            
            # Crea/aggiorna ordine completo
            ordine_completo = {
                'id': ordine_id,
                'info': info_ordine,
                'cibo': quantita_cibo,
                'stato': 'attivo',
                'ultimo_aggiornamento': datetime.now().isoformat()
            }
            
            # Salva ordine singolo
            self.ordini_singoli[ordine_id] = ordine_completo
            
            # Rigenera aggregazione
            self.rigenera_aggregazione()
            
            self.last_update = datetime.now()
            
            # ===== SYNC IMMEDIATO FIREBASE =====
            # ===== SYNC SELETTIVO FIREBASE =====
            if self.firebase.firebase_enabled:
                # Sync solo questo ordine specifico, non tutti
                ordini_da_syncronizzare = {ordine_id: ordine_completo}
                self.firebase.upload_ordini(ordini_da_syncronizzare, self.ordini_aggregati)
            
            print(f"✅ Ordine {ordine_id} aggiornato: {info_ordine['nome_cliente']} - {info_ordine['data_evento']}")
            print(f"   🍽️ {len(quantita_cibo)} prodotti trovati")
            
        except Exception as e:
            print(f"❌ Errore aggiornamento ordine {filepath}: {e}")

    # 🔥 AGGIUNGI QUESTA NUOVA FUNZIONE ALLA CLASSE ServerCentraleCloud:

# 🔥 AGGIUNGI import json ALL'INIZIO DEL FILE se non c'è già:
            
    def rigenera_aggregazione(self):
        """Rigenera aggregazione da ordini singoli"""
        self.ordini_aggregati = defaultdict(lambda: defaultdict(lambda: {"quantita": 0, "unita": ""}))
        
        for ordine_id, ordine in self.ordini_singoli.items():
            data_evento = ordine['info']['data_evento']
            if not data_evento:
                continue
                
            for key, info_cibo in ordine['cibo'].items():
                nome_prodotto = info_cibo['nome']
                quantita_str = info_cibo['quantita']
                
                # Estrai numero e unità
                numero_match = re.search(r'(\d+\.?\d*)', quantita_str)
                if numero_match:
                    numero = float(numero_match.group(1))
                    unita = quantita_str.replace(numero_match.group(1), '').strip()
                    
                    self.ordini_aggregati[data_evento][nome_prodotto]["quantita"] += numero
                    if not self.ordini_aggregati[data_evento][nome_prodotto]["unita"]:
                        self.ordini_aggregati[data_evento][nome_prodotto]["unita"] = unita
        
        print(f"📊 Aggregazione aggiornata per {len(self.ordini_aggregati)} date")
    
    # 🔥 SOSTITUISCI la funzione carica_ordini_esistenti nel tuo server_centrale_cloud.py

    def carica_ordini_esistenti(self):
        """Carica ordini esistenti dalla cartella (ESCLUDENDO _cancellati)"""
        if not os.path.exists(self.cartella_ordini):
            return
        
        print("📁 Caricamento ordini esistenti...")
        
        cartelle_da_controllare = [self.cartella_ordini]
        
        if hasattr(self, 'cartella_drive') and os.path.exists(self.cartella_drive):
            cartelle_da_controllare.append(self.cartella_drive)
            print(f"📂 Controllo anche: {self.cartella_drive}")
        
        for cartella in cartelle_da_controllare:
            print(f"📂 Controllo {cartella}...")
            
            try:
                for root, dirs, files in os.walk(cartella):
                    # 🔥 SALTA COMPLETAMENTE la cartella _cancellati
                    if "_cancellati" in root:
                        print(f"⏭️ SALTANDO cartella cancellati: {root}")
                        continue
                    
                    # 🔥 RIMUOVI _cancellati dalla lista delle directory da esplorare
                    if "_cancellati" in dirs:
                        dirs.remove("_cancellati")
                        print(f"⏭️ Rimossa _cancellati dalla lista directory da esplorare")
                    
                    for filename in files:
                        if filename.endswith('.docx') and not filename.startswith('~$'):
                            filepath = os.path.join(root, filename)
                            
                            try:
                                ordine_id = self.genera_id_ordine(filepath)
                                print(f"🔄 Elaborando ordine ID: {ordine_id} da {filepath}")
                                
                                info_ordine = self.estrai_info_ordine_da_filename(filename, filepath)
                                quantita_cibo = self.leggi_quantita_da_word(filepath)
                                
                                ordine_completo = {
                                    'id': ordine_id,
                                    'info': info_ordine,
                                    'cibo': quantita_cibo,
                                    'stato': 'attivo',
                                    'ultimo_aggiornamento': datetime.now().isoformat()
                                }
                                
                                self.ordini_singoli[ordine_id] = ordine_completo
                                print(f"📄 Caricato: {info_ordine['nome_cliente']} - {info_ordine['data_evento']}")
                                
                            except Exception as e:
                                print(f"❌ Errore caricamento {filepath}: {e}")
                                continue
                                
            except Exception as e:
                print(f"❌ Errore controllo cartella {cartella}: {e}")
                continue
        
        # Rigenera aggregazione
        self.rigenera_aggregazione()
        self.last_update = datetime.now()
        
        # Sync Firebase
        if self.firebase.firebase_enabled:
            print("🔥 Sync Firebase batch iniziale...")
            self.firebase.upload_ordini(self.ordini_singoli, self.ordini_aggregati)
        
        print(f"📁 Caricati {len(self.ordini_singoli)} ordini esistenti (esclusi _cancellati)")
    def avvia_file_watcher(self):
        """Avvia il monitoring dei file"""
        # Crea entrambe le cartelle se non esistono
        for cartella in [self.cartella_ordini, self.cartella_drive]:
            if not os.path.exists(cartella):
                os.makedirs(cartella)
                print(f"📁 Creata cartella: {cartella}")
        
        event_handler = WordFileHandler(self)
        observer = Observer()
        
        # Monitora ENTRAMBE le cartelle
        observer.schedule(event_handler, self.cartella_ordini, recursive=True)
        observer.schedule(event_handler, self.cartella_drive, recursive=True)
        
        print(f"👁️  Monitoring LAN: {self.cartella_ordini}")
        print(f"☁️  Monitoring Drive: {self.cartella_drive}")
        observer.start()
        
        print(f"👁️  Monitoring attivo su: {self.cartella_ordini}")
        
        # Mantieni observer in background
        def keep_observer_alive():
            try:
                while True:
                    time.sleep(1)
            except KeyboardInterrupt:
                observer.stop()
        
        thread = threading.Thread(target=keep_observer_alive, daemon=True)
        thread.start()
    
    def registra_routes(self):
        """Registra le route API Flask"""
        
        @self.app.route('/api/ordini', methods=['GET'])
        def get_lista_ordini():
            """Ritorna lista ordini singoli"""
            ordini_lista = []
            
            for ordine_id, ordine in self.ordini_singoli.items():
                info = ordine['info']
                cibo_count = len(ordine['cibo'])
                
                ordini_lista.append({
                    'id': ordine_id,
                    'nome_cliente': info['nome_cliente'],
                    'data_evento': info['data_evento'],
                    'filename': info['filename'],
                    'prodotti_count': cibo_count,
                    'stato': ordine['stato'],
                    'ultimo_aggiornamento': ordine['ultimo_aggiornamento']
                })
            
            # Ordina per data evento
            ordini_lista.sort(key=lambda x: x['data_evento'] or '9999-12-31')
            
            return jsonify({
                'ordini': ordini_lista,
                'totale_ordini': len(ordini_lista),
                'last_update': self.last_update.isoformat(),
                'firebase_status': self.firebase.firebase_enabled,
                'status': 'ok'
            })
        
        @self.app.route('/')
        def monitor_page():
            """Serve la pagina monitor.html"""
            try:
                with open('monitor.html', 'r', encoding='utf-8') as f:
                    return f.read()
            except FileNotFoundError:
                return "File monitor.html non trovato", 404

        @self.app.route('/monitor.html')
        def monitor_file():
            """Serve monitor.html come file statico"""
            try:
                with open('monitor.html', 'r', encoding='utf-8') as f:
                    return f.read()
            except FileNotFoundError:
                return "File monitor.html non trovato", 404
        
        @self.app.route('/api/ordini/<ordine_id>/dettagli', methods=['GET'])
        def get_dettagli_ordine(ordine_id):
            """Ritorna dettagli cibo per ordine specifico"""
            if ordine_id not in self.ordini_singoli:
                return jsonify({'error': 'Ordine non trovato'}), 404
            
            ordine = self.ordini_singoli[ordine_id]
            
            return jsonify({
                'ordine_id': ordine_id,
                'info': ordine['info'],
                'cibo': ordine['cibo'],
                'stato': ordine['stato'],
                'ultimo_aggiornamento': ordine['ultimo_aggiornamento']
            })
        
        @self.app.route('/api/ordini/aggregati', methods=['GET'])
        def get_ordini_aggregati():
            """Ritorna ordini aggregati"""
            return jsonify({
                'ordini': self.ordini_aggregati,
                'last_update': self.last_update.isoformat(),
                'status': 'ok'
            })
        
        @self.app.route('/api/status', methods=['GET'])
        def get_status():
            """Ritorna status del server"""
            return jsonify({
                'status': 'online',
                'ordini_singoli': len(self.ordini_singoli),
                'date_attive': list(self.ordini_aggregati.keys()),
                'last_update': self.last_update.isoformat(),
                'cartella_ordini': self.cartella_ordini,
                'firebase_enabled': self.firebase.firebase_enabled,
                'versione': '2.0-CLOUD con Firebase'
            })
        
        # ===== NUOVO: FORCE SYNC ENDPOINT =====
        @self.app.route('/api/firebase/sync', methods=['POST'])
        def force_firebase_sync():
            """Forza sincronizzazione immediata con Firebase"""
            if not self.firebase.firebase_enabled:
                return jsonify({
                    'success': False,
                    'error': 'Firebase non abilitato'
                }), 400
            
            success = self.firebase.upload_ordini(self.ordini_singoli, self.ordini_aggregati)
            
            return jsonify({
                'success': success,
                'message': 'Sync Firebase completato' if success else 'Errore sync Firebase',
                'timestamp': datetime.now().isoformat()
            })
        @self.app.route('/api/delete-orders', methods=['POST'])
        def delete_orders():
            """Cancella ordini specificati DA FIREBASE E DAI FILE FISICI"""
            try:
                data = request.get_json()
                order_ids = data.get('orderIds', [])
                
                if not order_ids:
                    return jsonify({'success': False, 'error': 'Nessun ordine specificato'}), 400
                
                deleted_count = 0
                deleted_files = []
                errors = []
                
                # 🔥 STEP 1: Cancella da Firebase PRIMA
                if self.firebase.firebase_enabled:
                    print("🔥 Cancellazione da Firebase...")
                    try:
                        firebase_admin.initialize_app()
                        db = firestore.client()
                        
                        for order_id in order_ids:
                            try:
                                doc_ref = db.collection('ordini_singoli').document(order_id)
                                doc_ref.delete()
                                print(f"🗑️ Firebase: cancellato {order_id}")
                            except Exception as e:
                                print(f"⚠️ Firebase: errore cancellazione {order_id}: {e}")
                                
                    except Exception as e:
                        print(f"⚠️ Errore connessione Firebase: {e}")
                
                # 🔥 STEP 2: Cancella da server locale E file fisici
                for order_id in order_ids:
                    try:
                        if order_id in self.ordini_singoli:
                            # Ottieni il path del file prima di cancellare
                            filepath = self.ordini_singoli[order_id]['info'].get('filepath', '')
                            
                            # Cancella dal dizionario server
                            del self.ordini_singoli[order_id]
                            deleted_count += 1
                            
                            # 🔥 CANCELLA ANCHE IL FILE FISICO
                            if filepath and os.path.exists(filepath):
                                try:
                                    # Sposta il file in una cartella "cancellati" invece di eliminarlo
                                    deleted_folder = os.path.join(os.path.dirname(filepath), "_cancellati")
                                    os.makedirs(deleted_folder, exist_ok=True)
                                    
                                    filename = os.path.basename(filepath)
                                    new_path = os.path.join(deleted_folder, filename)
                                    
                                    # Se esiste già, aggiungi timestamp
                                    if os.path.exists(new_path):
                                        name, ext = os.path.splitext(filename)
                                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                        new_path = os.path.join(deleted_folder, f"{name}_{timestamp}{ext}")
                                    
                                    shutil.move(filepath, new_path)
                                    deleted_files.append(new_path)
                                    print(f"📁 File spostato: {filepath} → {new_path}")
                                    
                                except Exception as e:
                                    print(f"⚠️ Errore spostamento file {filepath}: {e}")
                                    errors.append(f"File {filename}: {str(e)}")
                            
                            print(f"🗑️ Server: cancellato {order_id}")
                        else:
                            print(f"⚠️ Ordine non trovato nel server: {order_id}")
                            
                    except Exception as e:
                        error_msg = f"Errore cancellazione {order_id}: {str(e)}"
                        errors.append(error_msg)
                        print(f"❌ {error_msg}")
                
                # 🔥 STEP 3: Rigenera aggregazione
                if deleted_count > 0:
                    self.rigenera_aggregazione()
                    self.last_update = datetime.now()
                    print(f"📊 Aggregazione rigenerata dopo cancellazione di {deleted_count} ordini")
                
                return jsonify({
                    'success': True,
                    'deletedCount': deleted_count,
                    'deletedFiles': deleted_files,
                    'errors': errors,
                    'message': f'{deleted_count} ordini cancellati (server + file + Firebase)'
                })
                
            except Exception as e:
                print(f"❌ Errore cancellazione ordini: {e}")
                return jsonify({'success': False, 'error': str(e)}), 500
        
    def avvia_server(self, host='0.0.0.0', port=5001):
        """Avvia il server Flask"""
        print(f"🚀 Avvio Server Centrale CLOUD su {host}:{port}")
        print(f"🔥 Firebase: {'✅ ABILITATO' if self.firebase.firebase_enabled else '❌ DISABILITATO'}")
        print(f"📡 API disponibili:")
        print(f"   • GET /api/status - Status del server")
        print(f"   • GET /api/ordini - Lista ordini singoli")
        print(f"   • GET /api/ordini/<id>/dettagli - Dettagli cibo ordine")
        print(f"   • GET /api/ordini/aggregati - Tutti aggregati")
        print(f"   • POST /api/delete-orders - Cancella ordini")
        print(f"   • POST /api/firebase/sync - Forza sync Firebase")
        print(f"🌐 Monitor cucina: http://{host}:{port}")
        
        self.app.run(host=host, port=port, debug=False)


def main():
    """Funzione principale"""
    print("=" * 60)
    print("🍳 SERVER CENTRALE ORDINI CUCINA - VERSIONE CLOUD")
    print("🔥 Con sincronizzazione Firebase")
    print("=" * 60)
    
    server = ServerCentraleCloud()
    
    try:
        server.avvia_server()
    except KeyboardInterrupt:
        print("\n🛑 Server fermato dall'utente")
    except Exception as e:
        print(f"❌ Errore server: {e}")


if __name__ == "__main__":
    main()