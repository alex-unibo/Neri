#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Server Centrale con Cloud Bridge - VERSIONE FIREBASE
Sincronizza ordini con Firebase per accesso da reti diverse
"""

import os
import json
import time
import threading
import hashlib
import re
import shutil
import traceback
import logging
import sys
import contextlib
from datetime import datetime, timedelta
from collections import defaultdict
from pathlib import Path

# ===== LOGGING =====
# Funziona sia da script Python che da exe compilato
if getattr(sys, 'frozen', False):
    base_path = os.path.dirname(sys.executable)
else:
    base_path = os.path.dirname(os.path.abspath(__file__))

log_path = os.path.join(base_path, 'server_log.txt')
logging.basicConfig(
    filename=log_path,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    encoding='utf-8'   # FIX: evita crash su Windows con le emoji
)

# FIX: LogRedirect DOPO aver configurato logging
class LogRedirect:
    def write(self, msg):
        if msg.strip():
            logging.info(msg.strip())
    def flush(self):
        pass

sys.stdout = LogRedirect()
sys.stderr = LogRedirect()

# ===== FLASK =====
from flask import Flask, jsonify, request
try:
    from flask_cors import CORS
    CORS_AVAILABLE = True
except ImportError:
    CORS_AVAILABLE = False

# ===== WATCHDOG =====
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

# ===== FIREBASE =====
try:
    import firebase_admin
    from firebase_admin import credentials, firestore
    FIREBASE_AVAILABLE = True
    print("✅ Firebase disponibile")
except ImportError:
    FIREBASE_AVAILABLE = False
    print("⚠️ Firebase non disponibile - installare con: pip install firebase-admin")

# Flag globale per modalità debug (mette verbose output nel log)
DEBUG_PARSING = False


# ===========================
# FirebaseManager
# ===========================
class FirebaseManager:
    """Gestisce la sincronizzazione con Firebase"""

    def __init__(self):
        self.db = None
        self.firebase_enabled = False
        self._init_firebase()

    def _init_firebase(self):
        if not FIREBASE_AVAILABLE:
            print("❌ Firebase non disponibile")
            return

        try:
            cred_files = [
                "firebase-credentials.json",
                "serviceAccountKey.json",
                "firebase-adminsdk.json"
            ]
            cred_path = next((f for f in cred_files if os.path.exists(f)), None)

            if not cred_path:
                print("⚠️ File credenziali Firebase non trovato!")
                print("📋 Salva le credenziali come 'firebase-credentials.json'")
                return

            if not firebase_admin._apps:
                cred = credentials.Certificate(cred_path)
                firebase_admin.initialize_app(cred)

            self.db = firestore.client()
            self.firebase_enabled = True
            print(f"✅ Firebase inizializzato con {cred_path}")
            self._test_connection()

        except Exception as e:
            print(f"❌ Errore inizializzazione Firebase: {e}")
            self.firebase_enabled = False

    def _test_connection(self):
        try:
            self.db.collection('test').document('connection').set({
                'timestamp': datetime.now().isoformat(),
                'status': 'ok'
            })
            print("🔥 Test Firebase OK")
        except Exception as e:
            print(f"⚠️ Test Firebase fallito: {e}")

    def upload_ordini(self, ordini_singoli, ordini_aggregati):
        """Carica ordini su Firebase con batch chunking (max 500 doc/batch)."""
        if not self.firebase_enabled:
            return False

        try:
            items = list(ordini_singoli.items())
            chunk_size = 50

            for i in range(0, len(items), chunk_size):
                batch = self.db.batch()
                chunk = items[i:i + chunk_size]

                for ordine_id, ordine in chunk:
                    doc_ref = self.db.collection('ordini_singoli').document(ordine_id)
                    batch.set(doc_ref, {
                        'id': ordine['id'],
                        'info': ordine['info'],
                        'cibo': ordine['cibo'],
                        'stato': ordine['stato'],
                        'ultimo_aggiornamento': ordine['ultimo_aggiornamento'],
                        'cloud_sync': datetime.now().isoformat()
                    })

                # Metadata solo nell'ultimo pacchetto
                if i + chunk_size >= len(items):
                    meta_ref = self.db.collection('sistema').document('metadata')
                    batch.set(meta_ref, {
                        'totale_ordini': len(ordini_singoli),
                        'ultimo_sync': datetime.now().isoformat(),
                        'versione': '3.1'
                    })

                batch.commit()
                print(f"📦 Firebase: pacchetto {len(chunk)} ordini inviato")

            print(f"🔥 Firebase sync OK: {len(ordini_singoli)} ordini totali")
            return True

        except Exception as e:
            print(f"❌ Errore upload Firebase: {e}")
            traceback.print_exc()
            return False

    def delete_ordini(self, order_ids):
        """Cancella una lista di ordini da Firebase."""
        if not self.firebase_enabled or not order_ids:
            return
        try:
            batch = self.db.batch()
            for order_id in order_ids:
                ref = self.db.collection('ordini_singoli').document(order_id)
                batch.delete(ref)
            batch.commit()
            print(f"🗑️ Firebase: cancellati {len(order_ids)} ordini")
        except Exception as e:
            print(f"⚠️ Errore cancellazione Firebase: {e}")

    def cleanup_old_orders(self, max_hours=24):
        """No-op: la pulizia reale avviene in smart_firebase_sync_with_cleanup."""
        print("✅ cleanup_old_orders: gestito da smart_firebase_sync_with_cleanup")

    def start_auto_cleanup(self, interval_minutes=30):
        """Avvia cleanup automatico periodico (thread daemon)."""
        if not self.firebase_enabled:
            print("⚠️ Auto-cleanup disabilitato (Firebase non disponibile)")
            return

        def cleanup_worker():
            while True:
                try:
                    print(f"\n🕐 Avvio cleanup automatico...")
                    self.cleanup_old_orders(max_hours=24)
                    time.sleep(interval_minutes * 60)
                except Exception as e:
                    print(f"❌ Errore nel cleanup worker: {e}")
                    time.sleep(60)

        t = threading.Thread(target=cleanup_worker, daemon=True)
        t.start()
        print(f"✅ Auto-cleanup avviato (ogni {interval_minutes} minuti)")


# ===========================
# WordFileHandler
# ===========================
class WordFileHandler(FileSystemEventHandler):
    """Gestisce eventi di modifica sui file Word."""

    def __init__(self, server_instance):
        self.server = server_instance

    def _should_ignore(self, filepath):
        if "_cancellati" in filepath or "$Recycle.Bin" in filepath:
            return True
        filename = os.path.basename(filepath)
        return filename.startswith('~$') or not filepath.endswith('.docx')

    def _gestisci_eliminazione(self, filepath):
        ordine_id = self.server.genera_id_ordine(filepath)
        with self.server._lock:
            if ordine_id not in self.server.ordini_singoli:
                return
            del self.server.ordini_singoli[ordine_id]
            self.server._synced_timestamps.pop(ordine_id, None)

        self.server.rigenera_aggregazione()
        self.server.last_update = datetime.now()
        print(f"✅ Ordine {ordine_id} rimosso dalla memoria")

        self.server.firebase.delete_ordini([ordine_id])

    def on_modified(self, event):
        if event.is_directory or self._should_ignore(event.src_path):
            return
        print(f"📄 File modificato: {event.src_path}")
        self.server.aggiorna_ordine_da_file(event.src_path)

    def on_created(self, event):
        if event.is_directory or self._should_ignore(event.src_path):
            return
        print(f"📄 Nuovo file: {event.src_path}")
        time.sleep(1)  # Attendi che il file sia completamente scritto
        self.server.aggiorna_ordine_da_file(event.src_path)

    def on_deleted(self, event):
        if event.is_directory or self._should_ignore(event.src_path):
            return
        print(f"🗑️ File eliminato fisicamente: {event.src_path}")
        self._gestisci_eliminazione(event.src_path)

    def on_moved(self, event):
        if event.is_directory:
            return
        if "_cancellati" in event.dest_path or "Recycle.Bin" in event.dest_path:
            print(f"🗑️ File spostato nel cestino: {event.src_path}")
            self._gestisci_eliminazione(event.src_path)
        elif (("_cancellati" in event.src_path or "Recycle.Bin" in event.src_path)
              and not self._should_ignore(event.dest_path)):
            print(f"📄 File ripristinato: {event.dest_path}")
            time.sleep(1)
            self.server.aggiorna_ordine_da_file(event.dest_path)


# ===========================
# Parsing regex (centralizzato)
# ===========================
# Pattern ordinati per priorità; la prima lista ha match esclusivi.
# Ogni entry: (pattern_compilato, tipo)
#   tipo "kg_qty_name"   → gruppo 1=quantità, gruppo 2=nome
#   tipo "kg_name_qty"   → gruppo 1=nome, gruppo 2=quantità
#   tipo "paren_pz"      → gruppo 1=totale, gruppo 2=parziale, gruppo 3=nome
#   tipo "paren_kg"      → gruppo 1=totale, gruppo 2=quantità, gruppo 3=nome
#   tipo "paren_plain"   → gruppo 1=totale, gruppo 2=parziale, gruppo 3=nome
#   tipo "pz"            → gruppo 1=quantità, gruppo 2=nome
#   tipo "plain"         → gruppo 1=quantità, gruppo 2=nome

_PATTERNS = [
    # KG
    (re.compile(r'^\*{0,2}(\d+\.?\d*)\s*kg\s+(.+?)\*{0,2}$', re.I), 'kg_qty_name'),
    (re.compile(r'^\*{0,2}(.+?)\s+(\d+\.?\d*)\s*kg\*{0,2}$', re.I),  'kg_name_qty'),
    # Formato parentesi (intolleranze)
    (re.compile(r'^\((\d+)\)(\d+)\s*pz\s+(.+)$', re.I),               'paren_pz'),
    (re.compile(r'^\((\d+)\)(\d+\.?\d*)\s*kg\s+(.+)$', re.I),         'paren_kg'),
    (re.compile(r'^\((\d+)\)(\d+)\s+(.+)$', re.I),                    'paren_plain'),
    # Tradizionali
    (re.compile(r'^\*{0,2}(\d+)\s*pz\s+(.+?)\*{0,2}$', re.I),        'pz'),
    (re.compile(r'^\*{0,2}(\d+)\s+(.+?)\*{0,2}$', re.I),              'plain'),
]

_SKIP_KEYWORDS = {
    "nr progr", "all'attenzione", "presso", "aperitivo", "coffee break",
    "tea break", "lunch buffet", "persone", "allestimento", "pronti",
    "bevande", "sala", "lunedì", "martedì", "mercoledì", "giovedì",
    "venerdì", "sabato", "domenica", "gennaio", "febbraio", "marzo",
    "aprile", "maggio", "giugno", "luglio", "agosto", "settembre",
    "ottobre", "novembre", "dicembre", "totale", "pax", "euro", "€",
    "ore", "orario", "via", "telefono", "email", "fax", "intolleranti"
}

_PUNCTUATION_ONLY = {"•", "-", "*", "○", "●", "◦", "|", "+", "="}


def _parse_riga(riga_original):
    """
    Prova a fare match di una riga con i pattern definiti.
    Restituisce (nome, quantita_str) o None se non riconosciuta.
    """
    for pattern, tipo in _PATTERNS:
        m = pattern.match(riga_original)
        if not m:
            continue

        if tipo == 'kg_qty_name':
            qty, nome = float(m.group(1)), m.group(2)
            return _clean_name(nome), f"{qty} kg"

        elif tipo == 'kg_name_qty':
            nome, qty = m.group(1), float(m.group(2))
            return _clean_name(nome), f"{qty} kg"

        elif tipo == 'paren_pz':
            totale = int(m.group(1))
            nome = m.group(3)
            return _clean_name(nome), f"{totale} pezzi"

        elif tipo == 'paren_kg':
            qty = float(m.group(2))
            nome = m.group(3)
            return _clean_name(nome), f"{qty} kg"

        elif tipo == 'paren_plain':
            totale = int(m.group(1))
            nome = m.group(3)
            return _clean_name(nome), f"{totale} pezzi"

        elif tipo in ('pz', 'plain'):
            numero = int(m.group(1))
            nome = m.group(2)
            return _clean_name(nome), f"{numero} pezzi"

    return None


def _clean_name(nome):
    return " ".join(nome.replace('*', '').split())


def _accumula(prodotti, nome, quantita_str):
    """Aggiunge o somma un prodotto nel dizionario prodotti."""
    key = nome.lower().replace(' ', '_')
    if key not in prodotti:
        prodotti[key] = {'nome': nome, 'quantita': quantita_str}
        return

    prev = prodotti[key]['quantita']
    try:
        if 'kg' in quantita_str and 'kg' in prev:
            v1 = float(prev.replace(' kg', ''))
            v2 = float(quantita_str.replace(' kg', ''))
            prodotti[key]['quantita'] = f"{v1 + v2} kg"
        elif 'pezzi' in quantita_str and 'pezzi' in prev:
            v1 = int(prev.split()[0])
            v2 = int(quantita_str.split()[0])
            prodotti[key]['quantita'] = f"{v1 + v2} pezzi"
        else:
            # Unità incompatibili: sovrascrivi
            prodotti[key]['quantita'] = quantita_str
    except (ValueError, IndexError):
        prodotti[key]['quantita'] = quantita_str


# ===========================
# ServerCentraleCloud
# ===========================
class ServerCentraleCloud:
    """Server centrale con sincronizzazione Firebase."""

    def __init__(self):
        self._lock = threading.Lock()  # FIX: protezione accesso concorrente a ordini_singoli

        self.ordini_singoli = {}
        self.ordini_aggregati = {}
        self._synced_timestamps = {}

        self.cartella_ordini = "ordini_docx"
        self.cartella_drive = "ordini_drive"

        self.last_update = datetime.now()

        self.firebase = FirebaseManager()

        if self.firebase.firebase_enabled:
            print("\n🧹 Esecuzione cleanup iniziale ordini vecchi...")
            self.firebase.cleanup_old_orders(max_hours=24)
            self.firebase.start_auto_cleanup(interval_minutes=30)

        self.app = Flask(__name__)

        if CORS_AVAILABLE:
            CORS(self.app)
            print("✅ CORS abilitato")
        else:
            print("⚠️ CORS non disponibile - installare con: pip install flask-cors")

        self.registra_routes()
        self.avvia_file_watcher()
        self.carica_ordini_esistenti()
        self.avvia_sync_automatico()

    # -------- SYNC --------

    def avvia_sync_automatico(self):
        """Sync automatico INTELLIGENTE ogni 30 secondi."""
        def sync_loop():
            while True:
                try:
                    time.sleep(30)
                    if self.firebase.firebase_enabled:
                        self.smart_firebase_sync_with_cleanup()
                except Exception as e:
                    print(f"❌ Errore loop sync automatico: {e}")
                    traceback.print_exc()

        t = threading.Thread(target=sync_loop, daemon=True)
        t.start()
        print("🔄 Sync automatico Firebase avviato (ogni 30s)")

    def smart_firebase_sync_with_cleanup(self):
        """
        Sync intelligente OTTIMIZZATO - zero letture Firebase se nulla è cambiato.

        - Rimuove dalla memoria (e Firebase) ordini con evento >48h fa
        - Scrive su Firebase SOLO gli ordini nuovi o modificati
        """
        try:
            limite_scadenza = datetime.now() - timedelta(days=2)

            if not self.firebase.firebase_enabled:
                return

            # 1️⃣ Pulizia ordini scaduti
            orders_to_remove = []
            with self._lock:
                for order_id, order_data in list(self.ordini_singoli.items()):
                    data_evento_str = order_data.get('info', {}).get('data_evento')
                    if not data_evento_str:
                        continue
                    try:
                        parti = data_evento_str.split('/')
                        data_evento = datetime(int(parti[2]), int(parti[1]), int(parti[0]))
                        if data_evento < limite_scadenza:
                            orders_to_remove.append(order_id)
                            nome = order_data.get('info', {}).get('nome_cliente', order_id)
                            print(f"🧹 Ordine scaduto rimosso: {nome} ({data_evento_str})")
                    except Exception:
                        continue

                for order_id in orders_to_remove:
                    del self.ordini_singoli[order_id]
                    self._synced_timestamps.pop(order_id, None)

            if orders_to_remove:
                self.firebase.delete_ordini(orders_to_remove)
                self.rigenera_aggregazione()
                print(f"🗑️ Rimossi {len(orders_to_remove)} ordini scaduti")

            # 2️⃣ Sync selettivo: solo ordini nuovi o modificati
            to_sync = {}
            with self._lock:
                for order_id, order_data in self.ordini_singoli.items():
                    current_ts = order_data.get('ultimo_aggiornamento')
                    if current_ts != self._synced_timestamps.get(order_id):
                        to_sync[order_id] = order_data

            if to_sync:
                self.firebase.upload_ordini(to_sync, self.ordini_aggregati)
                with self._lock:
                    for order_id, order_data in to_sync.items():
                        self._synced_timestamps[order_id] = order_data.get('ultimo_aggiornamento')
                print(f"✅ Sync: {len(to_sync)} ordini aggiornati su Firebase")
            else:
                print("✅ Sync: nessuna modifica da inviare")

        except Exception as e:
            print(f"❌ Errore smart sync: {e}")
            traceback.print_exc()

    # -------- ORDINE ID --------

    def genera_id_ordine(self, filepath):
        """Genera ID ordine da nome file (MD5 troncato)."""
        filename = os.path.splitext(os.path.basename(filepath))[0]
        return hashlib.md5(filename.encode()).hexdigest()[:8]

    # -------- ESTRAZIONE INFO DA FILENAME --------

    def estrai_info_ordine_da_filename(self, filename, filepath):
        try:
            base_name = os.path.splitext(filename)[0]
            date_match = re.search(r'(\d{2}-\d{2}-\d{4})', base_name)
            if date_match:
                data_str = date_match.group(1).replace('-', '/')
                nome_cliente = base_name.replace(f"_{date_match.group(1)}", "").replace(date_match.group(1), "").strip('_')
            else:
                data_str = None
                nome_cliente = base_name

            stat = os.stat(filepath)
            return {
                'nome_cliente': nome_cliente,
                'data_evento': data_str,
                'filename': filename,
                'filepath': filepath,
                'data_creazione': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'data_modifica': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'dimensione_file': stat.st_size
            }
        except Exception as e:
            print(f"❌ Errore estrazione info da {filename}: {e}")
            return {
                'nome_cliente': os.path.splitext(filename)[0],
                'data_evento': None,
                'filename': filename,
                'filepath': filepath,
                'data_creazione': datetime.now().isoformat(),
                'data_modifica': datetime.now().isoformat(),
                'dimensione_file': 0
            }

    # -------- LETTURA WORD --------

    def leggi_quantita_da_word(self, path_docx):
        """
        Legge quantità da Word.
        Riconosce: kg, pz, (X)Y pz, (X)Y kg, numeri semplici.
        Ritorna: {chiave: {'nome': str, 'quantita': str}}
        """
        if not os.path.exists(path_docx):
            print(f"❌ File Word non trovato: {path_docx}")
            return {}

        try:
            from docx import Document
            doc = Document(path_docx)

            # Raccoglie tutto il testo (paragrafi + tabelle)
            all_text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            all_text += cell.text + "\n"

            righe = [r.strip() for r in all_text.split('\n') if r.strip()]

            prodotti = {}

            for riga in righe:
                if len(riga) < 3:
                    continue
                if riga in _PUNCTUATION_ONLY:
                    continue

                riga_lower = riga.lower()
                # Salta header/keywords, ma non se contiene "kg" (potrebbe essere un prodotto)
                if any(kw in riga_lower for kw in _SKIP_KEYWORDS) and 'kg' not in riga_lower:
                    continue

                result = _parse_riga(riga)
                if result:
                    nome, quantita = result
                    if len(nome) >= 3:
                        _accumula(prodotti, nome, quantita)
                        if DEBUG_PARSING:
                            print(f"   ✅ {nome}: {quantita}")
                else:
                    if DEBUG_PARSING and len(riga) > 5:
                        print(f"   ❓ NON RICONOSCIUTO: '{riga}'")

            print(f"🎯 {os.path.basename(path_docx)}: {len(prodotti)} prodotti trovati")
            return prodotti

        except Exception as e:
            print(f"❌ Errore lettura Word {path_docx}: {e}")
            traceback.print_exc()
            return {}

    # -------- AGGIORNA ORDINE --------

    def aggiorna_ordine_da_file(self, filepath):
        try:
            filename = os.path.basename(filepath)
            ordine_id = self.genera_id_ordine(filepath)
            info_ordine = self.estrai_info_ordine_da_filename(filename, filepath)
            quantita_cibo = self.leggi_quantita_da_word(filepath)

            ordine_completo = {
                'id': ordine_id,
                'info': info_ordine,
                'cibo': quantita_cibo,
                'stato': 'attivo',
                'ultimo_aggiornamento': datetime.now().isoformat()
            }

            with self._lock:
                self.ordini_singoli[ordine_id] = ordine_completo

            self.rigenera_aggregazione()
            self.last_update = datetime.now()

            if self.firebase.firebase_enabled:
                self.firebase.upload_ordini({ordine_id: ordine_completo}, self.ordini_aggregati)
                with self._lock:
                    self._synced_timestamps[ordine_id] = ordine_completo['ultimo_aggiornamento']

            print(f"✅ Ordine {ordine_id} aggiornato: {info_ordine['nome_cliente']} - {info_ordine['data_evento']}")

        except Exception as e:
            print(f"❌ Errore aggiornamento ordine {filepath}: {e}")
            traceback.print_exc()

    # -------- AGGREGAZIONE --------

    def rigenera_aggregazione(self):
        agg = defaultdict(lambda: defaultdict(lambda: {"quantita": 0, "unita": ""}))

        with self._lock:
            snapshot = dict(self.ordini_singoli)

        for ordine_id, ordine in snapshot.items():
            data_evento = ordine['info'].get('data_evento')
            if not data_evento:
                continue
            for key, info_cibo in ordine['cibo'].items():
                nome_prodotto = info_cibo['nome']
                quantita_str = info_cibo['quantita']
                m = re.search(r'(\d+\.?\d*)', quantita_str)
                if m:
                    numero = float(m.group(1))
                    unita = quantita_str.replace(m.group(1), '').strip()
                    agg[data_evento][nome_prodotto]["quantita"] += numero
                    if not agg[data_evento][nome_prodotto]["unita"]:
                        agg[data_evento][nome_prodotto]["unita"] = unita

        self.ordini_aggregati = agg
        print(f"📊 Aggregazione aggiornata per {len(self.ordini_aggregati)} date")

    # -------- CARICA ORDINI ESISTENTI --------

    def carica_ordini_esistenti(self):
        print("\n📁 Caricamento ordini esistenti (dal mese corrente in poi)...")

        cartelle = [self.cartella_ordini]
        if os.path.exists(self.cartella_drive):
            cartelle.append(self.cartella_drive)

        oggi = datetime.now()
        inizio_mese_corrente = datetime(oggi.year, oggi.month, 1)

        for cartella in cartelle:
            if not os.path.exists(cartella):
                continue
            try:
                for root, dirs, files in os.walk(cartella):
                    # Salta cartelle _cancellati
                    dirs[:] = [d for d in dirs if d != "_cancellati"]
                    if "_cancellati" in root:
                        continue

                    for filename in files:
                        if not filename.endswith('.docx') or filename.startswith('~$'):
                            continue
                        filepath = os.path.join(root, filename)
                        try:
                            info_ordine = self.estrai_info_ordine_da_filename(filename, filepath)

                            # Filtro temporale: ignora ordini di mesi passati
                            if info_ordine.get('data_evento'):
                                try:
                                    data_str = info_ordine['data_evento'].replace('-', '/')
                                    data_evento_obj = datetime.strptime(data_str, "%d/%m/%Y")
                                    if data_evento_obj < inizio_mese_corrente:
                                        continue
                                except ValueError:
                                    pass

                            ordine_id = self.genera_id_ordine(filepath)
                            quantita_cibo = self.leggi_quantita_da_word(filepath)

                            ordine_completo = {
                                'id': ordine_id,
                                'info': info_ordine,
                                'cibo': quantita_cibo,
                                'stato': 'attivo',
                                'ultimo_aggiornamento': datetime.now().isoformat()
                            }

                            with self._lock:
                                self.ordini_singoli[ordine_id] = ordine_completo

                            print(f"📄 Caricato: {info_ordine['nome_cliente']} - {info_ordine['data_evento']}")

                        except Exception as e:
                            print(f"❌ Errore caricamento {filepath}: {e}")

            except Exception as e:
                print(f"❌ Errore accesso cartella {cartella}: {e}")

        self.rigenera_aggregazione()
        self.last_update = datetime.now()

        if self.firebase.firebase_enabled:
            print("🔥 Sync Firebase batch iniziale...")
            with self._lock:
                snapshot = dict(self.ordini_singoli)
            self.firebase.upload_ordini(snapshot, self.ordini_aggregati)
            with self._lock:
                for order_id, order_data in self.ordini_singoli.items():
                    self._synced_timestamps[order_id] = order_data.get('ultimo_aggiornamento')
            print(f"📝 Registrati {len(self._synced_timestamps)} timestamp sync")

        with self._lock:
            count = len(self.ordini_singoli)
        print(f"✅ Caricati {count} ordini")

    # -------- FILE WATCHER --------

    def _crea_observer(self):
        """Crea e avvia un nuovo Observer su entrambe le cartelle."""
        event_handler = WordFileHandler(self)
        observer = Observer()
        observer.schedule(event_handler, self.cartella_ordini, recursive=True)
        observer.schedule(event_handler, self.cartella_drive, recursive=True)
        observer.start()
        return observer

    def avvia_file_watcher(self):
        for cartella in [self.cartella_ordini, self.cartella_drive]:
            os.makedirs(cartella, exist_ok=True)
            print(f"📁 Cartella pronta: {cartella}")

        self._observer = self._crea_observer()
        print(f"👁️  Monitoring LAN: {self.cartella_ordini}")
        print(f"☁️  Monitoring Drive: {self.cartella_drive}")

        def watcher_health_check():
            while True:
                time.sleep(60)
                try:
                    if not self._observer.is_alive():
                        print("⚠️ Watcher morto — riavvio in corso...")
                        try:
                            self._observer.stop()
                        except Exception:
                            pass
                        self._observer = self._crea_observer()
                        print("✅ Watcher riavviato correttamente")
                    elif not self._observer._watches:
                        print("⚠️ Watcher vivo ma senza watch attivi — riavvio...")
                        self._observer.stop()
                        self._observer = self._crea_observer()
                        print("✅ Watcher riavviato (watches persi)")
                except Exception as e:
                    print(f"❌ Errore health check watcher: {e}")

        t = threading.Thread(target=watcher_health_check, daemon=True)
        t.start()
        print("🔍 Watcher health check avviato (ogni 60s)")

    # -------- ROUTES --------

    def registra_routes(self):

        @self.app.route('/')
        @self.app.route('/monitor.html')
        def monitor_page():
            try:
                with open('monitor.html', 'r', encoding='utf-8') as f:
                    return f.read()
            except FileNotFoundError:
                return "File monitor.html non trovato", 404

        @self.app.route('/api/ordini', methods=['GET'])
        def get_lista_ordini():
            ordini_lista = []
            with self._lock:
                snapshot = dict(self.ordini_singoli)

            for ordine_id, ordine in snapshot.items():
                info = ordine['info']
                ordini_lista.append({
                    'id': ordine_id,
                    'nome_cliente': info['nome_cliente'],
                    'data_evento': info['data_evento'],
                    'filename': info['filename'],
                    'prodotti_count': len(ordine['cibo']),
                    'stato': ordine['stato'],
                    'ultimo_aggiornamento': ordine['ultimo_aggiornamento']
                })

            ordini_lista.sort(key=lambda x: x['data_evento'] or '9999-12-31')

            return jsonify({
                'ordini': ordini_lista,
                'totale_ordini': len(ordini_lista),
                'last_update': self.last_update.isoformat(),
                'firebase_status': self.firebase.firebase_enabled,
                'status': 'ok'
            })

        @self.app.route('/api/ordini/<ordine_id>/dettagli', methods=['GET'])
        def get_dettagli_ordine(ordine_id):
            with self._lock:
                ordine = self.ordini_singoli.get(ordine_id)
            if ordine is None:
                return jsonify({'error': 'Ordine non trovato'}), 404
            return jsonify({
                'ordine_id': ordine_id,
                'info': ordine['info'],
                'cibo': ordine['cibo'],
                'stato': ordine['stato'],
                'ultimo_aggiornamento': ordine['ultimo_aggiornamento']
            })

        @self.app.route('/api/ordini/aggregati', methods=['GET'])
        def get_ordini_aggregati():
            return jsonify({
                'ordini': self.ordini_aggregati,
                'last_update': self.last_update.isoformat(),
                'status': 'ok'
            })

        @self.app.route('/api/status', methods=['GET'])
        def get_status():
            with self._lock:
                n_ordini = len(self.ordini_singoli)
                date_attive = list(self.ordini_aggregati.keys())
            return jsonify({
                'status': 'online',
                'ordini_singoli': n_ordini,
                'date_attive': date_attive,
                'last_update': self.last_update.isoformat(),
                'cartella_ordini': self.cartella_ordini,
                'firebase_enabled': self.firebase.firebase_enabled,
                'versione': '3.1'
            })

        @self.app.route('/api/firebase/sync', methods=['POST'])
        def force_firebase_sync():
            if not self.firebase.firebase_enabled:
                return jsonify({'success': False, 'error': 'Firebase non abilitato'}), 400
            with self._lock:
                snapshot = dict(self.ordini_singoli)
            success = self.firebase.upload_ordini(snapshot, self.ordini_aggregati)
            return jsonify({
                'success': success,
                'message': 'Sync Firebase completato' if success else 'Errore sync Firebase',
                'timestamp': datetime.now().isoformat()
            })

        @self.app.route('/api/delete-orders', methods=['POST'])
        def delete_orders():
            """Cancella ordini da server locale, file fisici e Firebase."""
            try:
                data = request.get_json()
                order_ids = data.get('orderIds', [])

                if not order_ids:
                    return jsonify({'success': False, 'error': 'Nessun ordine specificato'}), 400

                deleted_count = 0
                deleted_files = []
                errors = []

                # STEP 1: Rimuovi dalla memoria locale e sposta i file fisici
                for order_id in order_ids:
                    with self._lock:
                        ordine = self.ordini_singoli.get(order_id)

                    if ordine is None:
                        print(f"⚠️ Ordine non trovato nel server: {order_id}")
                        continue

                    filepath = ordine['info'].get('filepath', '')

                    with self._lock:
                        del self.ordini_singoli[order_id]
                        self._synced_timestamps.pop(order_id, None)
                    deleted_count += 1

                    # Sposta il file fisico in _cancellati
                    if filepath and os.path.exists(filepath):
                        try:
                            deleted_folder = os.path.join(os.path.dirname(filepath), "_cancellati")
                            os.makedirs(deleted_folder, exist_ok=True)
                            filename = os.path.basename(filepath)
                            new_path = os.path.join(deleted_folder, filename)
                            if os.path.exists(new_path):
                                name, ext = os.path.splitext(filename)
                                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                                new_path = os.path.join(deleted_folder, f"{name}_{ts}{ext}")
                            shutil.move(filepath, new_path)
                            deleted_files.append(new_path)
                            print(f"📁 File spostato: {filepath} → {new_path}")
                        except Exception as e:
                            msg = f"File {os.path.basename(filepath)}: {e}"
                            errors.append(msg)
                            print(f"⚠️ Errore spostamento file: {msg}")

                    print(f"🗑️ Server: cancellato {order_id}")

                # STEP 2: Cancella da Firebase usando il client già inizializzato
                # FIX: usare self.firebase.db invece di reinizializzare firebase_admin
                if self.firebase.firebase_enabled:
                    self.firebase.delete_ordini(order_ids)

                # STEP 3: Rigenera aggregazione
                if deleted_count > 0:
                    self.rigenera_aggregazione()
                    self.last_update = datetime.now()
                    print(f"📊 Aggregazione rigenerata dopo cancellazione di {deleted_count} ordini")

                return jsonify({
                    'success': True,
                    'deletedCount': deleted_count,
                    'deletedFiles': deleted_files,
                    'errors': errors,
                    'message': f'{deleted_count} ordini cancellati'
                })

            except Exception as e:
                print(f"❌ Errore cancellazione ordini: {e}")
                traceback.print_exc()
                return jsonify({'success': False, 'error': str(e)}), 500

    # -------- AVVIO --------

    def avvia_server(self, host='0.0.0.0', port=5001):
        print(f"🚀 Avvio Server Centrale su {host}:{port}")
        print(f"🔥 Firebase: {'✅ ABILITATO' if self.firebase.firebase_enabled else '❌ DISABILITATO'}")
        print("📡 API disponibili:")
        print("   • GET  /api/status")
        print("   • GET  /api/ordini")
        print("   • GET  /api/ordini/<id>/dettagli")
        print("   • GET  /api/ordini/aggregati")
        print("   • POST /api/delete-orders")
        print("   • POST /api/firebase/sync")
        self.app.run(host=host, port=port, debug=False)


# ===========================
# Entry point
# ===========================
def main():
    print("=" * 60)
    print("🍳 SERVER CENTRALE ORDINI CUCINA")
    print("🔥 Con sincronizzazione Firebase")
    print("=" * 60)
    server = ServerCentraleCloud()
    try:
        server.avvia_server()
    except KeyboardInterrupt:
        print("\n🛑 Server fermato dall'utente")
    except Exception as e:
        print(f"❌ Errore server: {e}")
        traceback.print_exc()


if __name__ == '__main__':
    while True:
        try:
            print("\n🔄 Avvio/Riavvio del servizio principale...")
            main()
        except KeyboardInterrupt:
            print("\n🛑 Server spento manualmente.")
            break
        except Exception as e:
            print(f"\n💥 ERRORE CRITICO: {e}")
            print("⏳ Auto-ripristino tra 10 secondi...\n")
            time.sleep(10)
